#!/usr/bin/env python3
"""
Compare original and processed documents to see what was removed.
"""

import os
import sys
from docx import Document

def compare_documents():
    """Compare original and processed documents."""
    
    original_doc = "backend/test2_report_copy.docx"
    processed_doc = "backend/test2_report_copy_test.docx"
    
    if not os.path.exists(original_doc):
        print(f"❌ Original document not found: {original_doc}")
        return
        
    if not os.path.exists(processed_doc):
        print(f"❌ Processed document not found: {processed_doc}")
        return
    
    try:
        # Read both documents
        print("📖 Reading original document...")
        orig_doc = Document(original_doc)
        orig_paragraphs = [p.text.strip() for p in orig_doc.paragraphs if p.text.strip()]
        
        print("📖 Reading processed document...")
        proc_doc = Document(processed_doc)
        proc_paragraphs = [p.text.strip() for p in proc_doc.paragraphs if p.text.strip()]
        
        print(f"\n📊 Document comparison:")
        print(f"   Original: {len(orig_paragraphs)} paragraphs")
        print(f"   Processed: {len(proc_paragraphs)} paragraphs")
        print(f"   Removed: {len(orig_paragraphs) - len(proc_paragraphs)} paragraphs")
        
        # Find removed paragraphs
        removed_paragraphs = []
        for para in orig_paragraphs:
            if para not in proc_paragraphs:
                removed_paragraphs.append(para)
        
        print(f"\n🗑️ Removed paragraphs ({len(removed_paragraphs)}):")
        for i, para in enumerate(removed_paragraphs, 1):
            print(f"   {i}. '{para[:80]}{'...' if len(para) > 80 else ''}'")
        
        # Show first few paragraphs of each document
        print(f"\n📄 First 10 paragraphs of original document:")
        for i, para in enumerate(orig_paragraphs[:10], 1):
            print(f"   {i}. '{para[:60]}{'...' if len(para) > 60 else ''}'")
        
        print(f"\n📄 First 10 paragraphs of processed document:")
        for i, para in enumerate(proc_paragraphs[:10], 1):
            print(f"   {i}. '{para[:60]}{'...' if len(para) > 60 else ''}'")
            
    except Exception as e:
        print(f"❌ Error comparing documents: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    print("🔍 Comparing original and processed documents...")
    compare_documents()





