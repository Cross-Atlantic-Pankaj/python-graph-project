#!/usr/bin/env python3
"""
Compare original and aggressively cleaned documents.
"""

import os
import sys
from docx import Document

def compare_documents():
    """Compare original and aggressively cleaned documents."""
    
    original_doc = "backend/test2_report_copy.docx"
    cleaned_doc = "backend/test2_report_copy_aggressive_test.docx"
    
    if not os.path.exists(original_doc):
        print(f"❌ Original document not found: {original_doc}")
        return
        
    if not os.path.exists(cleaned_doc):
        print(f"❌ Cleaned document not found: {cleaned_doc}")
        return
    
    try:
        # Read both documents
        print("📖 Reading original document...")
        orig_doc = Document(original_doc)
        orig_paragraphs = [p.text.strip() for p in orig_doc.paragraphs if p.text.strip()]
        
        print("📖 Reading aggressively cleaned document...")
        cleaned_doc_obj = Document(cleaned_doc)
        cleaned_paragraphs = [p.text.strip() for p in cleaned_doc_obj.paragraphs if p.text.strip()]
        
        print(f"\n📊 Document comparison:")
        print(f"   Original: {len(orig_paragraphs)} paragraphs")
        print(f"   Cleaned: {len(cleaned_paragraphs)} paragraphs")
        print(f"   Removed: {len(orig_paragraphs) - len(cleaned_paragraphs)} paragraphs")
        
        print(f"\n📄 First 15 paragraphs of ORIGINAL document:")
        for i, para in enumerate(orig_paragraphs[:15], 1):
            print(f"   {i:2d}. '{para[:80]}{'...' if len(para) > 80 else ''}'")
        
        print(f"\n📄 First 15 paragraphs of CLEANED document:")
        for i, para in enumerate(cleaned_paragraphs[:15], 1):
            print(f"   {i:2d}. '{para[:80]}{'...' if len(para) > 80 else ''}'")
        
        # Check if TOC content is gone
        toc_indicators = ['table of contents', 'list of figures', 'list of tables', 'methodology', 'bnpl definitions']
        
        print(f"\n🔍 Checking for TOC indicators in cleaned document:")
        found_toc_content = False
        for indicator in toc_indicators:
            found_in_cleaned = any(indicator in para.lower() for para in cleaned_paragraphs[:20])  # Check first 20 paragraphs
            if found_in_cleaned:
                print(f"   ⚠️  '{indicator}' still found in cleaned document")
                found_toc_content = True
            else:
                print(f"   ✅ '{indicator}' successfully removed")
        
        if not found_toc_content:
            print(f"\n🎉 SUCCESS: All TOC/LOF/LOT content appears to be removed!")
        else:
            print(f"\n⚠️  Some TOC content may still remain")
            
    except Exception as e:
        print(f"❌ Error comparing documents: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    print("🔍 Comparing original and aggressively cleaned documents...")
    compare_documents()





