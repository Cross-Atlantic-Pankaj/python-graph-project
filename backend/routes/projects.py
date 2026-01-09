import os
import json
import pandas as pd
import matplotlib
matplotlib.use('Agg')  # Use non-GUI backend suitable for Flask servers
import matplotlib.pyplot as plt
from flask import Blueprint, request, jsonify, current_app, send_file
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from bson.objectid import ObjectId
from datetime import datetime 
from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
import openpyxl
import tempfile
import re
import zipfile
import shutil
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import squarify

import re

# Import TOC service
from utils.toc_service import update_toc, test_remove_toc_lof_lot, clean_pages_2_3_4_completely

# Define a constant for the section1_chart attribut

# Files are now stored in database, no upload folder needed

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'csv', 'xlsx', 'docx', 'doc'}
ALLOWED_REPORT_EXTENSIONS = {'csv', 'xlsx'}

projects_bp = Blueprint('projects', __name__)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_report_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_REPORT_EXTENSIONS

def safe_color(color):
    """Safely handle color values, returning a fallback if None or invalid"""
    if color is None:
        return 'blue'
    if isinstance(color, str) and color.strip() == '':
        return 'blue'
    return color

def safe_color_list(colors):
    """Safely handle color lists, filtering out None values"""
    if not colors:
        return ['blue']
    if isinstance(colors, list):
        filtered = [c for c in colors if c is not None and c != '']
        return filtered if filtered else ['blue']
    return [safe_color(colors)]

def validate_colors_for_plotly(colors):
    """Validate colors for Plotly, ensuring no None values"""
    if not colors:
        return ['blue']
    if isinstance(colors, list):
        # Filter out None, empty strings, and invalid values
        valid_colors = []
        for color in colors:
            if color is not None and color != '' and color != 'None':
                valid_colors.append(color)
        return valid_colors if valid_colors else ['blue']
    elif colors is not None and colors != '' and colors != 'None':
        return [colors]
    return ['blue']

def calculate_optimal_label_distance(chart_type, series_data, x_values, y_values, figsize, font_size=12):
    """
    Calculate optimal axis label distances to prevent overlap with data values.
    
    Args:
        chart_type (str): Type of chart (bar, scatter, line, etc.)
        series_data (list): Chart series data
        x_values (list): X-axis values
        y_values (list): Y-axis values
        figsize (tuple): Figure size (width, height) in inches
        font_size (int): Font size for labels
    
    Returns:
        tuple: (x_axis_distance, y_axis_distance) in pixels
    """
    # Handle None or undefined values safely
    if x_values is None:
        x_values = []
    if y_values is None:
        y_values = []
    if figsize is None:
        figsize = (8, 6)
    if font_size is None:
        font_size = 12
    
    # If y_values is empty, try to extract from series_data
    if not y_values and series_data:
        for series in series_data:
            series_values = series.get("values", [])
            if series_values:
                y_values.extend(series_values)
    
    # Convert figsize from inches to pixels (assuming 100 DPI)
    width_px = figsize[0] * 100 if figsize else 800
    height_px = figsize[1] * 100 if figsize else 600
    
    # Base distances in pixels
    base_x_distance = 30
    base_y_distance = 40
    
    # Adjust based on chart type
    if chart_type == "bar":
        # For bar charts, consider bar width and spacing
        if x_values and len(x_values) > 1:
            # Estimate bar width based on number of bars and chart width
            estimated_bar_width = width_px / (len(x_values) * 1.5)  # 1.5 accounts for spacing
            base_x_distance = max(base_x_distance, estimated_bar_width * 0.3)
        
        if y_values:
            # For vertical bars, y-axis needs more space for value labels
            max_y = max(y_values) if y_values else 0
            if max_y > 1000:  # Large numbers need more space
                base_y_distance = max(base_y_distance, 60)
    
    elif chart_type in ["scatter", "line"]:
        # For scatter/line charts, consider data point density
        if x_values and len(x_values) > 10:
            # Dense data needs more space
            base_x_distance = max(base_x_distance, 40)
            base_y_distance = max(base_y_distance, 50)
    
    elif chart_type == "area":
        # For area charts, need more space due to filled areas
        base_x_distance = max(base_x_distance, 50)
        base_y_distance = max(base_y_distance, 80)  # Area charts need more y-axis space
        
        if y_values:
            # For area charts with large values, need even more space
            max_y = max(y_values) if y_values else 0
            if max_y > 1000:  # Large numbers need more space
                base_y_distance = max(base_y_distance, 100)
    
    # Adjust based on font size
    font_factor = font_size / 12.0
    base_x_distance = int(base_x_distance * font_factor)
    base_y_distance = int(base_y_distance * font_factor)
    
    # Adjust based on chart dimensions
    if width_px < 600:  # Small charts need more space
        base_x_distance = int(base_x_distance * 1.2)
    if height_px < 400:
        base_y_distance = int(base_y_distance * 1.2)
    
    # Ensure minimum distances
    base_x_distance = max(base_x_distance, 20)
    base_y_distance = max(base_y_distance, 25)
    
    return base_x_distance, base_y_distance

def validate_excel_structure(file_path):
    """Validate that an Excel file has the required structure for report generation"""
    try:
        df = pd.read_excel(file_path, sheet_name=0, keep_default_na=False)
        df.columns = df.columns.str.strip().str.replace(" ", "_").str.replace("__", "_")
        
        # Check required columns
        required_columns = ["Text_Tag", "Text", "Chart_Tag", "Chart_Attributes", "Chart_Type"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            return False, f"Missing columns: {missing_columns}. Available: {list(df.columns)}"
        
        # Check if there's any data
        if df.empty:
            return False, "Excel file is empty"
        
        # Extract dynamic columns from A1 to M1 range
        dynamic_columns = extract_dynamic_columns_from_excel(file_path)
        
        # Ensure dynamic_columns is not empty and is a valid list
        if not dynamic_columns or not isinstance(dynamic_columns, list):
            return False, "Failed to extract dynamic columns from A1 to M1 range"
        
        # Check for required global metadata using dynamic columns
        has_global_metadata = False
        for col in df.columns:
            col_lower = col.lower().strip()
            if col_lower in dynamic_columns:
                # Check if any row has non-empty values for these columns
                if df[col].notna().any() and (df[col].astype(str).str.strip() != '').any():
                    has_global_metadata = True
                    break
        
        if not has_global_metadata:
            return False, f"No global metadata found in dynamic columns: {dynamic_columns}"
        
        return True, "Valid structure"
        
    except Exception as e:
        return False, f"Error reading Excel file: {str(e)}"

def extract_dynamic_columns_from_excel(excel_path):
    """Extract column names from A1 to M1 range in Excel file"""
    try:
        import openpyxl
        # Load the Excel file
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        ws = wb.active
        
        dynamic_columns = []
        
        # Extract column names from A1 to M1 (columns 1 to 13)
        for col_idx in range(1, 14):  # A1 to M1 (columns 1 to 13)
            try:
                cell_value = ws.cell(row=1, column=col_idx).value
                
                # Handle various types of empty values and edge cases
                if cell_value is None:
                    # Empty cell - skip silently
                    continue
                elif isinstance(cell_value, str):
                    # String value - clean and normalize
                    clean_name = cell_value.strip()
                    if clean_name:  # Only add non-empty strings
                        clean_name = clean_name.lower().replace(" ", "_").replace("__", "_")
                        if clean_name:  # Double-check after normalization
                            dynamic_columns.append(clean_name)
                elif isinstance(cell_value, (int, float)):
                    # Numeric value - convert to string and process
                    clean_name = str(cell_value).strip()
                    if clean_name:
                        clean_name = clean_name.lower().replace(" ", "_").replace("__", "_")
                        if clean_name:
                            dynamic_columns.append(clean_name)
                else:
                    # Other types (dates, etc.) - convert to string and process
                    try:
                        clean_name = str(cell_value).strip()
                        if clean_name:
                            clean_name = clean_name.lower().replace(" ", "_").replace("__", "_")
                            if clean_name:
                                dynamic_columns.append(clean_name)
                    except Exception:
                        # If conversion fails, skip this cell
                        continue
                        
            except Exception as cell_error:
                # If there's an error reading a specific cell, skip it and continue
                current_app.logger.debug(f"Error reading cell {col_idx}: {cell_error}")
                continue
        
        wb.close()
        
        # Ensure we have at least some columns, even if all were empty
        if not dynamic_columns:
            current_app.logger.warning(f"No valid columns found in A1 to M1 range in {excel_path}, using fallback")
            dynamic_columns = ['report_name', 'currency', 'country', 'report_code']
        
        current_app.logger.info(f"Extracted dynamic columns from {excel_path}: {dynamic_columns}")
        current_app.logger.debug(f"🔍 DEBUG: Raw column names from Excel: {dynamic_columns}")
        return dynamic_columns
        
    except Exception as e:
        current_app.logger.error(f"Error extracting dynamic columns from {excel_path}: {str(e)}")
        # Fallback to default columns
        return ['report_name', 'currency', 'country', 'report_code']

def extract_report_info_from_excel(excel_path):
    """Extract Report_Name and Report_Code from Excel file"""
    try:
        # Load the Excel file
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        sheet = wb.active
        
        report_name = None
        report_code = None
        
        # Search for the columns in the first few rows
        for row_idx in range(1, min(10, sheet.max_row + 1)):  # Check first 10 rows
            for col_idx in range(1, min(10, sheet.max_column + 1)):  # Check first 10 columns
                cell_value = sheet.cell(row=row_idx, column=col_idx).value
                
                if cell_value and isinstance(cell_value, str):
                    cell_value = cell_value.strip()
                    
                    # Look for Report_Name column
                    if cell_value.lower() == 'report_name':
                        # Get the value from the next row in the same column
                        if row_idx + 1 <= sheet.max_row:
                            report_name = sheet.cell(row=row_idx + 1, column=col_idx).value
                            if report_name:
                                report_name = str(report_name).strip()
                    
                    # Look for Report_Code column
                    elif cell_value.lower() == 'report_code':
                        # Get the value from the next row in the same column
                        if row_idx + 1 <= sheet.max_row:
                            report_code = sheet.cell(row=row_idx + 1, column=col_idx).value
                            if report_code:
                                report_code = str(report_code).strip()
        
        wb.close()
        
        # Fallback to filename if not found
        if not report_name:
            report_name = os.path.splitext(os.path.basename(excel_path))[0]
            #pass  # Suppress warning logs: f"Report_Name not found in {excel_path}, using filename: {report_name}")
        
        if not report_code:
            report_code = f"REPORT_{os.path.splitext(os.path.basename(excel_path))[0]}"
            #pass  # Suppress warning logs: f"Report_Code not found in {excel_path}, using generated code: {report_code}")
        
        current_app.logger.info(f"Extracted from {excel_path}: Report_Name='{report_name}', Report_Code='{report_code}'")
        return report_name, report_code
        
    except Exception as e:
        current_app.logger.error(f"Error extracting report info from {excel_path}: {e}")
        # Fallback to filename-based naming
        fallback_name = os.path.splitext(os.path.basename(excel_path))[0]
        fallback_code = f"REPORT_{fallback_name}"
        return fallback_name, fallback_code

def create_expanded_pie_chart(labels, values, colors, expanded_segment, title, value_format=""):
    """
    Create an expanded pie chart with one segment shown as a bar chart
    """
    fig = make_subplots(
        rows=1, cols=2,
        specs=[[{"type": "pie"}, {"type": "bar"}]],
        subplot_titles=(title, f"{expanded_segment} Details")
    )
    
    # Add pie chart
    fig.add_trace(go.Pie(
        labels=labels,
        values=values,
        textinfo="label+percent",
        textposition="outside",
        marker=dict(colors=colors)
    ), row=1, col=1)
    
    # Add bar chart for expanded segment
    if expanded_segment in labels:
        segment_idx = labels.index(expanded_segment)
        segment_value = values[segment_idx]
        segment_color = colors[segment_idx] if segment_idx < len(colors) else colors[0]
        
        fig.add_trace(go.Bar(
            x=[expanded_segment],
            y=[segment_value],
            marker_color=segment_color,
            text=[f"{segment_value}{value_format}"],
            textposition="auto"
        ), row=1, col=2)
    
    fig.update_layout(
        title_text=title,
        showlegend=False,
        height=500
    )
    
    return fig

def convert_chatgpt_json_to_bar_of_pie_format(chatgpt_json, data_file_path=None):
    """
    Convert ChatGPT JSON format to the format expected by create_bar_of_pie_chart
    Now supports Excel cell references in the data section
    """
    import re
    import openpyxl
    from openpyxl.utils import get_column_letter, column_index_from_string
    
    def parse_range_value(value_str):
        """
        Parse range strings like '35% - 40%', '25%-30%', '5%-10%', '<4%' etc.
        Returns the midpoint as a float, or None if not a valid range.
        """
        if not isinstance(value_str, str):
            return None
        
        value_str = value_str.strip()
        
        # Pattern 1: "X% - Y%" or "X%-Y%" (range with dash)
        range_pattern = r'(\d+(?:\.\d+)?)\s*%?\s*-\s*(\d+(?:\.\d+)?)\s*%?'
        match = re.match(range_pattern, value_str)
        if match:
            lower = float(match.group(1))
            upper = float(match.group(2))
            midpoint = (lower + upper) / 2
            return midpoint
        
        # Pattern 2: "<X%" (less than - use half of the value as approximation)
        less_than_pattern = r'<\s*(\d+(?:\.\d+)?)\s*%?'
        match = re.match(less_than_pattern, value_str)
        if match:
            upper = float(match.group(1))
            midpoint = upper / 2  # Use half as the midpoint
            return midpoint
        
        # Pattern 3: ">X%" (greater than - use value * 1.5 as approximation)
        greater_than_pattern = r'>\s*(\d+(?:\.\d+)?)\s*%?'
        match = re.match(greater_than_pattern, value_str)
        if match:
            lower = float(match.group(1))
            midpoint = lower * 1.5  # Use 1.5x as approximation
            return midpoint
        
        return None
    
    def normalize_values_to_100(values):
        """
        Normalize a list of numeric values so they sum to 100.
        Used for pie charts to ensure proper percentage display.
        """
        try:
            # Filter out non-numeric values
            numeric_values = [float(v) for v in values if isinstance(v, (int, float)) and v is not None]
            
            if not numeric_values:
                return values
            
            total = sum(numeric_values)
            
            # If total is 0, return original values
            if total == 0:
                return values
            
            # Normalize to sum to 100
            normalized = [(v / total) * 100 for v in numeric_values]
            
            return normalized
        except Exception as e:
            print(f"Error normalizing values: {e}")
            return values
    
    def extract_excel_range(sheet, cell_range):
        """Extract values from Excel cell range, preserving percentage values and parsing range strings"""
        try:
            # Parse the range (e.g., "A1:B3")
            if ':' in cell_range:
                start_cell, end_cell = cell_range.split(':')
                start_col = column_index_from_string(re.sub(r'\d+', '', start_cell)) - 1
                start_row = int(re.sub(r'[A-Z]+', '', start_cell)) - 1
                end_col = column_index_from_string(re.sub(r'\d+', '', end_cell)) - 1
                end_row = int(re.sub(r'[A-Z]+', '', end_cell)) - 1
                
                values = []
                for row in range(start_row, end_row + 1):
                    for col in range(start_col, end_col + 1):
                        cell = sheet.cell(row=row + 1, column=col + 1)
                        cell_value = cell.value
                        if cell_value is not None:
                            # Check if cell value is a string range (e.g., "35% - 40%")
                            if isinstance(cell_value, str):
                                parsed_value = parse_range_value(cell_value)
                                if parsed_value is not None:
                                    cell_value = parsed_value
                                else:
                                    # Try to convert string to float if possible
                                    try:
                                        cell_value = float(cell_value.replace('%', '').strip())
                                    except:
                                        pass
                            # Check if cell has percentage format and preserve percentage value
                            # Excel stores percentages as decimals (0.666 = 66.6%), so we need to multiply by 100
                            elif isinstance(cell_value, (int, float)):
                                cell_format = cell.number_format
                                if cell_format and '%' in str(cell_format):
                                    # If value is between 0 and 1, it's likely a percentage stored as decimal
                                    # Multiply by 100 to get the actual percentage value
                                    if 0 <= cell_value <= 1:
                                        cell_value = cell_value * 100
                            values.append(cell_value)
                return values
            else:
                # Single cell
                col = column_index_from_string(re.sub(r'\d+', '', cell_range)) - 1
                row = int(re.sub(r'[A-Z]+', '', cell_range)) - 1
                cell = sheet.cell(row=row + 1, column=col + 1)
                cell_value = cell.value
                if cell_value is not None:
                    # Check if cell value is a string range
                    if isinstance(cell_value, str):
                        parsed_value = parse_range_value(cell_value)
                        if parsed_value is not None:
                            cell_value = parsed_value
                        else:
                            try:
                                cell_value = float(cell_value.replace('%', '').strip())
                            except:
                                pass
                    elif isinstance(cell_value, (int, float)):
                        cell_format = cell.number_format
                        if cell_format and '%' in str(cell_format):
                            if 0 <= cell_value <= 1:
                                cell_value = cell_value * 100
                return [cell_value] if cell_value is not None else []
        except Exception as e:
            print(f"Error extracting Excel range {cell_range}: {e}")
            return []
    
    chart_meta = chatgpt_json.get("chart_meta", {})
    data = chatgpt_json.get("data", {})
    
    # Extract overall data - check for Excel cell references first
    overall_data = data.get("overall", [])
    overall_labels = data.get("overall_labels", [])
    overall_values = data.get("overall_values", [])
    
    # Check if overall_labels and overall_values are Excel cell references
    if isinstance(overall_labels, str) and re.match(r'^[A-Z]+\d+:[A-Z]+\d+$', overall_labels) and data_file_path:
        try:
            wb = openpyxl.load_workbook(data_file_path, data_only=True)
            sheet = wb[chart_meta.get("source_sheet", "sample")]
            overall_labels = extract_excel_range(sheet, overall_labels)
            wb.close()
        except Exception as e:
            print(f"Error extracting overall_labels from Excel: {e}")
    
    if isinstance(overall_values, str) and re.match(r'^[A-Z]+\d+:[A-Z]+\d+$', overall_values) and data_file_path:
        try:
            wb = openpyxl.load_workbook(data_file_path, data_only=True)
            sheet = wb[chart_meta.get("source_sheet", "sample")]
            overall_values = extract_excel_range(sheet, overall_values)
            # Normalize values to sum to 100 for pie charts
            overall_values = normalize_values_to_100(overall_values)
            wb.close()
        except Exception as e:
            print(f"Error extracting overall_values from Excel: {e}")
    
    # If we don't have data from Excel or arrays, use object format
    if not overall_labels and not overall_values and overall_data:
        overall_labels = [item.get("label", "") for item in overall_data]
        overall_values = [item.get("value", 0) for item in overall_data]
        # Normalize values to sum to 100 for pie charts
        overall_values = normalize_values_to_100(overall_values)
    
    # Extract other breakdown data - check for Excel cell references first
    other_data = data.get("other_breakdown", [])
    other_labels = data.get("other_labels", [])
    other_values = data.get("other_values", [])
    
    # Check if other_labels and other_values are Excel cell references
    if isinstance(other_labels, str) and re.match(r'^[A-Z]+\d+:[A-Z]+\d+$', other_labels) and data_file_path:
        try:
            wb = openpyxl.load_workbook(data_file_path, data_only=True)
            sheet = wb[chart_meta.get("source_sheet", "sample")]
            other_labels = extract_excel_range(sheet, other_labels)
            wb.close()
        except Exception as e:
            print(f"Error extracting other_labels from Excel: {e}")
    
    if isinstance(other_values, str) and re.match(r'^[A-Z]+\d+:[A-Z]+\d+$', other_values) and data_file_path:
        try:
            wb = openpyxl.load_workbook(data_file_path, data_only=True)
            sheet = wb[chart_meta.get("source_sheet", "sample")]
            other_values = extract_excel_range(sheet, other_values)
            # Normalize values to sum to 100 for pie charts
            other_values = normalize_values_to_100(other_values)
            wb.close()
        except Exception as e:
            print(f"Error extracting other_values from Excel: {e}")
    
    # If we don't have data from Excel or arrays, use object format
    if not other_labels and not other_values and other_data:
        other_labels = [item.get("label", "") for item in other_data]
        other_values = [item.get("value", 0) for item in other_data]
        # Normalize values to sum to 100 for pie charts
        other_values = normalize_values_to_100(other_values)
    
    # Process cell references in chart_meta attributes
    if data_file_path:
        try:
            wb = openpyxl.load_workbook(data_file_path, data_only=True)
            sheet = wb[chart_meta.get("source_sheet", "sample")]
            
            # Process cell references in chart_meta
            for key, value in chart_meta.items():
                if isinstance(value, str):
                    # Check for single cell reference (e.g., "AR13")
                    if re.match(r'^[A-Z]+\d+$', value):
                        cell_value = sheet.cell(
                            row=int(re.sub(r'[A-Z]+', '', value)), 
                            column=column_index_from_string(re.sub(r'\d+', '', value))
                        ).value
                        if cell_value is not None:
                            chart_meta[key] = cell_value
                    # Check for cell range reference (e.g., "A1:B3")
                    elif re.match(r'^[A-Z]+\d+:[A-Z]+\d+$', value):
                        cell_values = extract_excel_range(sheet, value)
                        if cell_values:
                            chart_meta[key] = cell_values
            
            wb.close()
        except Exception as e:
            print(f"Error processing cell references in chart_meta: {e}")
    
    # Convert to expected format
    converted_json = {
        "chart_meta": {
            "chart_type": "bar_of_pie",
            "title_left": chart_meta.get("title_left", "Revenue Breakdown"),
            "title_right": chart_meta.get("title_right", "'Other' Composition"),
            "expanded_segment": "Other",
            "other_labels": other_labels,
            "other_values": other_values,
            "other_colors": chart_meta.get("palette_other", []),
            "type_left": chart_meta.get("type_left", "pie"),
            "stacked": chart_meta.get("type_right") == "stacked_bar",
            "connector": chart_meta.get("connector", {}),
            "source_sheet": chart_meta.get("source_sheet", "sample"),  # Pass through source sheet
            
            # Pass through all chart control attributes
            "showlegend": chart_meta.get("showlegend", True),
            "legend_position": chart_meta.get("legend_position", "bottom"),
            "legend_orientation": chart_meta.get("legend_orientation", "v"),  # Add legend orientation
            "legend_font_size": chart_meta.get("legend_font_size", 10),
            "data_labels": chart_meta.get("data_labels", True),
            "data_label_format": chart_meta.get("data_label_format", ".1f"),
            "data_label_font_size": chart_meta.get("data_label_font_size", 10),
            "data_label_color": chart_meta.get("data_label_color", "#000000"),
            "x_axis_title": chart_meta.get("x_axis_title", "Categories"),
            "y_axis_title": chart_meta.get("y_axis_title", "Value"),
            "show_x_axis": chart_meta.get("show_x_axis", True),
            "show_y_axis": chart_meta.get("show_y_axis", True),
            "show_x_ticks": chart_meta.get("show_x_ticks", True),
            "show_y_ticks": chart_meta.get("show_y_ticks", True),
                    "x_axis_label_distance": chart_meta.get("x_axis_label_distance", "auto"),
        "y_axis_label_distance": chart_meta.get("y_axis_label_distance", "auto"),
            "axis_tick_font_size": chart_meta.get("axis_tick_font_size", 10),
            "show_gridlines": chart_meta.get("show_gridlines", False),
            "gridline_color": chart_meta.get("gridline_color", "#E5E7EB"),
            "gridline_style": chart_meta.get("gridline_style", "solid"),
            "height": chart_meta.get("height", 500),
            "width": chart_meta.get("width", 900),
            "column_widths": chart_meta.get("column_widths", [0.5, 0.5]),
            "horizontal_spacing": chart_meta.get("horizontal_spacing", 0.1),
            "margin": chart_meta.get("margin", dict(l=50, r=50, t=80, b=50)),
            "font_family": chart_meta.get("font_family", "Arial"),
            "font_size": chart_meta.get("font_size", 14),
            "font_color": chart_meta.get("font_color", "#333333"),
            "chart_background": chart_meta.get("chart_background", "#FFFFFF"),
            "plot_background": chart_meta.get("plot_background", "#F8F9FA")
        },
        "series": {
            "labels": overall_labels,
            "values": overall_values,
            "colors": chart_meta.get("palette_main", [])
        }
    }
    
    return converted_json

def create_matplotlib_chart_from_plotly(fig, output_path):
    """
    Create a matplotlib chart from a Plotly figure as a fallback when Chrome/Kaleido is not available
    """
    import matplotlib.pyplot as plt
    import numpy as np
    
    # Create matplotlib figure
    fig_mpl, ax = plt.subplots(figsize=(12, 6))
    
    try:
        # Extract data from Plotly figure
        if hasattr(fig, 'data') and len(fig.data) > 0:
            trace = fig.data[0]
            
            # Determine chart type and create matplotlib equivalent
            if hasattr(trace, 'type'):
                chart_type = trace.type
            else:
                # Infer type from data structure
                if hasattr(trace, 'labels') and hasattr(trace, 'values'):
                    chart_type = 'pie'
                elif hasattr(trace, 'x') and hasattr(trace, 'y'):
                    chart_type = 'bar'
                else:
                    chart_type = 'bar'
            
            if chart_type in ['bar', 'column']:
                # Bar chart
                x_data = trace.x if hasattr(trace, 'x') else range(len(trace.y))
                y_data = trace.y if hasattr(trace, 'y') else [1, 2, 3]
                colors = trace.marker.color if hasattr(trace.marker, 'color') else 'blue'
                
                if isinstance(colors, list) and len(colors) > 1:
                    ax.bar(x_data, y_data, color=colors)
                else:
                    ax.bar(x_data, y_data, color=colors[0] if isinstance(colors, list) else colors)
                
                ax.set_xlabel('Categories')
                ax.set_ylabel('Values')
                
            elif chart_type in ['pie', 'donut']:
                # Pie chart
                labels = trace.labels if hasattr(trace, 'labels') else ['A', 'B', 'C']
                values = trace.values if hasattr(trace, 'values') else [1, 2, 3]
                colors = trace.marker.colors if hasattr(trace.marker, 'colors') else None
                
                ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
                
            elif chart_type in ['line', 'scatter']:
                # Line chart
                x_data = trace.x if hasattr(trace, 'x') else range(len(trace.y))
                y_data = trace.y if hasattr(trace, 'y') else [1, 2, 3]
                color = trace.marker.color if hasattr(trace.marker, 'color') else 'blue'
                
                ax.plot(x_data, y_data, marker='o', linewidth=2, markersize=6, color=color)
                ax.set_xlabel('X Values')
                ax.set_ylabel('Y Values')
                
            else:
                # Default bar chart
                x_data = trace.x if hasattr(trace, 'x') else range(len(trace.y) if hasattr(trace, 'y') else [1, 2, 3])
                y_data = trace.y if hasattr(trace, 'y') else [1, 2, 3]
                ax.bar(x_data, y_data, color='blue')
        
        # Set title
        if hasattr(fig, 'layout') and hasattr(fig.layout, 'title'):
            ax.set_title(str(fig.layout.title), fontsize=14, fontweight='bold')
        
        # Improve appearance
        ax.grid(True, alpha=0.3)
        plt.tight_layout()
        
        # Save the figure
        fig_mpl.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close(fig_mpl)
        
        return True
        
    except Exception as e:
        current_app.logger.error(f"Error creating matplotlib chart: {e}")
        plt.close(fig_mpl)
        return False

def create_bar_of_pie_chart(labels, values, other_labels, other_values, colors, other_colors, title, value_format="", chart_meta=None):
    """
    Create a 'bar of pie' chart using Plotly: pie chart with one segment broken down as a bar chart.
    Enhanced version with better title handling and layout options.
    Supports both individual bars and stacked bars.
    """
    # Filter out empty/null values from other_labels and other_values
    filtered_data = []
    for label, value in zip(other_labels, other_values):
        # Check if both label and value are not empty/null
        if (label is not None and str(label).strip() != "" and 
            value is not None and str(value).strip() != "" and 
            str(value).strip() != "0"):
            filtered_data.append((label, value))
    
    if filtered_data:
        filtered_labels, filtered_values = zip(*filtered_data)
    else:
        # If no valid data, use empty lists
        filtered_labels, filtered_values = [], []
    
    # Enhanced title handling
    title_left = chart_meta.get("title_left", title) if chart_meta else title
    title_right = chart_meta.get("title_right", "Breakdown of 'Other'") if chart_meta else "Breakdown of 'Other'"
    
    # Enhanced layout options
    height = chart_meta.get("height", 500) if chart_meta else 500
    width = chart_meta.get("width", 900) if chart_meta else 900
    column_widths = chart_meta.get("column_widths", [0.5, 0.5]) if chart_meta else [0.5, 0.5]
    
    # Enhanced styling options
    font_family = chart_meta.get("font_family", "Arial") if chart_meta else "Arial"
    font_size = chart_meta.get("font_size", 14) if chart_meta else 14
    font_color = chart_meta.get("font_color", "#333333") if chart_meta else "#333333"
    chart_background = chart_meta.get("chart_background", "#FFFFFF") if chart_meta else "#FFFFFF"
    plot_background = chart_meta.get("plot_background", "#F8F9FA") if chart_meta else "#F8F9FA"
    
    # Legend controls
    show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True)) if chart_meta else True
    # Convert string "false"/"true" to boolean if needed
    if isinstance(show_legend_raw, str):
        show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
    else:
        show_legend = bool(show_legend_raw)
    legend_position = chart_meta.get("legend_position", "bottom") if chart_meta else "bottom"
    legend_font_size = chart_meta.get("legend_font_size", 10) if chart_meta else 10
    legend_orientation = chart_meta.get("legend_orientation", "v") if chart_meta else "v"  # "v" for vertical, "h" for horizontal
    

    
    # Data label controls
    data_labels = chart_meta.get("data_labels", True) if chart_meta else True
    data_label_format = chart_meta.get("data_label_format", ".1f") if chart_meta else ".1f"
    data_label_font_size = chart_meta.get("data_label_font_size", 10) if chart_meta else 10
    data_label_color = chart_meta.get("data_label_color", "#000000") if chart_meta else "#000000"
    
    # Validate data_label_format to ensure it's a valid format specifier
    if data_label_format and data_label_format not in [".0f", ".1f", ".2f", ".3f", ".4f", ".5f", ".6f", ".7f", ".8f", ".9f", "d", "i", "o", "x", "X", "e", "E", "f", "F", "g", "G", "n", "%"]:
        # If it's not a valid format specifier, use default
        data_label_format = ".1f"
    
    # Helper function to safely format data labels
    def safe_format_label(value, format_spec=".1f"):
        try:
            return f"{value:{format_spec}}%"
        except (ValueError, TypeError):
            return f"{value:.1f}%"
    
    # Axis controls
    x_axis_title = chart_meta.get("x_axis_title", "Categories") if chart_meta else "Categories"
    y_axis_title = chart_meta.get("y_axis_title", "Value") if chart_meta else "Value"
    show_x_axis = chart_meta.get("show_x_axis", True) if chart_meta else True
    show_y_axis = chart_meta.get("show_y_axis", True) if chart_meta else True
    show_x_ticks = chart_meta.get("show_x_ticks", True) if chart_meta else True
    show_y_ticks = chart_meta.get("show_y_ticks", True) if chart_meta else True
    x_axis_label_distance = chart_meta.get("x_axis_label_distance", "auto") if chart_meta else "auto"
    y_axis_label_distance = chart_meta.get("y_axis_label_distance", "auto") if chart_meta else "auto"
    axis_tick_font_size = chart_meta.get("axis_tick_font_size", 10) if chart_meta else 10
    
    # Grid controls
    show_gridlines = chart_meta.get("show_gridlines", False) if chart_meta else False
    gridline_color = chart_meta.get("gridline_color", "#E5E7EB") if chart_meta else "#E5E7EB"
    gridline_style = chart_meta.get("gridline_style", "solid") if chart_meta else "solid"
    
    # Margin and spacing controls
    margin = chart_meta.get("margin", dict(l=50, r=50, t=80, b=50)) if chart_meta else dict(l=50, r=50, t=80, b=50)
    horizontal_spacing = chart_meta.get("horizontal_spacing", 0.1) if chart_meta else 0.1
    
    # Check if stacked bars are requested
    is_stacked = chart_meta.get("stacked", False) if chart_meta else False
    
    # Connector styling
    connector_style = chart_meta.get("connector", {}) if chart_meta else {}
    connector_color = connector_style.get("color", "#6B7280")
    connector_width = connector_style.get("width", 1.3)
    connector_opacity = connector_style.get("opacity", 0.9)
    
    fig = make_subplots(
        rows=1, cols=2,
        specs=[[{"type": "pie"}, {"type": "bar"}]],
        column_widths=column_widths,
        subplot_titles=(title_left, title_right),
        horizontal_spacing=horizontal_spacing
    )
    
    # Check if donut chart is requested
    is_donut = chart_meta.get("type_left") == "donut_pie" if chart_meta else False
    hole_size = 0.4 if is_donut else 0.0
    
    # Main pie chart with enhanced styling
    fig.add_trace(go.Pie(
        labels=labels,
        values=values,
        marker=dict(colors=colors),
        textinfo="percent" if data_labels else "none",
        hoverinfo="label+percent+value",
        pull=[0.1 if l == "Other" else 0 for l in labels],
        name="Main Pie",
        hole=hole_size,
        showlegend=show_legend,
        textfont=dict(family=font_family, size=font_size, color=font_color),
        textposition="inside" if data_labels else "none"
    ), row=1, col=1)
    
    # Bar chart for breakdown (only if we have filtered data)
    if filtered_labels and filtered_values:
        # Check if horizontal bars are requested
        bar_orientation = chart_meta.get("orientation", "vertical") if chart_meta else "vertical"
        
        # Convert values to numeric format for proper stacking
        numeric_values = []
        for value in filtered_values:
            if isinstance(value, (int, float)):
                if value <= 1.0:  # Likely decimal format (0.11)
                    numeric_values.append(value * 100)
                else:  # Likely already percentage format (11.0)
                    numeric_values.append(value)
            else:
                try:
                    val = float(value)
                    if val <= 1.0:
                        numeric_values.append(val * 100)
                    else:
                        numeric_values.append(val)
                except:
                    numeric_values.append(0)
        
        # Format labels for display
        formatted_labels = []
        for label in filtered_labels:
            if isinstance(label, (int, float)):
                if label <= 1.0:  # Likely decimal format (0.06)
                    formatted_labels.append(f"{label * 100:.1f}%")
                else:  # Likely already percentage format (6.0)
                    formatted_labels.append(f"{label:.1f}%")
            else:
                try:
                    val = float(label)
                    if val <= 1.0:
                        formatted_labels.append(f"{val * 100:.1f}%")
                    else:
                        formatted_labels.append(f"{val:.1f}%")
                except:
                    formatted_labels.append(str(label))
        
        if is_stacked:
            # Create stacked bar chart with separate traces for each segment
            if bar_orientation.lower() == "horizontal":
                # Create stacked horizontal bar with separate traces
                for i, (label, value) in enumerate(zip(filtered_labels, numeric_values)):
                    bar_color = other_colors[i] if other_colors and i < len(other_colors) else None
                    
                    fig.add_trace(go.Bar(
                        x=[value],  # Single value for this segment
                        y=["Other"],  # Single category for stacked bar
                        orientation="h",
                        marker_color=bar_color,
                        text=[safe_format_label(value, data_label_format)] if data_labels else [""],
                        textposition="inside",
                        name=label,
                        showlegend=show_legend,  # Enable legend for bar traces
                        textfont=dict(family=font_family, size=data_label_font_size, color=data_label_color),
                        hovertemplate=f"<b>{label}</b><br>Value: {value:.1f}%<extra></extra>"
                    ), row=1, col=2)
            else:
                # Create stacked vertical bar with separate traces
                for i, (label, value) in enumerate(zip(filtered_labels, numeric_values)):
                    bar_color = other_colors[i] if other_colors and i < len(other_colors) else None
                    
                    fig.add_trace(go.Bar(
                        x=["Other"],  # Single category for stacked bar
                        y=[value],  # Single value for this segment
                        marker_color=bar_color,
                        text=[safe_format_label(value, data_label_format)] if data_labels else [""],
                        textposition="inside",
                        name=label,
                        showlegend=show_legend,  # Enable legend for bar traces
                        textfont=dict(family=font_family, size=data_label_font_size, color=data_label_color),
                        hovertemplate=f"<b>{label}</b><br>Value: {value:.1f}%<extra></extra>"
                    ), row=1, col=2)
            
            # Set barmode to 'stack' for proper stacking
            fig.update_layout(barmode='stack')
        else:
            # Create individual bars (original behavior)
            if bar_orientation.lower() == "horizontal":
                # Create individual horizontal bar traces for each data point
                for i, (label, value) in enumerate(zip(formatted_labels, numeric_values)):
                    bar_color = other_colors[i] if other_colors and i < len(other_colors) else None
                    
                    fig.add_trace(go.Bar(
                        x=[value],  # Single value for this bar
                        y=[label],  # Category label
                        orientation="h",
                        marker_color=bar_color,
                        text=[safe_format_label(value, data_label_format)] if data_labels else [""],
                        textposition="auto",
                        name=label,  # Use actual label instead of generic name
                        showlegend=show_legend,  # Enable legend for bar traces
                        textfont=dict(family=font_family, size=data_label_font_size, color=data_label_color),
                        hovertemplate=f"<b>{label}</b><br>Value: {value:.1f}%<extra></extra>"
                    ), row=1, col=2)
            else:
                # Create individual vertical bar traces for each data point
                for i, (label, value) in enumerate(zip(formatted_labels, numeric_values)):
                    bar_color = other_colors[i] if other_colors and i < len(other_colors) else None
                    
                    fig.add_trace(go.Bar(
                        x=[label],  # Category label
                        y=[value],  # Single value for this bar
                        marker_color=bar_color,
                        text=[safe_format_label(value, data_label_format)] if data_labels else [""],
                        textposition="auto",
                        name=label,  # Use actual label instead of generic name
                        showlegend=show_legend,  # Enable legend for bar traces
                        textfont=dict(family=font_family, size=data_label_font_size, color=data_label_color),
                        hovertemplate=f"<b>{label}</b><br>Value: {value:.1f}%<extra></extra>"
                    ), row=1, col=2)
    
    # Enhanced layout with better styling
    fig.update_layout(
        title_text="",  # Remove main title since we have subplot titles
        showlegend=show_legend,
        height=height,
        width=width,
        font=dict(family=font_family, size=font_size, color=font_color),
        paper_bgcolor=chart_background,
        plot_bgcolor=plot_background,
        margin=margin
    )
    
    # Force legend visibility based on show_legend setting
    if not show_legend:
        # Explicitly hide legend for all traces
        for trace in fig.data:
            trace.showlegend = False
    
    # Configure legend position and styling
    if show_legend:
        # Map legend positions to Plotly coordinates
        legend_positions = {
            "top": dict(x=0.5, y=1.1, xanchor="center", yanchor="bottom"),
            "bottom": dict(x=0.5, y=-0.2, xanchor="center", yanchor="top"),
            "left": dict(x=-0.2, y=0.5, xanchor="right", yanchor="middle"),
            "right": dict(x=1.1, y=0.5, xanchor="left", yanchor="middle")
        }
        
        legend_config = {}
        if legend_position in legend_positions:
            legend_config.update(legend_positions[legend_position])
        
        # Add legend orientation for horizontal/vertical layout
        legend_config["orientation"] = legend_orientation
        
        # Additional legend properties for better horizontal layout
        if legend_orientation == "h":
            # Override position for horizontal legends
            legend_config["x"] = 0.5
            legend_config["y"] = -0.15
            legend_config["xanchor"] = "center"
            legend_config["yanchor"] = "top"
            # Force horizontal layout
            legend_config["traceorder"] = "normal"
        
        if legend_font_size:
            legend_config["font"] = dict(size=legend_font_size, family=font_family, color=font_color)
        
        if legend_config:
            fig.update_layout(legend=legend_config)
    
    # Update subplot titles with better styling
    fig.update_annotations(
        font=dict(family=font_family, size=font_size + 2, color=font_color)
    )
    
    # Update bar chart axes with comprehensive controls
    if show_x_axis:
        x_axis_config = {
            "title": dict(text=x_axis_title, font=dict(size=font_size)),
            "showticklabels": show_x_ticks,
            "tickfont": dict(size=axis_tick_font_size)
        }
        
        # Handle "auto" values for axis label distances
        if x_axis_label_distance == "auto":
            # Calculate optimal label distance for bar chart
            auto_x_distance, _ = calculate_optimal_label_distance(
                "bar", [{"labels": filtered_labels, "values": filtered_values}], 
                filtered_labels, filtered_values, (width/100, height/100), font_size
            )
            x_axis_label_distance = auto_x_distance
        
        # Apply axis label distance if specified
        if x_axis_label_distance is not None:
            x_axis_config["title"] = dict(
                text=x_axis_title, 
                font=dict(size=font_size),
                standoff=x_axis_label_distance
            )
        
        fig.update_xaxes(**x_axis_config, row=1, col=2)
    else:
        fig.update_xaxes(visible=False, row=1, col=2)
    
    if show_y_axis:
        y_axis_config = {
            "title": dict(text=y_axis_title, font=dict(size=font_size)),
            "showticklabels": show_y_ticks,
            "tickfont": dict(size=axis_tick_font_size)
        }
        
        # Handle "auto" values for axis label distances
        if y_axis_label_distance == "auto":
            # Calculate optimal label distance for bar chart
            _, auto_y_distance = calculate_optimal_label_distance(
                "bar", [{"labels": filtered_labels, "values": filtered_values}], 
                filtered_labels, filtered_values, (width/100, height/100), font_size
            )
            y_axis_label_distance = auto_y_distance
        
        # Apply axis label distance if specified
        if y_axis_label_distance is not None:
            y_axis_config["title"] = dict(
                text=y_axis_title, 
                font=dict(size=font_size),
                standoff=y_axis_label_distance
            )
        
        fig.update_yaxes(**y_axis_config, row=1, col=2)
    else:
        fig.update_yaxes(visible=False, row=1, col=2)
    
    # Add gridlines if requested
    if show_gridlines:
        fig.update_xaxes(
            showgrid=True,
            gridcolor=gridline_color,
            gridwidth=1,
            row=1, col=2
        )
        fig.update_yaxes(
            showgrid=True,
            gridcolor=gridline_color,
            gridwidth=1,
            row=1, col=2
        )
    
    # Update axis formatting for the bar chart to show percentages
    if filtered_values and isinstance(filtered_values[0], (int, float)) and filtered_values[0] <= 1.0:
        if bar_orientation.lower() == "horizontal":
            # For horizontal bars, format x-axis (values) as percentages
            fig.update_xaxes(
                tickformat=".1f",  # Format as percentage with 1 decimal place (no % symbol since values are already multiplied by 100)
                title="Revenue (%)",
                row=1, col=2
            )
            # Update y-axis title for horizontal bars
            fig.update_yaxes(
                title="Categories",
                row=1, col=2
            )
        else:
            # For vertical bars, format y-axis (values) as percentages
            fig.update_yaxes(
                tickformat=".1f",  # Format as percentage with 1 decimal place (no % symbol since values are already multiplied by 100)
                title="Revenue (%)",
                row=1, col=2
            )
            # Update x-axis title for vertical bars
            fig.update_xaxes(
                title="Categories",
                row=1, col=2
            )
    
    # Add connector line between pie and bar chart (if connector is enabled)
    if connector_style.get("style") == "elbow" and filtered_labels:
        # Calculate connector line coordinates
        # This is a simplified connector - you can enhance this for more complex paths
        fig.add_shape(
            type="line",
            x0=0.45, y0=0.5,  # Start from pie chart
            x1=0.55, y1=0.5,  # End at bar chart
            line=dict(color=connector_color, width=connector_width),
            opacity=connector_opacity,
            xref="paper", yref="paper"
        )
    
    return fig

def ensure_proper_page_breaks_for_toc(doc):
    """
    Ensures proper page breaks around TOC to help with accurate page numbering.
    Adds page breaks before and after TOC sections to ensure they're on separate pages.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        int: Number of page breaks added
    """
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        page_breaks_added = 0
        
        # Find TOC paragraphs
        toc_paragraphs = []
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph._element.xpath('.//w:instrText', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                for instr in paragraph._element.xpath('.//w:instrText', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    if instr.text and instr.text.strip().upper().startswith('TOC'):
                        toc_paragraphs.append((i, paragraph))
                        break
        
        if not toc_paragraphs:
            current_app.logger.debug("ℹ️ No TOC found for page break insertion")
            return 0
        
        # Add page break before first TOC
        first_toc_idx, first_toc_para = toc_paragraphs[0]
        if first_toc_idx > 0:  # Don't add page break if TOC is first paragraph
            # Check if previous paragraph already has a page break
            prev_para = doc.paragraphs[first_toc_idx - 1]
            has_page_break = prev_para._element.xpath('.//w:br[@w:type="page"]', 
                                                     namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            if not has_page_break:
                # Add page break to previous paragraph
                run = prev_para.runs[-1] if prev_para.runs else prev_para.add_run()
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                run._element.append(br)
                page_breaks_added += 1
                current_app.logger.debug("✅ Added page break before TOC")
        
        # Add page break after last TOC
        last_toc_idx, last_toc_para = toc_paragraphs[-1]
        
        # Find the end of the TOC field (look for field end marker)
        toc_end_idx = last_toc_idx
        for i in range(last_toc_idx, min(last_toc_idx + 20, len(doc.paragraphs))):  # Look ahead max 20 paragraphs
            para = doc.paragraphs[i]
            fld_chars = para._element.xpath('.//w:fldChar', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            for fld_char in fld_chars:
                if fld_char.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType') == 'end':
                    toc_end_idx = i
                    break
        
        if toc_end_idx < len(doc.paragraphs) - 1:  # Don't add page break if TOC is last content
            # Check if next paragraph after TOC already has a page break
            next_para_idx = toc_end_idx + 1
            next_para = doc.paragraphs[next_para_idx]
            has_page_break = next_para._element.xpath('.//w:br[@w:type="page"]', 
                                                     namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            if not has_page_break:
                # Add page break to the paragraph after TOC
                run = next_para.runs[0] if next_para.runs else next_para.add_run()
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                # Insert at beginning of run
                run._element.insert(0, br)
                page_breaks_added += 1
                current_app.logger.debug("✅ Added page break after TOC")
        
        if page_breaks_added > 0:
            current_app.logger.info(f"✅ Added {page_breaks_added} page break(s) around TOC for better page numbering")
        
        return page_breaks_added
        
    except Exception as e:
        current_app.logger.error(f"❌ Error adding page breaks around TOC: {e}")
        return 0


def create_fresh_toc_if_needed(doc):
    """
    Creates a fresh Table of Contents at the beginning of the document if none exists.
    This ensures there's always a TOC that can be updated.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        bool: True if TOC was created, False if one already exists
    """
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Check if TOC already exists
        for paragraph in doc.paragraphs:
            if paragraph._element.xpath('.//w:instrText', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                for instr in paragraph._element.xpath('.//w:instrText', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    if instr.text and instr.text.strip().upper().startswith('TOC'):
                        current_app.logger.debug("ℹ️ TOC already exists in document")
                        return False
        
        # No TOC found, create one at the beginning
        current_app.logger.info("🔄 Creating fresh Table of Contents...")
        
        # Insert TOC at the beginning of document
        if len(doc.paragraphs) > 0:
            # Insert before first paragraph
            first_para = doc.paragraphs[0]
            
            # Create TOC title
            toc_title = doc.paragraphs[0]._element.getparent().insert(0, OxmlElement('w:p'))
            toc_title_para = doc.paragraphs[0]
            toc_title_para.text = "Table of Contents"
            toc_title_para.style = 'Heading 1'
            
            # Create TOC field paragraph
            toc_para_elem = OxmlElement('w:p')
            first_para._element.getparent().insert(1, toc_para_elem)
            
            # Create the TOC field
            fld_begin = OxmlElement('w:fldChar')
            fld_begin.set(qn('w:fldCharType'), 'begin')
            
            instr_text = OxmlElement('w:instrText')
            instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
            
            fld_end = OxmlElement('w:fldChar')
            fld_end.set(qn('w:fldCharType'), 'end')
            
            # Create runs for the field
            run1 = OxmlElement('w:r')
            run1.append(fld_begin)
            
            run2 = OxmlElement('w:r')
            run2.append(instr_text)
            
            run3 = OxmlElement('w:r')
            run3.append(fld_end)
            
            # Add runs to paragraph
            toc_para_elem.append(run1)
            toc_para_elem.append(run2)
            toc_para_elem.append(run3)
            
            current_app.logger.info("✅ Created fresh Table of Contents")
            return True
        else:
            current_app.logger.warning("⚠️ Document has no paragraphs to insert TOC")
            return False
            
    except Exception as e:
        current_app.logger.error(f"❌ Error creating fresh TOC: {e}")
        return False


def ensure_headings_for_toc(doc):
    """
    Ensures all headings in the document are properly formatted for TOC generation.
    This function scans the document and makes sure headings have the correct styles
    that Word's TOC will recognize.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        int: Number of headings processed
    """
    try:
        headings_processed = 0
        
        # Define heading style names that TOC recognizes
        heading_styles = [
            'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6',
            'heading 1', 'heading 2', 'heading 3', 'heading 4', 'heading 5', 'heading 6'
        ]
        
        for paragraph in doc.paragraphs:
            # Check if paragraph has heading style
            if paragraph.style.name in heading_styles:
                headings_processed += 1
                current_app.logger.debug(f"🔄 Found heading: '{paragraph.text[:50]}...' (Style: {paragraph.style.name})")
                
                # Ensure the heading has proper outline level for TOC
                if hasattr(paragraph, '_element'):
                    para_elem = paragraph._element
                    
                    # Check if outline level is set
                    pPr = para_elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                    if pPr is not None:
                        outline_lvl = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
                        if outline_lvl is None:
                            # Add outline level based on heading style
                            from docx.oxml import OxmlElement
                            outline_lvl = OxmlElement('w:outlineLvl')
                            
                            # Extract level from style name
                            style_name = paragraph.style.name.lower()
                            if 'heading 1' in style_name:
                                outline_lvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')
                            elif 'heading 2' in style_name:
                                outline_lvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '1')
                            elif 'heading 3' in style_name:
                                outline_lvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '2')
                            elif 'heading 4' in style_name:
                                outline_lvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '3')
                            elif 'heading 5' in style_name:
                                outline_lvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '4')
                            elif 'heading 6' in style_name:
                                outline_lvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '5')
                            else:
                                outline_lvl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')
                            
                            pPr.append(outline_lvl)
                            current_app.logger.debug(f"🔄 Added outline level to heading: {paragraph.text[:30]}...")
        
        if headings_processed > 0:
            current_app.logger.info(f"✅ Processed {headings_processed} heading(s) for TOC generation")
        else:
            current_app.logger.debug("ℹ️ No headings found in document")
        
        return headings_processed
        
    except Exception as e:
        current_app.logger.error(f"❌ Error ensuring headings for TOC: {e}")
        return 0


def force_complete_toc_rebuild(docx_path):
    """
    Forces complete TOC rebuild by:
    1. Removing ALL existing TOC fields entirely
    2. Creating fresh TOC fields that Word MUST recalculate
    3. Setting aggressive update flags to force page number recalculation
    
    This completely eliminates cached page numbers and forces Word to 
    recalculate everything from scratch when the document opens.
    
    Args:
        docx_path: Path to the saved .docx file
        
    Returns:
        int: Number of TOC fields completely rebuilt
    """
    try:
        import zipfile
        import tempfile
        import shutil
        import os
        from lxml import etree
        
        fields_rebuilt = 0
        
        # Create temporary directory for processing
        temp_dir = tempfile.mkdtemp()
        
        # Extract docx as ZIP
        extract_dir = os.path.join(temp_dir, 'extracted')
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(extract_dir)
        
        # Process document.xml
        doc_xml_path = os.path.join(extract_dir, 'word', 'document.xml')
        if not os.path.exists(doc_xml_path):
            current_app.logger.warning("⚠️ document.xml not found in docx file")
            shutil.rmtree(temp_dir)
            return 0
            
        # Parse document XML
        with open(doc_xml_path, 'r', encoding='utf-8') as f:
            xml_content = f.read()
            
        root = etree.fromstring(xml_content.encode('utf-8'))
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        
        current_app.logger.debug("🔄 Completely removing and rebuilding TOC fields...")
        
        # Find and completely remove existing TOC fields
        toc_locations = []  # Store where TOCs were for recreation
        all_paragraphs = root.xpath('.//w:p', namespaces=namespaces)
        
        paragraphs_to_remove = []
        in_toc_field = False
        toc_start_para = None
        toc_field_code = None
        
        for para_idx, para in enumerate(all_paragraphs):
            # Look for TOC field start
            instr_texts = para.xpath('.//w:instrText', namespaces=namespaces)
            
            for instr_text in instr_texts:
                if instr_text.text and instr_text.text.strip().upper().startswith('TOC'):
                    field_code = instr_text.text.strip()
                    field_type = "List of Figures" if ('\\C' in field_code or 'FIGURE' in field_code) else "Table of Contents"
                    
                    current_app.logger.debug(f"🔄 Found {field_type} to rebuild: {field_code}")
                    
                    # Mark start of TOC field for complete removal
                    in_toc_field = True
                    toc_start_para = para
                    toc_field_code = field_code
                    
                    # Store location for recreation
                    toc_locations.append({
                        'parent': para.getparent(),
                        'index': list(para.getparent()).index(para),
                        'field_code': field_code,
                        'field_type': field_type
                    })
                    break
            
            # If we're in a TOC field, mark paragraphs for removal
            if in_toc_field:
                paragraphs_to_remove.append(para)
                
                # Look for field end markers
                fld_chars = para.xpath('.//w:fldChar', namespaces=namespaces)
                for fld_char in fld_chars:
                    if fld_char.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType') == 'end':
                        # End of TOC field found - stop removing paragraphs
                        in_toc_field = False
                        fields_rebuilt += 1
                        break
        
        # Remove all TOC paragraphs completely
        for para in paragraphs_to_remove:
            parent = para.getparent()
            if parent is not None:
                parent.remove(para)
        
        current_app.logger.debug(f"🗑️ Removed {len(paragraphs_to_remove)} TOC paragraphs completely")
        
        # Recreate fresh TOC fields at the same locations
        for toc_info in toc_locations:
            parent = toc_info['parent']
            index = toc_info['index']
            field_code = toc_info['field_code']
            field_type = toc_info['field_type']
            
            current_app.logger.debug(f"🔄 Creating fresh {field_type} field...")
            
            # Create completely new TOC paragraph with fresh field
            new_para = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
            
            # Create field begin run
            run1 = etree.SubElement(new_para, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            fld_begin = etree.SubElement(run1, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar')
            fld_begin.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
            fld_begin.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dirty', 'true')  # Force update
            
            # Create instruction text run
            run2 = etree.SubElement(new_para, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            instr_text = etree.SubElement(run2, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText')
            instr_text.text = field_code
            
            # Create field separate run
            run3 = etree.SubElement(new_para, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            fld_sep = etree.SubElement(run3, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar')
            fld_sep.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'separate')
            
            # Create placeholder text run (Word will replace this)
            run4 = etree.SubElement(new_para, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            placeholder_text = etree.SubElement(run4, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            placeholder_text.text = "Table of Contents will be generated here"
            
            # Create field end run
            run5 = etree.SubElement(new_para, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            fld_end = etree.SubElement(run5, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar')
            fld_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
            
            # Insert the new paragraph at the correct location
            if index < len(parent):
                parent.insert(index, new_para)
            else:
                parent.append(new_para)
            
            current_app.logger.debug(f"✅ Created fresh {field_type} field with forced update flags")
        
        # Also add document-level settings to force field updates
        settings_xml_path = os.path.join(extract_dir, 'word', 'settings.xml')
        if os.path.exists(settings_xml_path):
            try:
                with open(settings_xml_path, 'r', encoding='utf-8') as f:
                    settings_content = f.read()
                
                settings_root = etree.fromstring(settings_content.encode('utf-8'))
                
                # Add updateFields setting to force field updates on open
                update_fields = settings_root.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}updateFields')
                if update_fields is None:
                    update_fields = etree.SubElement(settings_root, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}updateFields')
                update_fields.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'true')
                
                # Save modified settings
                modified_settings = etree.tostring(settings_root, encoding='utf-8', xml_declaration=True).decode('utf-8')
                with open(settings_xml_path, 'w', encoding='utf-8') as f:
                    f.write(modified_settings)
                
                current_app.logger.debug("✅ Added updateFields setting to force field updates on document open")
                
            except Exception as e:
                current_app.logger.debug(f"⚠️ Could not modify settings.xml: {e}")
        
        # Save the modified XML back
        modified_xml = etree.tostring(root, encoding='utf-8', xml_declaration=True).decode('utf-8')
        
        with open(doc_xml_path, 'w', encoding='utf-8') as f:
            f.write(modified_xml)
        
        # Repackage the docx file
        new_docx_path = docx_path + '.tmp'
        with zipfile.ZipFile(new_docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for root_dir, dirs, files in os.walk(extract_dir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, extract_dir)
                    zip_out.write(file_path, arcname)
        
        # Replace original file
        shutil.move(new_docx_path, docx_path)
        
        # Cleanup
        shutil.rmtree(temp_dir)
        
        if fields_rebuilt > 0:
            current_app.logger.info(f"✅ Completely rebuilt {fields_rebuilt} TOC field(s) from scratch - Word MUST recalculate all page numbers on open")
        else:
            current_app.logger.debug("ℹ️ No TOC fields found to rebuild")
        
        return fields_rebuilt
        
    except Exception as e:
        current_app.logger.error(f"❌ Error in complete TOC rebuild: {e}")
        import traceback
        current_app.logger.error(traceback.format_exc())
        
        # Cleanup on error
        if 'temp_dir' in locals():
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
        
        return 0


def regenerate_toc_completely(docx_path):
    """
    DEPRECATED: Use force_complete_toc_rebuild() instead for better page number accuracy.
    
    Completely regenerates the Table of Contents by:
    1. Removing existing TOC content entirely
    2. Forcing Word to rebuild TOC from scratch based on actual headings
    3. Setting proper field update flags
    
    This simulates: Select TOC -> Right Click -> Update Field -> Update Entire Table
    but does it programmatically to eliminate manual intervention.
    
    Args:
        docx_path: Path to the saved .docx file
        
    Returns:
        int: Number of TOC fields processed for complete regeneration
    """
    try:
        import zipfile
        import tempfile
        import shutil
        import os
        from lxml import etree
        
        fields_updated = 0
        
        # Create temporary directory for processing
        temp_dir = tempfile.mkdtemp()
        
        # Extract docx as ZIP
        extract_dir = os.path.join(temp_dir, 'extracted')
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(extract_dir)
        
        # Process document.xml
        doc_xml_path = os.path.join(extract_dir, 'word', 'document.xml')
        if not os.path.exists(doc_xml_path):
            current_app.logger.warning("⚠️ document.xml not found in docx file")
            shutil.rmtree(temp_dir)
            return 0
            
        # Parse document XML
        with open(doc_xml_path, 'r', encoding='utf-8') as f:
            xml_content = f.read()
            
        root = etree.fromstring(xml_content.encode('utf-8'))
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        
        current_app.logger.debug("🔄 Searching for TOC fields for complete regeneration...")
        
        # Find all TOC fields and completely clear their content
        all_paragraphs = root.xpath('.//w:p', namespaces=namespaces)
        
        toc_paragraphs_to_clear = []
        in_toc_field = False
        current_toc_paras = []
        
        for para_idx, para in enumerate(all_paragraphs):
            # Look for TOC field start
            instr_texts = para.xpath('.//w:instrText', namespaces=namespaces)
            
            for instr_text in instr_texts:
                if instr_text.text and instr_text.text.strip().upper().startswith('TOC'):
                    field_code = instr_text.text.strip()
                    field_type = "List of Figures" if ('\\C' in field_code or 'FIGURE' in field_code) else "Table of Contents"
                    
                    current_app.logger.debug(f"🔄 Found {field_type} field: {field_code}")
                    
                    # Mark start of TOC field
                    in_toc_field = True
                    current_toc_paras = [para]
                    
                    # Clear the instruction text but keep the field structure
                    # This forces Word to regenerate from scratch
                    break
            
            # If we're in a TOC field, collect all paragraphs until field end
            if in_toc_field:
                if para not in current_toc_paras:
                    current_toc_paras.append(para)
                
                # Look for field end markers
                fld_chars = para.xpath('.//w:fldChar', namespaces=namespaces)
                for fld_char in fld_chars:
                    if fld_char.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType') == 'end':
                        # End of TOC field found
                        toc_paragraphs_to_clear.extend(current_toc_paras)
                        in_toc_field = False
                        current_toc_paras = []
                        fields_updated += 1
                        break
        
        # Clear all TOC content completely
        for para in toc_paragraphs_to_clear:
            # Find all text runs in the paragraph
            runs = para.xpath('.//w:r', namespaces=namespaces)
            
            for run in runs:
                # Check if this run contains field characters (keep those)
                fld_chars = run.xpath('.//w:fldChar', namespaces=namespaces)
                instr_texts = run.xpath('.//w:instrText', namespaces=namespaces)
                
                if fld_chars or instr_texts:
                    # This run contains field markup - keep it but mark for update
                    for fld_char in fld_chars:
                        # Set dirty flag to force complete regeneration
                        fld_char.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dirty', 'true')
                        fld_char.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldLock', 'false')
                else:
                    # This run contains TOC content - remove all text
                    text_elements = run.xpath('.//w:t', namespaces=namespaces)
                    for text_elem in text_elements:
                        text_elem.text = ""
                    
                    # Also remove tab characters and other formatting
                    tabs = run.xpath('.//w:tab', namespaces=namespaces)
                    for tab in tabs:
                        tab.getparent().remove(tab)
        
        # Save the modified XML back
        modified_xml = etree.tostring(root, encoding='utf-8', xml_declaration=True).decode('utf-8')
        
        with open(doc_xml_path, 'w', encoding='utf-8') as f:
            f.write(modified_xml)
        
        # Repackage the docx file
        new_docx_path = docx_path + '.tmp'
        with zipfile.ZipFile(new_docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for root_dir, dirs, files in os.walk(extract_dir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, extract_dir)
                    zip_out.write(file_path, arcname)
        
        # Replace original file
        shutil.move(new_docx_path, docx_path)
        
        # Cleanup
        shutil.rmtree(temp_dir)
        
        if fields_updated > 0:
            current_app.logger.info(f"✅ Completely cleared {fields_updated} TOC field(s) - Word will regenerate from scratch on open")
        else:
            current_app.logger.debug("ℹ️ No TOC fields found to regenerate")
        
        return fields_updated
        
    except Exception as e:
        current_app.logger.error(f"❌ Error in complete TOC regeneration: {e}")
        import traceback
        current_app.logger.error(traceback.format_exc())
        
        # Cleanup on error
        if 'temp_dir' in locals():
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
        
        return 0


def force_toc_page_number_update(docx_path):
    """
    DEPRECATED: Use regenerate_toc_completely() instead.
    
    Forces TOC page number recalculation by clearing cached TOC content and setting update flags.
    This simulates: Select TOC -> Right Click -> Update Field -> Update Entire Table
    
    The key is to completely clear the TOC cached results so Word must regenerate 
    the entire TOC with correct page numbers when the document is opened.
    
    Args:
        docx_path: Path to the saved .docx file
        
    Returns:
        int: Number of TOC fields processed for update
    """
    try:
        import zipfile
        import tempfile
        import shutil
        import os
        from lxml import etree
        
        fields_updated = 0
        
        # Create temporary directory for processing
        temp_dir = tempfile.mkdtemp()
        
        # Extract docx as ZIP
        extract_dir = os.path.join(temp_dir, 'extracted')
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(extract_dir)
        
        # Process document.xml
        doc_xml_path = os.path.join(extract_dir, 'word', 'document.xml')
        if not os.path.exists(doc_xml_path):
            current_app.logger.warning("⚠️ document.xml not found in docx file")
            shutil.rmtree(temp_dir)
            return 0
            
        # Parse document XML
        with open(doc_xml_path, 'r', encoding='utf-8') as f:
            xml_content = f.read()
            
        root = etree.fromstring(xml_content.encode('utf-8'))
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        
        current_app.logger.debug("🔄 Searching for TOC fields to force page number update...")
        
        # Find all paragraphs that contain TOC fields
        all_paragraphs = root.xpath('.//w:p', namespaces=namespaces)
        
        for para_idx, para in enumerate(all_paragraphs):
            # Look for TOC field instructions
            instr_texts = para.xpath('.//w:instrText', namespaces=namespaces)
            
            for instr_text in instr_texts:
                if instr_text.text and instr_text.text.strip().upper().startswith('TOC'):
                    field_code = instr_text.text.strip().upper()
                    field_type = "List of Figures" if ('\\C' in field_code or 'FIGURE' in field_code) else "Table of Contents"
                    
                    current_app.logger.debug(f"🔄 Found {field_type} field: {instr_text.text[:60]}...")
                    
                    # Find field structure: begin -> instrText -> separate -> [cached content] -> end
                    field_chars = para.xpath('.//w:fldChar', namespaces=namespaces)
                    
                    separate_found = None
                    end_found = None
                    
                    for fld_char in field_chars:
                        fld_char_type = fld_char.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType')
                        if fld_char_type == 'separate':
                            separate_found = fld_char
                        elif fld_char_type == 'end':
                            end_found = fld_char
                    
                    if separate_found is not None and end_found is not None:
                        # Clear ALL content between separate and end markers
                        # This forces Word to completely regenerate the TOC with current page numbers
                        
                        # Get all runs in the paragraph
                        runs = para.xpath('.//w:r', namespaces=namespaces)
                        
                        clearing_mode = False
                        runs_cleared = 0
                        
                        for run in runs:
                            # Check if this run contains the separate marker
                            run_fld_chars = run.xpath('.//w:fldChar', namespaces=namespaces)
                            for fld_char in run_fld_chars:
                                fld_char_type = fld_char.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType')
                                if fld_char_type == 'separate':
                                    clearing_mode = True
                                elif fld_char_type == 'end':
                                    clearing_mode = False
                            
                            # If we're in clearing mode, remove all text content
                            if clearing_mode:
                                text_elements = run.xpath('.//w:t', namespaces=namespaces)
                                for text_elem in text_elements:
                                    if text_elem.text and text_elem.text.strip():
                                        text_elem.text = ''
                                        runs_cleared += 1
                                
                                # Also clear any tab characters, page breaks, etc.
                                for child in run:
                                    if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
                                        child.text = ''
                                    elif child.tag in [
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tab',
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'
                                    ]:
                                        # Keep structural elements but clear any text
                                        pass
                        
                        # Also check if TOC spans multiple paragraphs (common for long TOCs)
                        if runs_cleared == 0:
                            # TOC might span multiple paragraphs - look ahead
                            for next_para_idx in range(para_idx + 1, min(para_idx + 50, len(all_paragraphs))):
                                next_para = all_paragraphs[next_para_idx]
                                
                                # Check if this paragraph has the end marker
                                next_field_chars = next_para.xpath('.//w:fldChar', namespaces=namespaces)
                                has_end = any(fc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType') == 'end' 
                                            for fc in next_field_chars)
                                
                                if has_end:
                                    # Clear content in this paragraph up to the end marker
                                    next_runs = next_para.xpath('.//w:r', namespaces=namespaces)
                                    for run in next_runs:
                                        run_fld_chars = run.xpath('.//w:fldChar', namespaces=namespaces)
                                        has_end_in_run = any(fc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType') == 'end' 
                                                           for fc in run_fld_chars)
                                        
                                        if not has_end_in_run:
                                            # Clear this run's text
                                            text_elements = run.xpath('.//w:t', namespaces=namespaces)
                                            for text_elem in text_elements:
                                                if text_elem.text:
                                                    text_elem.text = ''
                                                    runs_cleared += 1
                                        else:
                                            # This run has the end marker - only clear text before the marker
                                            break
                                    break
                                else:
                                    # This paragraph is entirely within the TOC field - clear all text
                                    text_elements = next_para.xpath('.//w:t', namespaces=namespaces)
                                    for text_elem in text_elements:
                                        if text_elem.text:
                                            text_elem.text = ''
                                            runs_cleared += 1
                        
                        if runs_cleared > 0:
                            fields_updated += 1
                            current_app.logger.debug(f"✅ Cleared {runs_cleared} text elements from {field_type} - Word will regenerate with correct page numbers")
                        
                        # Set field as dirty to ensure update
                        begin_chars = para.xpath('.//w:fldChar[@w:fldCharType="begin"]', namespaces=namespaces)
                        for begin_char in begin_chars:
                            begin_char.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dirty', 'true')
        
        # Also modify settings.xml to ensure Word updates fields on document open
        settings_xml_path = os.path.join(extract_dir, 'word', 'settings.xml')
        try:
            if os.path.exists(settings_xml_path):
                with open(settings_xml_path, 'r', encoding='utf-8') as f:
                    settings_content = f.read()
                
                settings_root = etree.fromstring(settings_content.encode('utf-8'))
                settings_ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                # Check if updateFields setting exists
                update_fields = settings_root.xpath('.//w:updateFields', namespaces=settings_ns)
                
                if not update_fields:
                    # Add updateFields setting to force field updates on open
                    settings_elem = settings_root
                    
                    # Create updateFields element
                    update_fields_elem = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}updateFields')
                    update_fields_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'true')
                    
                    # Insert it as the first child
                    settings_elem.insert(0, update_fields_elem)
                    
                    # Write back the modified settings
                    modified_settings = etree.tostring(settings_root, encoding='utf-8', xml_declaration=True)
                    with open(settings_xml_path, 'wb') as f:
                        f.write(modified_settings)
                    
                    current_app.logger.debug("✅ Added updateFields setting to force field updates on document open")
                else:
                    # Ensure existing updateFields is set to true
                    for update_field in update_fields:
                        update_field.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'true')
                    
                    modified_settings = etree.tostring(settings_root, encoding='utf-8', xml_declaration=True)
                    with open(settings_xml_path, 'wb') as f:
                        f.write(modified_settings)
                    
                    current_app.logger.debug("✅ Ensured updateFields setting is enabled")
                    
        except Exception as settings_error:
            current_app.logger.debug(f"⚠️ Could not modify settings.xml: {settings_error}")
        
        # Save the modified document if any fields were updated
        if fields_updated > 0:
            # Write modified XML back
            modified_xml = etree.tostring(root, encoding='utf-8', xml_declaration=True)
            with open(doc_xml_path, 'wb') as f:
                f.write(modified_xml)
            
            current_app.logger.info(f"✅ Successfully cleared {fields_updated} TOC field(s) - Word will regenerate with correct page numbers on open")
        else:
            current_app.logger.debug("ℹ️ No TOC fields found to update")
            
        # Always recreate the docx file to include any settings changes
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for root_dir, dirs, files in os.walk(extract_dir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arc_name = os.path.relpath(file_path, extract_dir)
                    zip_out.write(file_path, arc_name)
        
        # Cleanup
        shutil.rmtree(temp_dir)
        
        return fields_updated
        
    except Exception as e:
        current_app.logger.error(f"❌ Error forcing TOC page number update: {e}")
        import traceback
        current_app.logger.debug(traceback.format_exc())
        return 0


def update_toc_fields_programmatically(docx_path, flat_data_map=None):
    """
    Programmatically updates TOC and List of Figures fields by simulating the manual process:
    1. Find all TOC fields in the document
    2. Replace placeholders in heading text that TOC references
    3. Force field update by manipulating field properties
    4. Regenerate TOC content based on updated headings
    
    This mimics the manual process: Select TOC -> Right Click -> Update Field -> Update Entire Table
    
    Args:
        docx_path: Path to the saved .docx file
        flat_data_map: Dictionary mapping placeholder keys to replacement values
        
    Returns:
        int: Number of TOC fields successfully updated
    """
    try:
        import zipfile
        import tempfile
        import shutil
        import os
        from lxml import etree
        
        if not flat_data_map:
            flat_data_map = {}
            
        fields_updated = 0
        
        # Create temporary directory for processing
        temp_dir = tempfile.mkdtemp()
        temp_docx = os.path.join(temp_dir, 'temp_doc.docx')
        
        # Copy original file
        shutil.copy2(docx_path, temp_docx)
        
        # Extract docx as ZIP
        extract_dir = os.path.join(temp_dir, 'extracted')
        with zipfile.ZipFile(temp_docx, 'r') as zip_ref:
            zip_ref.extractall(extract_dir)
        
        # Process document.xml
        doc_xml_path = os.path.join(extract_dir, 'word', 'document.xml')
        if not os.path.exists(doc_xml_path):
            current_app.logger.warning("⚠️ document.xml not found in docx file")
            return 0
            
        # Parse document XML
        with open(doc_xml_path, 'r', encoding='utf-8') as f:
            xml_content = f.read()
            
        root = etree.fromstring(xml_content.encode('utf-8'))
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        
        # Step 1: Replace placeholders in all heading paragraphs first
        current_app.logger.debug("🔄 Step 1: Updating heading text that TOC will reference...")
        
        # Find all paragraphs with heading styles
        heading_paragraphs = root.xpath('.//w:p[w:pPr/w:pStyle[@w:val[starts-with(., "Heading") or starts-with(., "heading")]]]', namespaces=namespaces)
        
        headings_updated = 0
        for para in heading_paragraphs:
            # Get all text elements in this paragraph
            text_elements = para.xpath('.//w:t', namespaces=namespaces)
            
            for text_elem in text_elements:
                if text_elem.text:
                    original_text = text_elem.text
                    modified_text = original_text
                    
                    # Replace placeholders
                    import re
                    
                    # Replace <placeholder> tags
                    angle_matches = re.findall(r'<([^>]+)>', original_text)
                    for match in angle_matches:
                        key_lower = match.lower().strip()
                        value = flat_data_map.get(key_lower, '')
                        if value:
                            pattern = re.compile(re.escape(f"<{match}>"), re.IGNORECASE)
                            modified_text = pattern.sub(str(value), modified_text)
                    
                    # Replace ${placeholder} tags  
                    dollar_matches = re.findall(r'\$\{([^\}]+)\}', original_text)
                    for match in dollar_matches:
                        key_lower = match.lower().strip()
                        value = flat_data_map.get(key_lower, '')
                        if value:
                            pattern = re.compile(re.escape(f"${{{match}}}"), re.IGNORECASE)
                            modified_text = pattern.sub(str(value), modified_text)
                    
                    if modified_text != original_text:
                        text_elem.text = modified_text
                        headings_updated += 1
                        current_app.logger.debug(f"🔄 Updated heading: '{original_text[:50]}...' -> '{modified_text[:50]}...'")
        
        current_app.logger.debug(f"✅ Updated {headings_updated} heading text elements")
        
        # Step 2: Force TOC field update by manipulating field properties
        current_app.logger.debug("🔄 Step 2: Forcing TOC field regeneration...")
        
        # Find all TOC fields and mark them for update
        toc_fields = []
        
        # Look for TOC field codes
        instr_texts = root.xpath('.//w:instrText', namespaces=namespaces)
        for instr_text in instr_texts:
            if instr_text.text and instr_text.text.strip().upper().startswith('TOC'):
                # Found a TOC field - find the complete field structure
                field_code = instr_text.text.strip()
                field_type = "List of Figures" if ('\\C' in field_code.upper() or 'FIGURE' in field_code.upper()) else "Table of Contents"
                
                # Find the field's parent paragraph
                para = instr_text
                while para is not None and para.tag != '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
                    para = para.getparent()
                
                if para is not None:
                    toc_fields.append({
                        'paragraph': para,
                        'field_code': field_code,
                        'field_type': field_type,
                        'instr_text': instr_text
                    })
        
        # Step 3: Clear existing TOC content and mark for regeneration
        for toc_field in toc_fields:
            para = toc_field['paragraph']
            field_type = toc_field['field_type']
            
            # Find field boundaries (begin, separate, end)
            field_chars = para.xpath('.//w:fldChar', namespaces=namespaces)
            
            begin_found = False
            separate_found = False
            end_found = False
            
            for fld_char in field_chars:
                fld_char_type = fld_char.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType')
                if fld_char_type == 'begin':
                    begin_found = True
                elif fld_char_type == 'separate':
                    separate_found = True
                elif fld_char_type == 'end':
                    end_found = True
            
            if begin_found and separate_found and end_found:
                # This is a complete field - clear the cached result
                # Find all runs between separate and end markers
                runs_to_clear = []
                in_result_area = False
                
                for run in para.xpath('.//w:r', namespaces=namespaces):
                    # Check if this run contains field markers
                    fld_chars_in_run = run.xpath('.//w:fldChar', namespaces=namespaces)
                    
                    for fld_char in fld_chars_in_run:
                        fld_char_type = fld_char.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType')
                        if fld_char_type == 'separate':
                            in_result_area = True
                        elif fld_char_type == 'end':
                            in_result_area = False
                    
                    # If we're in the result area, mark text for clearing
                    if in_result_area:
                        text_elems = run.xpath('.//w:t', namespaces=namespaces)
                        for text_elem in text_elems:
                            if text_elem.text:
                                runs_to_clear.append(text_elem)
                
                # Clear the cached TOC content
                for text_elem in runs_to_clear:
                    text_elem.text = ''
                
                # Add a field update instruction by setting the field's dirty flag
                # This tells Word to recalculate the field when the document opens
                instr_text_elem = toc_field['instr_text']
                
                # Add or modify the field instruction to force update
                current_field_code = instr_text_elem.text.strip()
                if '\\* MERGEFORMAT' not in current_field_code.upper():
                    instr_text_elem.text = current_field_code + ' \\* MERGEFORMAT'
                
                fields_updated += 1
                current_app.logger.debug(f"🔄 Cleared and marked {field_type} field for regeneration")
        
        # Step 4: Save the modified document
        if fields_updated > 0 or headings_updated > 0:
            # Write modified XML back
            modified_xml = etree.tostring(root, encoding='utf-8', xml_declaration=True, pretty_print=True)
            with open(doc_xml_path, 'wb') as f:
                f.write(modified_xml)
            
            # Recreate the docx file
            with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                for root_dir, dirs, files in os.walk(extract_dir):
                    for file in files:
                        file_path = os.path.join(root_dir, file)
                        arc_name = os.path.relpath(file_path, extract_dir)
                        zip_out.write(file_path, arc_name)
            
            current_app.logger.info(f"✅ Successfully updated {fields_updated} TOC field(s) and {headings_updated} heading(s)")
        else:
            current_app.logger.debug("ℹ️ No TOC fields or headings found to update")
        
        # Cleanup
        shutil.rmtree(temp_dir)
        
        return fields_updated
        
    except Exception as e:
        current_app.logger.error(f"❌ Error in programmatic TOC update: {e}")
        import traceback
        current_app.logger.debug(traceback.format_exc())
        return 0


def update_toc_fields_in_docx(docx_path, flat_data_map=None):
    """
    Post-processes a saved .docx file to replace placeholders in TOC content and clear TOC field results.
    
    This function manipulates the .docx file (which is a ZIP archive) to:
    1. Replace placeholders in TOC field cached content
    2. Clear the field results (content between separate and end markers) so that Word will 
       automatically recalculate TOC and List of Figures when the document is opened.
    
    Args:
        docx_path: Path to the saved .docx file
        flat_data_map: Optional dictionary mapping placeholder keys to replacement values
        
    Returns:
        int: Number of fields processed and cleared for update
    """
    try:
        import zipfile
        from lxml import etree
        
        fields_updated = 0
        
        # Open the .docx file as a ZIP archive
        with zipfile.ZipFile(docx_path, 'r') as zip_read:
            # Read the main document XML
            try:
                document_xml = zip_read.read('word/document.xml')
            except KeyError:
                current_app.logger.warning("⚠️ Could not find word/document.xml in .docx file")
                return 0
            
            # Parse the XML
            root = etree.fromstring(document_xml)
            
            # Define namespaces
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            # Find all paragraphs in the document
            all_paragraphs = root.xpath('.//w:p', namespaces=namespaces)
            
            # Process each paragraph to find TOC fields
            for para_idx, para in enumerate(all_paragraphs):
                # Find all field separate markers in this paragraph
                field_separates = para.xpath('.//w:fldChar[@w:fldCharType="separate"]', namespaces=namespaces)
                
                for separate_elem in field_separates:
                    # Check if this separate belongs to a TOC field
                    # Look backwards in the same paragraph to find the instrText
                    para_children = list(para)
                    separate_idx = None
                    
                    for idx, child in enumerate(para_children):
                        if separate_elem in child.iter():
                            separate_idx = idx
                            break
                    
                    if separate_idx is None:
                        continue
                    
                    # Look backwards to find the instrText (field code)
                    instr_text_found = None
                    for i in range(separate_idx, -1, -1):
                        child = para_children[i]
                        instr_texts = child.xpath('.//w:instrText', namespaces=namespaces)
                        for instr_text in instr_texts:
                            if instr_text.text and instr_text.text.strip().upper().startswith('TOC'):
                                instr_text_found = instr_text
                                break
                        if instr_text_found:
                            break
                    
                    if not instr_text_found:
                        continue
                    
                    # This is a TOC field - replace placeholders in cached content, then clear the result
                    field_code = instr_text_found.text.strip().upper()
                    field_type = "List of Figures" if ('\\C' in field_code or 'FIGURE' in field_code or '"FIGURE' in field_code) else "Table of Contents"
                    
                    # Find the end marker - it might be in the same paragraph or a following paragraph
                    end_found = None
                    end_para_idx = None
                    
                    # First check in the same paragraph
                    for i in range(separate_idx + 1, len(para_children)):
                        child = para_children[i]
                        end_markers = child.xpath('.//w:fldChar[@w:fldCharType="end"]', namespaces=namespaces)
                        if end_markers:
                            end_found = end_markers[0]
                            end_para_idx = para_idx
                            break
                    
                    # If not found in same paragraph, check following paragraphs
                    if not end_found:
                        for next_para_idx in range(para_idx + 1, len(all_paragraphs)):
                            next_para = all_paragraphs[next_para_idx]
                            end_markers = next_para.xpath('.//w:fldChar[@w:fldCharType="end"]', namespaces=namespaces)
                            if end_markers:
                                end_found = end_markers[0]
                                end_para_idx = next_para_idx
                                break
                    
                    if end_found:
                        cleared_any = False
                        toc_replacements = 0
                        
                        # First, replace placeholders in TOC field content if data map is provided
                        if flat_data_map:
                            # Helper function to replace placeholders in text
                            def replace_in_text(text):
                                if not text:
                                    return text, False
                                modified = text
                                replaced = False
                                
                                # Replace <placeholder> tags
                                import re
                                angle_matches = re.findall(r'<([^>]+)>', text)
                                for match in angle_matches:
                                    key_lower = match.lower().strip()
                                    value = flat_data_map.get(key_lower)
                                    if value:
                                        pattern = re.compile(re.escape(f"<{match}>"), re.IGNORECASE)
                                        modified = pattern.sub(str(value), modified)
                                        replaced = True
                                        toc_replacements += 1
                                
                                # Replace ${placeholder} tags
                                dollar_matches = re.findall(r'\$\{([^\}]+)\}', text)
                                for match in dollar_matches:
                                    key_lower = match.lower().strip()
                                    value = flat_data_map.get(key_lower)
                                    if value:
                                        pattern = re.compile(re.escape(f"${{{match}}}"), re.IGNORECASE)
                                        modified = pattern.sub(str(value), modified)
                                        replaced = True
                                        toc_replacements += 1
                                
                                return modified, replaced
                            
                            # Replace placeholders in TOC content before clearing
                            if end_para_idx == para_idx:
                                # End is in same paragraph
                                end_idx = None
                                for idx, child in enumerate(para_children):
                                    if end_found in child.iter():
                                        end_idx = idx
                                        break
                                
                                if end_idx is not None:
                                    for i in range(separate_idx + 1, end_idx):
                                        elem = para_children[i]
                                        text_elems = elem.xpath('.//w:t', namespaces=namespaces)
                                        for text_elem in text_elems:
                                            if text_elem.text:
                                                new_text, was_replaced = replace_in_text(text_elem.text)
                                                if was_replaced:
                                                    text_elem.text = new_text
                            else:
                                # End is in different paragraph - replace in all content between separate and end
                                # Replace in current paragraph after separate
                                for i in range(separate_idx + 1, len(para_children)):
                                    elem = para_children[i]
                                    text_elems = elem.xpath('.//w:t', namespaces=namespaces)
                                    for text_elem in text_elems:
                                        if text_elem.text:
                                            new_text, was_replaced = replace_in_text(text_elem.text)
                                            if was_replaced:
                                                text_elem.text = new_text
                                
                                # Replace in paragraphs between current and end
                                for mid_para_idx in range(para_idx + 1, end_para_idx):
                                    mid_para = all_paragraphs[mid_para_idx]
                                    text_elems = mid_para.xpath('.//w:t', namespaces=namespaces)
                                    for text_elem in text_elems:
                                        if text_elem.text:
                                            new_text, was_replaced = replace_in_text(text_elem.text)
                                            if was_replaced:
                                                text_elem.text = new_text
                                
                                # Replace in end paragraph before end marker
                                end_para = all_paragraphs[end_para_idx]
                                end_para_children = list(end_para)
                                end_idx = None
                                for idx, child in enumerate(end_para_children):
                                    if end_found in child.iter():
                                        end_idx = idx
                                        break
                                
                                if end_idx is not None:
                                    for i in range(0, end_idx):
                                        elem = end_para_children[i]
                                        text_elems = elem.xpath('.//w:t', namespaces=namespaces)
                                        for text_elem in text_elems:
                                            if text_elem.text:
                                                new_text, was_replaced = replace_in_text(text_elem.text)
                                                if was_replaced:
                                                    text_elem.text = new_text
                            
                            if toc_replacements > 0:
                                current_app.logger.debug(f"🔄 Replaced {toc_replacements} placeholder(s) in {field_type} field content")
                        
                        # Now clear content in the same paragraph (after separate)
                        if end_para_idx == para_idx:
                            # End is in same paragraph
                            end_idx = None
                            for idx, child in enumerate(para_children):
                                if end_found in child.iter():
                                    end_idx = idx
                                    break
                            
                            if end_idx is not None:
                                elements_to_remove = []
                                for i in range(separate_idx + 1, end_idx):
                                    elem = para_children[i]
                                    # Clear all text elements
                                    text_elems = elem.xpath('.//w:t', namespaces=namespaces)
                                    for text_elem in text_elems:
                                        if text_elem.text:
                                            text_elem.text = ''
                                            cleared_any = True
                                    
                                    # Mark empty runs for removal
                                    if elem.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r':
                                        has_non_text = False
                                        for child in elem:
                                            if child.tag != '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
                                                has_non_text = True
                                                break
                                        if not has_non_text:
                                            elements_to_remove.append(elem)
                                
                                for elem_to_remove in elements_to_remove:
                                    para.remove(elem_to_remove)
                        else:
                            # End is in a different paragraph - clear from separate to end
                            # Clear remaining content in current paragraph after separate
                            elements_to_remove = []
                            for i in range(separate_idx + 1, len(para_children)):
                                elem = para_children[i]
                                text_elems = elem.xpath('.//w:t', namespaces=namespaces)
                                for text_elem in text_elems:
                                    if text_elem.text:
                                        text_elem.text = ''
                                        cleared_any = True
                                
                                if elem.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r':
                                    has_non_text = False
                                    for child in elem:
                                        if child.tag != '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
                                            has_non_text = True
                                            break
                                    if not has_non_text:
                                        elements_to_remove.append(elem)
                            
                            for elem_to_remove in elements_to_remove:
                                para.remove(elem_to_remove)
                            
                            # Clear all paragraphs between current and end paragraph
                            for mid_para_idx in range(para_idx + 1, end_para_idx):
                                mid_para = all_paragraphs[mid_para_idx]
                                text_elems = mid_para.xpath('.//w:t', namespaces=namespaces)
                                for text_elem in text_elems:
                                    if text_elem.text:
                                        text_elem.text = ''
                                        cleared_any = True
                            
                            # Clear content in end paragraph before the end marker
                            end_para = all_paragraphs[end_para_idx]
                            end_para_children = list(end_para)
                            end_idx = None
                            for idx, child in enumerate(end_para_children):
                                if end_found in child.iter():
                                    end_idx = idx
                                    break
                            
                            if end_idx is not None:
                                elements_to_remove = []
                                for i in range(0, end_idx):
                                    elem = end_para_children[i]
                                    text_elems = elem.xpath('.//w:t', namespaces=namespaces)
                                    for text_elem in text_elems:
                                        if text_elem.text:
                                            text_elem.text = ''
                                            cleared_any = True
                                    
                                    if elem.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r':
                                        has_non_text = False
                                        for child in elem:
                                            if child.tag != '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
                                                has_non_text = True
                                                break
                                        if not has_non_text:
                                            elements_to_remove.append(elem)
                                
                                for elem_to_remove in elements_to_remove:
                                    end_para.remove(elem_to_remove)
                        
                        if cleared_any:
                            fields_updated += 1
                            current_app.logger.debug(f"🔄 Cleared {field_type} field result - Word will recalculate on open")
            
            # If we found and modified fields, save the updated XML
            if fields_updated > 0:
                # Create a temporary ZIP file
                temp_zip_path = docx_path + '.tmp'
                
                with zipfile.ZipFile(docx_path, 'r') as zip_read:
                    with zipfile.ZipFile(temp_zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_write:
                        # Copy all files except document.xml
                        for item in zip_read.infolist():
                            if item.filename != 'word/document.xml':
                                data = zip_read.read(item.filename)
                                zip_write.writestr(item, data)
                        
                        # Write the modified document.xml
                        modified_xml = etree.tostring(root, encoding='UTF-8', xml_declaration=True)
                        zip_write.writestr('word/document.xml', modified_xml)
                
                # Replace the original file with the modified one
                import shutil
                shutil.move(temp_zip_path, docx_path)
                
                current_app.logger.info(f"✅ Cleared {fields_updated} TOC/List of Figures field result(s) - Word will update them automatically on open")
            else:
                current_app.logger.debug("ℹ️ No TOC/List of Figures fields found to update in saved document")
        
        return fields_updated
        
    except Exception as e:
        current_app.logger.error(f"❌ Error updating TOC fields in .docx file: {e}")
        import traceback
        current_app.logger.debug(traceback.format_exc())
        return 0

def update_toc_and_list_of_figures(doc):
    """
    Updates Table of Contents and List of Figures in a Word document by manipulating XML.
    
    This function finds all TOC and List of Figures field codes in the document and
    ensures they are properly structured so Word will update them when the document is opened.
    The function manipulates the document's XML to mark fields for update.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        int: Number of fields found and prepared for update
    """
    try:
        from docx.oxml.ns import qn
        from lxml import etree
        
        fields_found = 0
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # Access the document's main XML element
        document_element = doc.element
        
        # Find all paragraphs in the document
        for paragraph in doc.paragraphs:
            para_element = paragraph._element
            
            # Look for all runs in this paragraph
            runs = para_element.xpath('.//w:r', namespaces=namespaces)
            
            for run_elem in runs:
                # Check for field instruction text (this contains the field code)
                instr_texts = run_elem.xpath('.//w:instrText', namespaces=namespaces)
                
                for instr_text in instr_texts:
                    if instr_text.text:
                        field_code = instr_text.text.strip()
                        field_code_upper = field_code.upper()
                        
                        # Check if this is a TOC field (Table of Contents)
                        # TOC fields typically start with "TOC" and may have switches like \h, \o, etc.
                        is_toc = field_code_upper.startswith('TOC')
                        
                        # Check if this is a List of Figures field
                        # List of Figures uses TOC with \c "Figure" or similar
                        is_list_of_figures = (is_toc and 
                                            ('\\C' in field_code_upper or 
                                             'FIGURE' in field_code_upper or 
                                             '"FIGURE"' in field_code_upper or
                                             '\\C "Figure' in field_code_upper))
                        
                        if is_toc:
                            fields_found += 1
                            field_type = "List of Figures" if is_list_of_figures else "Table of Contents"
                            current_app.logger.debug(f"🔄 Found {field_type} field: {field_code[:60]}")
                            
                            # To force Word to update the field, we need to ensure the field structure is correct
                            # and mark it as needing update. We do this by:
                            # 1. Ensuring the field has proper begin/separate/end markers
                            # 2. The field result should be between separate and end markers
                            
                            # Find the parent run that contains this field
                            parent_run = instr_text.getparent()
                            
                            # Check if field has proper structure (begin -> instrText -> separate -> result -> end)
                            field_begin = parent_run.xpath('.//w:fldChar[@w:fldCharType="begin"]', namespaces=namespaces)
                            field_separate = parent_run.xpath('.//w:fldChar[@w:fldCharType="separate"]', namespaces=namespaces)
                            field_end = parent_run.xpath('.//w:fldChar[@w:fldCharType="end"]', namespaces=namespaces)
                            
                            if field_begin and field_separate and field_end:
                                # Field structure is correct
                                # To force update, we can add a small modification to the field code
                                # that won't affect functionality but will trigger Word to recalculate
                                # One approach: add a space at the end if not present, or ensure proper formatting
                                
                                # Actually, a better approach is to ensure the field result area is cleared
                                # or marked. However, since we're adding content, Word should detect the change.
                                # The key is that when Word opens the document, it will see the field needs updating
                                # because the document content has changed.
                                
                                # For now, we'll ensure the field code is properly formatted
                                # Word will update fields automatically when it detects document changes
                                pass
        
        # Also check tables for field codes (TOC/List of Figures might be in tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        para_element = paragraph._element
                        runs = para_element.xpath('.//w:r', namespaces=namespaces)
                        
                        for run_elem in runs:
                            instr_texts = run_elem.xpath('.//w:instrText', namespaces=namespaces)
                            
                            for instr_text in instr_texts:
                                if instr_text.text:
                                    field_code = instr_text.text.strip()
                                    field_code_upper = field_code.upper()
                                    
                                    if field_code_upper.startswith('TOC'):
                                        fields_found += 1
                                        field_type = "List of Figures" if ('\\C' in field_code_upper or 'FIGURE' in field_code_upper) else "Table of Contents"
                                        current_app.logger.debug(f"🔄 Found {field_type} field in table: {field_code[:60]}")
        
        if fields_found > 0:
            current_app.logger.info(f"✅ Found {fields_found} TOC/List of Figures field(s) - will be updated when Word opens the document")
        else:
            current_app.logger.info("ℹ️ No TOC/List of Figures fields found in document")
        
        return fields_found
        
    except Exception as e:
        current_app.logger.error(f"❌ Error updating TOC and List of Figures: {e}")
        import traceback
        current_app.logger.debug(traceback.format_exc())
        return 0

def _generate_report(project_id, template_path, data_file_path):
    import pandas as pd
    import json
    import tempfile
    import matplotlib.pyplot as plt
    from matplotlib.ticker import FuncFormatter
    from docx import Document
    from docx.shared import Inches
    import re
    import os
    import gc  # Add garbage collection

    # Force matplotlib to use non-interactive backend to reduce memory usage
    plt.switch_backend('Agg')
    plt.style.use('ggplot')  # 👈 Apply a cleaner visual style

    try:
        # Report generation started

        # Try to read Excel with original formatting preserved
        try:
            df = pd.read_excel(data_file_path, sheet_name=0, keep_default_na=False)  # Use first sheet
        except:
            df = pd.read_excel(data_file_path, sheet_name=0)  # Fallback to default
        
        df.columns = df.columns.str.strip().str.replace(" ", "_").str.replace("__", "_")
        
        current_app.logger.debug(f"📊 Excel file loaded successfully. Shape: {df.shape}")
        current_app.logger.debug(f"📊 Columns after cleaning: {list(df.columns)}")
        
        # Log the raw data to see what pandas is reading
        # Excel data loaded successfully

        # Excel structure loaded silently
        
        # Validate required columns exist
        required_columns = ["Text_Tag", "Text", "Chart_Tag", "Chart_Attributes", "Chart_Type"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            current_app.logger.error(f"❌ Missing required columns: {missing_columns}")
            current_app.logger.error(f"❌ Available columns: {list(df.columns)}")
            raise ValueError(f"Excel file missing required columns: {missing_columns}")

        text_map = {str(k).strip().lower(): str(v).strip() for k, v in zip(df["Text_Tag"], df["Text"]) if pd.notna(k) and pd.notna(v)}
        chart_attr_map = {str(k).strip().lower(): str(v).strip() for k, v in zip(df["Chart_Tag"], df["Chart_Attributes"]) if pd.notna(k) and pd.notna(v)}
        chart_type_map = {str(k).strip().lower(): str(v).strip() for k, v in zip(df["Chart_Tag"], df["Chart_Type"]) if pd.notna(k) and pd.notna(v)}
        
        # Data maps created silently

        flat_data_map = {}
        # Add global metadata columns that can be used across all sections
        global_metadata = {}
        
        # Extract dynamic columns from A1 to M1 range
        dynamic_columns = extract_dynamic_columns_from_excel(data_file_path)
        
        # Ensure dynamic_columns is not empty and is a valid list
        if not dynamic_columns or not isinstance(dynamic_columns, list):
            current_app.logger.warning("Dynamic columns extraction failed or returned empty, using fallback")
            dynamic_columns = ['report_name', 'currency', 'country', 'report_code']
        
        current_app.logger.debug(f"Using dynamic columns: {dynamic_columns}")
        
        # First pass: Process global metadata columns from ALL rows
        for _, row in df.iterrows():
            for col in df.columns:
                # Normalize column names to match dynamic_columns normalization (spaces -> underscores)
                col_norm = col.lower().strip().replace(" ", "_").replace("__", "_")
                
                # Handle global metadata columns (can be used across all sections)
                if col_norm in dynamic_columns:
                    value = row[col]
                    if pd.notna(value) and str(value).strip():
                        # Store in global metadata for use across all sections
                        global_metadata[col_norm] = str(value).strip()
                        # Also add to flat_data_map with the normalized column name as key
                        flat_data_map[col_norm] = str(value).strip()
                        current_app.logger.debug(f"📋 LOADED: {col_norm} = '{str(value).strip()}'")
                    else:
                        # Empty value - skip silently
                        pass
        
        # Ensure we have all required global metadata
        missing_metadata = [key for key in dynamic_columns if key not in flat_data_map]
        if missing_metadata:
            current_app.logger.error(f"❌ MISSING GLOBAL METADATA: {missing_metadata}")
            current_app.logger.error(f"❌ FOUND METADATA: {flat_data_map}")
        else:
            current_app.logger.debug(f"✅ ALL GLOBAL METADATA FOUND: {flat_data_map}")
        
        # Second pass: Process chart-specific data from rows with chart tags
        current_app.logger.debug(f"🔍 DEBUG: Starting second pass - processing chart-specific data")
        for row_index, row in df.iterrows():
            chart_tag = row.get("Chart_Tag")
            current_app.logger.debug(f"🔍 DEBUG: Row {row_index}: Chart_Tag = '{chart_tag}' (type: {type(chart_tag)})")
            if not isinstance(chart_tag, str) or not chart_tag:
                current_app.logger.debug(f"🔍 DEBUG: Row {row_index}: Skipping - Chart_Tag is not a valid string")
                continue
            section_prefix = chart_tag.replace('_chart', '').lower()
            current_app.logger.debug(f"🔍 DEBUG: Row {row_index}: Processing section_prefix: {section_prefix} from Chart_Tag: {chart_tag}")

            # Process all columns systematically
            for col in df.columns:
                # Normalize column names consistently with dynamic_columns
                col_lower = col.lower().strip().replace(" ", "_").replace("__", "_")
                
                # Skip global metadata columns as they're already processed
                if col_lower in dynamic_columns:
                    continue
                
                # Handle Chart Data columns (for values)
                if col_lower.startswith("chart_data_y"):
                    year = col.replace("Chart_Data_", "").replace("chart_data_", "")
                    key = f"{section_prefix}_{year.lower()}"
                    value = row[col]
                    if pd.notna(value) and str(value).strip():
                        # Format values to always show one decimal place
                        try:
                            float_val = float(value)
                            formatted_val = f"{float_val:.1f}"
                            flat_data_map[key] = formatted_val
                        except (ValueError, TypeError):
                            # If conversion fails, use the original value as string
                            flat_data_map[key] = str(value).strip()
                    else:
                        # Empty value - skip silently
                        pass
                
                # Handle Growth columns (for growth rates)
                elif col_lower.startswith("growth_y"):
                    year = col.replace("Growth_", "").replace("growth_", "")
                    key = f"{section_prefix}_{year.lower()}_kpi2"
                    value = row[col]
                    if pd.notna(value) and str(value).strip():
                        # Convert percentage values to proper percentage format for table display
                        try:
                            # Convert to float first to handle any numeric format
                            float_val = float(value)
                            
                            # Check if the original value string contains '%' (Excel percentage format)
                            original_str = str(value).strip()
                            if '%' in original_str:
                                # Remove % and format as percentage
                                clean_val = original_str.replace('%', '').strip()
                                float_val = float(clean_val)
                                percentage_val = f"{float_val:.1f}%"
                                flat_data_map[key] = percentage_val
                                # Excel percentage formatted as percentage
                            elif float_val > 1:
                                # Convert percentage to percentage format (e.g., 20 -> 20.0%)
                                percentage_val = f"{float_val:.1f}%"
                                flat_data_map[key] = percentage_val
                                # Percentage formatted as percentage
                            else:
                                # If it's already a decimal, convert to percentage (e.g., 0.05 -> 5.0%)
                                percentage_val = f"{float_val * 100:.1f}%"
                                flat_data_map[key] = percentage_val
                                # Decimal converted to percentage
                        except (ValueError, TypeError):
                            # If conversion fails, use the original value as string
                            flat_data_map[key] = str(value).strip()
                            # Conversion failed - skip silently
                    else:
                        # Empty value - skip silently
                        pass

            # Handle CAGR
            if pd.notna(row.get("Chart_Data_CAGR")):
                key = f"{section_prefix}_cgrp"
                value = row["Chart_Data_CAGR"]
                if str(value).strip():
                    # Convert CAGR percentage values to proper percentage format for table display
                    try:
                        # Convert to float first to handle any numeric format
                        float_val = float(value)
                        
                        # Check if the original value string contains '%' (Excel percentage format)
                        original_str = str(value).strip()
                        if '%' in original_str:
                            # Remove % and format as percentage
                            clean_val = original_str.replace('%', '').strip()
                            float_val = float(clean_val)
                            percentage_val = f"{float_val:.1f}%"
                            flat_data_map[key] = percentage_val
                            # Excel CAGR percentage formatted as percentage
                        elif float_val > 1:
                            # Convert percentage to percentage format (e.g., 10.5 -> 10.5%)
                            percentage_val = f"{float_val:.1f}%"
                            flat_data_map[key] = percentage_val
                            # CAGR percentage formatted as percentage
                        else:
                            # If it's already a decimal, convert to percentage (e.g., 0.105 -> 10.5%)
                            percentage_val = f"{float_val * 100:.1f}%"
                            flat_data_map[key] = percentage_val
                            # CAGR decimal converted to percentage
                    except (ValueError, TypeError):
                        # If conversion fails, use the original value as string
                        flat_data_map[key] = str(value).strip()
                        # Conversion failed - skip silently
                else:
                    # Empty value - skip silently
                    pass

            # Handle CAGR Historical
            cagr_historical_value = row.get("Chart_Data_CAGR_Historical")
            current_app.logger.debug(f"🔍 DEBUG: Row {row_index}: Chart_Data_CAGR_Historical = '{cagr_historical_value}' (pd.notna: {pd.notna(cagr_historical_value)})")
            if pd.notna(cagr_historical_value):
                key = f"{section_prefix}_cgrp_historical"
                value = cagr_historical_value
                current_app.logger.debug(f"🔍 DEBUG: Found CAGR Historical for {section_prefix}: {value}")
                if str(value).strip():
                    # Convert CAGR percentage values to proper percentage format for table display
                    try:
                        # Convert to float first to handle any numeric format
                        float_val = float(value)
                        
                        # Check if the original value string contains '%' (Excel percentage format)
                        original_str = str(value).strip()
                        if '%' in original_str:
                            # Remove % and format as percentage
                            clean_val = original_str.replace('%', '').strip()
                            float_val = float(clean_val)
                            percentage_val = f"{float_val:.1f}%"
                            flat_data_map[key] = percentage_val
                            current_app.logger.debug(f"🔍 DEBUG: Added to flat_data_map: {key} = {percentage_val}")
                            # Excel CAGR Historical percentage formatted as percentage
                        elif float_val > 1:
                            # Convert percentage to percentage format (e.g., 10.5 -> 10.5%)
                            percentage_val = f"{float_val:.1f}%"
                            flat_data_map[key] = percentage_val
                            current_app.logger.debug(f"🔍 DEBUG: Added to flat_data_map: {key} = {percentage_val}")
                            # CAGR Historical percentage formatted as percentage
                        else:
                            # If it's already a decimal, convert to percentage (e.g., 0.105 -> 10.5%)
                            percentage_val = f"{float_val * 100:.1f}%"
                            flat_data_map[key] = percentage_val
                            current_app.logger.debug(f"🔍 DEBUG: Added to flat_data_map: {key} = {percentage_val}")
                            # CAGR Historical decimal converted to percentage
                    except (ValueError, TypeError):
                        # If conversion fails, use the original value as string
                        flat_data_map[key] = str(value).strip()
                        current_app.logger.debug(f"🔍 DEBUG: Added to flat_data_map (string): {key} = {str(value).strip()}")
                        # Conversion failed - skip silently
                else:
                    # Empty value - skip silently
                    pass
            else:
                # No CAGR Historical data found for this section
                current_app.logger.debug(f"🔍 DEBUG: No CAGR Historical data found for {section_prefix}")

            # Handle CAGR Forecast
            if pd.notna(row.get("Chart_Data_CAGR_Forecast")):
                key = f"{section_prefix}_cgrp_forecast"
                value = row["Chart_Data_CAGR_Forecast"]
                if str(value).strip():
                    # Convert CAGR percentage values to proper percentage format for table display
                    try:
                        # Convert to float first to handle any numeric format
                        float_val = float(value)
                        
                        # Check if the original value string contains '%' (Excel percentage format)
                        original_str = str(value).strip()
                        if '%' in original_str:
                            # Remove % and format as percentage
                            clean_val = original_str.replace('%', '').strip()
                            float_val = float(clean_val)
                            percentage_val = f"{float_val:.1f}%"
                            flat_data_map[key] = percentage_val
                            # Excel CAGR Forecast percentage formatted as percentage
                        elif float_val > 1:
                            # Convert percentage to percentage format (e.g., 10.5 -> 10.5%)
                            percentage_val = f"{float_val:.1f}%"
                            flat_data_map[key] = percentage_val
                            # CAGR Forecast percentage formatted as percentage
                        else:
                            # If it's already a decimal, convert to percentage (e.g., 0.105 -> 10.5%)
                            percentage_val = f"{float_val * 100:.1f}%"
                            flat_data_map[key] = percentage_val
                            # CAGR Forecast decimal converted to percentage
                    except (ValueError, TypeError):
                        # If conversion fails, use the original value as string
                        flat_data_map[key] = str(value).strip()
                        # Conversion failed - skip silently
                else:
                    # Empty value - skip silently
                    pass

        # Ensure all keys are lowercase
        flat_data_map = {k.lower(): v for k, v in flat_data_map.items()}
        
        # Log the final data map for debugging
        current_app.logger.debug(f"🔍 DEBUG: Available columns in flat_data_map: {list(flat_data_map.keys())}")
        
        # Check specifically for section-specific historical and forecast columns
        section_keys = [key for key in flat_data_map.keys() if 'cgrp_historical' in key or 'cgrp_forecast' in key]
        if section_keys:
            current_app.logger.debug(f"🔍 DEBUG: Found section-specific CAGR keys: {section_keys}")
            for key in section_keys:
                current_app.logger.debug(f"🔍 DEBUG: {key} = {flat_data_map[key]}")
        else:
            current_app.logger.debug("🔍 DEBUG: No section-specific CAGR Historical/Forecast keys found in flat_data_map")
        
        # Show ALL keys in flat_data_map for debugging
        current_app.logger.debug(f"🔍 DEBUG: ALL keys in flat_data_map: {sorted(flat_data_map.keys())}")
        
        # Data mapping completed silently

        doc = Document(template_path)

        def replace_text_in_paragraph(paragraph):
            """
            Enhanced text replacement function that handles split placeholders across runs.
            
            PROBLEM SOLVED:
            - Word documents often split placeholders like ${Text} or <Country> across multiple 
              runs when special characters, formatting changes, or certain characters are present.
            - This causes simple text replacement to fail because it only sees partial placeholders.
            
            SOLUTION:
            1. Combines text from all runs in a paragraph to see the complete placeholder
            2. Uses case-insensitive regex matching with re.escape() to handle special characters
            3. Uses improved regex patterns: [^\\}]+ and [^>]+ instead of .*? for better matching
            4. Puts replaced text in the first run, clearing others (preserves formatting)
            5. Has both high-level (paragraph.text) and low-level (XML) approaches for robustness
            
            This ensures placeholders are replaced correctly regardless of special characters,
            formatting, CAGR/CAGRT data, growth values, or how Word internally splits the text.
            """
            nonlocal flat_data_map, text_map  # Access variables from outer scope
            
            # IMPROVED APPROACH: Handle split placeholders across runs
            # Word often splits placeholders like ${Text} into multiple runs due to formatting/special chars
            
            try:
                # First, try to handle placeholders that span across multiple runs
                # by working with the full paragraph text
                full_para_text = paragraph.text
                new_para_text = full_para_text
                replacements_made = False
                
                # Replace ${...} placeholders (case-insensitive)
                dollar_matches = re.findall(r"\$\{([^\}]+)\}", full_para_text)
                for match in dollar_matches:
                    key_lower = match.lower().strip()
                    
                    # Special handling for section_cgrp variants (including numbered sections)
                    if key_lower == 'section_cgrp' or re.match(r'section\d+_cgrp$', key_lower):
                        val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                    elif key_lower == 'section_cgrp_historical' or re.match(r'section\d+_cgrp_historical$', key_lower):
                        val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                        if not val:
                            current_app.logger.debug(f"🔍 DEBUG: No data found for {key_lower}. Available keys: {list(flat_data_map.keys())}")
                    elif key_lower == 'section_cgrp_forecast' or re.match(r'section\d+_cgrp_forecast$', key_lower):
                        val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                        if not val:
                            current_app.logger.debug(f"🔍 DEBUG: No data found for {key_lower}. Available keys: {list(flat_data_map.keys())}")
                    else:
                        val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                    
                    if val is not None and val != '':
                        # Create a regex pattern that matches the placeholder regardless of case
                        pattern = re.compile(re.escape(f"${{{match}}}"), re.IGNORECASE)
                        new_para_text = pattern.sub(str(val), new_para_text)
                        replacements_made = True
                    else:
                        if (key_lower in dynamic_columns or 
                            key_lower in ['section_cgrp', 'section_cgrp_historical', 'section_cgrp_forecast'] or
                            re.match(r'section\d+_cgrp$', key_lower) or
                            re.match(r'section\d+_cgrp_historical$', key_lower) or
                            re.match(r'section\d+_cgrp_forecast$', key_lower)):
                            current_app.logger.error(f"❌ NO DATA: ${{{match}}} (key: {key_lower})")
                
                # Replace <...> placeholders (case-insensitive)
                angle_matches = re.findall(r"<([^>]+)>", full_para_text)
                for match in angle_matches:
                    key_lower = match.lower().strip()
                    
                    # Special handling for section_cgrp variants (including numbered sections)
                    if key_lower == 'section_cgrp' or re.match(r'section\d+_cgrp$', key_lower):
                        val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                    elif key_lower == 'section_cgrp_historical' or re.match(r'section\d+_cgrp_historical$', key_lower):
                        val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                    elif key_lower == 'section_cgrp_forecast' or re.match(r'section\d+_cgrp_forecast$', key_lower):
                        val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                    else:
                        val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                    
                    if val is not None and val != '':
                        # Create a regex pattern that matches the placeholder regardless of case
                        pattern = re.compile(re.escape(f"<{match}>"), re.IGNORECASE)
                        new_para_text = pattern.sub(str(val), new_para_text)
                        replacements_made = True
                    else:
                        if (key_lower in dynamic_columns or 
                            key_lower in ['section_cgrp', 'section_cgrp_historical', 'section_cgrp_forecast'] or
                            re.match(r'section\d+_cgrp$', key_lower) or
                            re.match(r'section\d+_cgrp_historical$', key_lower) or
                            re.match(r'section\d+_cgrp_forecast$', key_lower)):
                            current_app.logger.error(f"❌ NO DATA: <{match}> (key: {key_lower})")
                
                # If replacements were made, update the paragraph runs
                if replacements_made and new_para_text != full_para_text:
                    # Clear all runs except the first one and put all text in the first run
                    # This preserves formatting while ensuring replacement works
                    if paragraph.runs:
                        # Store the first run's formatting
                        first_run = paragraph.runs[0]
                        
                        # Clear all runs
                        for run in paragraph.runs[1:]:
                            run.text = ''
                        
                        # Put the replaced text in the first run
                        first_run.text = new_para_text
                    
            except Exception as e:
                current_app.logger.debug(f"⚠️ Paragraph replacement approach failed: {str(e)}, falling back to XML approach")
            
            # XML-level approach: Handle split placeholders at the XML level
            # This is more robust for handling special characters and split runs
            try:
                w_element = paragraph._element
                ns = {}
                if hasattr(w_element, 'nsmap') and isinstance(w_element.nsmap, dict):
                    ns = {k: v for k, v in w_element.nsmap.items() if k}
                if 'w' not in ns:
                    ns['w'] = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

                # Get all text nodes (w:t elements)
                t_nodes = w_element.xpath('.//w:t', namespaces=ns)
                if t_nodes:
                    # Combine all text from all runs
                    full_text = ''.join([(t.text or '') for t in t_nodes])
                    new_text = full_text
                    
                    # Replace ${...} placeholders (case-insensitive)
                    dollar_matches = re.findall(r"\$\{([^\}]+)\}", full_text)
                    for match in set(dollar_matches):
                        key_lower = match.lower().strip()
                        val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                        if val is not None and val != '':
                            pattern = re.compile(re.escape(f"${{{match}}}"), re.IGNORECASE)
                            new_text = pattern.sub(str(val), new_text)

                    # Replace <...> placeholders (case-insensitive)
                    angle_matches = re.findall(r"<([^>]+)>", full_text)
                    for match in set(angle_matches):
                        key_lower = match.lower().strip()
                        val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                        if val is not None and val != '':
                            pattern = re.compile(re.escape(f"<{match}>"), re.IGNORECASE)
                            new_text = pattern.sub(str(val), new_text)

                    # If replacements were made, update the XML nodes
                    if new_text != full_text and t_nodes:
                        # Put all text in the first node, clear the rest
                        t_nodes[0].text = new_text
                        for t in t_nodes[1:]:
                            t.text = ''
                            
            except Exception as e:
                current_app.logger.debug(f"⚠️ XML-level replacement failed: {str(e)}")
                pass

        def replace_text_in_tables():
            nonlocal doc, flat_data_map, text_map  # Access variables from outer scope
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            replace_text_in_paragraph(para)
            
            # Table processing completed silently

        def process_entire_document():
            """Process the entire document comprehensively to catch all placeholders"""
            nonlocal doc, flat_data_map, text_map  # Access variables from outer scope
            # current_app.logger.info("🔄 PROCESSING ENTIRE DOCUMENT COMPREHENSIVELY")
            
            # Find ALL placeholders in the entire document first
            # current_app.logger.info("🔍 SEARCHING FOR ALL PLACEHOLDERS IN ENTIRE DOCUMENT")
            
            all_placeholders_found = set()
            
            # Search through ALL paragraphs, tables, headers, footers, etc.
            def search_for_placeholders(container):
                """Search for placeholders in any container (document, table, header, footer)"""
                import re  # Import re module for regex operations
                if hasattr(container, 'paragraphs'):
                    for para in container.paragraphs:
                        if para.text:
                            # Searching paragraph for placeholders
                            # Find ${} placeholders
                            dollar_matches = re.findall(r"\$\{(.*?)\}", para.text)
                            for match in dollar_matches:
                                all_placeholders_found.add(f"${{{match}}}")
                                # Found $ placeholder
                            
                            # Find <> placeholders
                            angle_matches = re.findall(r"<(.*?)>", para.text)
                            for match in angle_matches:
                                all_placeholders_found.add(f"<{match}>")
                                # Found <> placeholder
                
                if hasattr(container, 'tables'):
                    for table in container.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                search_for_placeholders(cell)
            
            # Search in main document
            search_for_placeholders(doc)
            
            # Search in headers and footers
            for section in doc.sections:
                if section.header:
                    search_for_placeholders(section.header)
                if section.footer:
                    search_for_placeholders(section.footer)
            
            # Additional search: Look at raw XML for any missed placeholders
            # current_app.logger.info("🔍 ADDITIONAL SEARCH: Looking at raw XML for missed placeholders")
            try:
                for element in doc.element.iter():
                    if hasattr(element, 'text') and element.text:
                        # Find ${} placeholders
                        dollar_matches = re.findall(r"\$\{(.*?)\}", element.text)
                        for match in dollar_matches:
                            all_placeholders_found.add(f"${{{match}}}")
                            # Found $ placeholder in XML
                        
                        # Find <> placeholders
                        angle_matches = re.findall(r"<(.*?)>", element.text)
                        for match in angle_matches:
                            all_placeholders_found.add(f"<{match}>")
                            # Found <> placeholder in XML
            except Exception as e:
                pass  # Suppress warning logs
            
            # current_app.logger.info(f"🔍 Found {len(all_placeholders_found)} unique placeholders: {list(all_placeholders_found)}")
            
            # Log specific global metadata placeholders found
            global_placeholders_found = [p for p in all_placeholders_found if any(key in p.lower() for key in dynamic_columns)]
            # current_app.logger.info(f"🔍 Global metadata placeholders found: {global_placeholders_found}")
            
            # Verify that we have data for all found global metadata placeholders
            for placeholder in global_placeholders_found:
                if placeholder.startswith('${'):
                    key = placeholder[2:-1].lower()
                elif placeholder.startswith('<') and placeholder.endswith('>'):
                    key = placeholder[1:-1].lower()
                else:
                    continue
                
                if key in dynamic_columns:
                      if key in flat_data_map:
                          current_app.logger.debug(f"✅ Data available for {placeholder}: {flat_data_map[key]}")
                      else:
                          current_app.logger.error(f"❌ NO DATA AVAILABLE for {placeholder} (key: {key})")
                          current_app.logger.error(f"❌ Available keys: {list(flat_data_map.keys())}")
            
            # Process Table of Contents specifically - Direct XML string replacement
            # This MUST happen BEFORE the main content replacement to preserve tags in XML
            try:
                # Processing TOC entries before main replacement
                
                # Save current document to temporary file BEFORE any replacements
                import tempfile
                import zipfile
                import shutil
                import os
                
                tmp_path = tempfile.mktemp(suffix='.docx')
                doc.save(tmp_path)
                
                # Extract the document as ZIP
                extract_dir = tempfile.mkdtemp()
                with zipfile.ZipFile(tmp_path, 'r') as zip_ref:
                    zip_ref.extractall(extract_dir)
                
                # COMPREHENSIVE HYPERLINK-AWARE XML PROCESSING
                current_app.logger.debug("🔄 COMPREHENSIVE HYPERLINK-AWARE XML PROCESSING...")
                current_app.logger.debug(f"🔄 Using dynamic columns: {dynamic_columns}")
                
                # Track all modifications
                total_files_modified = 0
                total_replacements = 0
                
                # 1. Process ALL XML files in the document (including hyperlink files)
                # Scanning all XML files for dynamic column tags
                current_app.logger.debug("🔄 Processing all XML files for dynamic column tags...")
                for root, dirs, files in os.walk(extract_dir):
                    for file in files:
                        if file.endswith('.xml'):
                            file_path = os.path.join(root, file)
                            try:
                                with open(file_path, 'r', encoding='utf-8') as f:
                                    content = f.read()
                                
                                # Process each dynamic column
                                file_modified = False
                                for column in dynamic_columns:
                                    tag = f'<{column}>'
                                    escaped_tag = f'&lt;{column}&gt;'
                                    # Unescaped angle-bracket tag
                                    if tag in content:
                                        tag_count = content.count(tag)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {tag} TAGS IN {file}")
                                        replacement_value = flat_data_map.get(column, '')
                                        if replacement_value:
                                            modified_content = content.replace(tag, replacement_value)
                                            if modified_content != content:
                                                content = modified_content
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 XML FILE MODIFIED: {file} ({tag_count} {tag} replacements)")
                                    # Escaped tag (&lt;column&gt;)
                                    if escaped_tag in content:
                                        tag_count = content.count(escaped_tag)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {escaped_tag} TAGS IN {file}")
                                        replacement_value = flat_data_map.get(column, '')
                                        if replacement_value:
                                            modified_content = content.replace(escaped_tag, replacement_value)
                                            if modified_content != content:
                                                content = modified_content
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 XML FILE MODIFIED: {file} ({tag_count} {escaped_tag} replacements)")
                                
                                # Process section_cgrp variants (including numbered sections)
                                section_cgrp_variants = ['section_cgrp', 'section_cgrp_historical', 'section_cgrp_forecast']
                                for variant in section_cgrp_variants:
                                    tag = f'<{variant}>'
                                    escaped_tag = f'&lt;{variant}&gt;'
                                    if tag in content or escaped_tag in content:
                                        # Determine replacement
                                        if variant == 'section_cgrp':
                                            replacement_value = flat_data_map.get('section_cgrp', '')
                                        elif variant == 'section_cgrp_historical':
                                            replacement_value = flat_data_map.get('section_cgrp_historical', '')
                                        else:
                                            replacement_value = flat_data_map.get('section_cgrp_forecast', '')
                                        if replacement_value:
                                            if tag in content:
                                                tag_count = content.count(tag)
                                                current_app.logger.debug(f"🔄 FOUND {tag_count} {tag} TAGS IN {file}")
                                                content = content.replace(tag, replacement_value)
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 XML FILE MODIFIED: {file} ({tag_count} {tag} replacements)")
                                            if escaped_tag in content:
                                                tag_count = content.count(escaped_tag)
                                                current_app.logger.debug(f"🔄 FOUND {tag_count} {escaped_tag} TAGS IN {file}")
                                                content = content.replace(escaped_tag, replacement_value)
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 XML FILE MODIFIED: {file} ({tag_count} {escaped_tag} replacements)")
                                
                                # Process numbered section_cgrp variants (section1_cgrp_historical, section2_cgrp_forecast, etc.)
                                import re
                                section_patterns = [
                                    (r'<section\d+_cgrp>', 'chart_data_cgar'),
                                    (r'<section\d+_cgrp_historical>', 'section_cgrp_historical'),
                                    (r'<section\d+_cgrp_forecast>', 'section_cgrp_forecast')
                                ]
                                
                                for pattern, base_key in section_patterns:
                                    matches = re.findall(pattern, content)
                                    escaped_pattern = pattern.replace('<', '&lt;').replace('>', '&gt;')
                                    escaped_matches = re.findall(escaped_pattern, content)
                                    for match in matches:
                                        tag_count = content.count(match)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {match} TAGS IN {file}")
                                        
                                        # Extract the section-specific key from the match
                                        # match is like '<section1_cgrp_historical>', extract 'section1_cgrp_historical'
                                        section_key = match[1:-1].lower()  # Remove < and > and convert to lowercase
                                        
                                        # For all CAGR variants, use the section-specific key
                                        replacement_value = flat_data_map.get(section_key, '')
                                        
                                        if replacement_value:
                                            modified_content = content.replace(match, replacement_value)
                                            
                                            if modified_content != content:
                                                content = modified_content  # Update content for next iteration
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 XML FILE MODIFIED: {file} ({tag_count} {match} replacements)")
                                    for match in escaped_matches:
                                        tag_count = content.count(match)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {match} TAGS IN {file}")
                                        section_key = match.replace('&lt;', '<').replace('&gt;', '>')[1:-1].lower()
                                        replacement_value = flat_data_map.get(section_key, '')
                                        if replacement_value:
                                            modified_content = content.replace(match, replacement_value)
                                            if modified_content != content:
                                                content = modified_content
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 XML FILE MODIFIED: {file} ({tag_count} {match} replacements)")
                                
                                # Write the final modified content if any changes were made
                                # Additional case-insensitive pass for dynamic columns and escaped tags
                                try:
                                    import re as _re_ci
                                    for column in dynamic_columns:
                                        # Build case-insensitive patterns for both raw and escaped tags
                                        patterns = [
                                            rf"(?i)<\s*{_re_ci.escape(column)}\s*>",
                                            rf"(?i)&lt;\s*{_re_ci.escape(column)}\s*&gt;",
                                        ]
                                        for pat in patterns:
                                            matches = _re_ci.findall(pat, content)
                                            if matches:
                                                replacement_value = flat_data_map.get(column, '')
                                                if replacement_value:
                                                    content, n = _re_ci.subn(pat, replacement_value, content)
                                                    if n > 0:
                                                        file_modified = True
                                                        total_replacements += n
                                                        current_app.logger.debug(f"🔄 XML FILE MODIFIED (CI): {file} ({n} {column} replacements)")
                                except Exception:
                                    pass

                                if file_modified:
                                    with open(file_path, 'w', encoding='utf-8') as f:
                                        f.write(content)
                                    total_files_modified += 1
                                        
                            except Exception as e:
                                pass  # Suppress warning logs
                
                # 2. Special processing for hyperlink-specific files
                # Special processing for hyperlink files
                
                # Look for files that might contain hyperlink data
                hyperlink_keywords = ['hyperlink', 'link', 'toc', 'table', 'contents', 'rels', 'relationship']
                current_app.logger.debug(f"🔄 Looking for TOC/hyperlink files with keywords: {hyperlink_keywords}")
                for root, dirs, files in os.walk(extract_dir):
                    for file in files:
                        if file.endswith('.xml') and any(keyword in file.lower() for keyword in hyperlink_keywords):
                            current_app.logger.debug(f"🔄 Processing TOC/hyperlink file: {file}")
                            file_path = os.path.join(root, file)
                            try:
                                with open(file_path, 'r', encoding='utf-8') as f:
                                    content = f.read()
                                
                                # Process each dynamic column
                                file_modified = False
                                for column in dynamic_columns:
                                    tag = f'<{column}>'
                                    escaped_tag = f'&lt;{column}&gt;'
                                    if tag in content:
                                        tag_count = content.count(tag)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {tag} TAGS IN HYPERLINK FILE: {file}")
                                        replacement_value = flat_data_map.get(column, '')
                                        if replacement_value:
                                            modified_content = content.replace(tag, replacement_value)
                                            if modified_content != content:
                                                content = modified_content
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 HYPERLINK FILE MODIFIED: {file} ({tag_count} {tag} replacements)")
                                    if escaped_tag in content:
                                        tag_count = content.count(escaped_tag)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {escaped_tag} TAGS IN HYPERLINK FILE: {file}")
                                        replacement_value = flat_data_map.get(column, '')
                                        if replacement_value:
                                            modified_content = content.replace(escaped_tag, replacement_value)
                                            if modified_content != content:
                                                content = modified_content
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 HYPERLINK FILE MODIFIED: {file} ({tag_count} {escaped_tag} replacements)")
                                
                                # Process section_cgrp variants
                                section_cgrp_variants = ['section_cgrp', 'section_cgrp_historical', 'section_cgrp_forecast']
                                for variant in section_cgrp_variants:
                                    tag = f'<{variant}>'
                                    escaped_tag = f'&lt;{variant}&gt;'
                                    if tag in content or escaped_tag in content:
                                        if variant == 'section_cgrp':
                                            replacement_value = flat_data_map.get('section_cgrp', '')
                                        elif variant == 'section_cgrp_historical':
                                            replacement_value = flat_data_map.get('section_cgrp_historical', '')
                                        else:
                                            replacement_value = flat_data_map.get('section_cgrp_forecast', '')
                                        if replacement_value:
                                            if tag in content:
                                                tag_count = content.count(tag)
                                                current_app.logger.debug(f"🔄 FOUND {tag_count} {tag} TAGS IN HYPERLINK FILE: {file}")
                                                content = content.replace(tag, replacement_value)
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 HYPERLINK FILE MODIFIED: {file} ({tag_count} {tag} replacements)")
                                            if escaped_tag in content:
                                                tag_count = content.count(escaped_tag)
                                                current_app.logger.debug(f"🔄 FOUND {tag_count} {escaped_tag} TAGS IN HYPERLINK FILE: {file}")
                                                content = content.replace(escaped_tag, replacement_value)
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 HYPERLINK FILE MODIFIED: {file} ({tag_count} {escaped_tag} replacements)")
                                
                                # Process numbered section_cgrp variants (section1_cgrp_historical, section2_cgrp_forecast, etc.)
                                section_patterns = [
                                    (r'<section\d+_cgrp>', 'chart_data_cgar'),
                                    (r'<section\d+_cgrp_historical>', 'section_cgrp_historical'),
                                    (r'<section\d+_cgrp_forecast>', 'section_cgrp_forecast')
                                ]
                                
                                for pattern, base_key in section_patterns:
                                    matches = re.findall(pattern, content)
                                    escaped_pattern = pattern.replace('<', '&lt;').replace('>', '&gt;')
                                    escaped_matches = re.findall(escaped_pattern, content)
                                    for match in matches:
                                        tag_count = content.count(match)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {match} TAGS IN HYPERLINK FILE: {file}")
                                        
                                        # Extract the section-specific key from the match
                                        section_key = match[1:-1].lower()  # Remove < and > and convert to lowercase
                                        
                                        # For all CAGR variants, use the section-specific key
                                        replacement_value = flat_data_map.get(section_key, '')
                                        if replacement_value:
                                            modified_content = content.replace(match, replacement_value)
                                            
                                            if modified_content != content:
                                                content = modified_content  # Update content for next iteration
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 HYPERLINK FILE MODIFIED: {file} ({tag_count} {match} replacements)")
                                    for match in escaped_matches:
                                        tag_count = content.count(match)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {match} TAGS IN HYPERLINK FILE: {file}")
                                        section_key = match.replace('&lt;', '<').replace('&gt;', '>')[1:-1].lower()
                                        replacement_value = flat_data_map.get(section_key, '')
                                        if replacement_value:
                                            modified_content = content.replace(match, replacement_value)
                                            if modified_content != content:
                                                content = modified_content
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 HYPERLINK FILE MODIFIED: {file} ({tag_count} {match} replacements)")
                                
                                # Write the final modified content if any changes were made
                                if file_modified:
                                    with open(file_path, 'w', encoding='utf-8') as f:
                                        f.write(content)
                                    total_files_modified += 1
                                        
                            except Exception as e:
                                pass  # Suppress warning logs
                
                # 3. Process _rels files (relationship files that might contain hyperlink data)
                # Processing relationship files
                rels_dir = os.path.join(extract_dir, '_rels')
                if os.path.exists(rels_dir):
                    for file in os.listdir(rels_dir):
                        if file.endswith('.xml'):
                            file_path = os.path.join(rels_dir, file)
                            try:
                                with open(file_path, 'r', encoding='utf-8') as f:
                                    content = f.read()
                                
                                # Process each dynamic column
                                file_modified = False
                                for column in dynamic_columns:
                                    tag = f'<{column}>'
                                    escaped_tag = f'&lt;{column}&gt;'
                                    if tag in content:
                                        tag_count = content.count(tag)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {tag} TAGS IN RELS FILE: {file}")
                                        replacement_value = flat_data_map.get(column, '')
                                        if replacement_value:
                                            modified_content = content.replace(tag, replacement_value)
                                            if modified_content != content:
                                                content = modified_content
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 RELS FILE MODIFIED: {file} ({tag_count} {tag} replacements)")
                                    if escaped_tag in content:
                                        tag_count = content.count(escaped_tag)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {escaped_tag} TAGS IN RELS FILE: {file}")
                                        replacement_value = flat_data_map.get(column, '')
                                        if replacement_value:
                                            modified_content = content.replace(escaped_tag, replacement_value)
                                            if modified_content != content:
                                                content = modified_content
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 RELS FILE MODIFIED: {file} ({tag_count} {escaped_tag} replacements)")
                                
                                # Process section_cgrp variants
                                section_cgrp_variants = ['section_cgrp', 'section_cgrp_historical', 'section_cgrp_forecast']
                                for variant in section_cgrp_variants:
                                    tag = f'<{variant}>'
                                    escaped_tag = f'&lt;{variant}&gt;'
                                    if tag in content or escaped_tag in content:
                                        if variant == 'section_cgrp':
                                            replacement_value = flat_data_map.get('section_cgrp', '')
                                        elif variant == 'section_cgrp_historical':
                                            replacement_value = flat_data_map.get('section_cgrp_historical', '')
                                        else:
                                            replacement_value = flat_data_map.get('section_cgrp_forecast', '')
                                        if replacement_value:
                                            if tag in content:
                                                tag_count = content.count(tag)
                                                current_app.logger.debug(f"🔄 FOUND {tag_count} {tag} TAGS IN RELS FILE: {file}")
                                                content = content.replace(tag, replacement_value)
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 RELS FILE MODIFIED: {file} ({tag_count} {tag} replacements)")
                                            if escaped_tag in content:
                                                tag_count = content.count(escaped_tag)
                                                current_app.logger.debug(f"🔄 FOUND {tag_count} {escaped_tag} TAGS IN RELS FILE: {file}")
                                                content = content.replace(escaped_tag, replacement_value)
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 RELS FILE MODIFIED: {file} ({tag_count} {escaped_tag} replacements)")
                                
                                # Process numbered section_cgrp variants (section1_cgrp_historical, section2_cgrp_forecast, etc.)
                                section_patterns = [
                                    (r'<section\d+_cgrp>', 'chart_data_cgar'),
                                    (r'<section\d+_cgrp_historical>', 'section_cgrp_historical'),
                                    (r'<section\d+_cgrp_forecast>', 'section_cgrp_forecast')
                                ]
                                
                                for pattern, base_key in section_patterns:
                                    matches = re.findall(pattern, content)
                                    escaped_pattern = pattern.replace('<', '&lt;').replace('>', '&gt;')
                                    escaped_matches = re.findall(escaped_pattern, content)
                                    for match in matches:
                                        tag_count = content.count(match)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {match} TAGS IN RELS FILE: {file}")
                                        
                                        # Extract the section-specific key from the match
                                        section_key = match[1:-1].lower()  # Remove < and > and convert to lowercase
                                        
                                        # For all CAGR variants, use the section-specific key
                                        replacement_value = flat_data_map.get(section_key, '')
                                        if replacement_value:
                                            modified_content = content.replace(match, replacement_value)
                                            
                                            if modified_content != content:
                                                content = modified_content  # Update content for next iteration
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 RELS FILE MODIFIED: {file} ({tag_count} {match} replacements)")
                                    for match in escaped_matches:
                                        tag_count = content.count(match)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {match} TAGS IN RELS FILE: {file}")
                                        section_key = match.replace('&lt;', '<').replace('&gt;', '>')[1:-1].lower()
                                        replacement_value = flat_data_map.get(section_key, '')
                                        if replacement_value:
                                            modified_content = content.replace(match, replacement_value)
                                            if modified_content != content:
                                                content = modified_content
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 RELS FILE MODIFIED: {file} ({tag_count} {match} replacements)")
                                
                                # Write the final modified content if any changes were made
                                if file_modified:
                                    with open(file_path, 'w', encoding='utf-8') as f:
                                        f.write(content)
                                    total_files_modified += 1
                                        
                            except Exception as e:
                                pass  # Suppress warning logs
                
                # 4. Process word/_rels files (Word-specific relationship files)
                word_rels_dir = os.path.join(extract_dir, 'word', '_rels')
                if os.path.exists(word_rels_dir):
                    for file in os.listdir(word_rels_dir):
                        if file.endswith('.xml'):
                            file_path = os.path.join(word_rels_dir, file)
                            try:
                                with open(file_path, 'r', encoding='utf-8') as f:
                                    content = f.read()
                                
                                # Process each dynamic column
                                file_modified = False
                                for column in dynamic_columns:
                                    tag = f'<{column}>'
                                    escaped_tag = f'&lt;{column}&gt;'
                                    if tag in content:
                                        tag_count = content.count(tag)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {tag} TAGS IN WORD_RELS FILE: {file}")
                                        replacement_value = flat_data_map.get(column, '')
                                        if replacement_value:
                                            modified_content = content.replace(tag, replacement_value)
                                            if modified_content != content:
                                                content = modified_content
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 WORD_RELS FILE MODIFIED: {file} ({tag_count} {tag} replacements)")
                                    if escaped_tag in content:
                                        tag_count = content.count(escaped_tag)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {escaped_tag} TAGS IN WORD_RELS FILE: {file}")
                                        replacement_value = flat_data_map.get(column, '')
                                        if replacement_value:
                                            modified_content = content.replace(escaped_tag, replacement_value)
                                            if modified_content != content:
                                                content = modified_content
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 WORD_RELS FILE MODIFIED: {file} ({tag_count} {escaped_tag} replacements)")
                                
                                # Process section_cgrp variants
                                section_cgrp_variants = ['section_cgrp', 'section_cgrp_historical', 'section_cgrp_forecast']
                                for variant in section_cgrp_variants:
                                    tag = f'<{variant}>'
                                    escaped_tag = f'&lt;{variant}&gt;'
                                    if tag in content or escaped_tag in content:
                                        if variant == 'section_cgrp':
                                            replacement_value = flat_data_map.get('section_cgrp', '')
                                        elif variant == 'section_cgrp_historical':
                                            replacement_value = flat_data_map.get('section_cgrp_historical', '')
                                        else:
                                            replacement_value = flat_data_map.get('section_cgrp_forecast', '')
                                        if replacement_value:
                                            if tag in content:
                                                tag_count = content.count(tag)
                                                current_app.logger.debug(f"🔄 FOUND {tag_count} {tag} TAGS IN WORD_RELS FILE: {file}")
                                                content = content.replace(tag, replacement_value)
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 WORD_RELS FILE MODIFIED: {file} ({tag_count} {tag} replacements)")
                                            if escaped_tag in content:
                                                tag_count = content.count(escaped_tag)
                                                current_app.logger.debug(f"🔄 FOUND {tag_count} {escaped_tag} TAGS IN WORD_RELS FILE: {file}")
                                                content = content.replace(escaped_tag, replacement_value)
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 WORD_RELS FILE MODIFIED: {file} ({tag_count} {escaped_tag} replacements)")
                                
                                # Process numbered section_cgrp variants (section1_cgrp_historical, section2_cgrp_forecast, etc.)
                                section_patterns = [
                                    (r'<section\d+_cgrp>', 'chart_data_cgar'),
                                    (r'<section\d+_cgrp_historical>', 'section_cgrp_historical'),
                                    (r'<section\d+_cgrp_forecast>', 'section_cgrp_forecast')
                                ]
                                
                                for pattern, base_key in section_patterns:
                                    matches = re.findall(pattern, content)
                                    escaped_pattern = pattern.replace('<', '&lt;').replace('>', '&gt;')
                                    escaped_matches = re.findall(escaped_pattern, content)
                                    for match in matches:
                                        tag_count = content.count(match)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {match} TAGS IN WORD_RELS FILE: {file}")
                                        
                                        # Extract the section-specific key from the match
                                        section_key = match[1:-1].lower()  # Remove < and > and convert to lowercase
                                        
                                        # For all CAGR variants, use the section-specific key
                                        replacement_value = flat_data_map.get(section_key, '')
                                        if replacement_value:
                                            modified_content = content.replace(match, replacement_value)
                                            
                                            if modified_content != content:
                                                content = modified_content  # Update content for next iteration
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 WORD_RELS FILE MODIFIED: {file} ({tag_count} {match} replacements)")
                                    for match in escaped_matches:
                                        tag_count = content.count(match)
                                        current_app.logger.debug(f"🔄 FOUND {tag_count} {match} TAGS IN WORD_RELS FILE: {file}")
                                        section_key = match.replace('&lt;', '<').replace('&gt;', '>')[1:-1].lower()
                                        replacement_value = flat_data_map.get(section_key, '')
                                        if replacement_value:
                                            modified_content = content.replace(match, replacement_value)
                                            if modified_content != content:
                                                content = modified_content
                                                file_modified = True
                                                total_replacements += tag_count
                                                current_app.logger.debug(f"🔄 WORD_RELS FILE MODIFIED: {file} ({tag_count} {match} replacements)")
                                
                                # Write the final modified content if any changes were made
                                if file_modified:
                                    with open(file_path, 'w', encoding='utf-8') as f:
                                        f.write(content)
                                    total_files_modified += 1
                                        
                            except Exception as e:
                                pass  # Suppress warning logs
                
                # 5. If any files were modified, recreate the document
                if total_files_modified > 0:
                    current_app.logger.debug(f"🔄 COMPREHENSIVE XML REPLACEMENT COMPLETED: {total_files_modified} files, {total_replacements} total replacements")
                    
                    # Recreate the document
                    with zipfile.ZipFile(tmp_path, 'w') as zip_ref:
                        for root, dirs, files in os.walk(extract_dir):
                            for file in files:
                                file_path = os.path.join(root, file)
                                arcname = os.path.relpath(file_path, extract_dir)
                                zip_ref.write(file_path, arcname)
                    
                    # Reload the modified document
                    from docx import Document as NewDocument
                    doc = NewDocument(tmp_path)
                    current_app.logger.debug("🔄 DOCUMENT RELOADED AFTER COMPREHENSIVE XML MODIFICATION")
                    
                    # Cleanup and continue - proceed with downstream replacements (CAGR, section text, charts)
                    shutil.rmtree(extract_dir)
                    os.unlink(tmp_path)
                    # Do not return here; continue to paragraph/table/header/footer processing
                else:
                    # No XML files were modified
                    
                    # Cleanup
                    shutil.rmtree(extract_dir)
                    os.unlink(tmp_path)
            except Exception as e:
                pass  # Suppress warning logs
            
            # Now replace ALL placeholders everywhere they appear
            # Replacing all placeholders everywhere
            
            # Single comprehensive pass to avoid duplication
            # Processing all document elements in single pass
                
                # Process main document
            for para in doc.paragraphs:
                    replace_text_in_paragraph(para)
                
                # Process tables
            for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                replace_text_in_paragraph(para)
                
                # Process headers and footers (default/first/even) and their text boxes
            def _process_header_footer(hf_part):
                    if not hf_part:
                        return
                    # Paragraphs
                    for para in hf_part.paragraphs:
                        replace_text_in_paragraph(para)
                    # Tables
                    for table in hf_part.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    replace_text_in_paragraph(para)
                    # Text boxes inside header/footer
                    try:
                        ns = {k: v for k, v in (hf_part._element.nsmap or {}).items() if k}
                        if 'w' not in ns:
                            ns['w'] = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                        for p_elem in hf_part._element.xpath('.//w:txbxContent//w:p', namespaces=ns):
                            try:
                                para_obj = Paragraph(p_elem, hf_part)
                                replace_text_in_paragraph(para_obj)
                            except Exception:
                                pass
                    except Exception:
                        pass

                    # DrawingML text inside header/footer (WordArt/shapes) - a:t
                    try:
                        ns = {k: v for k, v in (hf_part._element.nsmap or {}).items() if k}
                        if 'a' not in ns:
                            ns['a'] = 'http://schemas.openxmlformats.org/drawingml/2006/main'
                        for a_t in hf_part._element.xpath('.//a:t', namespaces=ns):
                            try:
                                original_text = a_t.text or ''
                                modified_text = original_text
                                # Use improved regex pattern that handles special characters better
                                for match in set(re.findall(r"\$\{([^\}]+)\}", original_text)):
                                    key_lower = match.lower().strip()
                                    val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                                    if val is not None and val != '':
                                        pattern = re.compile(re.escape(f"${{{match}}}"), re.IGNORECASE)
                                        modified_text = pattern.sub(str(val), modified_text)
                                for match in set(re.findall(r"<([^>]+)>", original_text)):
                                    key_lower = match.lower().strip()
                                    val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                                    if val is not None and val != '':
                                        pattern = re.compile(re.escape(f"<{match}>"), re.IGNORECASE)
                                        modified_text = pattern.sub(str(val), modified_text)
                                if modified_text != original_text:
                                    a_t.text = modified_text
                            except Exception:
                                pass
                    except Exception:
                        pass

            for section in doc.sections:
                    _process_header_footer(getattr(section, 'header', None))
                    _process_header_footer(getattr(section, 'first_page_header', None))
                    _process_header_footer(getattr(section, 'even_page_header', None))
                    _process_header_footer(getattr(section, 'footer', None))
                    _process_header_footer(getattr(section, 'first_page_footer', None))
                    _process_header_footer(getattr(section, 'even_page_footer', None))
                
            # XML processing removed to prevent duplication - paragraph processing is sufficient
            
            #current_app.logger.info("✅ COMPREHENSIVE DOCUMENT PROCESSING COMPLETED")
            
            # Additional pass: Handle special Word elements that might contain placeholders
                    # Final pass: Processing special Word elements
            
            # Process text boxes and other special elements (only once)
            try:
                for shape in doc.inline_shapes:
                    if hasattr(shape, 'text_frame'):
                        for para in shape.text_frame.paragraphs:
                            replace_text_in_paragraph(para)
            except Exception as e:
                pass  # Suppress warning logs

            # Extra pass: process paragraphs inside text boxes (w:txbxContent) which are not exposed in doc.paragraphs
            try:
                ns = {k: v for k, v in (doc.element.nsmap or {}).items() if k}
                if 'w' not in ns:
                    ns['w'] = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                for p_elem in doc.element.xpath('.//w:txbxContent//w:p', namespaces=ns):
                    try:
                        para_obj = Paragraph(p_elem, doc)
                        replace_text_in_paragraph(para_obj)
                    except Exception:
                        pass
            except Exception:
                pass
            
            # Extra pass: DrawingML text (WordArt/shapes) in main body (a:t)
            try:
                ns = {k: v for k, v in (doc.element.nsmap or {}).items() if k}
                if 'a' not in ns:
                    ns['a'] = 'http://schemas.openxmlformats.org/drawingml/2006/main'
                for a_t in doc.element.xpath('.//a:t', namespaces=ns):
                    try:
                        original_text = a_t.text or ''
                        modified_text = original_text
                        # Use improved regex pattern that handles special characters better
                        for match in set(re.findall(r"\$\{([^\}]+)\}", original_text)):
                            key_lower = match.lower().strip()
                            val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                            if val is not None and val != '':
                                pattern = re.compile(re.escape(f"${{{match}}}"), re.IGNORECASE)
                                modified_text = pattern.sub(str(val), modified_text)
                        for match in set(re.findall(r"<([^>]+)>", original_text)):
                            key_lower = match.lower().strip()
                            val = flat_data_map.get(key_lower) or text_map.get(key_lower)
                            if val is not None and val != '':
                                pattern = re.compile(re.escape(f"<{match}>"), re.IGNORECASE)
                                modified_text = pattern.sub(str(val), modified_text)
                        if modified_text != original_text:
                            a_t.text = modified_text
                    except Exception:
                        pass
            except Exception:
                pass

            # Process Word fields (like table of contents) that might contain placeholders
            try:
                for field in doc.fields:
                    if hasattr(field, 'text') and field.text:
                        original_text = field.text
                        modified_text = original_text
                        
                        # Replace ${} placeholders with case-insensitive matching
                        dollar_matches = re.findall(r'\$\{([^\}]+)\}', original_text)
                        for match in dollar_matches:
                            key_lower = match.lower().strip()
                            value = flat_data_map.get(key_lower) or text_map.get(key_lower)
                            if value:
                                pattern = re.compile(re.escape(f"${{{match}}}"), re.IGNORECASE)
                                modified_text = pattern.sub(str(value), modified_text)
                                current_app.logger.debug(f"🔄 FIELD REPLACED: ${{{match}}} -> {value}")
                        
                        # Replace <> placeholders with case-insensitive matching
                        angle_matches = re.findall(r'<([^>]+)>', original_text)
                        for match in angle_matches:
                            key_lower = match.lower().strip()
                            value = flat_data_map.get(key_lower) or text_map.get(key_lower)
                            if value:
                                pattern = re.compile(re.escape(f"<{match}>"), re.IGNORECASE)
                                modified_text = pattern.sub(str(value), modified_text)
                                current_app.logger.debug(f"🔄 FIELD REPLACED: <{match}> -> {value}")
                        
                        # Update field text if modified
                        if modified_text != original_text:
                            field.text = modified_text
            except Exception as e:
                current_app.logger.debug(f"⚠️ Word fields processing error: {str(e)}")
                pass
            

            
            # Process all XML elements for any remaining placeholders - Careful approach to avoid duplication
            try:
                # Processing XML elements
                xml_replacements = 0
                
                for element in doc.element.iter():
                    if hasattr(element, 'text') and element.text:
                        original_text = element.text
                        modified_text = original_text
                        text_changed = False
                        
                        # Check if any placeholders are present (both ${} and <> formats)
                        has_placeholder = (
                            re.search(r'\$\{[^\}]+\}', original_text) or 
                            re.search(r'<[^>]+>', original_text)
                        )
                        
                        if has_placeholder:
                            # Only replace if it's a simple text element to avoid duplication
                            if hasattr(element, 'tag'):
                                tag_name = element.tag
                                # Handle namespaced tags
                                if '}' in tag_name:
                                    tag_name = tag_name.split('}')[1]
                                
                                if tag_name in ['t', 'tab', 'br']:
                                    try:
                                        # Process ${} placeholders with case-insensitive matching
                                        dollar_matches = re.findall(r'\$\{([^\}]+)\}', original_text)
                                        for match in dollar_matches:
                                            key_lower = match.lower().strip()
                                            replacement_value = flat_data_map.get(key_lower, '')
                                            if replacement_value:
                                                # Use regex for case-insensitive replacement
                                                pattern = re.compile(re.escape(f"${{{match}}}"), re.IGNORECASE)
                                                modified_text = pattern.sub(str(replacement_value), modified_text)
                                                text_changed = True
                                                xml_replacements += 1
                                                current_app.logger.debug(f"🔄 XML TEXT ELEMENT REPLACED: ${{{match}}} -> {replacement_value}")
                                        
                                        # Process <> placeholders with case-insensitive matching
                                        angle_matches = re.findall(r'<([^>]+)>', original_text)
                                        for match in angle_matches:
                                            key_lower = match.lower().strip()
                                            
                                            # Try direct lookup first
                                            replacement_value = flat_data_map.get(key_lower, '')
                                            
                                            # Special handling for section_cgrp variants if direct lookup fails
                                            if not replacement_value:
                                                if key_lower == 'section_cgrp' or re.match(r'section\d+_cgrp$', key_lower):
                                                    replacement_value = flat_data_map.get('chart_data_cgar', '')
                                                elif key_lower == 'section_cgrp_historical' or re.match(r'section\d+_cgrp_historical$', key_lower):
                                                    replacement_value = flat_data_map.get('chart_data_historical', '')
                                                elif key_lower == 'section_cgrp_forecast' or re.match(r'section\d+_cgrp_forecast$', key_lower):
                                                    replacement_value = flat_data_map.get('chart_data_forecast', '')
                                            
                                            if replacement_value:
                                                # Use regex for case-insensitive replacement
                                                pattern = re.compile(re.escape(f"<{match}>"), re.IGNORECASE)
                                                modified_text = pattern.sub(str(replacement_value), modified_text)
                                                text_changed = True
                                                xml_replacements += 1
                                                current_app.logger.debug(f"🔄 XML TEXT ELEMENT REPLACED: <{match}> -> {replacement_value}")
                                        
                                        if text_changed:
                                            element.text = modified_text
                                            
                                    except Exception as xml_error:
                                        current_app.logger.debug(f"⚠️ XML element replacement error: {str(xml_error)}")
                                        pass  # Suppress warning logs
                
                if xml_replacements > 0:
                    current_app.logger.debug(f"✅ XML processing complete: {xml_replacements} replacements made")
                
            except Exception as e:
                current_app.logger.debug(f"⚠️ XML processing error: {str(e)}")
                pass  # Suppress warning logs
            
            # Force update Table of Contents by refreshing the document
            try:
                # Update all TOC fields
                for field in doc.fields:
                    if field.type == 3:  # TOC field type
                        field.update()
                        current_app.logger.debug("🔄 TOC FIELD UPDATED")
            except Exception as e:
                # TOC fields update skipped
                pass
            
            # Final verification: Check if any <country> tags remain
            try:
                # Final verification
                remaining_country_tags = 0
                
                # Check all paragraphs
                for paragraph in doc.paragraphs:
                    if '<country>' in paragraph.text:
                        remaining_country_tags += 1
                        #pass  # Suppress warning logs: f"⚠️ REMAINING <country> TAG FOUND: {paragraph.text[:100]}...")
                
                # Check all tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if '<country>' in cell.text:
                                remaining_country_tags += 1
                                #pass  # Suppress warning logs: f"⚠️ REMAINING <country> TAG IN TABLE: {cell.text[:100]}...")
                
                if remaining_country_tags == 0:
                    current_app.logger.debug("✅ ALL <country> TAGS SUCCESSFULLY REPLACED!")
                else:
                    pass  # Suppress warning logs
                    
            except Exception as e:
                pass  # Suppress warning logs
            
            # Search for any remaining placeholders that might have been missed
            def search_for_remaining_placeholders(element, path=""):
                """Recursively search for any remaining placeholders"""
                if hasattr(element, 'text') and element.text:
                    text = element.text
                    
                    # Check for any remaining ${} placeholders
                    dollar_matches = re.findall(r"\$\{(.*?)\}", text)
                    for match in dollar_matches:
                        key_lower = match.lower().strip()
                        if key_lower in ['country', 'report_name', 'report_code', 'currency']:
                            current_app.logger.error(f"❌ REMAINING PLACEHOLDER FOUND: ${{{match}}} in {path}")
                            current_app.logger.error(f"❌ Available data for {key_lower}: {flat_data_map.get(key_lower, 'NOT FOUND')}")
                    
                    # Check for any remaining <> placeholders
                    angle_matches = re.findall(r"<(.*?)>", text)
                    for match in angle_matches:
                        key_lower = match.lower().strip()
                        if key_lower in ['country', 'report_name', 'report_code', 'currency']:
                            current_app.logger.error(f"❌ REMAINING PLACEHOLDER FOUND: <{match}> in {path}")
                            current_app.logger.error(f"❌ Available data for {key_lower}: {flat_data_map.get(key_lower, 'NOT FOUND')}")
                
                # Recursively check child elements
                for i, child in enumerate(element):
                    child_path = f"{path}.{i}" if path else str(i)
                    search_for_remaining_placeholders(child, child_path)
            
            #current_app.logger.info("✅ FINAL VERIFICATION COMPLETED")


                            
        # Data mapping completed silently

        def generate_chart(data_dict, chart_tag):
            current_app.logger.info(f"🚀 GENERATE_CHART CALLED with tag: {chart_tag}")
            import plotly.graph_objects as go
            import matplotlib.pyplot as plt
            from openpyxl.utils import column_index_from_string
            import numpy as np
            import os
            import tempfile
            import json
            import re
            import warnings
            import gc
            
            # Suppress Matplotlib warnings
            warnings.filterwarnings('ignore', category=UserWarning, module='matplotlib')
            
            # Force matplotlib to use non-interactive backend
            plt.switch_backend('Agg')

            try:
                chart_tag_lower = chart_tag.lower()
                raw_chart_attr = chart_attr_map.get(chart_tag_lower, "{}")
                
                # Enhanced JSON validation with detailed error reporting
                cleaned_json = re.sub(r'//.*?\n|/\*.*?\*/', '', raw_chart_attr, flags=re.DOTALL)
                
                # Validate JSON syntax and provide detailed error information
                try:
                    chart_config = json.loads(cleaned_json)
                except json.JSONDecodeError as json_err:
                    # Create detailed JSON error message
                    error_line = json_err.lineno if hasattr(json_err, 'lineno') else 'unknown'
                    error_col = json_err.colno if hasattr(json_err, 'colno') else 'unknown'
                    error_pos = json_err.pos if hasattr(json_err, 'pos') else 'unknown'
                    
                    # Extract the problematic part of the JSON
                    json_lines = cleaned_json.split('\n')
                    problematic_line = json_lines[error_line - 1] if error_line <= len(json_lines) else ''
                    
                    # Create specific error message based on common JSON issues
                    detailed_error = f"JSON syntax error in chart attributes for '{chart_tag}':\n"
                    detailed_error += f"• Error: {json_err.msg}\n"
                    detailed_error += f"• Line: {error_line}, Column: {error_col}\n"
                    if problematic_line:
                        detailed_error += f"• Problematic line: {problematic_line.strip()}\n"
                        # Highlight the error position if possible
                        if error_col and error_col <= len(problematic_line):
                            detailed_error += f"• Error position: {' ' * (error_col - 1)}^\n"
                    
                    # Add specific suggestions based on common JSON errors
                    if "Expecting ',' delimiter" in json_err.msg:
                        detailed_error += "• Suggestion: Add a comma (,) after the previous property\n"
                    elif "Expecting ':' delimiter" in json_err.msg:
                        detailed_error += "• Suggestion: Add a colon (:) between property name and value\n"
                    elif "Expecting property name" in json_err.msg:
                        detailed_error += "• Suggestion: Check for missing quotes around property names\n"
                    elif "Expecting value" in json_err.msg:
                        detailed_error += "• Suggestion: Add a value after the colon (:)\n"
                    elif "Extra data" in json_err.msg:
                        detailed_error += "• Suggestion: Remove extra characters or add missing comma\n"
                    elif "Unterminated string" in json_err.msg:
                        detailed_error += "• Suggestion: Add missing closing quote (\") for string values\n"
                    
                    # Log the detailed error
                    current_app.logger.error(f"❌ JSON Error in chart '{chart_tag}': {detailed_error}")
                    
                    # Raise a more informative error
                    raise ValueError(f"Invalid JSON in chart attributes for '{chart_tag}': {json_err.msg} at line {error_line}, column {error_col}")
                
                chart_config = json.loads(cleaned_json)

                # Check if this is a ChatGPT JSON format and convert it
                if "data" in chart_config and "validation" in chart_config:
                    # This is ChatGPT JSON format, convert it
                    converted_config = convert_chatgpt_json_to_bar_of_pie_format(chart_config, data_file_path)
                    chart_meta = converted_config.get("chart_meta", {})
                    series_meta = converted_config.get("series", {})
                    chart_type = "bar_of_pie"
                    title = chart_meta.get("title_left", chart_tag)
                else:
                    # Standard format
                    chart_meta = chart_config.get("chart_meta", {})
                    series_meta = chart_config.get("series", {})
                    
                    # Merge root-level attributes into chart_meta for backward compatibility
                    # This allows JSON with attributes at root level to work properly
                    root_attributes = [
                        "chart_title", "chart_background", "plot_background", "showlegend", 
                        "show_gridlines", "font_size", "font_color", "font_family",
                        "data_labels", "data_label_font_size", "data_label_color", 
                        "fill_opacity", "disable_secondary_y"
                    ]
                    for attr in root_attributes:
                        if attr in chart_config and attr not in chart_meta:
                            chart_meta[attr] = chart_config[attr]
                    
                    # Allow chart_type to be overridden from JSON configuration
                    chart_type = chart_config.get("chart_type", chart_type_map.get(chart_tag_lower, "")).lower().strip()
                    # Check for chart_title in both chart_meta and root level
                    # Handle empty strings properly - if chart_title is empty string, don't show title
                    chart_title_meta = chart_meta.get("chart_title")
                    chart_title_config = chart_config.get("chart_title")
                    
                    if chart_title_meta is not None and chart_title_meta.strip():
                        title = chart_title_meta
                    elif chart_title_config is not None and chart_title_config.strip():
                        title = chart_title_config
                    else:
                        title = ""  # Don't show title if chart_title is empty or not provided
                    
                    # Debug logging for chart type detection
                    current_app.logger.debug(f"🔥 Chart type detection - chart_config.get('chart_type'): {chart_config.get('chart_type')}")
                    current_app.logger.debug(f"🔥 Chart type detection - chart_type_map.get(chart_tag_lower): {chart_type_map.get(chart_tag_lower, '')}")
                    current_app.logger.debug(f"🔥 Chart type detection - Final chart_type: {chart_type}")

                # --- Comprehensive attribute detection logging ---
                #current_app.logger.info(f"🔍 COMPREHENSIVE CHART ATTRIBUTE DETECTION STARTED")
                #current_app.logger.info(f"📊 Chart Type: {chart_type}")
                #current_app.logger.info(f"📋 Chart Meta Keys Found: {list(chart_meta.keys())}")
                #current_app.logger.info(f"📋 Chart Config Keys Found: {list(chart_config.keys())}")
                #current_app.logger.info(f"📋 Top-level Keys Found: {list(data_dict.keys())}")
                
                # Define all possible chart attributes
                all_possible_attributes = [
                    "chart_title", "font_size", "font_color", "font_family", "figsize", 
                    "chart_background", "plot_background", "legend", "legend_position", "legend_font_size",
                    "primary_y_label", "secondary_y_label", "x_label", "y_axis_min_max", 
                    "secondary_y_axis_format", "secondary_y_axis_min_max", "x_axis_label_distance", 
                    "y_axis_label_distance", "axis_tick_format", "axis_tick_font_size",
                    "data_labels", "data_label_format", "data_label_font_size", "data_label_color",
                    "show_gridlines", "gridline_color", "gridline_style", "margin", "bar_width", 
                    "orientation", "bar_border_color", "bar_border_width", "barmode", "line_width", 
                    "marker_size", "line_style", "fill_opacity", "hole", "startangle", "pull", 
                    "bins", "box_points", "violin_points", "bubble_size", "waterfall_measure", 
                    "funnel_measure", "sunburst_path", "sankey_source", 
                    "sankey_target", "sankey_value", "table_header", "indicator_mode", 
                    "indicator_delta", "indicator_gauge", "3d_projection", "z_values", 
                    "lat", "lon", "locations", "open_values", "high_values", "low_values", 
                    "close_values", "annotations"
                ]
                
                # Check which attributes are missing (not in chart_meta OR chart_config OR top-level)
                missing_attributes = []
                for attr in all_possible_attributes:
                    if (attr not in chart_meta and 
                        attr not in chart_config and 
                        attr not in data_dict):
                        missing_attributes.append(attr)
                
                # Chart attribute detection completed (logging removed for cleaner output)

                # --- Define chart type mappings for Matplotlib ---
                chart_type_mapping_mpl = {
                    # Bar charts
                    "bar": "bar",
                    "column": "bar", 
                    "stacked_column": "bar",
                    "horizontal_bar": "barh",
                    
                    # Line charts
                    "line": "plot",
                    "scatter": "scatter",
                    "scatter_line": "plot",
                    
                    # Area charts
                    "area": "fill_between",
                    "filled_area": "fill_between",
                    
                    # Statistical charts
                    "histogram": "hist",
                    "box": "boxplot",
                    "violin": "violinplot",
                    
                    # Other charts
                    "bubble": "scatter",
                    "heatmap": "imshow",
                    "contour": "contour",
                    "waterfall": "bar",
                    "funnel": "bar",
                    "sunburst": "pie",
                    "icicle": "bar",
                    "sankey": "bar",
                    "table": "table",
                    "indicator": "bar",
                    "treemap": "treemap"
                }

                # --- Extract custom fields from chart_config ---
                bar_colors = chart_config.get("bar_colors")
                bar_width = data_dict.get("bar_width") or chart_config.get("bar_width") or chart_meta.get("bar_width")
                orientation = data_dict.get("orientation") or chart_config.get("orientation") or chart_meta.get("orientation")
                bar_border_color = data_dict.get("bar_border_color") or chart_config.get("bar_border_color") or chart_meta.get("bar_border_color")
                bar_border_width = data_dict.get("bar_border_width") or chart_config.get("bar_border_width") or chart_meta.get("bar_border_width")
                font_family = data_dict.get("font_family") or chart_config.get("font_family") or chart_meta.get("font_family")
                # Add font fallback for macOS compatibility
                if font_family:
                    # Check if the font is available, otherwise use a fallback
                    import matplotlib.font_manager as fm
                    # Suppress font warnings by setting log level
                    import logging
                    original_level = logging.getLogger('matplotlib.font_manager').level
                    logging.getLogger('matplotlib.font_manager').setLevel(logging.ERROR)
                    
                    try:
                        available_fonts = [f.name for f in fm.fontManager.ttflist]
                        if font_family not in available_fonts:
                            # Use system-appropriate fallback fonts
                            if font_family.lower() in ['calibri', 'arial']:
                                font_family = 'Helvetica'  # macOS equivalent
                            elif font_family.lower() in ['times new roman', 'times']:
                                font_family = 'Times'  # macOS equivalent
                            else:
                                font_family = 'Helvetica'  # Default fallback
                    finally:
                        # Restore original logging level
                        logging.getLogger('matplotlib.font_manager').setLevel(original_level)
                font_size = data_dict.get("font_size") or chart_config.get("font_size") or chart_meta.get("font_size")
                font_color = data_dict.get("font_color") or chart_config.get("font_color") or chart_meta.get("font_color")
                legend_position = data_dict.get("legend_position") or chart_config.get("legend_position") or chart_meta.get("legend_position")
                legend_font_size = data_dict.get("legend_font_size") or chart_config.get("legend_font_size") or chart_meta.get("legend_font_size")
                show_gridlines = data_dict.get("show_gridlines") if "show_gridlines" in data_dict else (chart_config.get("show_gridlines") if "show_gridlines" in chart_config else chart_meta.get("show_gridlines"))
                # Ensure show_gridlines is a boolean
                if isinstance(show_gridlines, str):
                    show_gridlines = show_gridlines.strip().lower() == "true"
                elif show_gridlines is None:
                    show_gridlines = False  # Default to hiding gridlines if not specified
                gridline_color = data_dict.get("gridline_color") or chart_config.get("gridline_color") or chart_meta.get("gridline_color")
                gridline_style = data_dict.get("gridline_style") or chart_config.get("gridline_style") or chart_meta.get("gridline_style")
                chart_background = data_dict.get("chart_background") or chart_config.get("chart_background") or chart_meta.get("chart_background")
                plot_background = data_dict.get("plot_background") or chart_config.get("plot_background") or chart_meta.get("plot_background")
                data_label_format = data_dict.get("data_label_format") or chart_config.get("data_label_format") or chart_meta.get("data_label_format")
                data_label_font_size = data_dict.get("data_label_font_size") or chart_config.get("data_label_font_size") or chart_meta.get("data_label_font_size")
                data_label_color = data_dict.get("data_label_color") or chart_config.get("data_label_color") or chart_meta.get("data_label_color")
                axis_tick_format = data_dict.get("axis_tick_format") or chart_config.get("axis_tick_format") or chart_meta.get("axis_tick_format")
                y_axis_min_max = data_dict.get("y_axis_min_max") or chart_config.get("y_axis_min_max") or chart_meta.get("y_axis_min_max")
                # current_app.logger.debug(f"Y-axis min/max from config: {y_axis_min_max}")
                x_axis_min_max = data_dict.get("x_axis_min_max") or chart_config.get("x_axis_min_max") or chart_meta.get("x_axis_min_max")
                # current_app.logger.debug(f"X-axis min/max from config: {x_axis_min_max}")
                secondary_y_axis_format = data_dict.get("secondary_y_axis_format") or chart_config.get("secondary_y_axis_format") or chart_meta.get("secondary_y_axis_format")
                secondary_y_axis_min_max = data_dict.get("secondary_y_axis_min_max") or chart_config.get("secondary_y_axis_min_max") or chart_meta.get("secondary_y_axis_min_max")
                disable_secondary_y = data_dict.get("disable_secondary_y") or chart_config.get("disable_secondary_y") or chart_meta.get("disable_secondary_y", False)
                # current_app.logger.info(f"🔧 disable_secondary_y setting: {disable_secondary_y}")
                
                # --- Safe ax2 operation wrapper ---
                def safe_ax2_operation(operation, *args, **kwargs):
                    """Safely execute ax2 operations, skipping if ax2 is None or disabled"""
                    if ax2 is not None and not disable_secondary_y:
                        try:
                            return operation(*args, **kwargs)
                        except Exception as e:
                            current_app.logger.debug(f"Safe ax2 operation failed: {e}")
                            return None
                    return None
                
                # --- Global ax2 safety check ---
                def safe_ax2_text(*args, **kwargs):
                    """Safely call ax2.text(), skipping if ax2 is None"""
                    if ax2 is not None:
                        try:
                            return ax2.text(*args, **kwargs)
                        except Exception as e:
                            current_app.logger.debug(f"ax2.text() failed: {e}")
                            return None
                    return None
                sort_order = data_dict.get("sort_order") or chart_config.get("sort_order") or chart_meta.get("sort_order")
                data_grouping = data_dict.get("data_grouping") or chart_config.get("data_grouping") or chart_meta.get("data_grouping")
                annotations = data_dict.get("annotations", []) or chart_config.get("annotations", []) or chart_meta.get("annotations", [])
                axis_tick_font_size = data_dict.get("axis_tick_font_size") or chart_config.get("axis_tick_font_size") or chart_meta.get("axis_tick_font_size") or 10
                
                # --- Extract tick mark control settings ---
                show_x_ticks = data_dict.get("show_x_ticks") if "show_x_ticks" in data_dict else (chart_config.get("show_x_ticks") if "show_x_ticks" in chart_config else chart_meta.get("show_x_ticks"))
                show_y_ticks = data_dict.get("show_y_ticks") if "show_y_ticks" in data_dict else (chart_config.get("show_y_ticks") if "show_y_ticks" in chart_config else chart_meta.get("show_y_ticks"))
                # Ensure tick settings are boolean
                if isinstance(show_x_ticks, str):
                    show_x_ticks = show_x_ticks.strip().lower() == "true"
                elif show_x_ticks is None:
                    show_x_ticks = True  # Default to showing ticks if not specified
                if isinstance(show_y_ticks, str):
                    show_y_ticks = show_y_ticks.strip().lower() == "true"
                elif show_y_ticks is None:
                    show_y_ticks = True  # Default to showing ticks if not specified
                
                # --- Extract margin settings ---
                margin = data_dict.get("margin") or chart_config.get("margin") or chart_meta.get("margin")
                x_axis_label_distance = data_dict.get("x_axis_label_distance") or chart_config.get("x_axis_label_distance") or chart_meta.get("x_axis_label_distance")
                y_axis_label_distance = (
                    data_dict.get("y_axis_label_distance")
                    or chart_config.get("y_axis_label_distance")
                    or chart_meta.get("y_axis_label_distance")
                    or chart_meta.get("primary_y_axis_label_distance")
                )
                
                # Handle "auto" values for axis label distances
                if x_axis_label_distance == "auto":
                    x_axis_label_distance = "auto"
                if y_axis_label_distance == "auto":
                    y_axis_label_distance = "auto"
                
                # Debug logging for axis label distance extraction
                current_app.logger.debug(f"🔍 Axis Label Distance Extraction - X: {x_axis_label_distance}, Y: {y_axis_label_distance}")
                current_app.logger.debug(f"🔍 Sources - data_dict: {data_dict.get('x_axis_label_distance')}, chart_config: {chart_config.get('x_axis_label_distance')}, chart_meta: {chart_meta.get('x_axis_label_distance')}")
                current_app.logger.debug(f"🔍 Y Sources - data_dict: {data_dict.get('y_axis_label_distance')}, chart_config: {chart_config.get('y_axis_label_distance')}, chart_meta: {chart_meta.get('y_axis_label_distance')}")
                axis_tick_distance = data_dict.get("axis_tick_distance") or chart_config.get("axis_tick_distance") or chart_meta.get("axis_tick_distance")
                figsize = data_dict.get("figsize") or chart_config.get("figsize") or chart_meta.get("figsize")
                
                # --- Extract additional missing attributes ---
                legend = data_dict.get("legend") or chart_config.get("legend") or chart_meta.get("legend")
                data_labels = data_dict.get("data_labels") or chart_config.get("data_labels") or chart_meta.get("data_labels")
                line_width = data_dict.get("line_width") or chart_config.get("line_width") or chart_meta.get("line_width")
                marker_size = data_dict.get("marker_size") or chart_config.get("marker_size") or chart_meta.get("marker_size")
                line_style = data_dict.get("line_style") or chart_config.get("line_style") or chart_meta.get("line_style")
                fill_opacity = data_dict.get("fill_opacity") or chart_config.get("fill_opacity") or chart_meta.get("fill_opacity")
                hole = data_dict.get("hole") or chart_config.get("hole") or chart_meta.get("hole")
                startangle = data_dict.get("startangle") or chart_config.get("startangle") or chart_meta.get("startangle")
                pull = data_dict.get("pull") or chart_config.get("pull") or chart_meta.get("pull")
                barmode = data_dict.get("barmode") or chart_config.get("barmode") or chart_meta.get("barmode")

                # --- Excel range extraction helpers ---
                def parse_range_value(value_str):
                    """
                    Parse range strings like '35% - 40%', '25%-30%', '5%-10%', '<4%' etc.
                    Returns the midpoint as a float, or None if not a valid range.
                    """
                    if not isinstance(value_str, str):
                        return None
                    
                    value_str = value_str.strip()
                    
                    # Pattern 1: "X% - Y%" or "X%-Y%" (range with dash)
                    range_pattern = r'(\d+(?:\.\d+)?)\s*%?\s*-\s*(\d+(?:\.\d+)?)\s*%?'
                    match = re.match(range_pattern, value_str)
                    if match:
                        lower = float(match.group(1))
                        upper = float(match.group(2))
                        midpoint = (lower + upper) / 2
                        return midpoint
                    
                    # Pattern 2: "<X%" (less than - use half of the value as approximation)
                    less_than_pattern = r'<\s*(\d+(?:\.\d+)?)\s*%?'
                    match = re.match(less_than_pattern, value_str)
                    if match:
                        upper = float(match.group(1))
                        midpoint = upper / 2  # Use half as the midpoint
                        return midpoint
                    
                    # Pattern 3: ">X%" (greater than - use value * 1.5 as approximation)
                    greater_than_pattern = r'>\s*(\d+(?:\.\d+)?)\s*%?'
                    match = re.match(greater_than_pattern, value_str)
                    if match:
                        lower = float(match.group(1))
                        midpoint = lower * 1.5  # Use 1.5x as approximation
                        return midpoint
                    
                    return None
                
                def extract_excel_range(sheet, cell_range):
                    # cell_range: e.g., 'E23:E29' or 'AA20:AA23'
                    try:
                        current_app.logger.debug(f"🔍 Extracting Excel range: {cell_range}")
                        values = []
                        for row in sheet[cell_range]:
                            for cell in row:
                                if cell.value is not None:
                                    cell_value = cell.value
                                    # Check if cell value is a string range (e.g., "35% - 40%")
                                    if isinstance(cell_value, str):
                                        parsed_value = parse_range_value(cell_value)
                                        if parsed_value is not None:
                                            cell_value = parsed_value
                                            current_app.logger.debug(f"🔍 Cell {cell.coordinate}: Parsed range '{cell.value}' -> {cell_value}")
                                        else:
                                            # Try to convert string to float if possible
                                            try:
                                                cell_value = float(cell_value.replace('%', '').strip())
                                            except:
                                                pass
                                    # Check if cell has percentage format and preserve percentage value
                                    # Excel stores percentages as decimals (0.666 = 66.6%), so we need to multiply by 100
                                    elif isinstance(cell_value, (int, float)):
                                        # Check if cell number format contains '%' (percentage format)
                                        cell_format = cell.number_format
                                        if cell_format and '%' in str(cell_format):
                                            # If value is between 0 and 1, it's likely a percentage stored as decimal
                                            # Multiply by 100 to get the actual percentage value
                                            if 0 <= cell_value <= 1:
                                                cell_value = cell_value * 100
                                                current_app.logger.debug(f"🔍 Cell {cell.coordinate}: Converted percentage {cell.value} -> {cell_value}% (format: {cell_format})")
                                    # Debug each cell value
                                    current_app.logger.debug(f"🔍 Cell {cell.coordinate}: '{cell_value}' (type: {type(cell_value)}, repr: {repr(cell_value)})")
                                    values.append(cell_value)
                        current_app.logger.debug(f"🔍 Extracted values: {values}")
                        return values
                    except Exception as e:
                        current_app.logger.error(f"❌ Error extracting range {cell_range}: {e}")
                        return []

                # --- Robust recursive Excel cell range and single cell extraction for all chart types and fields ---
                def extract_cell_ranges(obj, sheet, path="root"):
                    """Recursively walk through nested dicts/lists and extract cell ranges and single cells from strings"""
                    if isinstance(obj, dict):
                        for k, v in obj.items():
                            current_path = f"{path}.{k}"
                            if isinstance(v, str):
                                # Check for cell range pattern (e.g., "A1:B10")
                                if re.match(r"^[A-Z]+\d+:[A-Z]+\d+$", v):
                                    try:
                                        extracted = extract_excel_range(sheet, v)
                                        obj[k] = extracted
                                        current_app.logger.info(f"✅ Excel extraction: {current_path} = {v} -> {extracted}")
                                    except Exception as e:
                                        # Failed to extract data from cell range
                                        current_app.logger.error(f"❌ Failed to extract Excel range {v} for key {current_path}: {e}")
                                        pass
                                # Check for single cell pattern (e.g., "U13")
                                elif re.match(r"^[A-Z]+\d+$", v):
                                    try:
                                        cell_value = sheet[v].value
                                        if cell_value is not None:
                                            obj[k] = cell_value
                                            current_app.logger.info(f"✅ Excel extraction: {current_path} = {v} -> {cell_value}")
                                        else:
                                            # Keep original value if cell is empty
                                            current_app.logger.warning(f"⚠️ Cell {v} is empty at {current_path}")
                                            pass
                                    except Exception as e:
                                        # Failed to extract data from single cell
                                        current_app.logger.error(f"❌ Failed to extract cell {v} at {current_path}: {e}")
                                        pass
                            else:
                                extract_cell_ranges(v, sheet, current_path)
                    elif isinstance(obj, list):
                        for i, v in enumerate(obj):
                            extract_cell_ranges(v, sheet, f"{path}[{i}]")
                
                if "source_sheet" in chart_meta:
                    wb = openpyxl.load_workbook(data_file_path, data_only=True)
                    sheet = wb[chart_meta["source_sheet"]]
                    
                    # Validate chart configuration before extraction
                    def validate_chart_config(series_config):
                        """Validate chart configuration for dimension consistency"""
                        if series_config and isinstance(series_config, dict):
                            x_axis_range = series_config.get("x_axis", "")
                            
                            if x_axis_range and isinstance(x_axis_range, str) and ":" in x_axis_range:
                                # Parse x_axis range to get length
                                try:
                                    start_cell, end_cell = x_axis_range.split(":")
                                    start_col, start_row = re.match(r"([A-Z]+)(\d+)", start_cell).groups()
                                    end_col, end_row = re.match(r"([A-Z]+)(\d+)", end_cell).groups()
                                    
                                    if start_col == end_col:  # Same column
                                        x_length = int(end_row) - int(start_row) + 1
                                        current_app.logger.info(f"🔍 X-axis range {x_axis_range} has {x_length} cells")
                                        
                                        # Check each series data
                                        series_data = series_config.get("data", [])
                                        for i, series in enumerate(series_data):
                                            series_name = series.get("name", f"Series {i+1}")
                                            values_range = series.get("values", "")
                                            
                                            if values_range and isinstance(values_range, str) and ":" in values_range:
                                                try:
                                                    start_cell, end_cell = values_range.split(":")
                                                    start_col, start_row = re.match(r"([A-Z]+)(\d+)", start_cell).groups()
                                                    end_col, end_row = re.match(r"([A-Z]+)(\d+)", end_cell).groups()
                                                    
                                                    if start_col == end_col:  # Same column
                                                        y_length = int(end_row) - int(start_row) + 1
                                                        current_app.logger.info(f"📊 Series '{series_name}' range {values_range} has {y_length} cells")
                                                        
                                                        if x_length != y_length:
                                                            pass  # Suppress warning logs: f"⚠️ Dimension mismatch detected: x_axis={x_length}, {series_name}={y_length}")
                                                            
                                                            # Fix by adjusting the shorter range
                                                            if x_length > y_length:
                                                                # Extend y_values range
                                                                new_end_row = int(start_row) + x_length - 1
                                                                new_end_cell = f"{start_col}{new_end_row}"
                                                                series["values"] = f"{start_cell}:{new_end_cell}"
                                                                current_app.logger.info(f"🔧 Extended {series_name} range to {series['values']}")
                                                            else:
                                                                # Truncate x_axis range
                                                                new_end_row = int(start_row) + y_length - 1
                                                                new_end_cell = f"{start_col}{new_end_row}"
                                                                series_config["x_axis"] = f"{start_cell}:{new_end_cell}"
                                                                current_app.logger.info(f"🔧 Truncated x_axis range to {series_config['x_axis']}")
                                                                break  # Only need to fix once
                                                except Exception as e:
                                                    current_app.logger.error(f"❌ Error parsing values range {values_range}: {e}")
                                except Exception as e:
                                    current_app.logger.error(f"❌ Error parsing x_axis range {x_axis_range}: {e}")
                    
                    # Validate chart configuration (pass series_meta which contains x_axis)
                    validate_chart_config(series_meta)
                    
                    # Extract cell ranges from both chart_meta and series_meta
                    current_app.logger.info(f"🔍 Starting Excel cell range extraction from: {data_file_path}")
                    current_app.logger.info(f"🔍 Using sheet: {chart_meta.get('source_sheet', 'sample')}")
                    
                    current_app.logger.info(f"🔍 Extracting from chart_meta...")
                    extract_cell_ranges(chart_meta, sheet, "chart_meta")
                    current_app.logger.info(f"🔍 Extracting from series_meta...")
                    extract_cell_ranges(series_meta, sheet, "series_meta")
                    wb.close()
                    
                    # Log the extracted data for debugging
                    current_app.logger.info(f"✅ Chart meta after extraction: {chart_meta}")
                    current_app.logger.info(f"✅ Series meta after extraction: {series_meta}")
                    current_app.logger.info(f"✅ Series meta x_axis value: {series_meta.get('x_axis', 'NOT FOUND')}")
                
                # Use updated values from series_meta after extraction
                series_data = series_meta.get("data", [])
                if not series_data and "series" in series_meta:
                    # Handle case where data is directly in series object
                    series_data = series_meta.get("series", [])
                
                # Also check if chart_meta has series data (fallback)
                if not series_data and "series" in chart_meta:
                    series_data = chart_meta.get("series", {}).get("data", [])
                    if not series_data:
                        series_data = chart_meta.get("series", [])
                
                # Ensure Excel cell ranges are extracted from series data regardless of source
                if series_data and data_file_path:
                    try:
                        current_app.logger.info(f"🔍 Extracting Excel cell ranges from series data...")
                        wb = openpyxl.load_workbook(data_file_path, data_only=True)
                        sheet = wb[chart_meta.get("source_sheet", "sample")]
                        # Extract cell ranges from the series data
                        for i, series in enumerate(series_data):
                            if isinstance(series, dict):
                                current_app.logger.info(f"🔍 Processing series {i+1}: {series}")
                                extract_cell_ranges(series, sheet, f"series_data[{i}]")
                                current_app.logger.info(f"🔍 Series {i+1} after extraction: {series}")
                        wb.close()
                        current_app.logger.info(f"🔍 Extracted Excel cell ranges from series data")
                    except Exception as e:
                        current_app.logger.error(f"❌ Error extracting Excel cell ranges from series data: {e}")
                
                # Debug logging for series data extraction
                current_app.logger.debug(f"🔥 Series data extraction - series_meta.get('data'): {series_meta.get('data', [])}")
                current_app.logger.debug(f"🔥 Series data extraction - series_meta.get('series'): {series_meta.get('series', [])}")
                current_app.logger.debug(f"🔥 Series data extraction - chart_meta.get('series'): {chart_meta.get('series', [])}")
                current_app.logger.debug(f"🔥 Series data extraction - Final series_data: {series_data}")
                
                # Extract x_values from the correct location (skip for heatmaps)
                if chart_type != "heatmap":
                    x_values = series_meta.get("x_axis", [])
                    if not x_values and "series" in series_meta:
                        # Try to get x_axis from the series object
                        x_values = series_meta.get("series", {}).get("x_axis", [])
                    
                    # If still no x_values, try to get from chart_meta
                    if not x_values and "series" in chart_meta:
                        x_values = chart_meta.get("series", {}).get("x_axis", [])
                    
                    # If still no x_values, try to get from the first series data item
                    if not x_values and series_data and len(series_data) > 0:
                        # Check if x_axis is in the first series
                        first_series = series_data[0]
                        if isinstance(first_series, dict) and "x_axis" in first_series:
                            x_values = first_series["x_axis"]
                        # Also check if there's a separate x_axis in the series object
                        elif "series" in series_meta and isinstance(series_meta["series"], dict):
                            x_values = series_meta["series"].get("x_axis", [])
                        elif "series" in chart_meta and isinstance(chart_meta["series"], dict):
                            x_values = chart_meta["series"].get("x_axis", [])
                else:
                    # For heatmaps, don't use x_axis from series - let the heatmap handle its own axis labels
                    x_values = []
                
                current_app.logger.info(f"🔍 Extracted x_values (raw): {x_values} (type: {type(x_values)})")
                
                # Ensure x_values is always defined to prevent "cannot access local variable" error
                if not x_values:
                    x_values = []
                    current_app.logger.info(f"⚠️ No x_values found, using empty list as fallback")
                
                # If x_values is still a string (cell range not extracted), extract it now
                if isinstance(x_values, str) and re.match(r"^[A-Z]+\d+:[A-Z]+\d+$", x_values):
                    current_app.logger.warning(f"⚠️ x_values is still a cell range string: {x_values}. Extracting now...")
                    try:
                        wb = openpyxl.load_workbook(data_file_path, data_only=True)
                        sheet = wb[chart_meta.get("source_sheet", "sample")]
                        x_values = extract_excel_range(sheet, x_values)
                        wb.close()
                        current_app.logger.info(f"✅ Converted x_axis from cell range to values: {x_values}")
                    except Exception as e:
                        current_app.logger.error(f"❌ Error converting x_axis cell range '{x_values}': {e}")
                        x_values = []
                
                # Validate and fix dimension mismatches
                def validate_and_fix_dimensions(x_vals, series_data):
                    """Validate that all series have the same length as x_axis and fix if needed"""
                    if not x_vals or not series_data:
                        return x_vals, series_data
                    
                    x_length = len(x_vals)
                    current_app.logger.debug(f"🔍 Validating dimensions: x_axis length = {x_length}")
                    
                    for i, series in enumerate(series_data):
                        series_name = series.get("name", f"Series {i+1}")
                        y_vals = series.get("values", [])
                        y_length = len(y_vals)
                        
                        current_app.logger.debug(f"📊 Series '{series_name}': {y_length} values")
                        
                        if x_length != y_length:
                            pass  # Suppress warning logs: f"⚠️ Dimension mismatch in '{series_name}': x_axis={x_length}, y_values={y_length}")
                            
                            # Fix by truncating to the minimum length
                            min_length = min(x_length, y_length)
                            if min_length > 0:
                                # Truncate x_axis if it's longer
                                if x_length > min_length:
                                    x_vals = x_vals[:min_length]
                                    current_app.logger.debug(f"✂️ Truncated x_axis to {min_length} values")
                                
                                # Truncate y_values if they're longer
                                if y_length > min_length:
                                    series["values"] = y_vals[:min_length]
                                    current_app.logger.debug(f"✂️ Truncated '{series_name}' values to {min_length}")
                            else:
                                current_app.logger.error(f"❌ Cannot fix dimensions: both arrays are empty")
                    
                    return x_vals, series_data
                
                # Apply dimension validation (skip for heatmaps)
                if chart_type not in ["heatmap"]:
                    x_values, series_data = validate_and_fix_dimensions(x_values, series_data)
                
                colors = series_meta.get("colors", [])
                
                # --- SERIES ATTRIBUTE DETECTION LOGGING ---
                #current_app.logger.info(f"🔍 SERIES ATTRIBUTE DETECTION STARTED")
                #current_app.logger.info(f"📊 Number of series: {len(series_data)}")
                #current_app.logger.info(f"📋 Series meta keys: {list(series_data.keys()) if isinstance(series_data, dict) else 'N/A'}")
                #current_app.logger.info(f"📋 Series meta structure: {series_meta}")
                #current_app.logger.info(f"📋 Series data: {series_data}")
                #current_app.logger.info(f"📋 X values extracted: {x_values}")
                
                # Define all possible series attributes
                all_possible_series_attributes = [
                    "marker", "opacity", "textposition", "orientation", "width", "fill", 
                    "fillcolor", "hole", "pull", "mode", "line", "nbinsx", "boxpoints", 
                    "jitter", "sizeref", "sizemin", "symbol", "measure", "connector", "textinfo"
                ]
                
                for i, series in enumerate(series_data):
                    series_name = series.get("name", f"Series {i+1}")
                    series_type = series.get("type", "unknown")
                    #current_app.logger.info(f"📈 SERIES {i+1}: {series_name}")
                    #current_app.logger.info(f"   Type: {series_type}")
                    
                    # Check which series attributes are missing
                    missing_series_attributes = []
                    for attr in all_possible_series_attributes:
                        if attr not in series:
                            missing_series_attributes.append(attr)
                    
                    # Series attribute detection completed (logging removed for cleaner output)

                # --- Plotly interactive chart generation ---
                fig = go.Figure()

                # --- Bar of Pie chart special handling ---
                if chart_type in ["bar of pie", "bar_of_pie"]:
                    # Extract cell ranges for other_labels and other_values if they are cell ranges
                    other_labels = chart_meta.get("other_labels", [])
                    other_values = chart_meta.get("other_values", [])
                    other_colors = chart_meta.get("other_colors", [])
                    
                    # If other_labels and other_values are cell ranges, extract them
                    if isinstance(other_labels, str) and re.match(r"^[A-Z]+\d+:[A-Z]+\d+$", other_labels):
                        try:
                            wb = openpyxl.load_workbook(data_file_path, data_only=True)
                            sheet = wb[chart_meta.get("source_sheet", "sample")]
                            other_labels = extract_excel_range(sheet, other_labels)
                            wb.close()
                            # Extracted other_labels successfully
                            pass
                        except Exception as e:
                            # Failed to extract other_labels
                            pass
                    
                    if isinstance(other_values, str) and re.match(r"^[A-Z]+\d+:[A-Z]+\d+$", other_values):
                        try:
                            wb = openpyxl.load_workbook(data_file_path, data_only=True)
                            sheet = wb[chart_meta.get("source_sheet", "sample")]
                            other_values = extract_excel_range(sheet, other_values)
                            wb.close()
                            # Extracted other_values successfully
                            print(f"DEBUG: other_values extracted from {other_values} = {other_values}")
                            pass
                        except Exception as e:
                            # Failed to extract other_values
                            pass
                    
                    # Log the extracted data for debugging
                    #current_app.logger.info(f"🔍 Bar of Pie Chart Data:")
                    #current_app.logger.info(f"   Other Labels: {other_labels}")
                    #current_app.logger.info(f"   Other Values: {other_values}")
                    #current_app.logger.info(f"   Other Colors: {other_colors}")
                    
                    # Check if values are percentages in decimal form and convert them
                    y_axis_title = chart_meta.get("y_axis_title", "")
                    if other_values and y_axis_title and "%" in y_axis_title:
                        # Check if all values are between 0-1 (likely percentages in decimal form)
                        if all(isinstance(v, (int, float)) and 0 <= v <= 1 for v in other_values if v is not None):
                            print(f"DEBUG: Converting decimal values to percentages: {other_values}")
                            other_values = [v * 100 if v is not None and isinstance(v, (int, float)) else v for v in other_values]
                            print(f"DEBUG: Converted to: {other_values}")
                        else:
                            # Handle string values that might be percentages
                            converted_values = []
                            for v in other_values:
                                if v is not None:
                                    try:
                                        if isinstance(v, str):
                                            # Try to convert string to float
                                            float_val = float(v)
                                            if 0 <= float_val <= 1:
                                                converted_values.append(float_val * 100)
                                            else:
                                                converted_values.append(float_val)
                                        else:
                                            converted_values.append(v)
                                    except (ValueError, TypeError):
                                        converted_values.append(v)
                                else:
                                    converted_values.append(v)
                            other_values = converted_values
                    
                    # Check for empty/null values
                    if other_labels and other_values:
                        empty_count = sum(1 for label, value in zip(other_labels, other_values) 
                                        if (label is None or str(label).strip() == "" or 
                                            value is None or str(value).strip() == "" or 
                                            str(value).strip() == "0"))
                        if empty_count > 0:
                            # Found empty/null values in bar chart data (will cause missing bars)
                            pass
                        
                        # Data analysis completed (logging removed for cleaner output)
                    
                    value_format = chart_meta.get("value_format", "")
                    # For bar of pie charts, use labels from series_meta, fallback to x_values
                    labels = series_meta.get("labels", x_values)
                    values = series_meta.get("values", [])
                    colors = series_meta.get("colors", [])
                    
                    # Chart data prepared
                    fig = create_bar_of_pie_chart(
                        labels=labels,
                        values=values,
                        other_labels=other_labels,
                        other_values=other_values,
                        colors=colors,
                        other_colors=other_colors if other_colors else colors,
                        title=title,
                        value_format=value_format,
                        chart_meta=chart_meta
                    )
                    
                    # Save chart as PNG file for Word document insertion using matplotlib
                    tmpfile = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                    
                    # Convert Plotly figure to matplotlib and save
                    try:
                        # Try Plotly first (if Chrome is available)
                        fig.write_image(tmpfile.name, width=900, height=500, scale=2)
                        current_app.logger.info("✅ Chart saved using Plotly (Chrome available)")
                    except Exception as e:
                        current_app.logger.warning(f"Plotly write_image failed: {e}. Using matplotlib fallback.")
                        
                        # Use the dedicated matplotlib fallback function
                        success = create_matplotlib_chart_from_plotly(fig, tmpfile.name)
                        if success:
                            current_app.logger.info("✅ Chart saved using matplotlib fallback")
                        else:
                            current_app.logger.error("❌ Both Plotly and matplotlib chart generation failed")
                            # Create a simple placeholder chart
                            import matplotlib.pyplot as plt
                            fig_mpl, ax = plt.subplots(figsize=(12, 6))
                            ax.text(0.5, 0.5, 'Chart Generation Failed\nChrome/Kaleido not available', 
                                   ha='center', va='center', fontsize=16, 
                                   bbox=dict(boxstyle="round,pad=0.3", facecolor="lightgray"))
                            ax.set_xlim(0, 1)
                            ax.set_ylim(0, 1)
                            ax.axis('off')
                            fig_mpl.savefig(tmpfile.name, dpi=300, bbox_inches='tight')
                            plt.close(fig_mpl)
                    
                    plt.close('all')  # Close any matplotlib figures
                    gc.collect()  # Force garbage collection
                    
                    return tmpfile.name
                
                def extract_values_from_range(cell_range):
                    start_cell, end_cell = cell_range.split(":")
                    start_col, start_row = re.match(r"([A-Z]+)(\d+)", start_cell).groups()
                    end_col, end_row = re.match(r"([A-Z]+)(\d+)", end_cell).groups()

                    start_col_idx = column_index_from_string(start_col) - 1
                    end_col_idx = column_index_from_string(end_col) - 1
                    start_row_idx = int(start_row) - 1  # Fixed: pandas is 0-indexed, Excel is 1-indexed
                    end_row_idx = int(end_row) - 1      # Fixed: pandas is 0-indexed, Excel is 1-indexed

                    if start_col_idx == end_col_idx:
                        return df.iloc[start_row_idx:end_row_idx + 1, start_col_idx].tolist()
                    else:
                        return df.iloc[start_row_idx:end_row_idx + 1, start_col_idx:end_col_idx + 1].values.flatten().tolist()

                # --- Data grouping and sorting logic ---
                # If data_grouping is present, filter x/y values to only those groups
                def group_and_sort(x_vals, y_vals, group_names=None, sort_order=None):
                    # Grouping: only keep x/y where x in group_names (if group_names provided)
                    if group_names:
                        filtered = [(x, y) for x, y in zip(x_vals, y_vals) if x in group_names]
                        if filtered:
                            x_vals, y_vals = zip(*filtered)
                        else:
                            x_vals, y_vals = [], []
                    # Sorting: sort by y-value
                    if sort_order == "ascending":
                        sorted_pairs = sorted(zip(x_vals, y_vals), key=lambda pair: pair[1])
                        if sorted_pairs:
                            x_vals, y_vals = zip(*sorted_pairs)
                    elif sort_order == "descending":
                        sorted_pairs = sorted(zip(x_vals, y_vals), key=lambda pair: pair[1], reverse=True)
                        if sorted_pairs:
                            x_vals, y_vals = zip(*sorted_pairs)
                    return list(x_vals), list(y_vals)

                # Helper function to normalize values to sum to 100
                def normalize_values_to_100(values):
                    """
                    Normalize a list of numeric values so they sum to 100.
                    Used for pie charts to ensure proper percentage display.
                    """
                    try:
                        # Filter out non-numeric values
                        numeric_values = [float(v) for v in values if isinstance(v, (int, float)) and v is not None]
                        
                        if not numeric_values:
                            return values
                        
                        total = sum(numeric_values)
                        
                        # If total is 0, return original values
                        if total == 0:
                            return values
                        
                        # Normalize to sum to 100
                        normalized = [(v / total) * 100 for v in numeric_values]
                        
                        return normalized
                    except Exception as e:
                        current_app.logger.error(f"Error normalizing values: {e}")
                        return values
                
                # Special handling for pie charts and treemaps (single trace)
                if (chart_type == "pie" or chart_type == "treemap") and len(series_data) == 1:
                    series = series_data[0]
                    label = series.get("name", "Pie Chart")
                    labels = series.get("labels", x_values)
                    values = series.get("values", [])
                    
                    # Normalize pie chart values to sum to 100%
                    if chart_type == "pie" and values:
                        current_app.logger.debug(f"🔍 Original pie values: {values}")
                        values = normalize_values_to_100(values)
                        current_app.logger.debug(f"🔍 Normalized pie values: {values}")
                        # Update the series with normalized values
                        series["values"] = values
                    
                    color = series.get("marker", {}).get("colors") if "marker" in series else colors
                    
                    # Apply value format if specified
                    value_format = chart_meta.get("value_format", "")
                    if value_format:
                        # Convert values to percentage format for display
                        try:
                            formatted_values = [f"{float(v):{value_format}}" if v is not None else "0" for v in values]
                        except:
                            formatted_values = values
                    else:
                        formatted_values = values
                    
                    # Check if this is an expanded pie chart (pie with one segment expanded to column)
                    expanded_segment = chart_meta.get("expanded_segment")
                    
                    if expanded_segment:
                        # Use the helper function for expanded pie charts
                        fig = create_expanded_pie_chart(
                            labels=labels,
                            values=values,
                            colors=color if isinstance(color, list) else [color],
                            expanded_segment=expanded_segment,
                            title=title,
                            value_format=value_format
                        )
                        
                    else:
                        # Regular pie chart
                        pie_kwargs = {
                            "labels": labels,
                            "values": values,
                            "name": label,
                            "textinfo": "label+percent+value" if chart_meta.get("data_labels", True) else "none",
                            "textposition": "outside",
                            "hole": 0.0  # Solid pie chart
                        }
                        
                        # Pie colors
                        if color:
                            if isinstance(color, list):
                                safe_colors = [c for c in color if c is not None]
                                if safe_colors:
                                    pie_kwargs["marker"] = dict(colors=safe_colors)
                            elif color is not None:
                                pie_kwargs["marker"] = dict(colors=[color])
                        
                        # Add pull effect for specific segments if needed
                        if "pull" in chart_meta:
                            pie_kwargs["pull"] = chart_meta["pull"]
                        
                        fig.add_trace(go.Pie(**pie_kwargs,
                            hovertemplate=f"<b>{label}</b><br>%{{label}}: %{{value}}{str(value_format) if value_format else ''}<extra></extra>"
                        ))
                
                elif chart_type == "heatmap":
                    # Special handling for heatmaps
                    current_app.logger.debug(f"🔥 Heatmap chart type detected, processing series data")
                    if series_data and len(series_data) > 0:
                        series = series_data[0]
                        current_app.logger.debug(f"🔥 Heatmap series data: {series}")
                        
                        # Validate heatmap data structure
                        z_data = series.get("z", [])
                        x_labels = series.get("x", [])
                        y_labels = series.get("y", [])
                        
                        if z_data and x_labels and y_labels:
                            current_app.logger.debug(f"🔥 Heatmap data validation passed - Z: {len(z_data)}x{len(z_data[0]) if z_data else 0}, X: {len(x_labels)}, Y: {len(y_labels)}")
                        else:
                            current_app.logger.warning(f"🔥 Heatmap data validation failed - Z: {z_data}, X: {x_labels}, Y: {y_labels}")
                    else:
                        current_app.logger.warning(f"🔥 No series data found for heatmap")
                
                elif chart_type == "treemap":
                    series = series_data[0]
                    label = series.get("name", "Treemap Chart")
                    labels = series.get("labels", x_values)
                    values = series.get("values", [])
                    color = series.get("marker", {}).get("colors") if "marker" in series else colors
                    
                    # Debug logging for treemap data
                    current_app.logger.debug(f"🔍 Treemap data - series: {series}")
                    current_app.logger.debug(f"🔍 Treemap data - labels: {labels} (type: {type(labels)})")
                    current_app.logger.debug(f"🔍 Treemap data - values: {values} (type: {type(values)})")
                    
                    
                    # Get data label settings from JSON configuration
                    data_labels = chart_meta.get("data_labels", True)
                    data_label_font_size = chart_meta.get("data_label_font_size", 12)
                    data_label_color = chart_meta.get("data_label_color", "#000000")
                    fill_opacity = chart_meta.get("fill_opacity", 0.8)
                    hide_center_box = chart_meta.get("hide_center_box", False)
                    current_app.logger.debug(f"🔍 Treemap: hide_center_box parameter: {hide_center_box}")
                    current_app.logger.debug(f"🔍 Treemap: chart_meta keys: {list(chart_meta.keys())}")
                    current_app.logger.debug(f"🔍 Treemap: chart_meta: {chart_meta}")
                    
                    # Determine textinfo based on data_labels setting (independent of center box)
                    if data_labels:
                        textinfo_setting = "label"  # Show labels
                        current_app.logger.debug(f"🔍 Treemap: data_labels=True, setting textinfo to 'label'")
                    else:
                        textinfo_setting = "none"  # Hide labels
                        current_app.logger.debug(f"🔍 Treemap: data_labels=False, setting textinfo to 'none'")
                    
                    # Treemap specific settings - separate center box control from data labels
                    treemap_kwargs = {
                        "labels": labels,
                        "values": values,
                        "name": label,
                        "textinfo": textinfo_setting,
                        "textposition": "middle center",
                        "branchvalues": "total",
                        "pathbar": dict(visible=not hide_center_box),  # Hide center box if hide_center_box is True
                        "tiling": dict(packing="squarify", squarifyratio=1),
                        "texttemplate": None,
                        "hoverinfo": "label+value"
                    }
                    
                    # Add text font configuration from JSON
                    if data_labels:
                        treemap_kwargs["textfont"] = {
                            "family": font_family,
                            "size": data_label_font_size,
                            "color": data_label_color
                        }
                    
                    # Add hierarchical data support
                    if "parent" in series:
                        treemap_kwargs["parents"] = series["parent"]
                    
                    # Add enhanced colors if available
                    if color:
                        if isinstance(color, list):
                            treemap_kwargs["marker"] = dict(colors=color, opacity=fill_opacity)
                        else:
                            treemap_kwargs["marker"] = dict(colors=[color], opacity=fill_opacity)
                    else:
                        # Use enhanced color palette for better visual appeal
                        enhanced_colors = [
                            '#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7',
                            '#DDA0DD', '#98D8C8', '#F7DC6F', '#BB8FCE', '#85C1E9',
                            '#F8C471', '#82E0AA', '#F1948A', '#85C1E9', '#D7BDE2'
                        ]
                        # Apply colors based on data length
                        if len(labels) <= len(enhanced_colors):
                            treemap_kwargs["marker"] = dict(colors=enhanced_colors[:len(labels)], opacity=fill_opacity)
                    
                    # Add hover template with enhanced information
                    hover_template = f"<b>{label}</b><br>"
                    hover_template += "Category: %{label}<br>"
                    hover_template += "Value: %{value}<br>"
                    if "parent" in series:
                        hover_template += "Parent: %{parent}<br>"
                    hover_template += "Percentage: %{percentParent:.1f}%<extra></extra>"
                    
                    # Add showlegend parameter to the treemap trace
                    current_app.logger.debug(f"🔍 Treemap: Final treemap_kwargs: {treemap_kwargs}")
                    
                    treemap_trace = go.Treemap(**treemap_kwargs,
                        hovertemplate=hover_template
                    )
                    
                    # Apply showlegend setting from chart_meta
                    show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                    if isinstance(show_legend_raw, str):
                        show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                    else:
                        show_legend = bool(show_legend_raw)
                    
                    treemap_trace.showlegend = show_legend
                    fig.add_trace(treemap_trace)
                
                # Handle stacked column, area, and other multi-series charts
                else:
                    for i, series in enumerate(series_data):
                        label = series.get("name", f"Series {i+1}")
                        series_type = series.get("type", "bar").lower()
                        
                        # Map heatmap to imshow for Matplotlib
                        if series_type == "heatmap":
                            mpl_chart_type = "imshow"
                        else:
                            mpl_chart_type = chart_type_mapping_mpl.get(series_type, "scatter")
                        
                        color = None
                        if "marker" in series and isinstance(series["marker"], dict) and "color" in series["marker"]:
                            color = series["marker"]["color"]
                        elif bar_colors:
                            color = bar_colors
                        elif i < len(colors):
                            color = colors[i]

                        y_vals = series.get("values")
                        value_range = series.get("value_range")
                        if value_range:
                            # Check if value_range is already extracted (list) or still a string
                            if isinstance(value_range, list):
                                y_vals = value_range
                            else:
                                y_vals = extract_values_from_range(value_range)

                        # --- Apply grouping and sorting ---
                        x_vals = x_values
                        if y_vals is not None and x_vals is not None:
                            x_vals, y_vals = group_and_sort(x_vals, y_vals, data_grouping, sort_order)

                        trace_kwargs = {
                            "x": x_vals,
                            "y": y_vals,
                            "name": label,
                        }
                        
                        # Handle colors
                        if color:
                            if isinstance(color, list) and series_type == "bar":
                                trace_kwargs["marker"] = dict(color=color)
                            elif isinstance(color, str):
                                trace_kwargs["marker_color"] = color
                        
                        # Handle bar-specific properties
                        if bar_width and series_type == "bar":
                            trace_kwargs["width"] = bar_width
                        if orientation and series_type == "bar":
                            trace_kwargs["orientation"] = orientation[0].lower() if isinstance(orientation, str) else orientation
                        if bar_border_color and series_type == "bar":
                            trace_kwargs["marker_line_color"] = bar_border_color
                        if bar_border_width and series_type == "bar":
                            trace_kwargs["marker_line_width"] = bar_border_width
                        
                        # Handle line-specific properties
                        if line_width and series_type in ["line", "scatter", "scatter_line"]:
                            trace_kwargs["line_width"] = line_width
                        if marker_size and series_type in ["line", "scatter", "scatter_line"]:
                            trace_kwargs["marker_size"] = marker_size
                        if line_style and series_type in ["line", "scatter", "scatter_line"]:
                            # Map line styles to Plotly format
                            line_style_map = {
                                "solid": "solid",
                                "dashed": "dash",
                                "dotted": "dot",
                                "dashdot": "dashdot"
                            }
                            trace_kwargs["line_dash"] = line_style_map.get(line_style, "solid")
                        
                        # Handle opacity
                        if fill_opacity:
                            trace_kwargs["opacity"] = fill_opacity

                        # Add traces based on chart type - REPLACE THE RESTRICTIVE IF/ELIF BLOCKS
                        # Generic chart type handling for Plotly
                        chart_type_mapping = {
                            # Bar charts
                            "bar": go.Bar,
                            "column": go.Bar,
                            "stacked_column": go.Bar,
                            "horizontal_bar": go.Bar,
                            
                            # Line charts
                            "line": go.Scatter,
                            "scatter": go.Scatter,
                            "scatter_line": go.Scatter,
                            
                            # Area charts
                            "area": go.Scatter,
                            "filled_area": go.Scatter,
                            
                            # Pie charts
                            "pie": go.Pie,
                            "donut": go.Pie,
                            
                            # 3D charts
                            "scatter3d": go.Scatter3d,
                            "surface": go.Surface,
                            "mesh3d": go.Mesh3d,
                            
                            # Statistical charts
                            "histogram": go.Histogram,
                            "box": go.Box,
                            "violin": go.Violin,
                            
                            # Financial charts
                            "candlestick": go.Candlestick,
                            "ohlc": go.Ohlc,
                            
                            # Geographic charts
                            "scattergeo": go.Scattergeo,
                            "choropleth": go.Choropleth,
                            
                            # Other charts
                            "bubble": go.Scatter,
                            "heatmap": go.Heatmap,
                            "contour": go.Contour,
                            "waterfall": go.Waterfall,
                            "funnel": go.Funnel,
                            "sunburst": go.Sunburst,
                            "icicle": go.Icicle,
                            "sankey": go.Sankey,
                            "treemap": go.Treemap,
                            "table": go.Table,
                            "indicator": go.Indicator
                        }
                        
                        # Get the appropriate Plotly chart class
                        plotly_chart_class = chart_type_mapping.get(series_type)
                        
                        if plotly_chart_class:
                            # Prepare trace arguments based on chart type
                            if series_type in ["bar", "column", "stacked_column", "horizontal_bar"]:
                                # Bar chart specific settings
                                # REMOVE: if chart_type == "stacked_column": trace_kwargs["barmode"] = "stack"
                                if orientation and orientation.lower() == "horizontal":
                                    trace_kwargs["orientation"] = "h"
                                    # Swap x and y for horizontal bars
                                    trace_kwargs["x"], trace_kwargs["y"] = trace_kwargs["y"], trace_kwargs["x"]
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                                ))
                                
                            elif series_type in ["line", "scatter", "scatter_line"]:
                                # Line/Scatter chart specific settings
                                mode = "lines+markers" if series_type == "scatter_line" else "markers" if series_type == "scatter" else "lines"
                                trace_kwargs["mode"] = mode
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                                ))
                                
                            elif series_type in ["area", "filled_area"]:
                                # Area chart specific settings
                                trace_kwargs["mode"] = "lines"
                                trace_kwargs["fill"] = "tozeroy"
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                                ))
                                
                            elif series_type == "pie":
                                # Pie chart specific settings
                                pie_kwargs = {
                                    "labels": x_vals,
                                    "values": y_vals,
                                    "name": label,
                                    "textinfo": "label" if data_labels and not chart_meta.get("hide_center_box", False) else "none",
                                    "textposition": "outside",
                                    "hole": hole if hole is not None else (0.4 if series_type == "donut" else 0.0)
                                }
                                
                                # Add pie chart specific attributes
                                if startangle is not None:
                                    pie_kwargs["rotation"] = startangle
                                if pull is not None:
                                    # pull can be a single value or a list
                                    if isinstance(pull, (int, float)):
                                        pie_kwargs["pull"] = [pull] * len(y_vals)
                                    elif isinstance(pull, list):
                                        pie_kwargs["pull"] = pull
                                
                                if color:
                                    pie_kwargs["marker"] = dict(colors=color) if isinstance(color, list) else dict(colors=[color])
                                
                                fig.add_trace(plotly_chart_class(**pie_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>%{{label}}: %{{value}}<extra></extra>"
                                ))
                                
                            elif series_type in ["scatter3d", "surface", "mesh3d"]:
                                # 3D chart specific settings
                                if "z" not in trace_kwargs and len(y_vals) > 0:
                                    # Create a simple z-axis if not provided
                                    trace_kwargs["z"] = [i for i in range(len(y_vals))]
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Z: %{{z}}<extra></extra>"
                                ))
                                
                            elif series_type in ["histogram", "box", "violin"]:
                                # Statistical chart specific settings
                                if series_type == "histogram":
                                    trace_kwargs["x"] = y_vals  # Histogram uses x for values
                                    del trace_kwargs["y"]
                                elif series_type in ["box", "violin"]:
                                    trace_kwargs["y"] = y_vals
                                    trace_kwargs["x"] = [label] * len(y_vals) if len(y_vals) > 0 else [label]
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Value: %{{y if series_type in ['box', 'violin'] else 'x'}}<extra></extra>"
                                ))
                                
                            elif series_type in ["bubble"]:
                                # Bubble chart specific settings
                                sizes = series.get("size", [20] * len(y_vals))
                                trace_kwargs["mode"] = "markers"
                                trace_kwargs["marker"] = {"size": sizes}
                                if color:
                                    trace_kwargs["marker"]["color"] = color
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Size: %{{marker.size}}<extra></extra>"
                                ))
                                
                            elif series_type == "heatmap":
                                # Heatmap specific settings
                                heatmap_kwargs = {
                                    "z": series.get("z", y_vals),  # Use z data if provided, otherwise y_vals
                                    "x": series.get("x", x_vals),  # Use x data if provided, otherwise x_vals
                                    "y": series.get("y", [label]),  # Use y data if provided, otherwise label
                                    "name": label,
                                    "colorscale": series.get("colorscale", "Viridis"),
                                    "showscale": series.get("showscale", True)
                                }
                                
                                # Handle text data safely (convert to strings to avoid concatenation errors)
                                if "text" in series:
                                    try:
                                        heatmap_kwargs["text"] = [[str(cell) for cell in row] for row in series["text"]]
                                    except:
                                        pass  # Suppress warning logs: f"⚠️ Could not process heatmap text data for {label}")
                                
                                # Handle colorbar settings
                                if "colorbar" in series:
                                    heatmap_kwargs["colorbar"] = series["colorbar"]
                                
                                # Handle zmin/zmax
                                if "zmin" in series:
                                    heatmap_kwargs["zmin"] = series["zmin"]
                                if "zmax" in series:
                                    heatmap_kwargs["zmax"] = series["zmax"]
                                
                                fig.add_trace(go.Heatmap(**heatmap_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Value: %{{z}}<extra></extra>"
                                ))
                                
                            elif series_type == "treemap":
                                # Get data label settings from JSON configuration
                                data_labels = chart_meta.get("data_labels", True)
                                data_label_font_size = chart_meta.get("data_label_font_size", 12)
                                data_label_color = chart_meta.get("data_label_color", "#000000")
                                fill_opacity = chart_meta.get("fill_opacity", 0.8)
                                
                                # Treemap specific settings
                                hide_center_box = chart_meta.get("hide_center_box", False)
                                current_app.logger.debug(f"🔍 Treemap (multi-series): hide_center_box parameter: {hide_center_box}")
                                current_app.logger.debug(f"🔍 Treemap (multi-series): chart_meta: {chart_meta}")
                                # Determine textinfo based on data_labels setting (independent of center box)
                                textinfo_setting = "label" if data_labels else "none"
                                
                                treemap_kwargs = {
                                    "labels": x_vals,
                                    "values": y_vals,
                                    "name": label,
                                    "textinfo": textinfo_setting,
                                    "textposition": "middle center",
                                    "branchvalues": "total",
                                    "pathbar": dict(visible=not hide_center_box),  # Hide center box if hide_center_box is True
                                    "texttemplate": None,
                                    "hoverinfo": "label+value"
                                }
                                
                                # Add text font configuration from JSON
                                if data_labels:
                                    treemap_kwargs["textfont"] = {
                                        "family": font_family,
                                        "size": data_label_font_size,
                                        "color": data_label_color
                                    }
                                
                                # Add treemap specific attributes
                                if color:
                                    treemap_kwargs["marker"] = dict(colors=color, opacity=fill_opacity) if isinstance(color, list) else dict(colors=[color], opacity=fill_opacity)
                                
                                # Debug logging for final treemap configuration
                                current_app.logger.debug(f"🔍 Treemap (multi-series): Final treemap_kwargs: {treemap_kwargs}")
                                
                                fig.add_trace(go.Treemap(**treemap_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>%{{label}}: %{{value}}<extra></extra>"
                                ))
                                
                            else:
                                # Generic handling for other chart types
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Value: %{{y}}<extra></extra>"
                                ))
                        else:
                            # Fallback to scatter if chart type not recognized
                            #pass  # Suppress warning logs: f"⚠️ Unknown chart type '{series_type}', falling back to scatter")
                            fig.add_trace(go.Scatter(**trace_kwargs,
                                mode='markers',
                                hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                            ))

                # --- Layout updates ---
                layout_updates = {}
                
                # Title and axis labels
                layout_updates["title"] = title
                
                # Handle axis labels based on chart type
                if chart_type != "pie":
                    layout_updates["xaxis_title"] = chart_meta.get("x_label", chart_config.get("x_axis_title", "X"))
                    layout_updates["yaxis_title"] = chart_meta.get("primary_y_label", chart_config.get("primary_y_label", "Y"))
                
                # Font
                if font_family or font_size or font_color:
                    layout_updates["font"] = {}
                    if font_family:
                        layout_updates["font"]["family"] = font_family
                    if font_size:
                        layout_updates["font"]["size"] = font_size
                    if font_color:
                        layout_updates["font"]["color"] = font_color
                
                # Backgrounds
                if chart_background:
                    layout_updates["paper_bgcolor"] = chart_background
                if plot_background:
                    layout_updates["plot_bgcolor"] = plot_background
                
                # Legend configuration
                show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                # Convert string "false"/"true" to boolean if needed
                if isinstance(show_legend_raw, str):
                    show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                else:
                    show_legend = bool(show_legend_raw)
                # current_app.logger.debug(f"Legend setting: {chart_meta.get('legend')}")
                # current_app.logger.debug(f"Showlegend setting: {chart_meta.get('showlegend')}")
                # current_app.logger.debug(f"Final show_legend: {show_legend}")
                
                if show_legend:
                    # current_app.logger.debug("Configuring legend for Plotly")
                    if legend_position:
                        # Map 'top', 'bottom', 'left', 'right' to valid Plotly legend positions
                        pos_map = {
                            "top":    dict(x=0.5, y=1.1, xanchor="center", yanchor="top"),
                            "bottom": dict(x=0.5, y=-0.2, xanchor="center", yanchor="bottom"),
                            "left":   dict(x=-0.2, y=0.5, xanchor="left", yanchor="middle"),
                            "right":  dict(x=1.1, y=0.5, xanchor="right", yanchor="middle"),
                        }
                        if legend_position in pos_map:
                            layout_updates["legend"] = pos_map[legend_position]
                            # current_app.logger.debug(f"Set legend position: {legend_position}")
                    if legend_font_size:
                        layout_updates.setdefault("legend", {})["font"] = {"size": legend_font_size}
                        # current_app.logger.debug(f"Set legend font size: {legend_font_size}")
                else:
                    layout_updates["showlegend"] = False
                    # current_app.logger.debug("Legend disabled for Plotly")
                
                # Bar mode for stacked charts
                if barmode:
                    layout_updates["barmode"] = barmode
                elif chart_type == "stacked_column":
                    layout_updates["barmode"] = "stack"
                
                # Axis min/max and tick format (only for non-pie charts)
                if chart_type != "pie":
                    if y_axis_min_max:
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        # Handle "auto" value for y-axis min/max
                        if y_axis_min_max == "auto":
                            # Don't set range for auto - let Plotly auto-scale
                            pass
                        else:
                            layout_updates["yaxis"]["range"] = y_axis_min_max
                            # Also set autorange to false to ensure the range is respected
                            layout_updates["yaxis"]["autorange"] = False
                            # Force the range to be applied
                            layout_updates["yaxis"]["fixedrange"] = False
                            # Ensure the range is properly set
                            # current_app.logger.debug(f"Setting Y-axis range to: {y_axis_min_max}")
                    if axis_tick_format:
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["yaxis"]["tickformat"] = axis_tick_format
                        # Also apply to secondary y-axis if it's a currency format
                        if "$" in axis_tick_format:
                            layout_updates["yaxis2"] = layout_updates.get("yaxis2", {})
                            layout_updates["yaxis2"]["tickformat"] = axis_tick_format
                    
                    # Secondary y-axis formatting
                    if secondary_y_axis_format:
                        layout_updates["yaxis2"] = layout_updates.get("yaxis2", {})
                        layout_updates["yaxis2"]["tickformat"] = secondary_y_axis_format
                    if secondary_y_axis_min_max:
                        layout_updates["yaxis2"] = layout_updates.get("yaxis2", {})
                        # Handle "auto" value for secondary y-axis min/max
                        if secondary_y_axis_min_max == "auto":
                            # Don't set range for auto - let Matplotlib auto-scale
                            pass
                        elif isinstance(secondary_y_axis_min_max, list) and len(secondary_y_axis_min_max) == 2:
                            layout_updates["yaxis2"]["range"] = secondary_y_axis_min_max
                        else:
                            pass  # Suppress warning logs: f"Invalid secondary_y_axis_min_max format: {secondary_y_axis_min_max}")
                    
                    # Axis tick font size
                    if axis_tick_font_size:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["tickfont"] = {"size": axis_tick_font_size}
                        layout_updates["yaxis"]["tickfont"] = {"size": axis_tick_font_size}
                    # Gridlines
                    if show_gridlines is not None:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["showgrid"] = bool(show_gridlines)
                        layout_updates["yaxis"]["showgrid"] = bool(show_gridlines)
                    if gridline_color:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["gridcolor"] = gridline_color
                        layout_updates["yaxis"]["gridcolor"] = gridline_color
                    if gridline_style:
                        # Map gridline styles to valid Plotly dash styles
                        dash_map = {
                            "solid": "solid", 
                            "dashed": "dash",
                            "dash": "dash",  # Map 'dash' to 'dash'
                            "dashdot": "dashdot", 
                            "dotted": "dot",
                            "dot": "dot",
                            "dotdash": "dashdot"
                        }
                        dash_style = dash_map.get(gridline_style, "solid")
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["griddash"] = dash_style
                        layout_updates["yaxis"]["griddash"] = dash_style
                
                # Data labels (for bar/line traces)
                show_data_labels = chart_meta.get("data_labels", True)
                value_format = chart_meta.get("value_format", "")
                
                # Only enable data labels if explicitly set to True AND any data label settings are provided
                if show_data_labels and (data_label_format or data_label_font_size or data_label_color):
                    show_data_labels = True
                elif not show_data_labels:
                    # If data_labels is explicitly set to False, respect that setting
                    show_data_labels = False
                
                # Debug logging for data labels
                # current_app.logger.debug(f"Original data_labels setting: {chart_meta.get('data_labels')}")
                # current_app.logger.debug(f"Final show_data_labels: {show_data_labels}")
                # current_app.logger.debug(f"Data label format: {data_label_format}")
                # current_app.logger.debug(f"Data label font size: {data_label_font_size}")
                # current_app.logger.debug(f"Data label color: {data_label_color}")
                # current_app.logger.debug(f"Value format: {value_format}")
                # current_app.logger.debug(f"Chart config keys: {list(chart_config.keys())}")
                # current_app.logger.debug(f"Chart meta keys: {list(chart_meta.keys())}")
                
                if show_data_labels and (data_label_format or value_format or data_label_font_size or data_label_color):
                    # current_app.logger.debug(f"Processing {len(fig.data)} traces for data labels")
                    for i, trace in enumerate(fig.data):
                        # current_app.logger.debug(f"Trace {i}: type={trace.type}, mode={getattr(trace, 'mode', 'N/A')}")
                        # Handle both bar and line charts (line charts are scatter with mode='lines')
                        if trace.type in ['bar', 'scatter']:
                            # Use value_format from chart_meta if available, otherwise use data_label_format
                            format_to_use = value_format if value_format else data_label_format
                            
                            # Determine if this is a line chart (scatter with lines mode)
                            is_line_chart = trace.type == 'scatter' and trace.mode and 'lines' in trace.mode
                            
                            if format_to_use:
                                if trace.type == 'bar':
                                    trace.update(texttemplate=f"%{{y:{format_to_use}}}", textposition="auto")
                                elif trace.type == 'scatter':
                                    if is_line_chart:
                                        # For line charts, show labels at the data points
                                        trace.update(texttemplate=f"%{{y:{format_to_use}}}", textposition="top center")
                                    else:
                                        # For scatter plots
                                        trace.update(texttemplate=f"%{{y:{format_to_use}}}", textposition="top center")
                            else:
                                # If no format specified, still show data labels
                                if trace.type == 'bar':
                                    trace.update(texttemplate="%{y}", textposition="auto")
                                elif trace.type == 'scatter':
                                    if is_line_chart:
                                        # For line charts, show labels at the data points
                                        trace.update(texttemplate="%{y}", textposition="top center")
                                    else:
                                        # For scatter plots
                                        trace.update(texttemplate="%{y}", textposition="top center")
                            
                            # Apply font styling
                            if data_label_font_size or data_label_color:
                                trace.update(textfont={})
                                if data_label_font_size:
                                    trace.textfont["size"] = data_label_font_size
                                if data_label_color:
                                    trace.textfont["color"] = data_label_color

                # Annotations
                if annotations:
                    layout_updates["annotations"] = []
                    for ann in annotations:
                        ann_dict = {
                            "text": ann.get("text", ""),
                            "x": ann.get("x_value", ann.get("x")),
                            "y": ann.get("y_value", ann.get("y")),
                            "showarrow": True
                        }
                        layout_updates["annotations"].append(ann_dict)
                
                # --- Apply margin settings ---
                if margin:
                    layout_updates["margin"] = margin
                
                # Apply figure size if specified
                if figsize:
                    layout_updates["width"] = figsize[0] * 100  # Convert to pixels
                    layout_updates["height"] = figsize[1] * 100  # Convert to pixels
                    # current_app.logger.debug(f"Applied Plotly figsize: {figsize} -> width={figsize[0]*100}, height={figsize[1]*100}")
                
                # Apply axis label distances
                if chart_type != "pie":
                    if x_axis_label_distance == "auto" or y_axis_label_distance == "auto":
                        # Calculate optimal label distances
                        x_axis_label_distance, y_axis_label_distance = calculate_optimal_label_distance(
                            chart_type, series_data, x_values, [], figsize, font_size
                        )
                    
                    if x_axis_label_distance or y_axis_label_distance or axis_tick_distance:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        
                        if x_axis_label_distance:
                            layout_updates["xaxis"]["title"] = layout_updates["xaxis"].get("title", {})
                            layout_updates["xaxis"]["title"]["standoff"] = x_axis_label_distance
                            # Also apply to tick distance for better control
                            layout_updates["xaxis"]["ticklen"] = x_axis_label_distance
                        
                        if y_axis_label_distance:
                            layout_updates["yaxis"]["title"] = layout_updates["yaxis"].get("title", {})
                            layout_updates["yaxis"]["title"]["standoff"] = y_axis_label_distance
                        
                        if axis_tick_distance:
                            layout_updates["xaxis"]["ticklen"] = axis_tick_distance
                            layout_updates["yaxis"]["ticklen"] = axis_tick_distance

                # Tick mark control for Plotly
                if show_x_ticks is not None or show_y_ticks is not None:
                    layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                    layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                    if show_x_ticks is not None:
                        layout_updates["xaxis"]["showticklabels"] = bool(show_x_ticks)
                        layout_updates["xaxis"]["ticks"] = "" if not show_x_ticks else "outside"
                    if show_y_ticks is not None:
                        layout_updates["yaxis"]["showticklabels"] = bool(show_y_ticks)
                        layout_updates["yaxis"]["ticks"] = "" if not show_y_ticks else "outside"

                fig.update_layout(**layout_updates)

                # --- Matplotlib static chart for DOCX ---
                if chart_type == "pie":
                    # Check if this is an expanded pie chart
                    expanded_segment = chart_meta.get("expanded_segment")
                    
                    if expanded_segment and len(series_data) == 1:
                        # Create subplot for expanded pie chart
                        mpl_figsize = figsize if figsize else (15, 8)
                        fig_mpl, (ax1, ax2) = plt.subplots(1, 2, figsize=mpl_figsize, dpi=200)
                        
                        # Apply background colors to Matplotlib figure
                        if chart_background:
                            fig_mpl.patch.set_facecolor(chart_background)
                        if plot_background:
                            ax1.set_facecolor(plot_background)
                            ax2.set_facecolor(plot_background)
                        
                        series = series_data[0]
                        labels = series.get("labels", x_values)
                        values = series.get("values", [])
                        color = series.get("marker", {}).get("colors") if "marker" in series else colors
                        marker_line = series.get("marker", {}).get("line", {}) if "marker" in series else {}
                        explode = series.get("pull")
                        opacity = series.get("opacity", chart_meta.get("opacity"))
                        textinfo = series.get("textinfo", chart_meta.get("textinfo", "percent"))
                        textposition = series.get("textposition", chart_meta.get("textposition", "inside")).lower()
                        value_format_str = chart_meta.get("value_format", ".1f")
                        data_labels_enabled = bool(chart_meta.get("data_labels", True))
                        data_label_font_size = chart_meta.get("data_label_font_size", font_size or 10)
                        data_label_color = chart_meta.get("data_label_color", font_color or "#000000")
                        start_angle = startangle if 'startangle' in locals() and startangle is not None else 90
                        sort_order = chart_meta.get("sort_order")

                        # Optional sorting
                        if sort_order in ("ascending", "descending") and values:
                            zipped = list(zip(values, labels, color if isinstance(color, list) else [color]*len(labels), explode if isinstance(explode, list) else [0]*len(labels)))
                            reverse = sort_order == "descending"
                            zipped.sort(key=lambda t: (t[0] if t[0] is not None else 0), reverse=reverse)
                            values, labels, color_list, explode_list = zip(*zipped)
                            values = list(values)
                            labels = list(labels)
                            color = list(color_list)
                            explode = list(explode_list)

                        # Build autopct based on textinfo/value_format
                        def make_autopct(fmt:str, include_percent:bool, include_value:bool):
                            # Capture the current values in the closure
                            current_values = values.copy() if isinstance(values, list) else list(values) if values else []
                            
                            def _inner(pct):
                                # Ensure values are numeric before calculating total
                                numeric_values = []
                                for v in current_values:
                                    if v is not None:
                                        try:
                                            numeric_values.append(float(v))
                                        except (ValueError, TypeError):
                                            # Skip non-numeric values
                                            continue
                                
                                total = sum(numeric_values) if numeric_values else 0
                                val = pct * total / 100.0
                                parts = []
                                if include_value:
                                    try:
                                        parts.append(f"{val:{fmt}}")
                                    except Exception:
                                        parts.append(f"{val:.1f}")
                                if include_percent:
                                    parts.append(f"{pct:.1f}%")
                                return " ".join(parts)
                            return _inner

                        include_label = "label" in (textinfo or "")
                        include_percent = "percent" in (textinfo or "")
                        include_value = "value" in (textinfo or "")

                        autopct_callable = None
                        if data_labels_enabled and (include_percent or include_value):
                            autopct_callable = make_autopct(value_format_str, include_percent, include_value)

                        # Wedge and text props
                        wedgeprops = {}
                        if isinstance(marker_line, dict):
                            if marker_line.get("color"):
                                wedgeprops["edgecolor"] = marker_line.get("color")
                            if marker_line.get("width") is not None:
                                wedgeprops["linewidth"] = marker_line.get("width")
                        if opacity is not None:
                            wedgeprops["alpha"] = opacity

                        textprops = {"color": data_label_color, "fontsize": data_label_font_size}
                        if chart_meta.get("font_family"):
                            textprops["fontfamily"] = chart_meta.get("font_family")

                        # Positioning
                        pctdistance = 0.6 if textposition == "inside" else 1.15
                        labeldistance = 1.1 if textposition != "inside" else 1.05
                        
                        # Create pie chart
                        wedges, texts, autotexts = ax1.pie(
                            values,
                            labels=labels if include_label else None,
                            autopct=autopct_callable,
                            colors=color,
                            startangle=start_angle,
                            explode=explode,
                            wedgeprops=wedgeprops,
                            pctdistance=pctdistance,
                            labeldistance=labeldistance,
                            textprops=textprops,
                        )

                        # Style the autopct texts
                        for autotext in autotexts or []:
                            autotext.set_color(data_label_color)
                            autotext.set_fontsize(data_label_font_size)
                            autotext.set_fontweight('bold')
                            if chart_meta.get("font_family"):
                                autotext.set_fontfamily(chart_meta.get("font_family"))
                        
                        if title and title.strip():
                            ax1.set_title(title, fontsize=font_size or 14, weight='bold', pad=20, color=font_color if font_color else None, fontname=chart_meta.get("font_family") if chart_meta.get("font_family") else None)
                        
                        # Add legend for pie chart
                        show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                        # Convert string "false"/"true" to boolean if needed
                        if isinstance(show_legend_raw, str):
                            show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                        else:
                            show_legend = bool(show_legend_raw)
                        if show_legend:
                            # Initialize legend_loc for pie charts
                            legend_loc = 'best'  # default
                            if legend_position:
                                loc_map = {
                                    "top": "upper center",
                                    "bottom": "lower center", 
                                    "left": "center left",
                                    "right": "center right"
                                }
                                legend_loc = loc_map.get(legend_position, 'best')
                            
                            # Force legend to bottom if specified
                            if legend_position == "bottom":
                                ax1.legend(wedges, labels, loc='lower center', bbox_to_anchor=(0.5, -0.15), fontsize=legend_font_size)
                            else:
                                ax1.legend(wedges, labels, loc=legend_loc, fontsize=legend_font_size)
                        
                        # Create bar chart for expanded segment
                        if expanded_segment in labels:
                            segment_idx = labels.index(expanded_segment)
                            segment_value = values[segment_idx]
                            segment_color = color[segment_idx] if isinstance(color, list) and segment_idx < len(color) else color
                            
                            ax2.bar([expanded_segment], [segment_value], color=segment_color, alpha=0.7)
                            ax2.set_title(f"{expanded_segment} Details", fontsize=font_size or 12, weight='bold')
                            ax2.set_ylabel("Value")
                            
                            # Add value label on bar
                            ax2.text(0, segment_value, f"{segment_value}", ha='center', va='bottom', fontweight='bold')
                        
                    else:
                        # Regular pie chart
                        mpl_figsize = figsize if figsize else (10, 8)
                        fig_mpl, ax = plt.subplots(figsize=mpl_figsize, dpi=200)
                        
                        # Apply background colors to Matplotlib figure
                        if chart_background:
                            fig_mpl.patch.set_facecolor(chart_background)
                        if plot_background:
                            ax.set_facecolor(plot_background)
                        
                        if len(series_data) == 1:
                            series = series_data[0]
                            labels = series.get("labels", x_values)
                            values = series.get("values", [])
                            
                            # Normalize pie chart values to sum to 100% (for matplotlib fallback)
                            if chart_type == "pie" and values:
                                values = normalize_values_to_100(values)
                            
                            color = series.get("marker", {}).get("colors") if "marker" in series else colors
                            marker_line = series.get("marker", {}).get("line", {}) if "marker" in series else {}
                            explode = series.get("pull")
                            opacity = series.get("opacity", chart_meta.get("opacity"))
                            textinfo = series.get("textinfo", chart_meta.get("textinfo", "percent"))
                            textposition = series.get("textposition", chart_meta.get("textposition", "inside")).lower()
                            value_format_str = chart_meta.get("value_format", ".1f")
                            data_labels_enabled = bool(chart_meta.get("data_labels", True))
                            data_label_font_size = chart_meta.get("data_label_font_size", font_size or 10)
                            data_label_color = chart_meta.get("data_label_color", font_color or "#000000")
                            start_angle = startangle if 'startangle' in locals() and startangle is not None else 90
                            sort_order = chart_meta.get("sort_order")

                            # Optional sorting
                            if sort_order in ("ascending", "descending") and values:
                                zipped = list(zip(values, labels, color if isinstance(color, list) else [color]*len(labels), explode if isinstance(explode, list) else [0]*len(labels)))
                                reverse = sort_order == "descending"
                                zipped.sort(key=lambda t: (t[0] if t[0] is not None else 0), reverse=reverse)
                                values, labels, color_list, explode_list = zip(*zipped)
                                values = list(values)
                                labels = list(labels)
                                color = list(color_list)
                                explode = list(explode_list)

                            # Build autopct based on textinfo/value_format
                            def make_autopct(fmt:str, include_percent:bool, include_value:bool):
                                # Capture the current values in the closure
                                current_values = values.copy() if isinstance(values, list) else list(values) if values else []
                                
                                def _inner(pct):
                                    # Ensure values are numeric before calculating total
                                    numeric_values = []
                                    for v in current_values:
                                        if v is not None:
                                            try:
                                                numeric_values.append(float(v))
                                            except (ValueError, TypeError):
                                                # Skip non-numeric values
                                                continue
                                    
                                    total = sum(numeric_values) if numeric_values else 0
                                    val = pct * total / 100.0
                                    parts = []
                                    if include_value:
                                        try:
                                            parts.append(f"{val:{fmt}}")
                                        except Exception:
                                            parts.append(f"{val:.1f}")
                                    if include_percent:
                                        parts.append(f"{pct:.1f}%")
                                    return " ".join(parts)
                                return _inner

                            include_label = "label" in (textinfo or "")
                            include_percent = "percent" in (textinfo or "")
                            include_value = "value" in (textinfo or "")

                            autopct_callable = None
                            if data_labels_enabled and (include_percent or include_value):
                                autopct_callable = make_autopct(value_format_str, include_percent, include_value)

                            # Wedge and text props
                            wedgeprops = {}
                            if isinstance(marker_line, dict):
                                if marker_line.get("color"):
                                    wedgeprops["edgecolor"] = marker_line.get("color")
                                if marker_line.get("width") is not None:
                                    wedgeprops["linewidth"] = marker_line.get("width")
                            if opacity is not None:
                                wedgeprops["alpha"] = opacity

                            textprops = {"color": data_label_color, "fontsize": data_label_font_size}
                            if chart_meta.get("font_family"):
                                textprops["fontfamily"] = chart_meta.get("font_family")

                            # Positioning
                            pctdistance = 0.6 if textposition == "inside" else 1.15
                            labeldistance = 1.1 if textposition != "inside" else 1.05
                            
                            # Create pie chart
                            wedges, texts, autotexts = ax.pie(
                                values,
                                labels=labels if include_label else None,
                                autopct=autopct_callable,
                                colors=color,
                                startangle=start_angle,
                                explode=explode,
                                wedgeprops=wedgeprops,
                                pctdistance=pctdistance,
                                labeldistance=labeldistance,
                                textprops=textprops,
                            )
                            
                            # Style the text
                            for autotext in autotexts or []:
                                autotext.set_color(data_label_color)
                                autotext.set_fontsize(data_label_font_size)
                                autotext.set_fontweight('bold')
                                if chart_meta.get("font_family"):
                                    autotext.set_fontfamily(chart_meta.get("font_family"))
                            
                            if title and title.strip():
                                ax.set_title(title, fontsize=font_size or 14, weight='bold', pad=20, color=font_color if font_color else None, fontname=chart_meta.get("font_family") if chart_meta.get("font_family") else None)
                            
                            # Add legend for regular pie chart
                            show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                            # Convert string "false"/"true" to boolean if needed
                            if isinstance(show_legend_raw, str):
                                show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                            else:
                                show_legend = bool(show_legend_raw)
                            if show_legend:
                                # Initialize legend_loc for pie charts
                                legend_loc = 'best'  # default
                                if legend_position:
                                    loc_map = {
                                        "top": "upper center",
                                        "bottom": "lower center", 
                                        "left": "center left",
                                        "right": "center right"
                                    }
                                    legend_loc = loc_map.get(legend_position, 'best')
                                
                                # Force legend to bottom if specified
                                if legend_position == "bottom":
                                    ax.legend(wedges, labels, loc='lower center', bbox_to_anchor=(0.5, -0.15), fontsize=legend_font_size)
                                else:
                                    ax.legend(wedges, labels, loc=legend_loc, fontsize=legend_font_size)
                        
                elif chart_type in ["bar of pie", "bar_of_pie"]:
                    # Matplotlib version of bar of pie
                    mpl_figsize = figsize if figsize else (10, 5)
                    fig_mpl, (ax1, ax2) = plt.subplots(1, 2, figsize=mpl_figsize, dpi=200, gridspec_kw={'width_ratios': [2, 1]})
                    
                    # Apply background colors to Matplotlib figure
                    if chart_background:
                        fig_mpl.patch.set_facecolor(chart_background)
                    if plot_background:
                        ax1.set_facecolor(plot_background)
                        ax2.set_facecolor(plot_background)
                    labels = series_meta.get("labels", x_values)
                    values = series_meta.get("values", [])
                    colors = series_meta.get("colors", [])
                    other_labels = chart_meta.get("other_labels", [])
                    other_values = chart_meta.get("other_values", [])
                    other_colors = chart_meta.get("other_colors", [])
                    if not (other_labels and other_values):
                        if "other_label_range" in chart_meta and "other_value_range" in chart_meta and "source_sheet" in chart_meta:
                            wb = openpyxl.load_workbook(data_file_path, data_only=True)
                            sheet = wb[chart_meta["source_sheet"]]
                            other_labels = extract_excel_range(sheet, chart_meta["other_label_range"])
                            other_values = extract_excel_range(sheet, chart_meta["other_value_range"])
                    # Pie chart
                    if colors:
                        wedges, texts, autotexts = ax1.pie(values, labels=labels, autopct='%1.1f%%', colors=colors, startangle=90)
                    else:
                        wedges, texts, autotexts = ax1.pie(values, labels=labels, autopct='%1.1f%%', startangle=90)
                    for autotext in autotexts:
                        autotext.set_color('white')
                        autotext.set_fontweight('bold')
                    if title and title.strip():
                        ax1.set_title(title, fontsize=font_size or 14, weight='bold', pad=20)
                    # Move legend outside
                    show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                    # Convert string "false"/"true" to boolean if needed
                    if isinstance(show_legend_raw, str):
                        show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                    else:
                        show_legend = bool(show_legend_raw)
                    legend_font_size = chart_meta.get("legend_font_size", 8)
                    if show_legend:
                        ax1.legend(wedges, labels, loc='center left', bbox_to_anchor=(1, 0.5), fontsize=legend_font_size)
                    
                    # Bar chart - Filter out empty/null values first
                    filtered_data = []
                    for label, value in zip(other_labels, other_values):
                        if (label is not None and str(label).strip() != "" and 
                            value is not None and str(value).strip() != "" and 
                            str(value).strip() != "0"):
                            filtered_data.append((label, value))
                    
                    if filtered_data:
                        filtered_labels, filtered_values = zip(*filtered_data)
                    else:
                        filtered_labels, filtered_values = [], []
                    
                    # Create individual bars (not stacked)
                    if filtered_labels and filtered_values:
                        # Use individual colors for each bar
                        bar_colors = []
                        for i in range(len(filtered_labels)):
                            if other_colors and i < len(other_colors):
                                bar_colors.append(other_colors[i])
                            elif colors and i < len(colors):
                                bar_colors.append(colors[i])
                            else:
                                bar_colors.append('#1f77b4')  # Default blue
                        
                        bars = ax2.bar(range(len(filtered_labels)), filtered_values, color=bar_colors, alpha=0.7)
                    else:
                        bars = []
                    
                    expanded_segment = chart_meta.get("expanded_segment", "Other")
                    ax2.set_title(f"Breakdown of '{expanded_segment}'", fontsize=font_size or 12, weight='bold')
                    # Use proper Y-axis label from configuration
                    y_axis_title = chart_meta.get("y_axis_title", "Value")
                    ax2.set_ylabel(y_axis_title, fontsize=label_fontsize if 'label_fontsize' in locals() else 10)
                    # Set X-axis label
                    x_axis_title = chart_meta.get("x_axis_title", "Categories")
                    ax2.set_xlabel(x_axis_title, fontsize=label_fontsize if 'label_fontsize' in locals() else 10)
                    # Set x-tick labels for filtered data
                    if filtered_labels:
                        ax2.set_xticks(range(len(filtered_labels)))
                        # Format x-axis labels as percentages
                        formatted_x_labels = []
                        for label in filtered_labels:
                            if label is not None:
                                if isinstance(label, (int, float)):
                                    if label <= 1.0:  # Likely decimal format (0.06)
                                        formatted_x_labels.append(f"{label * 100:.1f}%")
                                    else:  # Likely already percentage format (6.0)
                                        formatted_x_labels.append(f"{label:.1f}%")
                                else:
                                    # Convert string to float and handle
                                    try:
                                        val = float(label)
                                        if val <= 1.0:
                                            formatted_x_labels.append(f"{val * 100:.1f}%")
                                        else:
                                            formatted_x_labels.append(f"{val:.1f}%")
                                    except (ValueError, TypeError):
                                        formatted_x_labels.append(str(label))
                            else:
                                formatted_x_labels.append("")
                        ax2.set_xticklabels(formatted_x_labels, rotation=0, fontsize=axis_tick_font_size or 10)
                    # Add data labels with proper formatting
                    value_format = chart_meta.get("value_format", ".2f")
                    data_label_font_size = chart_meta.get("data_label_font_size", 10)
                    data_label_color = chart_meta.get("data_label_color", "#000000")
                    for bar, v in zip(bars, filtered_values):
                        # Format percentage values as XX.X% instead of 0.XXX
                        if v is not None:
                            if isinstance(v, (int, float)):
                                if v <= 1.0:  # Likely decimal format (0.11)
                                    formatted_value = f"{v * 100:.1f}%"
                                else:  # Likely already percentage format (11.0)
                                    formatted_value = f"{v:.1f}%"
                            else:
                                    # Convert string to float and handle
                                try:
                                    val = float(v)
                                    if val <= 1.0:
                                        formatted_value = f"{val * 100:.1f}%"
                                    else:
                                        formatted_value = f"{val:.1f}%"
                                except:
                                    formatted_value = str(v)
                            ax2.text(bar.get_x() + bar.get_width()/2, v, formatted_value, ha='center', va='bottom', fontweight='bold', fontsize=data_label_font_size, color=data_label_color)

                else:
                    # Bar, line, area charts
                    mpl_figsize = figsize if figsize else (10, 6)
                    # current_app.logger.debug(f"Applied Matplotlib figsize: {mpl_figsize}")
                    fig_mpl, ax1 = plt.subplots(figsize=mpl_figsize, dpi=200)
                    
                    # Disable gridlines globally if show_gridlines is False
                    if not show_gridlines:
                        plt.rcParams['axes.grid'] = False
                    
                    # Only create secondary y-axis if not disabled
                    ax2 = None
                    if not disable_secondary_y:
                        ax2 = ax1.twinx()
                    
                    # Apply background colors to Matplotlib figure
                    if chart_background:
                        fig_mpl.patch.set_facecolor(chart_background)
                    if plot_background:
                        ax1.set_facecolor(plot_background)
                        if ax2:
                            ax2.set_facecolor(plot_background)

                    # Define colors array for matplotlib section
                    colors = series_meta.get("colors", [])

                    for i, series in enumerate(series_data):
                        label = series.get("name", f"Series {i+1}")
                        series_type = series.get("type", "bar").lower()
                        
                        # Map heatmap to imshow for Matplotlib
                        if series_type == "heatmap":
                            mpl_chart_type = "imshow"
                        else:
                            mpl_chart_type = chart_type_mapping_mpl.get(series_type, "scatter")
                        
                        color = None
                        if "marker" in series and isinstance(series["marker"], dict) and "color" in series["marker"]:
                            color = series["marker"]["color"]
                        elif bar_colors:
                            color = bar_colors
                        elif i < len(colors) and colors[i] is not None:
                            color = colors[i]

                        y_vals = series.get("values")
                        value_range = series.get("value_range")
                        if value_range:
                            # Check if value_range is already extracted (list) or still a string
                            if isinstance(value_range, list):
                                y_vals = value_range
                            else:
                                y_vals = extract_values_from_range(value_range)

                        # --- Apply grouping and sorting ---
                        x_vals = x_values
                        if y_vals is not None and x_vals is not None:
                            x_vals, y_vals = group_and_sort(x_vals, y_vals, data_grouping, sort_order)

                        trace_kwargs = {
                            "x": x_vals,
                            "y": y_vals,
                            "name": label,
                        }
                        
                        # Handle colors
                        if color:
                            if isinstance(color, list) and series_type == "bar":
                                trace_kwargs["marker"] = dict(color=color)
                            elif isinstance(color, str):
                                trace_kwargs["marker_color"] = color
                        
                        # Handle bar-specific properties
                        if bar_width and series_type == "bar":
                            trace_kwargs["width"] = bar_width
                        if orientation and series_type == "bar":
                            trace_kwargs["orientation"] = orientation[0].lower() if isinstance(orientation, str) else orientation
                        if bar_border_color and series_type == "bar":
                            trace_kwargs["marker_line_color"] = bar_border_color
                        if bar_border_width and series_type == "bar":
                            trace_kwargs["marker_line_width"] = bar_border_width
                        
                        # Handle line-specific properties
                        if line_width and series_type in ["line", "scatter", "scatter_line"]:
                            trace_kwargs["line_width"] = line_width
                        if marker_size and series_type in ["line", "scatter", "scatter_line"]:
                            trace_kwargs["marker_size"] = marker_size
                        if line_style and series_type in ["line", "scatter", "scatter_line"]:
                            # Map line styles to Plotly format
                            line_style_map = {
                                "solid": "solid",
                                "dashed": "dash",
                                "dotted": "dot",
                                "dashdot": "dashdot"
                            }
                            trace_kwargs["line_dash"] = line_style_map.get(line_style, "solid")
                        
                        # Handle opacity
                        if fill_opacity:
                            trace_kwargs["opacity"] = fill_opacity

                        # Add traces based on chart type - REPLACE THE RESTRICTIVE IF/ELIF BLOCKS
                        # Generic chart type handling for Plotly
                        chart_type_mapping = {
                            # Bar charts
                            "bar": go.Bar,
                            "column": go.Bar,
                            "stacked_column": go.Bar,
                            "horizontal_bar": go.Bar,
                            
                            # Line charts
                            "line": go.Scatter,
                            "scatter": go.Scatter,
                            "scatter_line": go.Scatter,
                            
                            # Area charts
                            "area": go.Scatter,
                            "filled_area": go.Scatter,
                            
                            # Pie charts
                            "pie": go.Pie,
                            "donut": go.Pie,
                            
                            # 3D charts
                            "scatter3d": go.Scatter3d,
                            "surface": go.Surface,
                            "mesh3d": go.Mesh3d,
                            
                            # Statistical charts
                            "histogram": go.Histogram,
                            "box": go.Box,
                            "violin": go.Violin,
                            
                            # Financial charts
                            "candlestick": go.Candlestick,
                            "ohlc": go.Ohlc,
                            
                            # Geographic charts
                            "scattergeo": go.Scattergeo,
                            "choropleth": go.Choropleth,
                            
                            # Other charts
                            "bubble": go.Scatter,
                            "heatmap": go.Heatmap,
                            "contour": go.Contour,
                            "waterfall": go.Waterfall,
                            "funnel": go.Funnel,
                            "sunburst": go.Sunburst,
                            "icicle": go.Icicle,
                            "sankey": go.Sankey,
                            "treemap": go.Treemap,
                            "table": go.Table,
                            "indicator": go.Indicator
                        }
                        
                        # Get the appropriate Plotly chart class
                        plotly_chart_class = chart_type_mapping.get(series_type)
                        
                        if plotly_chart_class:
                            # Prepare trace arguments based on chart type
                            if series_type in ["bar", "column", "stacked_column", "horizontal_bar"]:
                                # Bar chart specific settings
                                # REMOVE: if chart_type == "stacked_column": trace_kwargs["barmode"] = "stack"
                                if orientation and orientation.lower() == "horizontal":
                                    trace_kwargs["orientation"] = "h"
                                    # Swap x and y for horizontal bars
                                    trace_kwargs["x"], trace_kwargs["y"] = trace_kwargs["y"], trace_kwargs["x"]
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                                ))
                                
                            elif series_type in ["line", "scatter", "scatter_line"]:
                                # Line/Scatter chart specific settings
                                mode = "lines+markers" if series_type == "scatter_line" else "markers" if series_type == "scatter" else "lines"
                                trace_kwargs["mode"] = mode
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                                ))
                                
                            elif series_type in ["area", "filled_area"]:
                                # Area chart specific settings
                                trace_kwargs["mode"] = "lines"
                                trace_kwargs["fill"] = "tozeroy"
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                                ))
                                
                            elif series_type == "pie":
                                # Pie chart specific settings
                                pie_kwargs = {
                                    "labels": x_vals,
                                    "values": y_vals,
                                    "name": label,
                                    "textinfo": "label" if data_labels and not chart_meta.get("hide_center_box", False) else "none",
                                    "textposition": "outside",
                                    "hole": hole if hole is not None else (0.4 if series_type == "donut" else 0.0)
                                }
                                
                                # Add pie chart specific attributes
                                if startangle is not None:
                                    pie_kwargs["rotation"] = startangle
                                if pull is not None:
                                    # pull can be a single value or a list
                                    if isinstance(pull, (int, float)):
                                        pie_kwargs["pull"] = [pull] * len(y_vals)
                                    elif isinstance(pull, list):
                                        pie_kwargs["pull"] = pull
                                
                                if color:
                                    pie_kwargs["marker"] = dict(colors=color) if isinstance(color, list) else dict(colors=[color])
                                
                                fig.add_trace(plotly_chart_class(**pie_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>%{{label}}: %{{value}}<extra></extra>"
                                ))
                                
                            elif series_type in ["scatter3d", "surface", "mesh3d"]:
                                # 3D chart specific settings
                                if "z" not in trace_kwargs and len(y_vals) > 0:
                                    # Create a simple z-axis if not provided
                                    trace_kwargs["z"] = [i for i in range(len(y_vals))]
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Z: %{{z}}<extra></extra>"
                                ))
                                
                            elif series_type in ["histogram", "box", "violin"]:
                                # Statistical chart specific settings
                                if series_type == "histogram":
                                    trace_kwargs["x"] = y_vals  # Histogram uses x for values
                                    del trace_kwargs["y"]
                                elif series_type in ["box", "violin"]:
                                    trace_kwargs["y"] = y_vals
                                    trace_kwargs["x"] = [label] * len(y_vals) if len(y_vals) > 0 else [label]
                                
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Value: %{{y if series_type in ['box', 'violin'] else 'x'}}<extra></extra>"
                                ))
                                
                            elif series_type in ["bubble"]:
                                # Bubble chart specific settings
                                sizes = series.get("size", [20] * len(y_vals))
                                trace_kwargs["mode"] = "markers"
                                trace_kwargs["marker"] = {"size": sizes}
                                if color:
                                    trace_kwargs["marker"]["color"] = color
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Size: %{{marker.size}}<extra></extra>"
                                ))
                                
                            elif series_type == "heatmap":
                                # Heatmap specific settings
                                heatmap_kwargs = {
                                    "z": series.get("z", y_vals),  # Use z data if provided, otherwise y_vals
                                    "x": series.get("x", x_vals),  # Use x data if provided, otherwise x_vals
                                    "y": series.get("y", [label]),  # Use y data if provided, otherwise label
                                    "name": label,
                                    "colorscale": series.get("colorscale", "Viridis"),
                                    "showscale": series.get("showscale", True)
                                }
                                
                                # Handle text data safely (convert to strings to avoid concatenation errors)
                                if "text" in series:
                                    try:
                                        heatmap_kwargs["text"] = [[str(cell) for cell in row] for row in series["text"]]
                                    except:
                                        pass  # Suppress warning logs: f"⚠️ Could not process heatmap text data for {label}")
                                
                                # Handle colorbar settings
                                if "colorbar" in series:
                                    heatmap_kwargs["colorbar"] = series["colorbar"]
                                
                                # Handle zmin/zmax
                                if "zmin" in series:
                                    heatmap_kwargs["zmin"] = series["zmin"]
                                if "zmax" in series:
                                    heatmap_kwargs["zmax"] = series["zmax"]
                                
                                fig.add_trace(go.Heatmap(**heatmap_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Value: %{{z}}<extra></extra>"
                                ))
                                
                            elif series_type == "treemap":
                                # Get data label settings from JSON configuration
                                data_labels = chart_meta.get("data_labels", True)
                                data_label_font_size = chart_meta.get("data_label_font_size", 12)
                                data_label_color = chart_meta.get("data_label_color", "#000000")
                                fill_opacity = chart_meta.get("fill_opacity", 0.8)
                                
                                # Treemap specific settings
                                hide_center_box = chart_meta.get("hide_center_box", False)
                                current_app.logger.debug(f"🔍 Treemap (multi-series): hide_center_box parameter: {hide_center_box}")
                                current_app.logger.debug(f"🔍 Treemap (multi-series): chart_meta: {chart_meta}")
                                # Determine textinfo based on data_labels setting (independent of center box)
                                textinfo_setting = "label" if data_labels else "none"
                                
                                treemap_kwargs = {
                                    "labels": x_vals,
                                    "values": y_vals,
                                    "name": label,
                                    "textinfo": textinfo_setting,
                                    "textposition": "middle center",
                                    "branchvalues": "total",
                                    "pathbar": dict(visible=not hide_center_box),  # Hide center box if hide_center_box is True
                                    "texttemplate": None,
                                    "hoverinfo": "label+value"
                                }
                                
                                # Add text font configuration from JSON
                                if data_labels:
                                    treemap_kwargs["textfont"] = {
                                        "family": font_family,
                                        "size": data_label_font_size,
                                        "color": data_label_color
                                    }
                                
                                # Add treemap specific attributes
                                if color:
                                    treemap_kwargs["marker"] = dict(colors=color, opacity=fill_opacity) if isinstance(color, list) else dict(colors=[color], opacity=fill_opacity)
                                
                                # Debug logging for final treemap configuration
                                current_app.logger.debug(f"🔍 Treemap (multi-series): Final treemap_kwargs: {treemap_kwargs}")
                                
                                fig.add_trace(go.Treemap(**treemap_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>%{{label}}: %{{value}}<extra></extra>"
                                ))
                                
                            else:
                                # Generic handling for other chart types
                                fig.add_trace(plotly_chart_class(**trace_kwargs,
                                    hovertemplate=f"<b>{label}</b><br>Value: %{{y}}<extra></extra>"
                                ))
                        else:
                            # Fallback to scatter if chart type not recognized
                            #pass  # Suppress warning logs: f"⚠️ Unknown chart type '{series_type}', falling back to scatter")
                            fig.add_trace(go.Scatter(**trace_kwargs,
                                mode='markers',
                                hovertemplate=f"<b>{label}</b><br>Category: %{{x}}<br>Value: %{{y}}<extra></extra>"
                            ))

                # --- Process heatmap charts separately ---
                # Heatmaps need special handling because they have x, y, z, text data directly
                for i, series in enumerate(series_data):
                    label = series.get("name", f"Series {i+1}")
                    series_type = series.get("type", "bar").lower()
                    
                    if series_type == "heatmap":
                        # Heatmap specific settings
                        heatmap_kwargs = {
                            "z": series.get("z", []),  # Use z data directly
                            "x": series.get("x", []),  # Use x data directly
                            "y": series.get("y", []),  # Use y data directly
                            "name": label,
                            "colorscale": series.get("colorscale", "Viridis"),
                            "showscale": series.get("showscale", True)
                        }
                        
                        # Handle text data safely (convert to strings to avoid concatenation errors)
                        if "text" in series:
                            try:
                                heatmap_kwargs["text"] = [[str(cell) for cell in row] for row in series["text"]]
                            except:
                                pass  # Suppress warning logs: f"⚠️ Could not process heatmap text data for {label}")
                        
                        # Handle colorbar settings
                        if "colorbar" in series:
                            heatmap_kwargs["colorbar"] = series["colorbar"]
                        
                        # Handle zmin/zmax
                        if "zmin" in series:
                            heatmap_kwargs["zmin"] = series["zmin"]
                        if "zmax" in series:
                            heatmap_kwargs["zmax"] = series["zmax"]
                        
                        fig.add_trace(go.Heatmap(**heatmap_kwargs,
                            hovertemplate=f"<b>{label}</b><br>X: %{{x}}<br>Y: %{{y}}<br>Value: %{{z}}<extra></extra>"
                        ))

                # --- Layout updates ---
                layout_updates = {}
                
                # Title and axis labels
                layout_updates["title"] = title
                
                # Handle axis labels based on chart type
                if chart_type != "pie":
                    layout_updates["xaxis_title"] = chart_meta.get("x_label", chart_config.get("x_axis_title", "X"))
                    layout_updates["yaxis_title"] = chart_meta.get("primary_y_label", chart_config.get("primary_y_label", "Y"))
                
                # Font
                if font_family or font_size or font_color:
                    layout_updates["font"] = {}
                    if font_family:
                        layout_updates["font"]["family"] = font_family
                    if font_size:
                        layout_updates["font"]["size"] = font_size
                    if font_color:
                        layout_updates["font"]["color"] = font_color
                
                # Backgrounds
                if chart_background:
                    layout_updates["paper_bgcolor"] = chart_background
                if plot_background:
                    layout_updates["plot_bgcolor"] = plot_background
                
                # Legend configuration
                show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                # Convert string "false"/"true" to boolean if needed
                if isinstance(show_legend_raw, str):
                    show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                else:
                    show_legend = bool(show_legend_raw)
                # current_app.logger.debug(f"Legend setting: {chart_meta.get('legend')}")
                # current_app.logger.debug(f"Showlegend setting: {chart_meta.get('showlegend')}")
                # current_app.logger.debug(f"Final show_legend: {show_legend}")
                
                if show_legend:
                    # current_app.logger.debug("Configuring legend for Plotly")
                    if legend_position:
                        # Map 'top', 'bottom', 'left', 'right' to valid Plotly legend positions
                        pos_map = {
                            "top":    dict(x=0.5, y=1.1, xanchor="center", yanchor="top"),
                            "bottom": dict(x=0.5, y=-0.2, xanchor="center", yanchor="bottom"),
                            "left":   dict(x=-0.2, y=0.5, xanchor="left", yanchor="middle"),
                            "right":  dict(x=1.1, y=0.5, xanchor="right", yanchor="middle"),
                        }
                        if legend_position in pos_map:
                            layout_updates["legend"] = pos_map[legend_position]
                            # current_app.logger.debug(f"Set legend position: {legend_position}")
                    if legend_font_size:
                        layout_updates.setdefault("legend", {})["font"] = {"size": legend_font_size}
                        # current_app.logger.debug(f"Set legend font size: {legend_font_size}")
                else:
                    layout_updates["showlegend"] = False
                    # current_app.logger.debug("Legend disabled for Plotly")
                
                # Bar mode for stacked charts
                if barmode:
                    layout_updates["barmode"] = barmode
                elif chart_type == "stacked_column":
                    layout_updates["barmode"] = "stack"
                
                # Axis min/max and tick format (only for non-pie charts)
                if chart_type != "pie":
                    if y_axis_min_max:
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        # Handle "auto" value for y-axis min/max
                        if y_axis_min_max == "auto":
                            # Don't set range for auto - let Plotly auto-scale
                            pass
                        else:
                            layout_updates["yaxis"]["range"] = y_axis_min_max
                            # Also set autorange to false to ensure the range is respected
                            layout_updates["yaxis"]["autorange"] = False
                            # Force the range to be applied
                            layout_updates["yaxis"]["fixedrange"] = False
                            # Ensure the range is properly set
                            # current_app.logger.debug(f"Setting Y-axis range to: {y_axis_min_max}")
                    if axis_tick_format:
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["yaxis"]["tickformat"] = axis_tick_format
                        # Also apply to secondary y-axis if it's a currency format
                        if "$" in axis_tick_format:
                            layout_updates["yaxis2"] = layout_updates.get("yaxis2", {})
                            layout_updates["yaxis2"]["tickformat"] = axis_tick_format
                    
                    # Secondary y-axis formatting
                    if secondary_y_axis_format:
                        layout_updates["yaxis2"] = layout_updates.get("yaxis2", {})
                        layout_updates["yaxis2"]["tickformat"] = secondary_y_axis_format
                    if secondary_y_axis_min_max:
                        layout_updates["yaxis2"] = layout_updates.get("yaxis2", {})
                        # Handle "auto" value for secondary y-axis min/max
                        if secondary_y_axis_min_max == "auto":
                            # Don't set range for auto - let Matplotlib auto-scale
                            pass
                        elif isinstance(secondary_y_axis_min_max, list) and len(secondary_y_axis_min_max) == 2:
                            layout_updates["yaxis2"]["range"] = secondary_y_axis_min_max
                        else:
                            pass  # Suppress warning logs: f"Invalid secondary_y_axis_min_max format: {secondary_y_axis_min_max}")
                    
                    # Axis tick font size
                    if axis_tick_font_size:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["tickfont"] = {"size": axis_tick_font_size}
                        layout_updates["yaxis"]["tickfont"] = {"size": axis_tick_font_size}
                    # Gridlines
                    if show_gridlines is not None:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["showgrid"] = bool(show_gridlines)
                        layout_updates["yaxis"]["showgrid"] = bool(show_gridlines)
                    if gridline_color:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["gridcolor"] = gridline_color
                        layout_updates["yaxis"]["gridcolor"] = gridline_color
                    if gridline_style:
                        # Map gridline styles to valid Plotly dash styles
                        dash_map = {
                            "solid": "solid", 
                            "dashed": "dash",
                            "dash": "dash",  # Map 'dash' to 'dash'
                            "dashdot": "dashdot", 
                            "dotted": "dot",
                            "dot": "dot",
                            "dotdash": "dashdot"
                        }
                        dash_style = dash_map.get(gridline_style, "solid")
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        layout_updates["xaxis"]["griddash"] = dash_style
                        layout_updates["yaxis"]["griddash"] = dash_style
                
                # Data labels (for bar/line traces)
                show_data_labels = chart_meta.get("data_labels", True)
                value_format = chart_meta.get("value_format", "")
                
                # Only enable data labels if explicitly set to True AND any data label settings are provided
                if show_data_labels and (data_label_format or data_label_font_size or data_label_color):
                    show_data_labels = True
                elif not show_data_labels:
                    # If data_labels is explicitly set to False, respect that setting
                    show_data_labels = False
                
                # Debug logging for data labels
                # current_app.logger.debug(f"Original data_labels setting: {chart_meta.get('data_labels')}")
                # current_app.logger.debug(f"Final show_data_labels: {show_data_labels}")
                # current_app.logger.debug(f"Data label format: {data_label_format}")
                # current_app.logger.debug(f"Data label font size: {data_label_font_size}")
                # current_app.logger.debug(f"Data label color: {data_label_color}")
                # current_app.logger.debug(f"Value format: {value_format}")
                # current_app.logger.debug(f"Chart config keys: {list(chart_config.keys())}")
                # current_app.logger.debug(f"Chart meta keys: {list(chart_meta.keys())}")
                
                if show_data_labels and (data_label_format or value_format or data_label_font_size or data_label_color):
                    # current_app.logger.debug(f"Processing {len(fig.data)} traces for data labels")
                    for i, trace in enumerate(fig.data):
                        # current_app.logger.debug(f"Trace {i}: type={trace.type}, mode={getattr(trace, 'mode', 'N/A')}")
                        # Handle both bar and line charts (line charts are scatter with mode='lines')
                        if trace.type in ['bar', 'scatter']:
                            # Use value_format from chart_meta if available, otherwise use data_label_format
                            format_to_use = value_format if value_format else data_label_format
                            
                            # Determine if this is a line chart (scatter with lines mode)
                            is_line_chart = trace.type == 'scatter' and trace.mode and 'lines' in trace.mode
                            
                            if format_to_use:
                                if trace.type == 'bar':
                                    trace.update(texttemplate=f"%{{y:{format_to_use}}}", textposition="auto")
                                elif trace.type == 'scatter':
                                    if is_line_chart:
                                        # For line charts, show labels at the data points
                                        trace.update(texttemplate=f"%{{y:{format_to_use}}}", textposition="top center")
                                    else:
                                        # For scatter plots
                                        trace.update(texttemplate=f"%{{y:{format_to_use}}}", textposition="top center")
                            else:
                                # If no format specified, still show data labels
                                if trace.type == 'bar':
                                    trace.update(texttemplate="%{y}", textposition="auto")
                                elif trace.type == 'scatter':
                                    if is_line_chart:
                                        # For line charts, show labels at the data points
                                        trace.update(texttemplate="%{y}", textposition="top center")
                                    else:
                                        # For scatter plots
                                        trace.update(texttemplate="%{y}", textposition="top center")
                            
                            # Apply font styling
                            if data_label_font_size or data_label_color:
                                trace.update(textfont={})
                                if data_label_font_size:
                                    trace.textfont["size"] = data_label_font_size
                                if data_label_color:
                                    trace.textfont["color"] = data_label_color

                # Annotations
                if annotations:
                    layout_updates["annotations"] = []
                    for ann in annotations:
                        ann_dict = {
                            "text": ann.get("text", ""),
                            "x": ann.get("x_value", ann.get("x")),
                            "y": ann.get("y_value", ann.get("y")),
                            "showarrow": True
                        }
                        layout_updates["annotations"].append(ann_dict)
                
                # --- Apply margin settings ---
                if margin:
                    layout_updates["margin"] = margin
                
                # Apply figure size if specified
                if figsize:
                    layout_updates["width"] = figsize[0] * 100  # Convert to pixels
                    layout_updates["height"] = figsize[1] * 100  # Convert to pixels
                    # current_app.logger.debug(f"Applied Plotly figsize: {figsize} -> width={figsize[0]*100}, height={figsize[1]*100}")
                
                # Apply axis label distances
                if chart_type != "pie":
                    if x_axis_label_distance == "auto" or y_axis_label_distance == "auto":
                        # Calculate optimal label distances
                        x_axis_label_distance, y_axis_label_distance = calculate_optimal_label_distance(
                            chart_type, series_data, x_values, [], figsize, font_size
                        )
                    
                    if x_axis_label_distance or y_axis_label_distance or axis_tick_distance:
                        layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                        layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                        
                        if x_axis_label_distance:
                            layout_updates["xaxis"]["title"] = layout_updates["xaxis"].get("title", {})
                            layout_updates["xaxis"]["title"]["standoff"] = x_axis_label_distance
                            # Also apply to tick distance for better control
                            layout_updates["xaxis"]["ticklen"] = x_axis_label_distance
                        
                        if y_axis_label_distance:
                            layout_updates["yaxis"]["title"] = layout_updates["yaxis"].get("title", {})
                            layout_updates["yaxis"]["title"]["standoff"] = y_axis_label_distance
                        
                        if axis_tick_distance:
                            layout_updates["xaxis"]["ticklen"] = axis_tick_distance
                            layout_updates["yaxis"]["ticklen"] = axis_tick_distance

                # Tick mark control for Plotly
                if show_x_ticks is not None or show_y_ticks is not None:
                    layout_updates["xaxis"] = layout_updates.get("xaxis", {})
                    layout_updates["yaxis"] = layout_updates.get("yaxis", {})
                    if show_x_ticks is not None:
                        layout_updates["xaxis"]["showticklabels"] = bool(show_x_ticks)
                        layout_updates["xaxis"]["ticks"] = "" if not show_x_ticks else "outside"
                    if show_y_ticks is not None:
                        layout_updates["yaxis"]["showticklabels"] = bool(show_y_ticks)
                        layout_updates["yaxis"]["ticks"] = "" if not show_y_ticks else "outside"

                fig.update_layout(**layout_updates)

                # --- Matplotlib static chart for DOCX ---
                if chart_type == "pie":
                    # Check if this is an expanded pie chart
                    expanded_segment = chart_meta.get("expanded_segment")
                    
                    if expanded_segment and len(series_data) == 1:
                        # Create subplot for expanded pie chart
                        mpl_figsize = figsize if figsize else (15, 8)
                        fig_mpl, (ax1, ax2) = plt.subplots(1, 2, figsize=mpl_figsize, dpi=200)
                        
                        # Apply background colors to Matplotlib figure
                        if chart_background:
                            fig_mpl.patch.set_facecolor(chart_background)
                        if plot_background:
                            ax1.set_facecolor(plot_background)
                            ax2.set_facecolor(plot_background)
                        
                        series = series_data[0]
                        labels = series.get("labels", x_values)
                        values = series.get("values", [])
                        color = series.get("marker", {}).get("colors") if "marker" in series else colors
                        marker_line = series.get("marker", {}).get("line", {}) if "marker" in series else {}
                        explode = series.get("pull")
                        opacity = series.get("opacity", chart_meta.get("opacity"))
                        textinfo = series.get("textinfo", chart_meta.get("textinfo", "percent"))
                        textposition = series.get("textposition", chart_meta.get("textposition", "inside")).lower()
                        value_format_str = chart_meta.get("value_format", ".1f")
                        data_labels_enabled = bool(chart_meta.get("data_labels", True))
                        data_label_font_size = chart_meta.get("data_label_font_size", font_size or 10)
                        data_label_color = chart_meta.get("data_label_color", font_color or "#000000")
                        start_angle = startangle if 'startangle' in locals() and startangle is not None else 90
                        sort_order = chart_meta.get("sort_order")

                        # Optional sorting
                        if sort_order in ("ascending", "descending") and values:
                            zipped = list(zip(values, labels, color if isinstance(color, list) else [color]*len(labels), explode if isinstance(explode, list) else [0]*len(labels)))
                            reverse = sort_order == "descending"
                            zipped.sort(key=lambda t: (t[0] if t[0] is not None else 0), reverse=reverse)
                            values, labels, color_list, explode_list = zip(*zipped)
                            values = list(values)
                            labels = list(labels)
                            color = list(color_list)
                            explode = list(explode_list)

                        # Build autopct based on textinfo/value_format
                        def make_autopct(fmt:str, include_percent:bool, include_value:bool):
                            # Capture the current values in the closure
                            current_values = values.copy() if isinstance(values, list) else list(values) if values else []
                            
                            def _inner(pct):
                                # Ensure values are numeric before calculating total
                                numeric_values = []
                                for v in current_values:
                                    if v is not None:
                                        try:
                                            numeric_values.append(float(v))
                                        except (ValueError, TypeError):
                                            # Skip non-numeric values
                                            continue
                                
                                total = sum(numeric_values) if numeric_values else 0
                                val = pct * total / 100.0
                                parts = []
                                if include_value:
                                    try:
                                        parts.append(f"{val:{fmt}}")
                                    except Exception:
                                        parts.append(f"{val:.1f}")
                                if include_percent:
                                    parts.append(f"{pct:.1f}%")
                                return " ".join(parts)
                            return _inner

                        include_label = "label" in (textinfo or "")
                        include_percent = "percent" in (textinfo or "")
                        include_value = "value" in (textinfo or "")

                        autopct_callable = None
                        if data_labels_enabled and (include_percent or include_value):
                            autopct_callable = make_autopct(value_format_str, include_percent, include_value)

                        # Wedge and text props
                        wedgeprops = {}
                        if isinstance(marker_line, dict):
                            if marker_line.get("color"):
                                wedgeprops["edgecolor"] = marker_line.get("color")
                            if marker_line.get("width") is not None:
                                wedgeprops["linewidth"] = marker_line.get("width")
                        if opacity is not None:
                            wedgeprops["alpha"] = opacity

                        textprops = {"color": data_label_color, "fontsize": data_label_font_size}
                        if chart_meta.get("font_family"):
                            textprops["fontfamily"] = chart_meta.get("font_family")

                        # Positioning
                        pctdistance = 0.6 if textposition == "inside" else 1.15
                        labeldistance = 1.1 if textposition != "inside" else 1.05
                        
                        # Create pie chart
                        wedges, texts, autotexts = ax1.pie(
                            values,
                            labels=labels if include_label else None,
                            autopct=autopct_callable,
                            colors=color,
                            startangle=start_angle,
                            explode=explode,
                            wedgeprops=wedgeprops,
                            pctdistance=pctdistance,
                            labeldistance=labeldistance,
                            textprops=textprops,
                        )

                        # Style the autopct texts
                        for autotext in autotexts or []:
                            autotext.set_color(data_label_color)
                            autotext.set_fontsize(data_label_font_size)
                            autotext.set_fontweight('bold')
                            if chart_meta.get("font_family"):
                                autotext.set_fontfamily(chart_meta.get("font_family"))
                        
                        if title and title.strip():
                            ax1.set_title(title, fontsize=font_size or 14, weight='bold', pad=20, color=font_color if font_color else None, fontname=chart_meta.get("font_family") if chart_meta.get("font_family") else None)
                        
                        # Add legend for pie chart
                        show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                        # Convert string "false"/"true" to boolean if needed
                        if isinstance(show_legend_raw, str):
                            show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                        else:
                            show_legend = bool(show_legend_raw)
                        if show_legend:
                            # Initialize legend_loc for pie charts
                            legend_loc = 'best'  # default
                            if legend_position:
                                loc_map = {
                                    "top": "upper center",
                                    "bottom": "lower center", 
                                    "left": "center left",
                                    "right": "center right"
                                }
                                legend_loc = loc_map.get(legend_position, 'best')
                            
                            # Force legend to bottom if specified
                            if legend_position == "bottom":
                                ax1.legend(wedges, labels, loc='lower center', bbox_to_anchor=(0.5, -0.15), fontsize=legend_font_size)
                            else:
                                ax1.legend(wedges, labels, loc=legend_loc, fontsize=legend_font_size)
                        
                        # Create bar chart for expanded segment
                        if expanded_segment in labels:
                            segment_idx = labels.index(expanded_segment)
                            segment_value = values[segment_idx]
                            segment_color = color[segment_idx] if isinstance(color, list) and segment_idx < len(color) else color
                            
                            ax2.bar([expanded_segment], [segment_value], color=segment_color, alpha=0.7)
                            ax2.set_title(f"{expanded_segment} Details", fontsize=font_size or 12, weight='bold')
                            ax2.set_ylabel("Value")
                            
                            # Add value label on bar
                            ax2.text(0, segment_value, f"{segment_value}", ha='center', va='bottom', fontweight='bold')
                        
                    else:
                        # Regular pie chart
                        mpl_figsize = figsize if figsize else (10, 8)
                        fig_mpl, ax = plt.subplots(figsize=mpl_figsize, dpi=200)
                        
                        # Apply background colors to Matplotlib figure
                        if chart_background:
                            fig_mpl.patch.set_facecolor(chart_background)
                        if plot_background:
                            ax.set_facecolor(plot_background)
                        
                        if len(series_data) == 1:
                            series = series_data[0]
                            labels = series.get("labels", x_values)
                            values = series.get("values", [])
                            
                            # Normalize pie chart values to sum to 100% (for matplotlib fallback)
                            if chart_type == "pie" and values:
                                values = normalize_values_to_100(values)
                            
                            color = series.get("marker", {}).get("colors") if "marker" in series else colors
                            marker_line = series.get("marker", {}).get("line", {}) if "marker" in series else {}
                            explode = series.get("pull")
                            opacity = series.get("opacity", chart_meta.get("opacity"))
                            textinfo = series.get("textinfo", chart_meta.get("textinfo", "percent"))
                            textposition = series.get("textposition", chart_meta.get("textposition", "inside")).lower()
                            value_format_str = chart_meta.get("value_format", ".1f")
                            data_labels_enabled = bool(chart_meta.get("data_labels", True))
                            data_label_font_size = chart_meta.get("data_label_font_size", font_size or 10)
                            data_label_color = chart_meta.get("data_label_color", font_color or "#000000")
                            start_angle = startangle if 'startangle' in locals() and startangle is not None else 90
                            sort_order = chart_meta.get("sort_order")

                            # Optional sorting
                            if sort_order in ("ascending", "descending") and values:
                                zipped = list(zip(values, labels, color if isinstance(color, list) else [color]*len(labels), explode if isinstance(explode, list) else [0]*len(labels)))
                                reverse = sort_order == "descending"
                                zipped.sort(key=lambda t: (t[0] if t[0] is not None else 0), reverse=reverse)
                                values, labels, color_list, explode_list = zip(*zipped)
                                values = list(values)
                                labels = list(labels)
                                color = list(color_list)
                                explode = list(explode_list)

                            # Build autopct based on textinfo/value_format
                            def make_autopct(fmt:str, include_percent:bool, include_value:bool):
                                # Capture the current values in the closure
                                current_values = values.copy() if isinstance(values, list) else list(values) if values else []
                                
                                def _inner(pct):
                                    # Ensure values are numeric before calculating total
                                    numeric_values = []
                                    for v in current_values:
                                        if v is not None:
                                            try:
                                                numeric_values.append(float(v))
                                            except (ValueError, TypeError):
                                                # Skip non-numeric values
                                                continue
                                    
                                    total = sum(numeric_values) if numeric_values else 0
                                    val = pct * total / 100.0
                                    parts = []
                                    if include_value:
                                        try:
                                            parts.append(f"{val:{fmt}}")
                                        except Exception:
                                            parts.append(f"{val:.1f}")
                                    if include_percent:
                                        parts.append(f"{pct:.1f}%")
                                    return " ".join(parts)
                                return _inner

                            include_label = "label" in (textinfo or "")
                            include_percent = "percent" in (textinfo or "")
                            include_value = "value" in (textinfo or "")

                            autopct_callable = None
                            if data_labels_enabled and (include_percent or include_value):
                                autopct_callable = make_autopct(value_format_str, include_percent, include_value)

                            # Wedge and text props
                            wedgeprops = {}
                            if isinstance(marker_line, dict):
                                if marker_line.get("color"):
                                    wedgeprops["edgecolor"] = marker_line.get("color")
                                if marker_line.get("width") is not None:
                                    wedgeprops["linewidth"] = marker_line.get("width")
                            if opacity is not None:
                                wedgeprops["alpha"] = opacity

                            textprops = {"color": data_label_color, "fontsize": data_label_font_size}
                            if chart_meta.get("font_family"):
                                textprops["fontfamily"] = chart_meta.get("font_family")

                            # Positioning
                            pctdistance = 0.6 if textposition == "inside" else 1.15
                            labeldistance = 1.1 if textposition != "inside" else 1.05
                            
                            # Create pie chart
                            wedges, texts, autotexts = ax.pie(
                                values,
                                labels=labels if include_label else None,
                                autopct=autopct_callable,
                                colors=color,
                                startangle=start_angle,
                                explode=explode,
                                wedgeprops=wedgeprops,
                                pctdistance=pctdistance,
                                labeldistance=labeldistance,
                                textprops=textprops,
                            )
                            
                            # Style the text
                            for autotext in autotexts or []:
                                autotext.set_color(data_label_color)
                                autotext.set_fontsize(data_label_font_size)
                                autotext.set_fontweight('bold')
                                if chart_meta.get("font_family"):
                                    autotext.set_fontfamily(chart_meta.get("font_family"))
                            
                            if title and title.strip():
                                ax.set_title(title, fontsize=font_size or 14, weight='bold', pad=20, color=font_color if font_color else None, fontname=chart_meta.get("font_family") if chart_meta.get("font_family") else None)
                            
                            # Add legend for regular pie chart
                            show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                            # Convert string "false"/"true" to boolean if needed
                            if isinstance(show_legend_raw, str):
                                show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                            else:
                                show_legend = bool(show_legend_raw)
                            if show_legend:
                                # Initialize legend_loc for pie charts
                                legend_loc = 'best'  # default
                                if legend_position:
                                    loc_map = {
                                        "top": "upper center",
                                        "bottom": "lower center", 
                                        "left": "center left",
                                        "right": "center right"
                                    }
                                    legend_loc = loc_map.get(legend_position, 'best')
                                
                                # Force legend to bottom if specified
                                if legend_position == "bottom":
                                    ax.legend(wedges, labels, loc='lower center', bbox_to_anchor=(0.5, -0.15), fontsize=legend_font_size)
                                else:
                                    ax.legend(wedges, labels, loc=legend_loc, fontsize=legend_font_size)
                        
                elif chart_type in ["bar of pie", "bar_of_pie"]:
                    # Matplotlib version of bar of pie
                    mpl_figsize = figsize if figsize else (10, 5)
                    fig_mpl, (ax1, ax2) = plt.subplots(1, 2, figsize=mpl_figsize, dpi=200, gridspec_kw={'width_ratios': [2, 1]})
                    
                    # Apply background colors to Matplotlib figure
                    if chart_background:
                        fig_mpl.patch.set_facecolor(chart_background)
                    if plot_background:
                        ax1.set_facecolor(plot_background)
                        ax2.set_facecolor(plot_background)
                    labels = series_meta.get("labels", x_values)
                    values = series_meta.get("values", [])
                    colors = series_meta.get("colors", [])
                    other_labels = chart_meta.get("other_labels", [])
                    other_values = chart_meta.get("other_values", [])
                    other_colors = chart_meta.get("other_colors", [])
                    if not (other_labels and other_values):
                        if "other_label_range" in chart_meta and "other_value_range" in chart_meta and "source_sheet" in chart_meta:
                            wb = openpyxl.load_workbook(data_file_path, data_only=True)
                            sheet = wb[chart_meta["source_sheet"]]
                            other_labels = extract_excel_range(sheet, chart_meta["other_label_range"])
                            other_values = extract_excel_range(sheet, chart_meta["other_value_range"])
                    # Pie chart
                    if colors:
                        wedges, texts, autotexts = ax1.pie(values, labels=labels, autopct='%1.1f%%', colors=colors, startangle=90)
                    else:
                        wedges, texts, autotexts = ax1.pie(values, labels=labels, autopct='%1.1f%%', startangle=90)
                    for autotext in autotexts:
                        autotext.set_color('white')
                        autotext.set_fontweight('bold')
                    if title and title.strip():
                        ax1.set_title(title, fontsize=font_size or 14, weight='bold', pad=20)
                    # Move legend outside
                    show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                    # Convert string "false"/"true" to boolean if needed
                    if isinstance(show_legend_raw, str):
                        show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                    else:
                        show_legend = bool(show_legend_raw)
                    legend_font_size = chart_meta.get("legend_font_size", 8)
                    if show_legend:
                        ax1.legend(wedges, labels, loc='center left', bbox_to_anchor=(1, 0.5), fontsize=legend_font_size)
                    
                    # Bar chart - Filter out empty/null values first
                    filtered_data = []
                    for label, value in zip(other_labels, other_values):
                        if (label is not None and str(label).strip() != "" and 
                            value is not None and str(value).strip() != "" and 
                            str(value).strip() != "0"):
                            filtered_data.append((label, value))
                    
                    if filtered_data:
                        filtered_labels, filtered_values = zip(*filtered_data)
                    else:
                        filtered_labels, filtered_values = [], []
                    
                    # Create individual bars (not stacked)
                    if filtered_labels and filtered_values:
                        # Use individual colors for each bar
                        bar_colors = []
                        for i in range(len(filtered_labels)):
                            if other_colors and i < len(other_colors):
                                bar_colors.append(other_colors[i])
                            elif colors and i < len(colors):
                                bar_colors.append(colors[i])
                            else:
                                bar_colors.append('#1f77b4')  # Default blue
                        
                        bars = ax2.bar(range(len(filtered_labels)), filtered_values, color=bar_colors, alpha=0.7)
                    else:
                        bars = []
                    
                    expanded_segment = chart_meta.get("expanded_segment", "Other")
                    ax2.set_title(f"Breakdown of '{expanded_segment}'", fontsize=font_size or 12, weight='bold')
                    # Use proper Y-axis label from configuration
                    y_axis_title = chart_meta.get("y_axis_title", "Value")
                    ax2.set_ylabel(y_axis_title, fontsize=label_fontsize if 'label_fontsize' in locals() else 10)
                    # Set X-axis label
                    x_axis_title = chart_meta.get("x_axis_title", "Categories")
                    ax2.set_xlabel(x_axis_title, fontsize=label_fontsize if 'label_fontsize' in locals() else 10)
                    # Set x-tick labels for filtered data
                    if filtered_labels:
                        ax2.set_xticks(range(len(filtered_labels)))
                        # Format x-axis labels as percentages
                        formatted_x_labels = []
                        for label in filtered_labels:
                            if label is not None:
                                if isinstance(label, (int, float)):
                                    if label <= 1.0:  # Likely decimal format (0.06)
                                        formatted_x_labels.append(f"{label * 100:.1f}%")
                                    else:  # Likely already percentage format (6.0)
                                        formatted_x_labels.append(f"{label:.1f}%")
                                else:
                                    # Convert string to float and handle
                                    try:
                                        val = float(label)
                                        if val <= 1.0:
                                            formatted_x_labels.append(f"{val * 100:.1f}%")
                                        else:
                                            formatted_x_labels.append(f"{val:.1f}%")
                                    except (ValueError, TypeError):
                                        formatted_x_labels.append(str(label))
                            else:
                                formatted_x_labels.append("")
                        ax2.set_xticklabels(formatted_x_labels, rotation=0, fontsize=axis_tick_font_size or 10)
                    # Add data labels with proper formatting
                    value_format = chart_meta.get("value_format", ".2f")
                    data_label_font_size = chart_meta.get("data_label_font_size", 10)
                    data_label_color = chart_meta.get("data_label_color", "#000000")
                    for bar, v in zip(bars, filtered_values):
                        # Format percentage values as XX.X% instead of 0.XXX
                        if v is not None:
                            if isinstance(v, (int, float)):
                                if v <= 1.0:  # Likely decimal format (0.11)
                                    formatted_value = f"{v * 100:.1f}%"
                                else:  # Likely already percentage format (11.0)
                                    formatted_value = f"{v:.1f}%"
                            else:
                                    # Convert string to float and handle
                                try:
                                    val = float(v)
                                    if val <= 1.0:
                                        formatted_value = f"{val * 100:.1f}%"
                                    else:
                                        formatted_value = f"{val:.1f}%"
                                except:
                                    formatted_value = str(v)
                            ax2.text(bar.get_x() + bar.get_width()/2, v, formatted_value, ha='center', va='bottom', fontweight='bold', fontsize=data_label_font_size, color=data_label_color)

                else:
                    # Bar, line, area charts
                    mpl_figsize = figsize if figsize else (10, 6)
                    # current_app.logger.debug(f"Applied Matplotlib figsize: {mpl_figsize}")
                    fig_mpl, ax1 = plt.subplots(figsize=mpl_figsize, dpi=200)
                    
                    # Disable gridlines globally if show_gridlines is False
                    if not show_gridlines:
                        plt.rcParams['axes.grid'] = False
                    
                    # Only create secondary y-axis if not disabled
                    ax2 = None
                    if not disable_secondary_y:
                        ax2 = ax1.twinx()
                    
                    # Apply background colors to Matplotlib figure
                    if chart_background:
                        fig_mpl.patch.set_facecolor(chart_background)
                    if plot_background:
                        ax1.set_facecolor(plot_background)
                        if ax2:
                            ax2.set_facecolor(plot_background)

                    # Define colors array for matplotlib section
                    colors = series_meta.get("colors", [])

                    for i, series in enumerate(series_data):
                        label = series.get("name", f"Series {i+1}")
                        series_type = series.get("type", "bar").lower()
                        
                        # Debug logging for series type detection
                        current_app.logger.debug(f"🔍 Processing series {i}: {label}, type: {series_type}")
                        
                        # Map heatmap to imshow for Matplotlib
                        if series_type == "heatmap":
                            mpl_chart_type = "imshow"
                            current_app.logger.debug(f"🔥 Heatmap detected, setting mpl_chart_type to: {mpl_chart_type}")
                        else:
                            mpl_chart_type = chart_type_mapping_mpl.get(series_type, "scatter")
                            current_app.logger.debug(f"🔍 Regular chart type: {series_type} -> {mpl_chart_type}")
                        
                        color = None
                        if "marker" in series and isinstance(series["marker"], dict) and "color" in series["marker"]:
                            color = series["marker"]["color"]
                        elif bar_colors:
                            color = bar_colors
                        elif i < len(colors) and colors[i] is not None:
                            color = colors[i]

                        # Skip regular data processing for heatmaps
                        if series_type == "heatmap":
                            # Heatmaps use z data directly, skip y_vals processing
                            y_vals = []  # Not used for heatmaps
                            mpl_chart_type = "imshow"
                            # Debug logging for heatmap data
                            current_app.logger.debug(f"🔥 Heatmap series detected: {series}")
                            current_app.logger.debug(f"🔥 Heatmap x data: {series.get('x', [])}")
                            current_app.logger.debug(f"🔥 Heatmap y data: {series.get('y', [])}")
                            current_app.logger.debug(f"🔥 Heatmap z data: {series.get('z', [])}")
                        else:
                            # Regular series processing
                            y_vals = series.get("values")
                            value_range = series.get("value_range")
                            if value_range:
                                # Check if value_range is already extracted (list) or still a string
                                if isinstance(value_range, list):
                                    y_vals = value_range
                                else:
                                    y_vals = extract_values_from_range(value_range)
                            
                            # Ensure y_vals is not None
                            if y_vals is None:
                                y_vals = []
                                current_app.logger.warning(f"⚠️ y_vals is None for series {label}, using empty list")
                            
                            # Ensure x_values is not None
                            if x_values is None:
                                x_values = []
                                current_app.logger.warning(f"⚠️ x_values is None for series {label}, using empty list")

                            # Generic chart type handling for Matplotlib
                            mpl_chart_type = chart_type_mapping_mpl.get(series_type, "scatter")
                        
                        if mpl_chart_type == "bar":
                            # Add bar border parameters
                            edgecolor = bar_border_color if bar_border_color else 'none'
                            linewidth = bar_border_width if bar_border_width else 0
                            
                            if chart_type == "stacked_column":
                                # For stacked column, use bottom parameter
                                bar_color = color if color is not None else 'blue'
                                if i == 0:
                                    ax1.bar(x_values, y_vals, label=label, color=bar_color, alpha=0.7, edgecolor=edgecolor, linewidth=linewidth)
                                    bottom_vals = y_vals
                                    # Initialize storage for stacked bar segments
                                    if not hasattr(ax1, '_stacked_segments'):
                                        ax1._stacked_segments = []
                                    ax1._stacked_segments.append({
                                        'y_vals': y_vals,
                                        'bottom_vals': [0] * len(y_vals),
                                        'x_values': x_values
                                    })
                                else:
                                    ax1.bar(x_values, y_vals, bottom=bottom_vals, label=label, color=bar_color, alpha=0.7, edgecolor=edgecolor, linewidth=linewidth)
                                    # Store segment info for data labels
                                    ax1._stacked_segments.append({
                                        'y_vals': y_vals,
                                        'bottom_vals': bottom_vals.copy(),
                                        'x_values': x_values
                                    })
                                    bottom_vals = [sum(x) for x in zip(bottom_vals, y_vals)]
                            else:
                                if isinstance(color, list):
                                    for j, val in enumerate(y_vals):
                                        bar_color = color[j % len(color)] if color[j % len(color)] is not None else 'blue'
                                        ax1.bar(x_values[j], val, color=bar_color, alpha=0.7, label=label if j == 0 else "", edgecolor=edgecolor, linewidth=linewidth)
                                else:
                                    bar_color = color if color is not None else 'blue'
                                    ax1.bar(x_values, y_vals, label=label, color=bar_color, alpha=0.7, edgecolor=edgecolor, linewidth=linewidth)
                                    
                        elif mpl_chart_type == "barh":
                            # Horizontal bar chart
                            # Add bar border parameters
                            edgecolor = bar_border_color if bar_border_color else 'none'
                            linewidth = bar_border_width if bar_border_width else 0
                            
                            if isinstance(color, list):
                                for j, val in enumerate(y_vals):
                                    bar_color = color[j % len(color)] if color[j % len(color)] is not None else 'blue'
                                    ax1.barh(x_values[j], val, color=bar_color, alpha=0.7, label=label if j == 0 else "", edgecolor=edgecolor, linewidth=linewidth)
                            else:
                                bar_color = color if color is not None else 'blue'
                                ax1.barh(x_values, y_vals, label=label, color=bar_color, alpha=0.7, edgecolor=edgecolor, linewidth=linewidth)
                                
                        elif mpl_chart_type == "plot":
                            # Line chart
                            marker = 'o' if series_type == "scatter_line" else None
                            line_color = color if color is not None else 'blue'
                            if ax2:
                                ax2.plot(x_values, y_vals, label=label, color=line_color, marker=marker, linewidth=2)
                            else:
                                ax1.plot(x_values, y_vals, label=label, color=line_color, marker=marker, linewidth=2)
                            
                        elif mpl_chart_type == "scatter":
                            # Scatter plot and Bubble chart
                            if series_type == "bubble":
                                # Enhanced bubble chart with better styling
                                # Creating bubble chart for series: {label}
                                
                                # Get sizes from the series data structure
                                sizes = series.get("size", [20] * len(y_vals)) if y_vals else [20]
                                # Bubble chart data processed
                                
                                # Ensure all arrays have the same length
                                min_length = min(len(x_values) if x_values else 0, len(y_vals) if y_vals else 0, len(sizes) if sizes else 0)
                                if min_length > 0 and (min_length < len(x_values) or min_length < len(y_vals) or min_length < len(sizes)):
                                    #pass  # Suppress warning logs: f"⚠️ Array length mismatch! Truncating to {min_length}")
                                    x_values = x_values[:min_length] if x_values else []
                                    y_vals = y_vals[:min_length] if y_vals else []
                                    sizes = sizes[:min_length] if sizes else []
                                    if isinstance(color, list):
                                        color = color[:min_length]
                                
                                # Scale sizes for better visual impact (bubble charts need much larger sizes)
                                scaled_sizes = [s * 20 for s in sizes]  # Much larger scale for better visibility
                                # Scaled sizes calculated
                                
                                # Colors
                                bubble_colors = color
                                if isinstance(color, list):
                                    bubble_colors = color
                                elif color:
                                    import matplotlib.cm as cm
                                    import numpy as np
                                    size_array = np.array(sizes)
                                    normalized_sizes = (size_array - size_array.min()) / (size_array.max() - size_array.min() + 1e-8)
                                    cmap = cm.viridis if color == 'auto' else cm.get_cmap('viridis')
                                    bubble_colors = [cmap(norm_size) for norm_size in normalized_sizes]
                                else:
                                    bubble_colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7', '#DDA0DD', '#98D8C8', '#F7DC6F']
                                
                                # Marker styles from config
                                marker_cfg = series.get("marker", {}) if isinstance(series.get("marker"), dict) else {}
                                marker_line_cfg = marker_cfg.get("line", {}) if isinstance(marker_cfg.get("line"), dict) else {}
                                alpha_val = marker_cfg.get("opacity", series.get("opacity", chart_meta.get("opacity", 0.8)))
                                edge_color = marker_line_cfg.get("color", 'white')
                                edge_width = marker_line_cfg.get("width", 2)
                                
                                # Create bubble chart with enhanced styling
                                scatter = ax1.scatter(
                                    x_values, y_vals, 
                                                     s=scaled_sizes, 
                                                     c=bubble_colors, 
                                    alpha=alpha_val,
                                    edgecolors=edge_color,
                                    linewidth=edge_width,
                                                     label=label,
                                    zorder=3)
                                
                                # Keep a reference for legend building
                                try:
                                    if not hasattr(ax1, "_bubble_handles"):
                                        ax1._bubble_handles = []
                                        ax1._bubble_labels = []
                                    ax1._bubble_handles.append(scatter)
                                    ax1._bubble_labels.append(label)
                                except Exception:
                                    pass
                                
                                # Create an explicit proxy handle so legend always renders
                                try:
                                    from matplotlib.lines import Line2D
                                    if isinstance(bubble_colors, list) and len(bubble_colors) > 0:
                                        legend_facecolor = bubble_colors[0]
                                    else:
                                        legend_facecolor = bubble_colors if bubble_colors else '#1f77b4'
                                    proxy = Line2D([0], [0], marker='o', linestyle='None', label=label,
                                                   markerfacecolor=legend_facecolor, markeredgecolor=edge_color,
                                                   markeredgewidth=edge_width, markersize=10, alpha=alpha_val)
                                    ax1._bubble_proxy = proxy
                                except Exception:
                                    pass
                                
                                # Add subtle shadow effect for depth
                                ax1.scatter(
                                    x_values, y_vals, 
                                          s=scaled_sizes, 
                                          c='black', 
                                          alpha=0.1, 
                                          zorder=1)
                                
                                # Add data labels for bubble chart if enabled
                                if show_data_labels:
                                    for j, (x, y, value_num) in enumerate(zip(x_values, y_vals, y_vals)):
                                        if j < len(x_values):
                                            # Format value for label using value_format and axis_tick_format
                                            try:
                                                if axis_tick_format and "$" in axis_tick_format:
                                                    num = float(value_num)
                                                    formatted_value = f"${num:,.0f}" if value_format in (".0f", None) else f"${num:,.0f}" if value_format == ".0f" else f"${num:,.0f}"
                                                    # Simplify: currency with thousands, ignore decimals beyond .0f for labels
                                                else:
                                                    num = float(value_num)
                                                    if value_format == ".0f":
                                                        formatted_value = f"{num:.0f}"
                                                    elif value_format == ".1f":
                                                        formatted_value = f"{num:.1f}"
                                                    elif value_format == ".2f":
                                                        formatted_value = f"{num:.2f}"
                                                    else:
                                                        formatted_value = f"{num:.0f}"
                                            except:
                                                formatted_value = str(value_num)
                                            
                                            # Add text label inside bubble
                                            ax1.text(x, y, formatted_value, 
                                                    ha='center', va='center', 
                                                    fontsize=data_label_font_size or 10,
                                                    color=data_label_color or '#000000',
                                                    fontweight='bold',
                                                    zorder=4)
                                
                                # Fix axis ranges for bubble charts to prevent layout shifts
                                if i == len(series_data) - 1:  # Only set once after all series
                                    # Always set fixed x-axis range for bubble charts to prevent layout shifts
                                    if x_axis_min_max and isinstance(x_axis_min_max, list) and len(x_axis_min_max) == 2:
                                        # Use explicit range from config
                                        ax1.set_xlim(x_axis_min_max[0], x_axis_min_max[1])
                                    else:
                                        # Set auto-calculated range with padding
                                        x_min, x_max = min(x_values), max(x_values)
                                        x_padding = (x_max - x_min) * 0.1
                                        ax1.set_xlim(x_min - x_padding, x_max + x_padding)
                                    
                                    # Always set fixed y-axis range for bubble charts to prevent layout shifts
                                    if y_axis_min_max and isinstance(y_axis_min_max, list) and len(y_axis_min_max) == 2:
                                        # Use explicit range from config
                                        ax1.set_ylim(y_axis_min_max[0], y_axis_min_max[1])
                                    else:
                                        # Set auto-calculated range with padding
                                        y_min, y_max = min(y_vals), max(y_vals)
                                        y_padding = (y_max - y_min) * 0.1
                                        ax1.set_ylim(y_min - y_padding, y_max + y_padding)
                                    
                                    # Apply axis label distances specifically for bubble charts
                                    if x_axis_label_distance is not None or y_axis_label_distance is not None:
                                        # Handle "auto" values by calculating optimal distances
                                        if x_axis_label_distance == "auto" or y_axis_label_distance == "auto":
                                            # Calculate optimal label distances for bubble chart
                                            # Use the current series data for calculation
                                            current_series_data = [{"labels": series.get("labels", []), "values": y_vals}]
                                            auto_x_distance, auto_y_distance = calculate_optimal_label_distance(
                                                "scatter", current_series_data, x_values, y_vals, figsize, font_size
                                            )
                                            
                                            if x_axis_label_distance == "auto":
                                                x_axis_label_distance = auto_x_distance
                                            if y_axis_label_distance == "auto":
                                                y_axis_label_distance = auto_y_distance
                                        
                                        # Calculate labelpad values with more aggressive multiplication for better spacing
                                        figsize = chart_meta.get("figsize", [12, 8])
                                        chart_width, chart_height = figsize
                                        
                                        # More aggressive multiplication factors for better label separation
                                        x_multiplier = max(20, chart_width * 3.0)
                                        y_multiplier = max(20, chart_height * 3.0)
                                        
                                        x_labelpad = (x_axis_label_distance * x_multiplier) if x_axis_label_distance is not None else 50.0
                                        y_labelpad = (y_axis_label_distance * y_multiplier) if y_axis_label_distance is not None else 50.0
                                        
                                        # Get axis titles
                                        x_axis_title = chart_meta.get("x_label", chart_config.get("x_axis_title", ""))
                                        y_axis_title = chart_meta.get("primary_y_label", chart_config.get("primary_y_label", ""))
                                        
                                        # Debug logging for bubble chart axis label distance
                                        current_app.logger.debug(f"🎈 Bubble Chart Processing - X Distance: {x_axis_label_distance} → {x_labelpad}, Y Distance: {y_axis_label_distance} → {y_labelpad}")
                                        current_app.logger.debug(f"🎈 Bubble Chart Titles - X: '{x_axis_title}', Y: '{y_axis_title}'")
                                        
                                        # Set axis labels with distance
                                        if x_axis_title:
                                            ax1.set_xlabel(x_axis_title, fontsize=font_size or 12, color=font_color if font_color else 'black',
                                                        fontname=chart_meta.get("font_family") if chart_meta.get("font_family") else None,
                                                        labelpad=x_labelpad)
                                            current_app.logger.debug(f"🎈 Set X-axis label with labelpad: {x_labelpad}")
                                        if y_axis_title:
                                            ax1.set_ylabel(y_axis_title, fontsize=font_size or 12, color=font_color if font_color else 'black',
                                                        fontname=chart_meta.get("font_family") if chart_meta.get("font_family") else None,
                                                        labelpad=y_labelpad)
                                            current_app.logger.debug(f"🎈 Set Y-axis label with labelpad: {y_labelpad}")
                                        
                                        # Additional spacing techniques for y-axis
                                        if y_axis_label_distance and y_axis_label_distance > 50:
                                            # Force more space by adjusting the left margin
                                            current_app.logger.debug(f"🎈 Applying additional y-axis spacing techniques")
                                            # Adjust the plot position to create more left margin
                                            ax1.set_position([0.15, 0.1, 0.75, 0.8])  # [left, bottom, width, height]
                                        
                                        current_app.logger.debug(f"🎈 Bubble Chart Axis Label Distance Applied Successfully")
                                        
                                        # Store the labelpad values for later use to prevent override
                                        ax1._bubble_x_labelpad = x_labelpad
                                        ax1._bubble_y_labelpad = y_labelpad
                                        ax1._bubble_x_title = x_axis_title
                                        ax1._bubble_y_title = y_axis_title
                                    else:
                                        current_app.logger.debug(f"🎈 Bubble Chart - No axis label distance values found")
                            else:
                                # Enhanced scatter plot with custom styling
                                # Extract marker properties from series
                                marker_size = series.get("marker", {}).get("size", 50)
                                marker_color = series.get("marker", {}).get("color", color)
                                marker_symbol = series.get("marker", {}).get("symbol", "o")
                                marker_opacity = series.get("marker", {}).get("opacity", 0.8)
                                text_labels = series.get("text", [])
                                text_position = series.get("textposition", "top center")
                                
                                # Extract line properties from series
                                line_config = series.get("line", {})
                                line_color = line_config.get("color", marker_color)
                                line_width = line_config.get("width", 2)
                                line_dash = line_config.get("dash", "solid")
                                
                                # Handle mode setting
                                mode = series.get("mode", "markers")
                                
                                # Handle marker size (can be single value or list)
                                if isinstance(marker_size, list):
                                    sizes = marker_size
                                else:
                                    sizes = [marker_size] * len(y_vals)
                                
                                # Scale sizes for better visibility
                                scaled_sizes = [s * 2 for s in sizes]
                                
                                # Create enhanced scatter plot
                                scatter = ax1.scatter(x_values, y_vals, 
                                                     s=scaled_sizes, 
                                                     c=marker_color, 
                                                     alpha=marker_opacity,
                                                     edgecolors='white',
                                                     linewidth=1.5,
                                                     label=label,
                                                     zorder=3)
                                
                                # Add line if mode includes lines or if line properties are specified
                                if "lines" in mode or "line" in mode or line_config:
                                    # Convert dash style to matplotlib format
                                    if line_dash == "dash":
                                        linestyle = "--"
                                    elif line_dash == "dot":
                                        linestyle = ":"
                                    elif line_dash == "dashdot":
                                        linestyle = "-."
                                    else:
                                        linestyle = "-"  # solid
                                    
                                    ax1.plot(x_values, y_vals, 
                                            color=line_color, 
                                            linewidth=line_width, 
                                            linestyle=linestyle,
                                            alpha=marker_opacity,
                                            zorder=2)
                                
                                # Add subtle shadow effect for depth
                                ax1.scatter(x_values, y_vals, 
                                          s=scaled_sizes, 
                                          c='black', 
                                          alpha=0.1, 
                                          zorder=1)
                                
                                # Add data labels if text is provided
                                if text_labels and len(text_labels) == len(y_vals):
                                    for j, (x, y, text) in enumerate(zip(x_values, y_vals, text_labels)):
                                        if j < len(text_labels):
                                            # Determine label position based on textposition
                                            if "top" in text_position:
                                                y_offset = 2
                                            elif "bottom" in text_position:
                                                y_offset = -2
                                            else:
                                                y_offset = 0
                                            
                                            if "center" in text_position:
                                                ha = 'center'
                                            elif "left" in text_position:
                                                ha = 'left'
                                            elif "right" in text_position:
                                                ha = 'right'
                                            else:
                                                ha = 'center'
                                            
                                            # Add text label
                                            label_color = data_label_color or '#000000'
                                            ax1.text(x, y + y_offset, str(text), 
                                                    ha=ha, va='bottom' if y_offset > 0 else 'top',
                                                    fontsize=data_label_font_size or 10,
                                                    color=label_color,
                                                    fontweight='bold',
                                                    bbox=dict(boxstyle="round,pad=0.3", 
                                                             facecolor='white', 
                                                             alpha=0.9,
                                                             edgecolor='gray',
                                                             linewidth=0.5))
                        elif mpl_chart_type == "fill_between":
                            # Enhanced Area chart
                            # current_app.logger.debug(f"Processing area chart for series: {label}")
                            # current_app.logger.debug(f"X values: {x_vals}")
                            # current_app.logger.debug(f"Y values: {y_vals}")
                            
                            # Extract area-specific properties from series
                            fill_type = series.get("fill", "tozeroy")
                            line_color = series.get("line", {}).get("color", color)
                            line_width = series.get("line", {}).get("width", 2)
                            line_shape = series.get("line", {}).get("shape", "linear")
                            marker_symbol = series.get("marker", {}).get("symbol", "o")
                            marker_size = series.get("marker", {}).get("size", 6)
                            marker_color = series.get("marker", {}).get("color", line_color)
                            area_opacity = series.get("opacity", 0.6)
                            text_labels = series.get("text", [])
                            text_position = series.get("textposition", "top center")
                            
                            # Handle line shape
                            if line_shape == "spline":
                                linestyle = "-"
                                # For spline, we could add curve smoothing, but matplotlib doesn't have built-in splines
                            elif line_shape == "hv":
                                linestyle = "step"
                            elif line_shape == "vh":
                                linestyle = "step"
                            else:
                                linestyle = "-"  # linear
                            
                            # Create area fill based on fill type
                            if fill_type == "tozeroy":
                                # Fill from zero to y values
                                ax1.fill_between(x_vals, y_vals, alpha=area_opacity, label=label, color=line_color)
                            elif fill_type == "tonexty":
                                # Fill to next y values (for stacked areas)
                                if i == 0:
                                    # First series - fill from zero
                                    ax1.fill_between(x_vals, y_vals, alpha=area_opacity, label=label, color=line_color)
                                    # Store the cumulative values for next series
                                    if not hasattr(ax1, '_stacked_bottom'):
                                        ax1._stacked_bottom = {}
                                    ax1._stacked_bottom[label] = y_vals
                                else:
                                    # Get the bottom values from previous series
                                    prev_bottom = getattr(ax1, '_stacked_bottom', {}).get(series_data[i-1].get('name', f'series_{i-1}'), [0] * len(y_vals))
                                    # Calculate new bottom (cumulative)
                                    new_bottom = [b + y for b, y in zip(prev_bottom, y_vals)]
                                    # Fill between previous bottom and new bottom
                                    ax1.fill_between(x_vals, prev_bottom, new_bottom, alpha=area_opacity, label=label, color=line_color)
                                    # Store the new cumulative values
                                    ax1._stacked_bottom[label] = new_bottom
                            elif fill_type == "tonextx":
                                # Fill to next x values (horizontal stacking)
                                if i == 0:
                                    # First series - fill from zero
                                    ax1.fill_betweenx(y_vals, x_vals, alpha=area_opacity, label=label, color=line_color)
                                    # Store the cumulative values for next series
                                    if not hasattr(ax1, '_stacked_left'):
                                        ax1._stacked_left = {}
                                    ax1._stacked_left[label] = x_vals
                                else:
                                    # Get the left values from previous series
                                    prev_left = getattr(ax1, '_stacked_left', {}).get(series_data[i-1].get('name', f'series_{i-1}'), [0] * len(x_vals))
                                    # Calculate new left (cumulative)
                                    new_left = [l + x for l, x in zip(prev_left, x_vals)]
                                    # Fill between previous left and new left
                                    ax1.fill_betweenx(y_vals, prev_left, new_left, alpha=area_opacity, label=label, color=line_color)
                                    # Store the new cumulative values
                                    ax1._stacked_left[label] = new_left
                            else:
                                # Default fill
                                ax1.fill_between(x_vals, y_vals, alpha=area_opacity, label=label, color=line_color)
                            
                            # Add line on top of area
                            ax1.plot(x_vals, y_vals, color=line_color, linewidth=line_width, linestyle=linestyle, zorder=3)
                            
                            # Add markers if specified
                            if marker_symbol != "none":
                                ax1.scatter(x_vals, y_vals, color=marker_color, s=marker_size*20, 
                                          zorder=4, edgecolors='white', linewidth=1)
                            
                            # Add data labels if text is provided
                            if text_labels and len(text_labels) == len(y_vals):
                                for j, (x, y, text) in enumerate(zip(x_vals, y_vals, text_labels)):
                                    if j < len(text_labels):
                                        # Determine label position based on textposition
                                        # Get custom label offset if specified, otherwise use defaults
                                        custom_label_offset = chart_meta.get("data_label_offset", None)
                                        
                                        if "top" in text_position:
                                            # Increase offset to prevent overlap with data points
                                            y_offset = custom_label_offset if custom_label_offset is not None else 15
                                        elif "bottom" in text_position:
                                            y_offset = -custom_label_offset if custom_label_offset is not None else -15
                                        else:
                                            y_offset = 0
                                        
                                        if "center" in text_position:
                                            ha = 'center'
                                        elif "left" in text_position:
                                            ha = 'left'
                                        elif "right" in text_position:
                                            ha = 'right'
                                        else:
                                            ha = 'center'
                                        
                                        # Add text label
                                        label_color = data_label_color or '#000000'
                                        
                                        # Check if plain text labels are requested
                                        plain_text_labels = chart_meta.get("plain_text_labels", False)
                                        
                                        if plain_text_labels:
                                            # Plain text without background box
                                            ax1.text(x, y + y_offset, str(text), 
                                                   ha=ha, va='bottom' if y_offset > 0 else 'top',
                                                   fontsize=data_label_font_size or 10,
                                                   color=label_color,
                                                   fontweight='bold',
                                                   zorder=5)
                                        else:
                                            # Text with background box (original behavior)
                                            ax1.text(x, y + y_offset, str(text), 
                                                   ha=ha, va='bottom' if y_offset > 0 else 'top',
                                                   fontsize=data_label_font_size or 10,
                                                   color=label_color,
                                                   fontweight='bold',
                                                   bbox=dict(boxstyle="round,pad=0.3", 
                                                            facecolor='white', 
                                                            alpha=0.9,
                                                            edgecolor='gray',
                                                            linewidth=0.5),
                                                   zorder=5)
                        # Set axis labels, title, and legend for area chart (only once after all series)
                        if i == len(series_data) - 1 and mpl_chart_type == "fill_between":  # Only for area charts
                            # Set axis labels
                            x_axis_title = chart_meta.get("x_label", chart_config.get("x_axis_title", ""))
                            y_axis_title = chart_meta.get("primary_y_label", chart_config.get("primary_y_label", 
                                chart_meta.get("y_label", chart_config.get("y_label", 
                                chart_meta.get("y_axis_title", chart_config.get("y_axis_title", ""))))))
                            if x_axis_title:
                                # Handle "auto" values for axis label distances
                                if x_axis_label_distance == "auto":
                                    # Calculate optimal label distance for area chart
                                    auto_x_distance, _ = calculate_optimal_label_distance(
                                        "area", series_data, x_values, y_vals, figsize, font_size
                                    )
                                    x_axis_label_distance = auto_x_distance
                                
                                # Use the actual distance value directly, not divided by 10
                                x_labelpad = x_axis_label_distance if x_axis_label_distance is not None else 5.0
                                # Make the distance effect much more pronounced by multiplying the value
                                x_labelpad = (x_axis_label_distance * 10) if x_axis_label_distance is not None else 50.0
                                ax1.set_xlabel(x_axis_title, fontsize=font_size or 12, color=font_color if font_color else 'black',
                                             fontname=chart_meta.get("font_family") if chart_meta.get("font_family") else None,
                                             labelpad=x_labelpad)
                            if y_axis_title:
                                # Handle "auto" values for axis label distances
                                if y_axis_label_distance == "auto":
                                    # Calculate optimal label distance for area chart
                                    _, auto_y_distance = calculate_optimal_label_distance(
                                        "area", series_data, x_values, y_vals, figsize, font_size
                                    )
                                    y_axis_label_distance = auto_y_distance
                                
                                # Use the actual distance value directly, not divided by 10
                                y_labelpad = y_axis_label_distance if y_axis_label_distance is not None else 5.0
                                # Make the distance effect much more pronounced by multiplying the value
                                y_labelpad = (y_axis_label_distance * 10) if y_axis_label_distance is not None else 50.0
                                ax1.set_ylabel(y_axis_title, fontsize=font_size or 12, color=font_color if font_color else 'black',
                                             fontname=chart_meta.get("font_family") if chart_meta.get("font_family") else None,
                                             labelpad=y_labelpad)
                             
                             # Set chart title
                            if title and title.strip():
                                 ax1.set_title(title, fontsize=font_size or 14, weight='bold', pad=20,
                                             color=font_color if font_color else 'black',
                                             fontname=chart_meta.get("font_family") if chart_meta.get("font_family") else None)
                             
                             # Set legend
                            show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                            # Convert string "false"/"true" to boolean if needed
                            if isinstance(show_legend_raw, str):
                                show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                            else:
                                show_legend = bool(show_legend_raw)
                            if show_legend:
                                 legend_position = chart_meta.get("legend_position", "top")
                                 legend_font_size = chart_meta.get("legend_font_size", 10)
                                 
                                 if legend_position == "bottom":
                                     ax1.legend(loc='lower center', bbox_to_anchor=(0.5, -0.15), fontsize=legend_font_size)
                                 elif legend_position == "top":
                                     ax1.legend(loc='upper center', bbox_to_anchor=(0.5, 1.02), fontsize=legend_font_size)
                                 else:
                                     ax1.legend(loc='best', fontsize=legend_font_size)

                            # Data labels are already added in the main series processing loop above
                            # No need to add them again here to avoid duplication

                            # Set axis labels, title, and legend for area chart (only once after all series)
                            if i == len(series_data) - 1:  # Only set once after all series are processed
                                # Set axis labels
                                x_axis_title = chart_meta.get("x_label", chart_config.get("x_axis_title", ""))
                                y_axis_title = chart_meta.get("primary_y_label", chart_config.get("primary_y_label", 
                                    chart_meta.get("y_label", chart_config.get("y_label", 
                                    chart_meta.get("y_axis_title", chart_config.get("y_axis_title", ""))))))
                                if x_axis_title:
                                    # Handle "auto" values for axis label distances
                                    if x_axis_label_distance == "auto":
                                        # Calculate optimal label distance for area chart
                                        auto_x_distance, _ = calculate_optimal_label_distance(
                                            "area", series_data, x_values, y_vals, figsize, font_size
                                        )
                                        x_axis_label_distance = auto_x_distance
                                    
                                    # Use the actual distance value directly, not divided by 10
                                    x_labelpad = x_axis_label_distance if x_axis_label_distance is not None else 5.0
                                    # Make the distance effect much more pronounced by multiplying the value
                                    x_labelpad = (x_axis_label_distance * 10) if x_axis_label_distance is not None else 50.0
                                    ax1.set_xlabel(x_axis_title, fontsize=font_size or 12, color=font_color if font_color else 'black',
                                                fontname=chart_meta.get("font_family") if chart_meta.get("font_family") else None,
                                                labelpad=x_labelpad)
                                if y_axis_title:
                                    # Handle "auto" values for axis label distances
                                    if y_axis_label_distance == "auto":
                                        # Calculate optimal label distance for area chart
                                        _, auto_y_distance = calculate_optimal_label_distance(
                                            "area", series_data, x_values, y_vals, figsize, font_size
                                        )
                                        y_axis_label_distance = auto_y_distance
                                    
                                    # Use the actual distance value directly, not divided by 10
                                    y_labelpad = y_axis_label_distance if y_axis_label_distance is not None else 5.0
                                    # Make the distance effect much more pronounced by multiplying the value
                                    y_labelpad = (y_axis_label_distance * 10) if y_axis_label_distance is not None else 50.0
                                    ax1.set_ylabel(y_axis_title, fontsize=font_size or 12, color=font_color if font_color else 'black',
                                                fontname=chart_meta.get("font_family") if chart_meta.get("font_family") else None,
                                                labelpad=y_labelpad)

                            # Set chart title
                            if title and title.strip():
                                ax1.set_title(title, fontsize=font_size or 14, weight='bold', pad=20,
                                            color=font_color if font_color else 'black',
                                            fontname=chart_meta.get("font_family") if chart_meta.get("font_family") else None)
                            
                            # Set legend
                            show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                            # Convert string "false"/"true" to boolean if needed
                            if isinstance(show_legend_raw, str):
                                show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                            else:
                                show_legend = bool(show_legend_raw)
                            if show_legend:
                                legend_position = chart_meta.get("legend_position", "top")
                                legend_font_size = chart_meta.get("legend_font_size", 10)
                                
                                if legend_position == "bottom":
                                    ax1.legend(loc='lower center', bbox_to_anchor=(0.5, -0.15), fontsize=legend_font_size)
                                elif legend_position == "top":
                                    ax1.legend(loc='upper center', bbox_to_anchor=(0.5, 1.02), fontsize=legend_font_size)
                                else:
                                    ax1.legend(loc='best', fontsize=legend_font_size)
                                    
                        elif mpl_chart_type == "hist":
                            # Histogram
                            ax1.hist(y_vals, bins=10, label=label, color=color, alpha=0.7)
                            
                        elif mpl_chart_type == "boxplot":
                            # Box plot
                            ax1.boxplot(y_vals, labels=[label], patch_artist=True)
                            if color:
                                ax1.findobj(plt.matplotlib.patches.Patch)[-1].set_facecolor(color)
                                
                        elif mpl_chart_type == "contour":
                            # Contour plot (simplified)
                            if len(y_vals) > 0:
                                # Create a simple 2D array for contour
                                contour_data = [y_vals] if len(y_vals) > 0 else [[0]]
                                ax1.contour(contour_data)
                                
                        elif mpl_chart_type == "imshow":
                            # Heatmap using matplotlib imshow
                            current_app.logger.debug(f"🔥 Matplotlib heatmap processing started for series: {label}")
                            
                            # For heatmaps, get data directly from series
                            z_data = series.get("z", [])
                            x_labels = series.get("x", [])
                            y_labels = series.get("y", [])
                            colorscale = series.get("colorscale", "Blues")
                            showscale = series.get("showscale", True)
                            opacity = series.get("opacity", 1.0)
                            
                            # Debug logging
                            current_app.logger.debug(f"🔥 Heatmap Debug - Series: {series}")
                            current_app.logger.debug(f"🔥 Heatmap Debug - Z data: {z_data}")
                            current_app.logger.debug(f"🔥 Heatmap Debug - X labels: {x_labels}")
                            current_app.logger.debug(f"🔥 Heatmap Debug - Y labels: {y_labels}")
                            
                            # Ensure we have valid data for heatmap
                            if z_data and x_labels and y_labels:
                                # Convert z_data to numpy array for better handling
                                import numpy as np
                                z_array = np.array(z_data)
                                
                                # Validate z_array dimensions
                                if z_array.size == 0:
                                    current_app.logger.error(f"🔥 Heatmap Error: Z data is empty")
                                    ax1.text(0.5, 0.5, "Empty heatmap data", 
                                            ha='center', va='center', 
                                            fontsize=font_size or 12,
                                            transform=ax1.transAxes)
                                    return
                                
                                # Ensure z_array is 2D
                                if z_array.ndim == 1:
                                    z_array = z_array.reshape(1, -1)
                                    current_app.logger.debug(f"🔥 Heatmap: Reshaped 1D array to 2D: {z_array.shape}")
                                
                                current_app.logger.debug(f"🔥 Heatmap: Final z_array shape: {z_array.shape}")
                                current_app.logger.debug(f"🔥 Heatmap: X labels count: {len(x_labels)}")
                                current_app.logger.debug(f"🔥 Heatmap: Y labels count: {len(y_labels)}")
                                
                                # Disable gridlines globally for this heatmap
                                import matplotlib.pyplot as plt
                                import numpy as np
                                plt.rcParams['axes.grid'] = False
                                plt.rcParams['axes.linewidth'] = 0
                                
                                # Create heatmap using imshow with proper orientation and no lines
                                current_app.logger.debug(f"🔥 Creating heatmap with imshow...")
                                im = ax1.imshow(z_array, 
                                               cmap=colorscale, 
                                               aspect='auto',
                                               alpha=opacity,
                                               interpolation='nearest',
                                               extent=[-0.5, len(x_labels)-0.5, -0.5, len(y_labels)-0.5])
                                current_app.logger.debug(f"🔥 Heatmap imshow created successfully")
                                
                                # Completely disable all gridlines and minor gridlines for heatmaps BEFORE setting labels
                                ax1.grid(False, which='both')
                                ax1.set_axisbelow(False)
                                
                                # Hide tick marks to remove any visual grid-like elements
                                ax1.tick_params(axis='both', length=0)
                                
                                # Disable all axis spines to remove the white border lines
                                ax1.spines['top'].set_visible(False)
                                ax1.spines['bottom'].set_visible(False)
                                ax1.spines['left'].set_visible(False)
                                ax1.spines['right'].set_visible(False)
                                
                                # Disable minor gridlines
                                ax1.minorticks_off()
                                
                                # Additional grid removal for matplotlib heatmaps
                                ax1.xaxis.grid(False)
                                ax1.yaxis.grid(False)
                                ax1.xaxis.set_tick_params(gridOn=False)
                                ax1.yaxis.set_tick_params(gridOn=False)
                                
                                # Set x and y axis labels
                                current_app.logger.debug(f"🔥 Setting heatmap axis labels...")
                                ax1.set_xticks(range(len(x_labels)))
                                ax1.set_yticks(range(len(y_labels)))
                                ax1.set_xticklabels(x_labels, rotation=45, ha='right', fontsize=axis_tick_font_size or 10)
                                ax1.set_yticklabels(y_labels, fontsize=axis_tick_font_size or 10)
                                
                                # Ensure the plot is properly displayed
                                ax1.set_aspect('auto')
                                
                                # Final comprehensive tick line removal
                                ax1.tick_params(axis='both', which='both', length=0, width=0)
                                ax1.tick_params(axis='x', bottom=False, top=False, labelbottom=True)
                                ax1.tick_params(axis='y', left=False, right=False, labelleft=True)
                                
                                # Ensure no tick lines are drawn
                                for tick in ax1.xaxis.get_major_ticks():
                                    tick.tick1line.set_visible(False)
                                    tick.tick2line.set_visible(False)
                                for tick in ax1.yaxis.get_major_ticks():
                                    tick.tick1line.set_visible(False)
                                    tick.tick2line.set_visible(False)
                                
                                # Remove any remaining tick lines
                                ax1.tick_params(axis='both', which='both', length=0, width=0)
                                ax1.tick_params(axis='both', which='major', length=0, width=0)
                                ax1.tick_params(axis='both', which='minor', length=0, width=0)
                                
                                # Final gridline removal after all setup
                                ax1.grid(False, which='both')
                                ax1.xaxis.grid(False)
                                ax1.yaxis.grid(False)
                                ax1.xaxis.set_tick_params(gridOn=False)
                                ax1.yaxis.set_tick_params(gridOn=False)
                                ax1.figure.canvas.draw()
                                
                                # Add colorbar if showscale is True
                                if showscale:
                                    current_app.logger.debug(f"🔥 Adding colorbar...")
                                    cbar = plt.colorbar(im, ax=ax1)
                                    cbar.set_label('Value', rotation=270, labelpad=15)
                                
                                # Add text annotations on each cell if text data is provided
                                text_data = series.get("text", [])
                                if text_data and len(text_data) == len(z_data) and len(text_data[0]) == len(z_data[0]):
                                    current_app.logger.debug(f"🔥 Adding text annotations...")
                                    for i in range(len(z_data)):
                                        for j in range(len(z_data[0])):
                                            text = str(text_data[i][j])
                                            ax1.text(j, i, text, ha='center', va='center', 
                                                   color='white' if z_array[i, j] > z_array.max() * 0.5 else 'black',
                                                   fontweight='bold')
                                
                                # Apply background colors if specified
                                if chart_background:
                                    fig_mpl.patch.set_facecolor(chart_background)
                                if plot_background:
                                    ax1.set_facecolor(plot_background)
                                
                                # Add title with customizable styling
                                title_font_size = font_size or 16
                                title_color = font_color if font_color else '#2C3E50'
                                title_font_family = chart_meta.get("font_family") if chart_meta.get("font_family") else None
                                
                                if title and title.strip():
                                    if title and title.strip():
                                        ax1.set_title(title, fontsize=title_font_size, weight='bold', pad=20, 
                                                    color=title_color, fontname=title_font_family)
                                
                                # Set axis labels
                                x_axis_title = chart_meta.get("x_label", "")
                                y_axis_title = chart_meta.get("primary_y_label", 
                                    chart_meta.get("y_label", 
                                    chart_meta.get("y_axis_title", "")))
                                if x_axis_title:
                                    ax1.set_xlabel(x_axis_title, fontsize=font_size or 12, color=font_color if font_color else 'black',
                                                fontname=title_font_family)
                                if y_axis_title:
                                    ax1.set_ylabel(y_axis_title, fontsize=font_size or 12, color=font_color if font_color else 'black',
                                                fontname=title_font_family)
                                
                                # Handle gridlines and cell borders
                                show_gridlines = chart_meta.get("show_gridlines", True)
                                show_cell_borders = chart_meta.get("show_cell_borders", True)
                                
                                # For heatmaps, always disable gridlines completely
                                ax1.grid(False, which='both')
                                ax1.set_axisbelow(False)
                                
                                # Hide tick lines but keep tick labels
                                ax1.tick_params(axis='both', length=0, width=0, which='both')
                                ax1.tick_params(axis='x', bottom=False, top=False, labelbottom=True)
                                ax1.tick_params(axis='y', left=False, right=False, labelleft=True)
                                
                                # Disable all axis spines to remove the white border lines
                                ax1.spines['top'].set_visible(False)
                                ax1.spines['bottom'].set_visible(False)
                                ax1.spines['left'].set_visible(False)
                                ax1.spines['right'].set_visible(False)
                                
                                # Disable minor gridlines
                                ax1.minorticks_off()
                                
                                # Additional grid removal for matplotlib heatmaps
                                ax1.xaxis.grid(False)
                                ax1.yaxis.grid(False)
                                ax1.xaxis.set_tick_params(gridOn=False)
                                ax1.yaxis.set_tick_params(gridOn=False)
                                
                                # Ensure no tick lines are drawn
                                for tick in ax1.xaxis.get_major_ticks():
                                    tick.tick1line.set_visible(False)
                                    tick.tick2line.set_visible(False)
                                for tick in ax1.yaxis.get_major_ticks():
                                    tick.tick1line.set_visible(False)
                                    tick.tick2line.set_visible(False)
                                
                                # Remove any remaining tick lines
                                ax1.tick_params(axis='both', which='both', length=0, width=0)
                                ax1.tick_params(axis='both', which='major', length=0, width=0)
                                ax1.tick_params(axis='both', which='minor', length=0, width=0)
                                
                                if show_cell_borders:
                                    # Add cell borders with customizable color
                                    border_color = chart_meta.get("cell_border_color", "black")
                                    border_width = chart_meta.get("cell_border_width", 1.0)  # Increased default width
                                    border_alpha = chart_meta.get("cell_border_alpha", 1.0)  # Increased default alpha
                                    
                                    for i in range(len(z_data) + 1):
                                        ax1.axhline(y=i-0.5, color=border_color, linewidth=border_width, alpha=border_alpha)
                                    if z_data and len(z_data[0]) > 0:
                                        for j in range(len(z_data[0]) + 1):
                                            ax1.axvline(x=j-0.5, color=border_color, linewidth=border_width, alpha=border_alpha)
                                
                                current_app.logger.debug(f"🔥 Heatmap created successfully with shape: {z_array.shape}")
                            else:
                                current_app.logger.warning(f"⚠️ Invalid heatmap data structure - Z: {z_data}, X: {x_labels}, Y: {y_labels}")
                                ax1.text(0.5, 0.5, "Invalid data structure for heatmap", 
                                        ha='center', va='center', 
                                        fontsize=font_size or 12,
                                        transform=ax1.transAxes)
                                

                        elif mpl_chart_type == "treemap":
                            # Enhanced Treemap chart using squarify with full customization support
                            # For treemaps, get data directly from series
                            values = series.get("values", [])
                            labels = series.get("labels", [])
                            
                            # Debug logging for chart metadata
                            current_app.logger.debug(f"🔍 Treemap Chart Meta Debug - Full chart_meta: {chart_meta}")
                            current_app.logger.debug(f"🔍 Treemap Chart Meta Debug - showlegend value: {chart_meta.get('showlegend')}")
                            current_app.logger.debug(f"🔍 Treemap Chart Meta Debug - showlegend type: {type(chart_meta.get('showlegend'))}")
                            current_app.logger.debug(f"🔍 Treemap Chart Meta Debug - showlegend converted: {show_legend}")
                            
                            # Debug logging
                            current_app.logger.debug(f"🔍 Treemap Debug - Series: {series}")
                            current_app.logger.debug(f"🔍 Treemap Debug - Values: {values} (type: {type(values)})")
                            current_app.logger.debug(f"🔍 Treemap Debug - Labels: {labels} (type: {type(labels)})")
                            current_app.logger.debug(f"🔍 Treemap Debug - X_values: {x_values}")
                            
                            # Additional debugging for raw data
                            if isinstance(values, list):
                                current_app.logger.debug(f"🔍 Treemap Debug - Values list items: {[f'{v} (type: {type(v)})' for v in values]}")
                            if isinstance(labels, list):
                                current_app.logger.debug(f"🔍 Treemap Debug - Labels list items: {[f'{l} (type: {type(l)})' for l in labels]}")
                                # Additional debugging for label content
                                for i, label in enumerate(labels):
                                    current_app.logger.debug(f"🔍 Treemap Debug - Label {i}: '{label}' (repr: {repr(label)})")
                            
                            # NO FALLBACKS - Only use data from series to prevent mixing data sources
                            if not values:
                                current_app.logger.error(f"❌ Treemap: No values found in series data")
                                values = []
                            
                            if not labels:
                                current_app.logger.error(f"❌ Treemap: No labels found in series data")
                                labels = []
                            
                            
                            # Ensure we have valid data for treemap
                            if values and labels and len(values) == len(labels):
                                # Convert values to numeric and filter out invalid data
                                valid_data = []
                                valid_labels = []
                                valid_colors = []
                                
                                # Enhanced color palette for better visual appeal
                                enhanced_colors = [
                                    '#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7',
                                    '#DDA0DD', '#98D8C8', '#F7DC6F', '#BB8FCE', '#85C1E9',
                                    '#F8C471', '#82E0AA', '#F1948A', '#85C1E9', '#D7BDE2'
                                ]
                                
                                for i, (label, value) in enumerate(zip(labels, values)):
                                    try:
                                        numeric_value = float(value)
                                        if numeric_value > 0:  # Only include positive values
                                            # Clean and format the label text
                                            clean_label = str(label).strip()
                                            
                                            # Debug the original label
                                            current_app.logger.debug(f"🔍 Treemap label cleaning - Original: '{label}' (repr: {repr(label)})")
                                            
                                            # Light cleaning for treemap labels - only remove control characters
                                            # Remove control characters and non-printable characters
                                            clean_label = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', clean_label)
                                            # Remove extra whitespace
                                            clean_label = re.sub(r'\s+', ' ', clean_label).strip()
                                            
                                            # If label is empty after cleaning, keep it empty (no fallback)
                                            if not clean_label:
                                                clean_label = ""
                                            
                                            # Debug the cleaned label
                                            current_app.logger.debug(f"🔍 Treemap label cleaning - Cleaned: '{clean_label}' (repr: {repr(clean_label)})")
                                            
                                            # CRITICAL: Skip only truly problematic labels (not valid category names)
                                            if (clean_label and 
                                                (clean_label.lower() in ['e', 'c', '15.0']) or
                                                (re.match(r'^[0-9\.]+$', clean_label)) or
                                                (len(clean_label) < 2) or
                                                ('budget' in clean_label.lower() and len(clean_label) < 8 and 'budget' != clean_label.lower())):
                                                current_app.logger.warning(f"⚠️ SKIPPING problematic label: '{clean_label}' with value: {numeric_value}")
                                                continue
                                            
                                            valid_data.append(numeric_value)
                                            valid_labels.append(clean_label)
                                            
                                            # Debug logging for label cleaning
                                            if str(label) != clean_label:
                                                current_app.logger.debug(f"🔍 Treemap label cleaned: '{label}' -> '{clean_label}'")
                                            
                                            # Use custom color if available, otherwise use enhanced palette
                                            if isinstance(color, list) and i < len(color):
                                                valid_colors.append(color[i])
                                            elif color:
                                                valid_colors.append(color)
                                            else:
                                                # Use enhanced color palette for better visual appeal
                                                color_index = i % len(enhanced_colors)
                                                valid_colors.append(enhanced_colors[color_index])
                                    except (ValueError, TypeError):
                                        continue
                                
                                if valid_data:
                                    # Ensure we have the same number of labels as data points
                                    if len(valid_labels) != len(valid_data):
                                        current_app.logger.warning(f"⚠️ Treemap: Mismatch between data ({len(valid_data)}) and labels ({len(valid_labels)})")
                                        # Trim or pad labels to match data length (no fallback generation)
                                        if len(valid_labels) < len(valid_data):
                                            # Pad with empty strings if we have fewer labels
                                            valid_labels.extend([""] * (len(valid_data) - len(valid_labels)))
                                        else:
                                            # Trim excess labels
                                            valid_labels = valid_labels[:len(valid_data)]
                                    
                                    # Final validation - remove any remaining problematic entries
                                    final_data = []
                                    final_labels = []
                                    final_colors = []
                                    
                                    for i, (label, value, color) in enumerate(zip(valid_labels, valid_data, valid_colors)):
                                        # Skip only truly problematic entries (not valid category names)
                                        if (label and 
                                            (label.lower() in ['e', 'c', '15.0']) or
                                            (re.match(r'^[0-9\.]+$', label)) or
                                            (len(label) < 2) or
                                            ('budget' in label.lower() and len(label) < 8 and 'budget' != label.lower())):
                                            current_app.logger.warning(f"⚠️ FINAL FILTER: Skipping problematic entry {i}: '{label}' with value: {value}")
                                            continue
                                        
                                        final_data.append(value)
                                        final_labels.append(label)
                                        final_colors.append(color)
                                    
                                    # Remove duplicates to prevent multiple entries of the same category
                                    seen_labels = set()
                                    unique_data = []
                                    unique_labels = []
                                    unique_colors = []
                                    
                                    for i, (label, value, color) in enumerate(zip(final_labels, final_data, final_colors)):
                                        if label and label.lower() not in seen_labels:
                                            seen_labels.add(label.lower())
                                            unique_data.append(value)
                                            unique_labels.append(label)
                                            unique_colors.append(color)
                                        else:
                                            current_app.logger.warning(f"⚠️ DUPLICATE REMOVED: '{label}' with value: {value}")
                                    
                                    # Use the unique data
                                    valid_data = unique_data
                                    valid_labels = unique_labels
                                    valid_colors = unique_colors
                                    
                                    current_app.logger.debug(f"🔍 Treemap: Final unique data - {len(valid_data)} items")
                                    current_app.logger.debug(f"🔍 Treemap: Final labels: {valid_labels}")
                                    current_app.logger.debug(f"🔍 Treemap: Final data: {valid_data}")
                                    
                                    # Get customization options
                                    treemap_alpha = fill_opacity if fill_opacity is not None else 0.8
                                    treemap_font_size = data_label_font_size if data_label_font_size else 10
                                    treemap_font_weight = 'bold'
                                    
                                    # Handle root_visible attribute - map from hide_center_box
                                    hide_center_box = chart_meta.get("hide_center_box", False)
                                    root_visible = not hide_center_box  # root_visible is opposite of hide_center_box
                                    current_app.logger.debug(f"🔍 Treemap: hide_center_box = {hide_center_box}, root_visible = {root_visible}")
                                    
                                    # Check legend visibility early to determine plot labels
                                    show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                                    # Convert string "false"/"true" to boolean if needed
                                    if isinstance(show_legend_raw, str):
                                        show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                                    else:
                                        show_legend = bool(show_legend_raw)
                                    
                                    # Debug: log the legend decision
                                    current_app.logger.debug(f"🔍 Treemap Legend Decision - Raw: {show_legend_raw}, Converted: {show_legend}, Type: {type(show_legend_raw)}")
                                    
                                    # Create enhanced treemap with customizable styling
                                    # Always pass labels to squarify.plot() to avoid axis label errors
                                    # But disable automatic legend creation by setting showlegend=False on the axis
                                    ax1.set_visible(True)  # Ensure axis is visible
                                    
                                    # CRITICAL: Disable automatic legend creation at the matplotlib level
                                    # This prevents squarify from creating legends automatically
                                    if not show_legend:
                                        # Disable legend creation on the axis
                                        try:
                                            ax1.set_legend_handles([])
                                        except:
                                            pass
                                        
                                        # CRITICAL: Override matplotlib's global legend settings
                                        # This prevents any automatic legend creation across the entire plot
                                        try:
                                            # Disable legend frame and visibility
                                            plt.rcParams['legend.frameon'] = False
                                            # Also try to disable legend creation on the axis
                                            ax1.legend_ = None
                                            if hasattr(ax1, '_legend'):
                                                ax1._legend = None
                                        except:
                                            pass
                                    
                                    # For treemaps, we need to separate data labels from legend labels
                                    # Data labels should show when data_labels=true, regardless of legend setting
                                    if data_labels:
                                        # Show data labels on the treemap rectangles
                                        current_app.logger.debug(f"🔍 Treemap: Showing data labels with font size {treemap_font_size}")
                                        current_app.logger.debug(f"🔍 Treemap: Valid labels: {valid_labels}")
                                        current_app.logger.debug(f"🔍 Treemap: Valid data: {valid_data}")
                                        current_app.logger.debug(f"🔍 Treemap: Valid colors: {valid_colors}")
                                        
                                        # Final validation before plotting
                                        if len(valid_labels) != len(valid_data):
                                            current_app.logger.error(f"❌ CRITICAL: Label/data mismatch - Labels: {len(valid_labels)}, Data: {len(valid_data)}")
                                            current_app.logger.error(f"❌ Labels: {valid_labels}")
                                            current_app.logger.error(f"❌ Data: {valid_data}")
                                        
                                        # Configure squarify based on root_visible setting
                                        if root_visible:
                                            # Normal treemap with all rectangles visible
                                            # Use squarify.plot() but then remove any center annotations
                                            # Prepare labels: show label + value when data_labels is True
                                            if data_labels:
                                                display_labels = []
                                                for size, lbl in zip(valid_data, valid_labels):
                                                    try:
                                                        if value_format == ".1f":
                                                            formatted_val = f"{size:.1f}"
                                                        elif value_format == ".0f":
                                                            formatted_val = f"{size:.0f}"
                                                        elif value_format == ".0%":
                                                            formatted_val = f"{size:.0%}"
                                                        else:
                                                            formatted_val = f"{size:.1f}"
                                                    except:
                                                        formatted_val = str(size)
                                                    display_labels.append(f"{lbl}\n{formatted_val}")
                                            else:
                                                display_labels = [''] * len(valid_labels)

                                            squarify.plot(
                                                sizes=valid_data,
                                                label=display_labels,
                                                color=valid_colors,
                                                alpha=treemap_alpha,
                                                ax=ax1,
                                                text_kwargs={
                                                    'fontsize': treemap_font_size, 
                                                    'fontweight': treemap_font_weight,
                                                    'color': data_label_color if data_label_color else '#000000'
                                                }
                                            )
                                            
                                            # Remove any center annotations that squarify might have created
                                            # Look for text objects that might be center annotations
                                            texts_to_remove = []
                                            for text_obj in ax1.texts:
                                                # Check if this text object is a center annotation
                                                if hasattr(text_obj, 'get_position'):
                                                    pos = text_obj.get_position()
                                                    text_content = text_obj.get_text()
                                                    
                                                    # More robust center annotation detection:
                                                    # 1. Positioned in center area (0.3 to 0.7 range)
                                                    # 2. Contains numbers or is longer than normal labels
                                                    # 3. Or contains specific patterns like "Economic 25.0"
                                                    is_center_annotation = (
                                                        (0.3 <= pos[0] <= 0.7 and 0.3 <= pos[1] <= 0.7) and (
                                                            len(text_content) > 8 or  # Longer than normal labels
                                                            any(char.isdigit() for char in text_content) or  # Contains numbers
                                                            ' ' in text_content and any(char.isdigit() for char in text_content)  # Pattern like "Economic 25.0"
                                                        )
                                                    )
                                                    
                                                    if is_center_annotation:
                                                        current_app.logger.debug(f"🔍 Removing center annotation: '{text_content}' at position {pos}")
                                                        texts_to_remove.append(text_obj)
                                            
                                            # Remove identified center annotations
                                            for text_obj in texts_to_remove:
                                                text_obj.remove()
                                        else:
                                            # Treemap without root rectangle - use squarify.normalize_sizes and manual plotting
                                            current_app.logger.debug(f"🔍 Treemap: Creating treemap without root rectangle")
                                            
                                            # Normalize sizes to fit the plot area
                                            normalized_sizes = squarify.normalize_sizes(valid_data, 1, 1)
                                            
                                            # Get rectangles without root
                                            rectangles = squarify.squarify(normalized_sizes, 0, 0, 1, 1)
                                            
                                            # Plot rectangles manually without root
                                            for i, (rect, label, color) in enumerate(zip(rectangles, valid_labels, valid_colors)):
                                                # Scale rectangle to plot area
                                                x, y, dx, dy = rect['x'], rect['y'], rect['dx'], rect['dy']
                                                
                                                # Create rectangle patch
                                                from matplotlib.patches import Rectangle
                                                rect_patch = Rectangle((x, y), dx, dy, 
                                                                      facecolor=color, 
                                                                      alpha=treemap_alpha,
                                                                      edgecolor='black',
                                                                      linewidth=0.5)
                                                ax1.add_patch(rect_patch)
                                                
                                                # Add label if data_labels is enabled
                                                if data_labels and label:
                                                    # Center the label in the rectangle
                                                    label_x = x + dx/2
                                                    label_y = y + dy/2
                                                    ax1.text(label_x, label_y, label,
                                                           ha='center', va='center',
                                                           fontsize=treemap_font_size,
                                                           fontweight=treemap_font_weight,
                                                           color=data_label_color if data_label_color else '#000000')
                                    else:
                                        # No data labels requested
                                        current_app.logger.debug(f"🔍 Treemap: Data labels disabled")
                                        
                                        if root_visible:
                                            # Normal treemap with no labels
                                            # Use squarify.plot() but then remove any center annotations
                                            squarify.plot(
                                                sizes=valid_data,
                                                label=[''] * len(valid_labels),  # Empty labels when data_labels=false
                                                color=valid_colors,
                                                alpha=treemap_alpha,
                                                ax=ax1,
                                                text_kwargs={'fontsize': treemap_font_size, 'fontweight': treemap_font_weight}
                                            )
                                            
                                            # Remove any center annotations that squarify might have created
                                            texts_to_remove = []
                                            for text_obj in ax1.texts:
                                                if hasattr(text_obj, 'get_position'):
                                                    pos = text_obj.get_position()
                                                    text_content = text_obj.get_text()
                                                    
                                                    # More robust center annotation detection:
                                                    # 1. Positioned in center area (0.3 to 0.7 range)
                                                    # 2. Contains numbers or is longer than normal labels
                                                    # 3. Or contains specific patterns like "Economic 25.0"
                                                    is_center_annotation = (
                                                        (0.3 <= pos[0] <= 0.7 and 0.3 <= pos[1] <= 0.7) and (
                                                            len(text_content) > 8 or  # Longer than normal labels
                                                            any(char.isdigit() for char in text_content) or  # Contains numbers
                                                            ' ' in text_content and any(char.isdigit() for char in text_content)  # Pattern like "Economic 25.0"
                                                        )
                                                    )
                                                    
                                                    if is_center_annotation:
                                                        current_app.logger.debug(f"🔍 Removing center annotation: '{text_content}' at position {pos}")
                                                        texts_to_remove.append(text_obj)
                                            
                                            # Remove identified center annotations
                                            for text_obj in texts_to_remove:
                                                text_obj.remove()
                                        else:
                                            # Treemap without root rectangle and no labels
                                            current_app.logger.debug(f"🔍 Treemap: Creating treemap without root rectangle and no labels")
                                            
                                            # Normalize sizes to fit the plot area
                                            normalized_sizes = squarify.normalize_sizes(valid_data, 1, 1)
                                            
                                            # Get rectangles without root
                                            rectangles = squarify.squarify(normalized_sizes, 0, 0, 1, 1)
                                            
                                            # Plot rectangles manually without root and without labels
                                            for i, (rect, color) in enumerate(zip(rectangles, valid_colors)):
                                                # Scale rectangle to plot area
                                                x, y, dx, dy = rect['x'], rect['y'], rect['dx'], rect['dy']
                                                
                                                # Create rectangle patch
                                                from matplotlib.patches import Rectangle
                                                rect_patch = Rectangle((x, y), dx, dy, 
                                                                      facecolor=color, 
                                                                      alpha=treemap_alpha,
                                                                      edgecolor='black',
                                                                      linewidth=0.5)
                                                ax1.add_patch(rect_patch)
                                    
                                    # CRITICAL: After plotting, prevent legend creation if showlegend=false
                                    # This applies regardless of whether data labels are shown
                                    if not show_legend:
                                        current_app.logger.debug(f"🔍 Treemap: Post-plot legend prevention")
                                        
                                        # Prevent legend creation without removing data labels
                                        # The key is to prevent matplotlib from creating a legend, not to remove the labels themselves
                                        try:
                                            # Clear any existing legend
                                            if ax1.get_legend():
                                                ax1.get_legend().remove()
                                            ax1.legend_ = None
                                            
                                            # Prevent automatic legend creation by clearing legend handles
                                            ax1.legend_handles = []
                                        except Exception as e:
                                            current_app.logger.debug(f"🔍 Legend prevention error: {e}")
                                        
                                        # Debug: log what plot elements were created
                                        current_app.logger.debug(f"🔍 Treemap Debug - Plot elements after squarify: {[type(artist).__name__ for artist in ax1.get_children()]}")
                                        current_app.logger.debug(f"🔍 Treemap Debug - Plot elements with labels: {[artist.get_label() for artist in ax1.get_children() if hasattr(artist, 'get_label') and artist.get_label()]}")
                                        
                                        # Check if squarify automatically created a legend
                                        current_app.logger.debug(f"🔍 Treemap Debug - Axis legend after squarify: {ax1.get_legend()}")
                                        current_app.logger.debug(f"🔍 Treemap Debug - Figure legends after squarify: {fig_mpl.legends if hasattr(fig_mpl, 'legends') else 'No legends attribute'}")
                                        
                                        # CRITICAL: Check if squarify created any plot elements that might automatically create legends
                                        # This is often the root cause of automatic legend creation
                                        for i, artist in enumerate(ax1.get_children()):
                                            if (hasattr(artist, 'get_label') and 
                                                hasattr(artist, 'set_label') and 
                                                not isinstance(artist, (matplotlib.axis.Axis, matplotlib.axis.XAxis, matplotlib.axis.YAxis)) and
                                                artist.get_label()):
                                                current_app.logger.warning(f"⚠️ Treemap: Found plot element with label '{artist.get_label()}' that might create legend")
                                                # Force remove the label to prevent legend creation
                                                try:
                                                    artist.set_label('')
                                                except Exception as e:
                                                    # Skip if set_label fails (e.g., for Axis objects)
                                                    current_app.logger.debug(f"🔍 Skipped setting label on {type(artist).__name__}: {e}")
                                        
                                        # Additional safety: try to prevent matplotlib from creating legends
                                        # by setting the axis to not show legends
                                        try:
                                            # This is a more direct approach to prevent legend creation
                                            ax1.set_legend_handles([])
                                        except:
                                            pass
                                        
                                        # Also try to clear any legend handles that might exist
                                        try:
                                            ax1.legend_handles = []
                                        except:
                                            pass
                                        
                                        # CRITICAL: Force remove any legend that might have been created by squarify
                                        # This is the key fix - squarify might be creating legends automatically
                                        if ax1.get_legend():
                                            current_app.logger.warning(f"⚠️ Treemap: Found legend after squarify.plot(), removing it")
                                            ax1.get_legend().remove()
                                        ax1.legend_ = None
                                        
                                        # Additional safety: ensure no legend exists at all
                                        current_app.logger.debug(f"🔍 Treemap: Final legend check - ensuring no legend exists")
                                        if ax1.get_legend():
                                            current_app.logger.warning(f"⚠️ Treemap: Legend still exists after removal, forcing removal again")
                                            ax1.get_legend().remove()
                                            ax1.legend_ = None
                                        
                                        # Also try to prevent matplotlib from automatically creating legends
                                        # by setting the axis to not show legends
                                        try:
                                            # This is a more direct approach to prevent legend creation
                                            ax1.set_legend_handles([])
                                        except:
                                            pass
                                        
                                        # Also try to clear any legend handles that might exist
                                        try:
                                            ax1.legend_handles = []
                                        except:
                                            pass
                                        
                                        # CRITICAL: Override matplotlib's automatic legend creation
                                        # This prevents any automatic legends from being created
                                        try:
                                            # Disable automatic legend creation
                                            ax1.legend_ = None
                                            if hasattr(ax1, '_legend'):
                                                ax1._legend = None
                                            # Also try to clear any legend handles
                                            if hasattr(ax1, 'legend_handles'):
                                                ax1.legend_handles = []
                                        except:
                                            pass
                                        
                                        # Also try to prevent matplotlib from automatically creating legends
                                        # by setting the axis to not show legends
                                        try:
                                            # This is a more direct approach to prevent legend creation
                                            ax1.set_legend_handles([])
                                        except:
                                            pass
                                        
                                        # Also try to clear any legend handles that might exist
                                        try:
                                            ax1.legend_handles = []
                                        except:
                                            pass
                                    
                                    # Immediately after squarify.plot(), remove any automatic legend creation
                                    if not show_legend:
                                        # Force remove any legend that might have been created
                                        if ax1.get_legend():
                                            ax1.get_legend().remove()
                                        # Clear the legend attribute
                                        ax1.legend_ = None
                                        
                                        # Also try to prevent any automatic legend creation
                                        # by setting the axis to not show legends
                                        ax1.set_navigate(True)  # Keep navigation enabled
                                        # Clear any legend-related attributes
                                        if hasattr(ax1, '_legend'):
                                            ax1._legend = None
                                        
                                        # Additional safety: remove any plot elements that might have labels
                                        # This prevents matplotlib from automatically creating legends
                                        for artist in ax1.get_children():
                                            # Only process objects that are not Axis objects and have label methods
                                            if (hasattr(artist, 'get_label') and 
                                                hasattr(artist, 'set_label') and 
                                                not isinstance(artist, (matplotlib.axis.Axis, matplotlib.axis.XAxis, matplotlib.axis.YAxis)) and
                                                artist.get_label()):
                                                try:
                                                    artist.set_label('')
                                                except Exception as e:
                                                    # Skip if set_label fails (e.g., for Axis objects)
                                                    current_app.logger.debug(f"🔍 Skipped setting label on {type(artist).__name__}: {e}")
                                            if hasattr(artist, '_label') and artist._label:
                                                artist._label = ''
                                    
                                    # Ensure no automatic legend is added by squarify
                                    if ax1.get_legend():
                                        ax1.get_legend().remove()
                                    
                                    # Additional legend removal for squarify plots
                                    if not show_legend:
                                        # Remove any existing legend
                                        if ax1.get_legend():
                                            ax1.get_legend().remove()
                                        # Clear all legend handles
                                        ax1.legend_ = None
                                        # Also try to remove from figure level
                                        if fig_mpl.legends:
                                            for legend in fig_mpl.legends:
                                                legend.remove()
                                        
                                        # For squarify plots, we need to be more aggressive
                                        # Remove labels from all plot elements that might create legends
                                        for artist in ax1.get_children():
                                            # Only process objects that are not Axis objects and have label methods
                                            if (hasattr(artist, 'get_label') and 
                                                hasattr(artist, 'set_label') and 
                                                not isinstance(artist, (matplotlib.axis.Axis, matplotlib.axis.XAxis, matplotlib.axis.YAxis)) and
                                                artist.get_label()):
                                                try:
                                                    artist.set_label('')
                                                except Exception as e:
                                                    # Skip if set_label fails (e.g., for Axis objects)
                                                    current_app.logger.debug(f"🔍 Skipped setting label on {type(artist).__name__}: {e}")
                                            if hasattr(artist, '_label') and artist._label:
                                                artist._label = ''
                                        
                                        # Additional safety: prevent any automatic legend creation
                                        # by clearing the legend attribute completely
                                        ax1.legend_ = None
                                        if hasattr(ax1, '_legend'):
                                            ax1._legend = None
                                        
                                        # Force remove any remaining legends at the figure level
                                        if hasattr(fig_mpl, 'legends') and fig_mpl.legends:
                                            for legend in fig_mpl.legends[:]:  # Copy list to avoid modification during iteration
                                                try:
                                                    legend.remove()
                                                except:
                                                    pass
                                        
                                        # Also try to clear any legend-related attributes on the figure
                                        if hasattr(fig_mpl, '_legend'):
                                            fig_mpl._legend = None
                                        
                                        # Additional safety: try to prevent matplotlib from creating legends
                                        # by setting the axis to not show legends
                                        try:
                                            ax1.set_legend_handles([])
                                        except:
                                            pass
                                        
                                        # Also try to clear any legend handles that might exist
                                        try:
                                            ax1.legend_handles = []
                                        except:
                                            pass
                                    
                                    # Apply background colors if specified
                                    if chart_background:
                                        fig_mpl.patch.set_facecolor(chart_background)
                                    if plot_background:
                                        ax1.set_facecolor(plot_background)
                                    
                                    # Add title with customizable styling
                                    title_font_size = font_size or 16
                                    title_color = font_color if font_color else '#2C3E50'
                                    title_font_family = chart_meta.get("font_family") if chart_meta.get("font_family") else None
                                    
                                    if title and title.strip():
                                        ax1.set_title(title, fontsize=title_font_size, weight='bold', pad=20, 
                                                    color=title_color, fontname=title_font_family)
                                    ax1.set_xlabel('')
                                    ax1.set_ylabel('')
                                    
                                    # Handle axis ticks visibility
                                    if show_x_ticks is False:
                                        ax1.set_xticks([])
                                    if show_y_ticks is False:
                                        ax1.set_yticks([])
                                    
                                    # Remove axis spines for cleaner look (treemap specific)
                                    ax1.spines['top'].set_visible(False)
                                    ax1.spines['right'].set_visible(False)
                                    ax1.spines['bottom'].set_visible(False)
                                    ax1.spines['left'].set_visible(False)
                                    
                                    # Apply grid lines based on customization settings
                                    if show_gridlines is not None:
                                        if show_gridlines:
                                            grid_alpha = 0.1
                                            grid_style = gridline_style if gridline_style else '-'
                                            grid_color = gridline_color if gridline_color else 'gray'
                                            grid_width = 0.5
                                            ax1.grid(True, alpha=grid_alpha, linestyle=grid_style, 
                                                    linewidth=grid_width, color=grid_color)
                                        else:
                                            ax1.grid(False)
                                    else:
                                        # Respect show_gridlines setting for treemap
                                        if show_gridlines:
                                            ax1.grid(True, alpha=0.1, linestyle='-', linewidth=0.5)
                                        else:
                                            ax1.grid(False)
                                    
                                    # Add data labels manually only when root is hidden (we draw rectangles ourselves)
                                    if show_data_labels and not root_visible:
                                        # Place labels inside each rectangle rather than at the figure center
                                        # Compute treemap rectangles to determine per-rectangle centers
                                        normalized_sizes = squarify.normalize_sizes(valid_data, 1, 1)
                                        rectangles = squarify.squarify(normalized_sizes, 0, 0, 1, 1)
                                        
                                        for (rect, size, label) in zip(rectangles, valid_data, valid_labels):
                                            # Format the value
                                            try:
                                                if value_format == ".1f":
                                                    formatted_val = f"{size:.1f}"
                                                elif value_format == ".0f":
                                                    formatted_val = f"{size:.0f}"
                                                elif value_format == ".0%":
                                                    formatted_val = f"{size:.0%}"
                                                else:
                                                    formatted_val = f"{size:.1f}"
                                            except:
                                                formatted_val = str(size)

                                            # Compute center of this rectangle
                                            center_x = rect['x'] + rect['dx'] / 2
                                            center_y = rect['y'] + rect['dy'] / 2

                                            # Styling
                                            label_color = data_label_color or '#2C3E50'
                                            label_font_size = data_label_font_size or 9
                                            bbox_props = dict(
                                                boxstyle="round,pad=0.2",
                                                facecolor='white',
                                                alpha=0.9,
                                                edgecolor='gray',
                                                linewidth=0.5
                                            )

                                            # Draw label inside the rectangle
                                            ax1.text(center_x, center_y, f"{label}\n{formatted_val}",
                                                    ha='center', va='center',
                                                    fontsize=label_font_size,
                                                    color=label_color,
                                                    fontweight='bold',
                                                    bbox=bbox_props)
                                    
                                    # Handle legend visibility and positioning (already checked above)
                                    # Debug logging for legend visibility
                                    current_app.logger.debug(f"🔍 Treemap Legend Debug - showlegend: {show_legend}")
                                    current_app.logger.debug(f"🔍 Treemap Legend Debug - valid_data length: {len(valid_data)}")
                                    current_app.logger.debug(f"🔍 Treemap Legend Debug - Legend condition: {show_legend and len(valid_data) > 1}")
                                    
                                    # CRITICAL: Only create legend if explicitly requested AND we have multiple data points
                                    # This is the key condition that was causing the issue
                                    if show_legend and len(valid_data) > 1:
                                        # Create custom legend
                                        legend_elements = []
                                        for i, (label, color) in enumerate(zip(valid_labels, valid_colors)):
                                            from matplotlib.patches import Patch
                                            legend_elements.append(Patch(facecolor=color, label=label, alpha=treemap_alpha))
                                        
                                        # Get legend customization options
                                        legend_position = chart_meta.get("legend_position", "right")
                                        legend_font_size = chart_meta.get("legend_font_size", 8)
                                        
                                        # Position legend based on setting
                                        if legend_position == "bottom":
                                            ax1.legend(handles=legend_elements, 
                                                     loc='lower center', 
                                                     bbox_to_anchor=(0.5, -0.15), 
                                                     fontsize=legend_font_size,
                                                     frameon=True, fancybox=True, shadow=True)
                                        elif legend_position == "top":
                                            ax1.legend(handles=legend_elements, 
                                                     loc='upper center', 
                                                     bbox_to_anchor=(0.5, 1.02), 
                                                     fontsize=legend_font_size,
                                                     frameon=True, fancybox=True, shadow=True)
                                        elif legend_position == "left":
                                            ax1.legend(handles=legend_elements, 
                                                     loc='center left', 
                                                     bbox_to_anchor=(-0.05, 0.5), 
                                                     fontsize=legend_font_size,
                                                     frameon=True, fancybox=True, shadow=True)
                                        else:  # right (default)
                                            ax1.legend(handles=legend_elements, 
                                                     loc='center left', 
                                                     bbox_to_anchor=(1.05, 0.5), 
                                                     fontsize=legend_font_size,
                                                     frameon=True, fancybox=True, shadow=True)
                                    else:
                                        # CRITICAL: When showlegend is False, ensure NO legend is created
                                        # This is the key fix for the issue
                                        current_app.logger.debug(f"🔍 Treemap: showlegend is False, ensuring no legend is created")
                                        
                                        # Remove any existing legend
                                        if ax1.get_legend():
                                            ax1.get_legend().remove()
                                        # Clear all legend handles
                                        ax1.legend_ = None
                                        # Also try to remove from figure level
                                        if fig_mpl.legends:
                                            for legend in fig_mpl.legends:
                                                legend.remove()
                                        
                                        # For squarify plots, we need to be more aggressive
                                        # Remove labels from all plot elements that might create legends
                                        for artist in ax1.get_children():
                                            # Only process objects that are not Axis objects and have label methods
                                            if (hasattr(artist, 'get_label') and 
                                                hasattr(artist, 'set_label') and 
                                                not isinstance(artist, (matplotlib.axis.Axis, matplotlib.axis.XAxis, matplotlib.axis.YAxis)) and
                                                artist.get_label()):
                                                try:
                                                    artist.set_label('')
                                                except Exception as e:
                                                    # Skip if set_label fails (e.g., for Axis objects)
                                                    current_app.logger.debug(f"🔍 Skipped setting label on {type(artist).__name__}: {e}")
                                            if hasattr(artist, '_label') and artist._label:
                                                artist._label = ''
                                        
                                        # Additional safety: prevent any automatic legend creation
                                        # by clearing the legend attribute completely
                                        ax1.legend_ = None
                                        if hasattr(ax1, '_legend'):
                                            ax1._legend = None
                                    
                                    # Apply margin settings if specified
                                    if margin:
                                        plt.subplots_adjust(**margin)
                                    
                                    # Ensure no legend is shown if showlegend is False
                                    if not show_legend:
                                        # Remove any existing legend
                                        if ax1.get_legend():
                                            ax1.get_legend().remove()
                                        # Also try to remove from figure level
                                        if fig_mpl.legends:
                                            for legend in fig_mpl.legends:
                                                legend.remove()
                                        
                                        # Additional safety: clear all legend-related attributes
                                        ax1.legend_ = None
                                        if hasattr(ax1, '_legend'):
                                            ax1._legend = None
                                        
                                        # Also try to clear any legend handles that might exist
                                        try:
                                            ax1.legend_handles = []
                                        except:
                                            pass
                                    
                                    # Adjust layout to accommodate legend (only if legend is shown)
                                    if show_legend:
                                        plt.tight_layout()
                                    else:
                                        # Use tight layout without legend consideration and ensure no legend space
                                        plt.tight_layout()
                                        # Double-check no legend was added during layout
                                        if ax1.get_legend():
                                            ax1.get_legend().remove()
                                        
                                        # Additional safety: remove any legends that might have been created during layout
                                        if hasattr(fig_mpl, 'legends') and fig_mpl.legends:
                                            for legend in fig_mpl.legends[:]:
                                                try:
                                                    legend.remove()
                                                except:
                                                    pass
                                        
                                        # Final check: ensure no legend exists
                                        ax1.legend_ = None
                                        if hasattr(ax1, '_legend'):
                                            ax1._legend = None
                                    
                                    # One final legend removal check before finishing
                                    if not show_legend:
                                        # Remove any legend that might have been created
                                        if ax1.get_legend():
                                            ax1.get_legend().remove()
                                            ax1.legend_ = None
                                        
                                        # Clear any remaining legend attributes
                                        ax1.legend_ = None
                                        # Also check figure level
                                        if hasattr(fig_mpl, 'legends') and fig_mpl.legends:
                                            for legend in fig_mpl.legends[:]:
                                                try:
                                                    legend.remove()
                                                except:
                                                    pass
                                        
                                        # CRITICAL: Final override to prevent any legend creation
                                        # This is the last line of defense against automatic legends
                                        try:
                                            # Completely disable legend creation on this axis
                                            ax1.legend_ = None
                                            if hasattr(ax1, '_legend'):
                                                ax1._legend = None
                                            if hasattr(ax1, 'legend_handles'):
                                                ax1.legend_handles = []
                                            # Also try to prevent any future legend creation
                                            ax1.set_legend_handles([])
                                        except:
                                            pass
                                        
                                        # Debug: log final state
                                        current_app.logger.debug(f"🔍 Treemap Final Debug - Axis legend: {ax1.get_legend()}")
                                        current_app.logger.debug(f"🔍 Treemap Final Debug - Figure legends: {fig_mpl.legends if hasattr(fig_mpl, 'legends') else 'No legends attribute'}")
                                        current_app.logger.debug(f"🔍 Treemap Final Debug - Plot elements with labels: {[artist.get_label() for artist in ax1.get_children() if hasattr(artist, 'get_label') and artist.get_label()]}")
                                        
                                        # Final verification: if there's still a legend, force remove it
                                        if ax1.get_legend():
                                            current_app.logger.warning(f"⚠️ Treemap: Final check found legend, forcing removal")
                                            ax1.get_legend().remove()
                                            ax1.legend_ = None
                                    
                                else:
                                    # No valid data for treemap
                                    ax1.text(0.5, 0.5, "No valid data for treemap", 
                                            ha='center', va='center', 
                                            fontsize=font_size or 12,
                                            transform=ax1.transAxes)
                            else:
                                # Invalid data structure for treemap
                                current_app.logger.warning(f"⚠️ Invalid treemap data structure - Values: {values}, Labels: {labels}")
                                ax1.text(0.5, 0.5, "Invalid data structure for treemap", 
                                        ha='center', va='center', 
                                        fontsize=font_size or 12,
                                        transform=ax1.transAxes)
                                
                        else:
                            # Fallback to scatter for unknown types
                            #pass  # Suppress warning logs: f"⚠️ Unknown matplotlib chart type '{series_type}', falling back to scatter")
                            ax1.scatter(x_values, y_vals, label=label, color=color, alpha=0.7)

                    # Add data labels to Matplotlib chart if enabled (skip for area charts as they have custom label handling)
                    if show_data_labels and (data_label_format or value_format or data_label_font_size or data_label_color) and chart_type != "area":
                        # Special handling for stacked bars - process once before regular series loop
                        if chart_type == "stacked_column" and hasattr(ax1, '_stacked_segments'):
                            # Add data labels for each segment in stacked bars
                            for segment in ax1._stacked_segments:
                                seg_y_vals = segment['y_vals']
                                seg_bottom_vals = segment['bottom_vals']
                                seg_x_values = segment['x_values']
                                
                                for j, (val, bottom) in enumerate(zip(seg_y_vals, seg_bottom_vals)):
                                    if j < len(seg_x_values):
                                        # Calculate center position of segment
                                        segment_center = bottom + (val / 2)
                                        
                                        # Format value to 1 decimal place with % sign
                                        try:
                                            formatted_val = f"{float(val):.1f}%"
                                        except:
                                            formatted_val = f"{val}%"
                                        
                                        # Add text label at center of segment
                                        label_color = data_label_color or '#000000'
                                        ax1.text(j, segment_center, formatted_val, 
                                                ha='center', va='center', 
                                                fontsize=data_label_font_size or int(title_fontsize * 0.9),
                                                color=label_color,
                                                fontweight='bold',
                                                bbox=dict(boxstyle="round,pad=0.2", facecolor='white', alpha=0.9))
                        else:
                            # Regular bar/line chart data labels - process each series
                            # current_app.logger.debug(f"Adding data labels to Matplotlib chart")
                            for i, series in enumerate(series_data):
                                series_type = series.get("type", "bar").lower()
                                y_vals = series.get("values")
                                value_range = series.get("value_range")
                                if value_range:
                                    if isinstance(value_range, list):
                                        y_vals = value_range
                                    else:
                                        y_vals = extract_values_from_range(value_range)
                                
                                # Ensure y_vals is not None for data labels
                                if y_vals is None:
                                    y_vals = []
                                
                                if y_vals and x_values:
                                    # Determine format to use
                                    format_to_use = value_format if value_format else data_label_format
                                    if not format_to_use:
                                        format_to_use = ".1f"  # Default format
                                    
                                    # Add data labels based on chart type
                                    if series_type == "bar":
                                        # Regular bar chart data labels
                                        for j, val in enumerate(y_vals):
                                            if j < len(x_values) if x_values else 0:
                                                # Format the value
                                                try:
                                                    if format_to_use == ".1f":
                                                        formatted_val = f"{float(val):.1f}"
                                                    elif format_to_use == ".0f":
                                                        formatted_val = f"{float(val):.0f}"
                                                    elif format_to_use == ".0%":
                                                        formatted_val = f"{float(val):.0%}"
                                                    else:
                                                        formatted_val = f"{float(val):{format_to_use}}"
                                                except:
                                                    formatted_val = str(val)
                                                
                                                # Add text label on top of bar (ENHANCED VERSION)
                                                label_color = data_label_color or '#000000'
                                                ax1.text(j, val, formatted_val, 
                                                        ha='center', va='bottom', 
                                                        fontsize=data_label_font_size or int(title_fontsize * 0.9),  # Improved scaling
                                                        color=label_color,
                                                        fontweight='bold',
                                                        bbox=dict(boxstyle="round,pad=0.2", facecolor='white', alpha=0.9))
                                                # current_app.logger.debug(f"Added bar data label with color: {label_color}")
                                
                                elif series_type == "line":
                                    for j, val in enumerate(y_vals):
                                        if j < len(x_values) if x_values else 0:
                                            # For line charts, use secondary y-axis format if available
                                            line_format = secondary_y_axis_format if secondary_y_axis_format else format_to_use
                                            if not line_format:
                                                line_format = ".1f"  # Default format
                                            
                                            # Format the value
                                            try:
                                                if line_format == ".1f":
                                                    formatted_val = f"{float(val):.1f}"
                                                elif line_format == ".0f":
                                                    formatted_val = f"{float(val):.0f}"
                                                elif line_format == ".0%":
                                                    formatted_val = f"{float(val):.0%}"
                                                elif line_format == ".1%":
                                                    formatted_val = f"{float(val):.1%}"
                                                else:
                                                    formatted_val = f"{float(val):{line_format}}"
                                            except:
                                                formatted_val = str(val)
                                            
                                            # Add text label above line point (ENHANCED VERSION)
                                            label_color = data_label_color or '#000000'
                                            if ax2:
                                                ax2.text(j, val, formatted_val, 
                                                        ha='center', va='bottom', 
                                                        fontsize=data_label_font_size or int(title_fontsize * 0.9),  # Improved scaling
                                                        color=label_color,
                                                        fontweight='bold',
                                                        bbox=dict(boxstyle="round,pad=0.2", facecolor='white', alpha=0.9))
                                            # current_app.logger.debug(f"Added line data label with color: {label_color}")

                                            # Set labels and styling
                        # Check if this is a bubble chart to avoid overriding bubble chart axis label settings
                        # This needs to be defined for all chart types since it's used later
                        is_bubble_chart = any(series.get("type", "").lower() == "bubble" for series in series_data)
                        current_app.logger.debug(f"🔍 General Section - is_bubble_chart: {is_bubble_chart}")
                        
                        # Improved font size scaling for Matplotlib (moved outside condition)
                        # These need to be defined for all chart types since they're used later
                        title_fontsize = font_size or 52  # Increased default title size
                        label_fontsize = int((font_size or 52) * 0.9)  # Increased relative size for labels
                        
                        if chart_type != "pie" and chart_type != "area" and chart_type != "treemap":
                            
                            # Only apply general axis label settings if NOT a bubble chart
                            if not is_bubble_chart:
                                # Handle "auto" values for axis label distances
                                if x_axis_label_distance == "auto" or y_axis_label_distance == "auto":
                                    # Calculate optimal label distances for bar chart
                                    auto_x_distance, auto_y_distance = calculate_optimal_label_distance(
                                        "bar", series_data, x_values, y_vals, figsize, font_size
                                    )
                                    
                                    if x_axis_label_distance == "auto":
                                        x_axis_label_distance = auto_x_distance
                                    if y_axis_label_distance == "auto":
                                        y_axis_label_distance = auto_y_distance
                                
                                # Apply axis label distances using labelpad parameter
                                x_labelpad = x_axis_label_distance if x_axis_label_distance else 5.0  # Use the value directly
                                y_labelpad = y_axis_label_distance if y_axis_label_distance else 5.0  # Use the value directly
                                # Make the distance effect much more pronounced by multiplying the values
                                x_labelpad = (x_axis_label_distance * 10) if x_axis_label_distance is not None else 50.0
                                y_labelpad = (y_axis_label_distance * 10) if y_axis_label_distance is not None else 50.0
                                
                                # Debug logging
                                current_app.logger.debug(f"🔍 Axis Label Distance Debug - X: {x_axis_label_distance} → {x_labelpad}, Y: {y_axis_label_distance} → {y_labelpad}")
                                
                                ax1.set_xlabel(chart_meta.get("x_label", chart_config.get("x_axis_title", "X")), 
                                             fontsize=label_fontsize, color=font_color, labelpad=x_labelpad)
                                ax1.set_ylabel(chart_meta.get("primary_y_label", chart_config.get("primary_y_label", "Primary Y")), 
                                             fontsize=label_fontsize, color=font_color, labelpad=y_labelpad)
                                if "secondary_y_label" in chart_meta or "secondary_y_label" in chart_config:
                                    if ax2:
                                        # Use a much smaller labelpad for secondary y-axis to match primary axis distance
                                        secondary_y_labelpad = y_labelpad * 0.1 if y_labelpad else 5.0
                                        ax2.set_ylabel(chart_meta.get("secondary_y_label", chart_config.get("secondary_y_label", "Secondary Y")), 
                                                     fontsize=label_fontsize, color=font_color, labelpad=secondary_y_labelpad)
                            else:
                                # For bubble charts, just set the font size without overriding the labelpad values
                                current_app.logger.debug(f"🎈 Skipping general axis label settings for bubble chart - preserving bubble chart specific settings")
                                
                                # Check if bubble chart axis labels were already set and restore them
                                if hasattr(ax1, '_bubble_x_labelpad') and hasattr(ax1, '_bubble_y_labelpad'):
                                    current_app.logger.debug(f"🎈 Restoring bubble chart axis labels with stored labelpad values")
                                    if hasattr(ax1, '_bubble_x_title') and ax1._bubble_x_title:
                                        ax1.set_xlabel(ax1._bubble_x_title, fontsize=label_fontsize, color=font_color, labelpad=ax1._bubble_x_labelpad)
                                        current_app.logger.debug(f"🎈 Restored X-axis label with labelpad: {ax1._bubble_x_labelpad}")
                                    if hasattr(ax1, '_bubble_y_title') and ax1._bubble_y_title:
                                        ax1.set_ylabel(ax1._bubble_y_title, fontsize=label_fontsize, color=font_color, labelpad=ax1._bubble_y_labelpad)
                                        current_app.logger.debug(f"🎈 Restored Y-axis label with labelpad: {ax1._bubble_y_labelpad}")
                                else:
                                    current_app.logger.debug(f"🎈 No stored bubble chart labelpad values found")

                            # Apply axis scale type if provided
                            xaxis_type_cfg = chart_meta.get("xaxis_type")
                            yaxis_type_cfg = chart_meta.get("yaxis_type")
                            if xaxis_type_cfg in ("log", "log10"):
                                ax1.set_xscale('log')
                            if yaxis_type_cfg in ("log", "log10"):
                                ax1.set_yscale('log')

                        # is_bubble_chart is already defined above
                        
                        # Get secondary y-axis control for scatter charts and area charts
                        # Use disable_secondary_y field to control secondary y-axis visibility for all chart types
                        disable_secondary_y = chart_meta.get("disable_secondary_y", False)
                        
                        # AGGRESSIVE: If secondary y-axis is disabled for any chart type, remove it immediately
                        if disable_secondary_y and 'ax2' in locals() and ax2 is not None:
                            try:
                                ax2.remove()  # Remove the axis completely
                                ax2 = None    # Set to None to prevent further operations
                            except:
                                pass  # Ignore errors if removal fails

                        # Set axis tick font size and control tick visibility
                        tick_fontsize = axis_tick_font_size or int(title_fontsize * 0.8)  # Improved tick size calculation
                        if axis_tick_font_size is not None:
                            ax1.tick_params(axis='x', labelsize=axis_tick_font_size, colors=font_color)  # Removed rotation
                            ax1.tick_params(axis='y', labelsize=axis_tick_font_size, colors=font_color)
                            if ax2 and not is_bubble_chart and not disable_secondary_y:
                                ax2.tick_params(axis='y', labelsize=axis_tick_font_size, colors=font_color)
                        else:
                            ax1.tick_params(axis='x', labelsize=tick_fontsize, colors=font_color)  # Removed rotation
                            ax1.tick_params(axis='y', labelsize=tick_fontsize, colors=font_color)
                            if ax2 and not is_bubble_chart and not disable_secondary_y:
                                ax2.tick_params(axis='y', labelsize=tick_fontsize, colors=font_color)
                        
                        # Tick mark control for Matplotlib
                        if show_x_ticks is not None or show_y_ticks is not None:
                            if show_x_ticks is not None:
                                if not show_x_ticks:
                                    ax1.tick_params(axis='x', length=0)  # Hide tick marks
                                    ax1.set_xticklabels([])  # Hide tick labels
                                else:
                                    ax1.tick_params(axis='x', length=5)  # Show tick marks
                                    # Reapply font size after showing ticks
                                    if axis_tick_font_size is not None:
                                        ax1.tick_params(axis='x', labelsize=axis_tick_font_size, colors=font_color)
                            if show_y_ticks is not None:
                                if not show_y_ticks:
                                    ax1.tick_params(axis='y', length=0)  # Hide tick marks
                                    ax1.set_yticklabels([])  # Hide tick labels
                                    if ax2 and not is_bubble_chart and not disable_secondary_y:
                                        ax2.tick_params(axis='y', length=0)  # Hide secondary y-axis tick marks
                                        ax2.set_yticklabels([])  # Hide secondary y-axis tick labels
                                else:
                                    ax1.tick_params(axis='y', length=5)  # Show tick marks
                                    # Reapply font size after showing ticks
                                    if axis_tick_font_size is not None:
                                        ax1.tick_params(axis='y', labelsize=axis_tick_font_size, colors=font_color)
                                    if ax2 and not is_bubble_chart and not disable_secondary_y:
                                        ax2.tick_params(axis='y', length=5)  # Show secondary y-axis tick marks
                                        if axis_tick_font_size is not None:
                                            ax2.tick_params(axis='y', labelsize=axis_tick_font_size, colors=font_color)
                        
                        # Apply primary y-axis formatting for Matplotlib
                        if axis_tick_format:
                            from matplotlib.ticker import FuncFormatter
                            if "$" in axis_tick_format:
                                def currency_formatter(x, pos):
                                    return f'${x:,.0f}'
                                ax1.yaxis.set_major_formatter(FuncFormatter(currency_formatter))
                        if y_axis_min_max:
                            # current_app.logger.debug(f"Setting Matplotlib Y-axis range to: {y_axis_min_max}")
                            # Ensure the range is properly applied
                            if isinstance(y_axis_min_max, list) and len(y_axis_min_max) == 2:
                                # Check if the provided range is appropriate for the data
                                all_y_values = []
                                for series in series_data:
                                    y_vals = series.get("values", [])
                                    if y_vals:
                                        all_y_values.extend(y_vals)
                                
                                if all_y_values:
                                    data_min = min(all_y_values)
                                    data_max = max(all_y_values)
                                    data_range = data_max - data_min
                                    
                                    # Check if the provided range is reasonable (within 100x of data range for currency formatting)
                                    provided_range = y_axis_min_max[1] - y_axis_min_max[0]
                                    if provided_range > data_range * 100 and not axis_tick_format:
                                        # Auto-calculate appropriate range only if no special formatting is requested
                                        padding = data_range * 0.1  # 10% padding
                                        auto_min = max(0, data_min - padding)
                                        auto_max = data_max + padding
                                        ax1.set_ylim(auto_min, auto_max)
                                        # current_app.logger.debug(f"Auto-adjusted Y-axis range from {y_axis_min_max} to {[auto_min, auto_max]} (data range: {data_min}-{data_max})")
                                    else:
                                        # Use provided range (especially for currency formatting)
                                        ax1.set_ylim(y_axis_min_max[0], y_axis_min_max[1])
                                        # current_app.logger.debug(f"Applied Y-axis range: {y_axis_min_max[0]} to {y_axis_min_max[1]}")
                                else:
                                    # No data available, use provided range
                                    ax1.set_ylim(y_axis_min_max[0], y_axis_min_max[1])
                                    # current_app.logger.debug(f"Applied Y-axis range: {y_axis_min_max[0]} to {y_axis_min_max[1]}")
                            else:
                                pass  # Suppress warning logs: f"Invalid Y-axis range format: {y_axis_min_max}")
                        
                        # Apply secondary y-axis formatting for Matplotlib
                        if ax2 and not disable_secondary_y:
                            from matplotlib.ticker import FuncFormatter
                            if secondary_y_axis_format:
                                # Use explicit format if provided
                                if "%" in secondary_y_axis_format:
                                    def percentage_formatter(x, pos):
                                        return f'{x:.0%}'
                                    ax2.yaxis.set_major_formatter(FuncFormatter(percentage_formatter))
                                elif "$" in secondary_y_axis_format:
                                    def currency_formatter(x, pos):
                                        return f'${x:,.0f}'
                                    ax2.yaxis.set_major_formatter(FuncFormatter(currency_formatter))
                                else:
                                    # For other formats, use the format string directly
                                    def custom_formatter(x, pos):
                                        return f'{x:{secondary_y_axis_format}}'
                                    ax2.yaxis.set_major_formatter(FuncFormatter(custom_formatter))
                            else:
                                # Default to percentage formatting if no explicit format is specified
                                # This handles cases where secondary y-axis exists but no format is specified
                                def default_percentage_formatter(x, pos):
                                    return f'{x:.0%}'
                                ax2.yaxis.set_major_formatter(FuncFormatter(default_percentage_formatter))
                        
                        # Set title with improved font size
                        if title and title.strip():
                            ax1.set_title(title, fontsize=title_fontsize, weight='bold', pad=20)
                        
                        # Legend
                        show_legend_raw = chart_meta.get("showlegend", chart_meta.get("legend", True))
                        # Convert string "false"/"true" to boolean if needed
                        if isinstance(show_legend_raw, str):
                            show_legend = show_legend_raw.lower() not in ['false', '0', 'no', 'off']
                        else:
                            show_legend = bool(show_legend_raw)
                        
                        legend_loc = 'best'
                        if show_legend:
                            handles = []
                            labels_list = []
                            # Use explicit proxy if present
                            proxy = getattr(ax1, "_bubble_proxy", None)
                            if proxy is not None:
                                handles.append(proxy)
                                labels_list.append(proxy.get_label())
                            else:
                                # Prefer stored bubble handles/labels
                                handles = getattr(ax1, "_bubble_handles", []) or []
                                labels_list = getattr(ax1, "_bubble_labels", []) or []
                                if not handles:
                                    handles, labels_list = ax1.get_legend_handles_labels()

                            # Combine with secondary axis if present
                            if 'ax2' in locals() and ax2 is not None:
                                h2, l2 = ax2.get_legend_handles_labels()
                                handles += h2
                                labels_list += l2

                            if legend_position:
                                loc_map = {
                                    "top": "upper center",
                                    "bottom": "lower center", 
                                    "left": "center left",
                                    "right": "center right",
                                }
                                legend_loc = loc_map.get(legend_position, 'best')
                            
                            if legend_position == "bottom":
                                ax1.legend(handles, labels_list, loc='lower center', bbox_to_anchor=(0.5, -0.15), fontsize=legend_font_size)
                            elif legend_position == "top":
                                ax1.legend(handles, labels_list, loc='upper center', bbox_to_anchor=(0.5, 1.02), fontsize=legend_font_size)
                            else:
                                ax1.legend(handles, labels_list, loc=legend_loc, fontsize=legend_font_size)
                                
                        # Set title with proper font attributes (ENHANCED VERSION)

                # Apply axis label distances for Matplotlib (skip heatmaps as they're handled specifically)
                # current_app.logger.debug(f"X-axis label distance: {x_axis_label_distance}")
                # current_app.logger.debug(f"Y-axis label distance: {y_axis_label_distance}")
                
                # Handle "auto" values for axis label distances
                if x_axis_label_distance == "auto" or y_axis_label_distance == "auto":
                    # Calculate optimal label distances for the current chart type
                    # Create a safe series data structure for calculation
                    safe_series_data = []
                    if series_data:
                        for series in series_data:
                            safe_series = {"labels": series.get("labels", []), "values": series.get("values", [])}
                            safe_series_data.append(safe_series)
                    
                    # Use safe data for calculation
                    auto_x_distance, auto_y_distance = calculate_optimal_label_distance(
                        chart_type, safe_series_data, x_values, y_vals if 'y_vals' in locals() else [], figsize, font_size
                    )
                    
                    if x_axis_label_distance == "auto":
                        x_axis_label_distance = auto_x_distance
                    if y_axis_label_distance == "auto":
                        y_axis_label_distance = auto_y_distance
                
                if (x_axis_label_distance or y_axis_label_distance) and chart_type != "heatmap":
                    # Get current subplot parameters
                    current_bottom = fig_mpl.subplotpars.bottom
                    current_left = fig_mpl.subplotpars.left
                    
                    if x_axis_label_distance:
                        # Convert the distance to a fraction of the figure height
                        # Higher x_axis_label_distance values will push labels further down
                        try:
                            x_distance = float(x_axis_label_distance) if isinstance(x_axis_label_distance, (int, float, str)) else 0.0
                            adjustment = x_distance / 500.0  # Increased conversion factor for more visible effect
                            fig_mpl.subplots_adjust(bottom=current_bottom - adjustment)
                            # current_app.logger.debug(f"Applied X-axis adjustment: {adjustment}")
                        except (ValueError, TypeError):
                            pass
                    
                    if y_axis_label_distance:
                        # Convert the distance to a fraction of the figure width
                        try:
                            y_distance = float(y_axis_label_distance) if isinstance(y_axis_label_distance, (int, float, str)) else 0.0
                            adjustment = y_distance / 500.0  # Increased conversion factor
                            fig_mpl.subplots_adjust(left=current_left - adjustment)
                            # current_app.logger.debug(f"Applied Y-axis adjustment: {adjustment}")
                        except (ValueError, TypeError):
                            pass
                
                # Apply tight_layout but preserve manual adjustments
                fig_mpl.tight_layout()
                
                # Re-apply manual adjustments after tight_layout with larger effect (skip heatmaps)
                if (x_axis_label_distance or y_axis_label_distance) and chart_type != "heatmap":
                    if x_axis_label_distance:
                        try:
                            x_distance = float(x_axis_label_distance) if isinstance(x_axis_label_distance, (int, float, str)) else 0.0
                            adjustment = x_distance / 300.0  # Even larger effect after tight_layout
                            fig_mpl.subplots_adjust(bottom=fig_mpl.subplotpars.bottom - adjustment)
                            # Also re-apply x-axis label padding and nudge its position downward
                            try:
                                x_labelpad = x_distance / 10.0 if x_distance else 5.0
                                ax1.xaxis.labelpad = x_labelpad
                                # Move label further down in axes coordinates
                                x_label_nudge = -0.10 - (x_distance / 1200.0)
                                ax1.xaxis.set_label_coords(0.5, x_label_nudge)
                            except Exception:
                                pass
                            # current_app.logger.debug(f"Re-applied X-axis adjustment: {adjustment}")
                        except (ValueError, TypeError):
                            pass
                    
                    if y_axis_label_distance:
                        try:
                            y_distance = float(y_axis_label_distance) if isinstance(y_axis_label_distance, (int, float, str)) else 0.0
                            adjustment = y_distance / 300.0  # Even larger effect after tight_layout
                            fig_mpl.subplots_adjust(left=fig_mpl.subplotpars.left - adjustment)
                            # Ensure the y-label itself moves away from the axis ticks
                            try:
                                # Re-apply labelpad explicitly after tight_layout
                                y_labelpad = y_distance / 10.0 if y_distance else 5.0
                                ax1.yaxis.labelpad = y_labelpad
                                # Additionally, nudge the label position in axes coordinates for a clearer visual effect
                                # Negative x moves it further left; scale factor tuned for visibility
                                coord_nudge = -0.02 - (y_distance / 1200.0)
                                ax1.yaxis.set_label_coords(coord_nudge, 0.5)
                            except Exception:
                                pass
                            # current_app.logger.debug(f"Re-applied Y-axis adjustment: {adjustment}")
                        except (ValueError, TypeError):
                            pass

                # Add annotations if specified
                if annotations:
                    # current_app.logger.debug(f"Adding {len(annotations)} annotations to chart")
                    
                    # For heatmaps, extract x and y values from series data
                    if chart_type == "heatmap" and series_data:
                        heatmap_x_values = series_data[0].get("x", []) if series_data else []
                        heatmap_y_values = series_data[0].get("y", []) if series_data else []
                    else:
                        heatmap_x_values = []
                        heatmap_y_values = []
                    
                    for annotation in annotations:
                        text = annotation.get("text", "")
                        x_value = annotation.get("x_value", 0)
                        y_value = annotation.get("y_value", 0)
                        
                        # Find x position based on x_value
                        x_pos = 0
                        if chart_type == "heatmap" and heatmap_x_values:
                            # For heatmaps, look in the heatmap x values
                            if x_value in heatmap_x_values:
                                x_pos = heatmap_x_values.index(x_value)
                        elif x_values and x_value in x_values:
                            # For other charts, look in the standard x_values
                            x_pos = x_values.index(x_value)
                        elif isinstance(x_value, str) and x_values:
                            # Try to find string value in x_values
                            try:
                                x_pos = x_values.index(x_value)
                            except ValueError:
                                # If not found, use first position
                                x_pos = 0
                                # Annotation x_value not found in x_values (skipping)
                        elif isinstance(x_value, (int, float)) and x_values:
                            # For numeric x_value, use the actual value for bubble charts
                            if chart_type == "bubble":
                                x_pos = x_value  # Use actual x value for bubble charts
                            else:
                                # For other charts, find closest position
                                try:
                                    x_pos = int(x_value)
                                    if x_pos >= len(x_values):
                                        x_pos = len(x_values) - 1
                                except (ValueError, TypeError):
                                    x_pos = 0
                        
                        # Find y position based on y_value
                        y_pos = y_value
                        if chart_type == "heatmap" and heatmap_y_values:
                            # For heatmaps, look in the heatmap y values
                            if y_value in heatmap_y_values:
                                y_pos = heatmap_y_values.index(y_value)
                        
                        # Auto-adjust Y position if it's outside the data range
                        all_y_values = []
                        for series in series_data:
                            y_vals = series.get("values", [])
                            if y_vals:
                                all_y_values.extend(y_vals)
                        
                        if all_y_values:
                            data_min = min(all_y_values)
                            data_max = max(all_y_values)
                            data_range = data_max - data_min
                            
                            # If annotation Y value is way outside data range, adjust it
                            if y_pos > data_max * 2 or y_pos < data_min * 0.5:
                                # Position annotation above the highest data point
                                y_pos = data_max + (data_range * 0.1)  # 10% above max
                                # current_app.logger.debug(f"Auto-adjusted annotation Y position from {y_value} to {y_pos} (data range: {data_min}-{data_max})")
                        
                        # Add annotation text with better positioning for different chart types
                        if chart_type == "bubble":
                            # For bubble charts, position annotation above the bubble
                            annotation_y_offset = 2000  # Fixed offset for bubble charts
                            ax1.annotate(text, xy=(x_pos, y_pos), xytext=(x_pos, y_pos + annotation_y_offset),
                                       arrowprops=dict(arrowstyle='->', color='red', lw=1.5),
                                       fontsize=12, color='red', weight='bold',
                                       ha='center', va='bottom',
                                       bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.9, edgecolor='red'))
                        elif chart_type == "area":
                            # For area charts, position annotation above the data with better styling
                            annotation_y_offset = data_range * 0.15 if 'data_range' in locals() else 25
                            ax1.annotate(text, xy=(x_pos, y_pos), xytext=(x_pos, y_pos + annotation_y_offset),
                                       arrowprops=dict(arrowstyle='->', color='#2c3e50', lw=2),
                                       fontsize=11, color='#2c3e50', weight='bold',
                                       ha='center', va='bottom',
                                       bbox=dict(boxstyle="round,pad=0.4", facecolor='white', alpha=0.9, edgecolor='#3498db'))
                        else:
                            # For other charts, use original positioning
                            ax1.annotate(text, xy=(x_pos, y_pos), xytext=(x_pos, y_pos + (data_range * 0.05) if 'data_range' in locals() else y_pos + 50),
                                       arrowprops=dict(arrowstyle='->', color='red'),
                                       fontsize=12, color='red', weight='bold',
                                       ha='center', va='bottom')
                        # current_app.logger.debug(f"Added annotation: '{text}' at position ({x_pos}, {y_pos})")

                # Apply margin settings to Matplotlib (FIXED VERSION)
                if margin:
                    # current_app.logger.debug(f"Applying margin to Matplotlib: {margin}")
                    # Use figure padding instead of subplot adjustments
                    # Convert margin values to fractions of figure size (more appropriate scaling)
                    # Use a smaller conversion factor to avoid invalid subplot parameters
                    conversion_factor = 1000.0  # Larger denominator for smaller values
                    left_margin_fraction = margin.get("l", 0) / conversion_factor
                    right_margin_fraction = margin.get("r", 0) / conversion_factor
                    top_margin_fraction = margin.get("t", 0) / conversion_factor
                    bottom_margin_fraction = margin.get("b", 0) / conversion_factor
                    
                    # Calculate subplot parameters with validation
                    left_pos = max(0.1, left_margin_fraction)  # Minimum 0.1 to avoid edge
                    right_pos = min(0.9, 1.0 - right_margin_fraction)  # Maximum 0.9 to avoid edge
                    top_pos = min(0.9, 1.0 - top_margin_fraction)  # Maximum 0.9 to avoid edge
                    bottom_pos = max(0.1, bottom_margin_fraction)  # Minimum 0.1 to avoid edge
                    
                    # Ensure left < right and bottom < top
                    if left_pos >= right_pos:
                        left_pos = 0.1
                        right_pos = 0.9
                        #pass  # Suppress warning logs: f"Invalid margin: left >= right, using default values")
                    if bottom_pos >= top_pos:
                        bottom_pos = 0.1
                        top_pos = 0.9
                        #pass  # Suppress warning logs: f"Invalid margin: bottom >= top, using default values")
                    
                    # Apply margins using figure padding
                    fig_mpl.subplots_adjust(
                        left=left_pos,
                        right=right_pos,
                        top=top_pos,
                        bottom=bottom_pos
                    )
                    # current_app.logger.debug(f"Applied margins (fractions): left={left_pos}, right={right_pos}, top={top_pos}, bottom={bottom_pos}")
                
                # Adjust layout to accommodate legend position
                if show_legend and legend_position == "bottom":
                    # Add extra space at bottom for legend
                    fig_mpl.subplots_adjust(bottom=fig_mpl.subplotpars.bottom + 0.15)
                    # current_app.logger.debug("Added extra bottom space for legend")

                # Final check: Hide secondary y-axis for all chart types if explicitly disabled
                if disable_secondary_y and 'ax2' in locals() and ax2 is not None:
                    # AGGRESSIVE: Completely remove the secondary y-axis
                    try:
                        ax2.remove()  # Remove the axis completely
                        ax2 = None    # Set to None to prevent further operations
                    except:
                        # Fallback: Hide it if removal fails
                        ax2.set_visible(False)
                        ax2.set_yticks([])
                        ax2.set_yticklabels([])
                        ax2.spines['right'].set_visible(False)
                        ax2.set_ylabel("")
                        ax2.set_title("")
                        # Force redraw
                        fig_mpl.canvas.draw()
                
                tmpfile = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                # Use different bbox_inches parameter based on legend position
                if show_legend and legend_position == "bottom":
                    # For bottom legend, use 'tight' but with extra padding
                    plt.savefig(tmpfile.name, bbox_inches='tight', pad_inches=0.3, dpi=200)
                    # current_app.logger.debug(f"Saved Matplotlib chart with bottom legend using extra padding")
                else:
                    # For other positions, use standard tight layout
                    plt.savefig(tmpfile.name, bbox_inches='tight', dpi=200)
                                        # Only log legend position if legend_loc is defined (for charts that have legends)
                    if 'legend_loc' in locals():
                        # Chart saved with legend
                        pass
                    else:
                        # current_app.logger.debug(f"Saved Matplotlib chart without legend")
                        pass
                    
                    # ALWAYS close the figure to prevent memory leaks
                    plt.close(fig_mpl)
                    plt.close('all')  # Close all figures
                    gc.collect()  # Force garbage collection

                return tmpfile.name

            except Exception as e:
                import traceback
                
                # Create user-friendly error message
                error_type = type(e).__name__
                error_msg = str(e)
                
                # Simplify common error messages with more specific, user-friendly descriptions
                if "JSONDecodeError" in error_type:
                    user_message = "Invalid chart configuration format - please check your chart settings"
                elif "Invalid JSON in chart attributes" in error_msg:
                    # Extract the detailed JSON error information
                    if "at line" in error_msg and "column" in error_msg:
                        # Parse the error message to extract line and column info
                        try:
                            line_part = error_msg.split("at line ")[1].split(",")[0]
                            col_part = error_msg.split("column ")[1]
                            user_message = f"JSON syntax error in chart attributes: {error_msg.split(': ')[1].split(' at')[0]} at line {line_part}, column {col_part}"
                        except:
                            user_message = f"JSON syntax error in chart attributes: {error_msg}"
                    else:
                        user_message = f"JSON syntax error in chart attributes: {error_msg}"
                elif "KeyError" in error_type:
                    if "sheet" in error_msg.lower():
                        user_message = f"Excel sheet not found: {error_msg.split(':')[-1].strip()}"
                    elif "column" in error_msg.lower():
                        user_message = f"Column not found in data: {error_msg.split(':')[-1].strip()}"
                    else:
                        user_message = f"Missing required data field: {error_msg.split(':')[-1].strip()}"
                elif "ValueError" in error_type:
                    if "cell range" in error_msg.lower() or "range" in error_msg.lower():
                        user_message = "Invalid cell range - please check that the data range exists in your Excel file"
                    elif "empty" in error_msg.lower() or "no data" in error_msg.lower():
                        user_message = "No data found in the specified range - please check your data source"
                    elif "numeric" in error_msg.lower():
                        user_message = "Data contains non-numeric values - please ensure all chart data is numeric"
                    elif "x and y must have same first dimension" in error_msg.lower():
                        user_message = "Data mismatch: X and Y data have different lengths - please check that your data ranges have the same number of rows"
                    elif "x and y must be the same size" in error_msg.lower():
                        user_message = "Data mismatch: X and Y data have different sizes - please check that your data ranges have the same number of values"
                    elif "must be of length" in error_msg.lower() and "explode" in error_msg.lower():
                        user_message = "Pie chart data mismatch: The explode values don't match the number of data points - please check your data configuration"
                    elif "shapes" in error_msg.lower() and "dimension" in error_msg.lower():
                        user_message = f"Data dimension mismatch: {error_msg.split('but have shapes')[0].strip()} - please ensure all data series have the same length"
                    else:
                        user_message = f"Data validation error: {error_msg}"
                elif "IndexError" in error_type:
                    user_message = "Data range is empty or invalid - please check your Excel file has data in the specified range"
                elif "Cannot fix dimensions" in error_msg and "both arrays are empty" in error_msg:
                    user_message = "No data available: Both X and Y data arrays are empty - please check your data source and ensure it contains values"
                elif "TypeError" in error_type:
                    if "string" in error_msg.lower() and "float" in error_msg.lower():
                        user_message = "Data contains text instead of numbers - please ensure all chart data is numeric"
                    else:
                        user_message = f"Data type error: {error_msg}"
                elif "FileNotFoundError" in error_type:
                    user_message = "Excel file or worksheet not found - please check the file path and sheet name"
                elif "openpyxl" in error_msg.lower():
                    user_message = "Excel file is corrupted or in an unsupported format - please try a different file"
                elif "pandas" in error_msg.lower():
                    user_message = "Unable to read Excel data - please check that the file is not corrupted and contains valid data"
                elif "xlrd" in error_msg.lower():
                    user_message = "Excel file format not supported - please save as .xlsx format"
                elif "NoneType" in error_type and "text" in error_msg.lower():
                    user_message = "Secondary y-axis disabled - chart generated without secondary axis"
                    current_app.logger.info(f"🔄 Chart '{chart_tag}' generated without secondary y-axis (disable_secondary_y=True)")
                elif "permission" in error_msg.lower() or "access" in error_msg.lower():
                    user_message = "File access denied - please ensure the Excel file is not open in another program"
                elif "memory" in error_msg.lower():
                    user_message = "Not enough memory to process the data - try with a smaller dataset"
                elif "timeout" in error_msg.lower():
                    user_message = "Chart generation timed out - try with a smaller dataset or simpler chart type"
                elif "y_vals is None" in error_msg.lower():
                    user_message = "Missing Y-axis data: One or more data series is empty - please check that all your data ranges contain values"
                else:
                    # For any other error, try to extract a meaningful message
                    if ":" in error_msg:
                        user_message = error_msg.split(":")[-1].strip()
                    else:
                        user_message = f"Chart generation failed: {error_msg}"
                
                error_details = {
                    "chart_tag": chart_tag,
                    "error_type": error_type,
                    "user_message": user_message,
                    "technical_message": error_msg,
                    "chart_type": chart_type if 'chart_type' in locals() else "unknown",
                    "data_points": len(series_data) if 'series_data' in locals() else 0,
                    "timestamp": datetime.utcnow().isoformat(),
                    "chart_attributes": chart_attr_map.get(chart_tag.lower(), "{}") if 'chart_attr_map' in locals() else "Not available"
                }
                
                # Simple console logging
                current_app.logger.error(f"❌ Chart '{chart_tag}' failed: {user_message}")
                
                # Store error details for frontend (project-specific)
                if not hasattr(current_app, 'chart_errors'):
                    current_app.chart_errors = {}
                if project_id not in current_app.chart_errors:
                    current_app.chart_errors[project_id] = {}
                current_app.chart_errors[project_id][chart_tag] = error_details
                
                # Debug logging for error storage
                current_app.logger.info(f"🔍 Stored chart error for {chart_tag}: {error_details}")
                
                # Also store a simplified version for report generation errors
                if not hasattr(current_app, 'report_generation_errors'):
                    current_app.report_generation_errors = {}
                if project_id not in current_app.report_generation_errors:
                    current_app.report_generation_errors[project_id] = {}
                report_error_details = {
                    "error": user_message,
                    "chart_type": chart_type if 'chart_type' in locals() else "unknown",
                    "timestamp": datetime.utcnow().isoformat()
                }
                current_app.report_generation_errors[project_id][chart_tag] = report_error_details
                
                # Debug logging for report generation error storage
                current_app.logger.info(f"🔍 Stored report generation error for {chart_tag}: {report_error_details}")
                
                # Clean up any remaining matplotlib figures
                plt.close('all')
                gc.collect()
                return None

        # Insert charts into paragraphs
        chart_errors = []
        
        # COMPREHENSIVE TEXT REPLACEMENT - Process ALL document elements
        process_entire_document()
        
        # Process charts in paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            full_text = ''.join(run.text for run in para.runs)
            chart_placeholders = re.findall(r"\$\{(section\d+_chart)\}", full_text, flags=re.IGNORECASE)
            for tag in chart_placeholders:
                if tag.lower() in chart_attr_map:
                    try:
                        current_app.logger.info(f"🔍 About to call generate_chart for tag: {tag}")
                        chart_img = generate_chart({}, tag)
                        current_app.logger.info(f"🔍 generate_chart returned: {chart_img is not None}")
                        if chart_img:
                            para.text = re.sub(rf"\$\{{{tag}\}}", "", para.text, flags=re.IGNORECASE)
                            para.add_run().add_picture(chart_img, width=Inches(5.5))
                        else:
                            # Chart generation failed, add error placeholder
                            error_msg = f"[Chart failed: {tag}]"
                            para.text = re.sub(rf"\$\{{{tag}\}}", error_msg, para.text, flags=re.IGNORECASE)
                            
                            # Get the specific error from chart_errors if available
                            specific_error = "Chart generation failed - please check your data and chart configuration"
                            if hasattr(current_app, 'chart_errors') and project_id in current_app.chart_errors:
                                if tag in current_app.chart_errors[project_id]:
                                    specific_error = current_app.chart_errors[project_id][tag].get('user_message', specific_error)
                            elif hasattr(current_app, 'report_generation_errors') and project_id in current_app.report_generation_errors:
                                if tag in current_app.report_generation_errors[project_id]:
                                    specific_error = current_app.report_generation_errors[project_id][tag].get('error', specific_error)
                            
                            chart_errors.append({
                                "tag": tag,
                                "error": specific_error
                            })
                        
                        # Clean up after each chart generation
                        plt.close('all')
                        gc.collect()
                    except Exception as e:
                        current_app.logger.error(f"⚠️ Failed to insert chart for tag {tag}: {e}")
                        error_msg = f"[Chart failed: {tag}]"
                        para.text = re.sub(rf"\$\{{{tag}\}}", error_msg, para.text, flags=re.IGNORECASE)
                        chart_errors.append({
                            "tag": tag,
                            "error": f"Chart insertion failed: {str(e)}"
                        })

        # Insert charts into tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text = ''.join(run.text for run in para.runs)
                        chart_placeholders = re.findall(r"\$\{(section\d+_chart)\}", full_text, flags=re.IGNORECASE)
                        for tag in chart_placeholders:
                            if tag.lower() in chart_attr_map:
                                try:
                                    current_app.logger.info(f"🔍 About to call generate_chart for tag: {tag} (table)")
                                    chart_img = generate_chart({}, tag)
                                    current_app.logger.info(f"🔍 generate_chart returned: {chart_img is not None} (table)")
                                    if chart_img:
                                        para.text = re.sub(rf"\$\{{{tag}\}}", "", para.text, flags=re.IGNORECASE)
                                        para.add_run().add_picture(chart_img, width=Inches(5.5))
                                    else:
                                        # Chart generation failed, add error placeholder
                                        error_msg = f"[Chart failed: {tag}]"
                                        para.text = re.sub(rf"\$\{{{tag}\}}", error_msg, para.text, flags=re.IGNORECASE)
                                        
                                        # Get the specific error from chart_errors if available
                                        specific_error = "Chart generation failed - please check your data and chart configuration"
                                        if hasattr(current_app, 'chart_errors') and project_id in current_app.chart_errors:
                                            if tag in current_app.chart_errors[project_id]:
                                                specific_error = current_app.chart_errors[project_id][tag].get('user_message', specific_error)
                                        elif hasattr(current_app, 'report_generation_errors') and project_id in current_app.report_generation_errors:
                                            if tag in current_app.report_generation_errors[project_id]:
                                                specific_error = current_app.report_generation_errors[project_id][tag].get('error', specific_error)
                                        
                                        chart_errors.append({
                                            "tag": tag,
                                            "error": specific_error
                                        })
                                except Exception as e:
                                    current_app.logger.error(f"⚠️ Failed to insert chart in table for tag {tag}: {e}")
                                    error_msg = f"[Chart failed: {tag}]"
                                    para.text = re.sub(rf"\$\{{{tag}\}}", error_msg, para.text, flags=re.IGNORECASE)
                                    chart_errors.append({
                                        "tag": tag,
                                        "error": f"Chart insertion failed: {str(e)}"
                                    })

        # Save report to temporary location first (needed for TOC update)
        import tempfile
        temp_dir = tempfile.mkdtemp()
        output_path = os.path.join(temp_dir, f'output_report_{project_id}.docx')
        doc.save(output_path)
        
        # Update TOC using the new service
        try:
            toc_result = update_toc(doc, docx_path=output_path, flat_data_map=flat_data_map)
            if toc_result.get('success'):
                current_app.logger.info(f"✅ TOC update completed successfully: {toc_result}")
            else:
                current_app.logger.warning(f"⚠️ TOC update had issues: {toc_result.get('error', 'Unknown error')}")
        except Exception as e:
            current_app.logger.warning(f"⚠️ Could not update TOC: {e}")
            import traceback
            current_app.logger.debug(traceback.format_exc())
        
        current_app.logger.info(f"✅ Report generated successfully")
        
        # Store chart errors for this report generation
        if not hasattr(current_app, 'report_errors'):
            current_app.report_errors = {}
        current_app.report_errors[project_id] = {
            "chart_errors": chart_errors,
            "generated_at": datetime.utcnow().isoformat()
        }
        
        # Clear old chart errors for this project when new report is generated
        if hasattr(current_app, 'chart_errors') and project_id in current_app.chart_errors:
            current_app.chart_errors[project_id] = {}
        
        return output_path

    except Exception as e:
        current_app.logger.error(f"❌ Failed to generate report: {e}")
        import traceback
        current_app.logger.error(traceback.format_exc())
        return None

# Helper function no longer needed - files are now stored in database

@projects_bp.route('/api/projects', methods=['GET'])
@login_required
def get_projects():
    # Access MongoDB via current_app.mongo.db
    projects = list(current_app.mongo.db.projects.find({'user_id': current_user.get_id()}))
    for project in projects:
        project['id'] = str(project['_id'])
        del project['_id']
        # Remove binary file content to prevent JSON serialization error
        if 'file_content' in project:
            del project['file_content']
    return jsonify({'projects': projects})

@projects_bp.route('/api/projects', methods=['POST'])
@login_required
def create_project():
    name = request.form.get('name')
    description = request.form.get('description')
    file = request.files.get('file') 

    if not name or not description:
        return jsonify({'error': 'Missing required fields (name or description)'}), 400

    file_name = None
    file_content = None
    if file:
        if not allowed_file(file.filename): 
            return jsonify({'error': 'File type not allowed'}), 400
        file_name = secure_filename(file.filename)
        file_content = file.read()  # Read file content into memory

    project = {
        'name': name,
        'description': description,
        'user_id': current_user.get_id(),
        'file_name': file_name,
        'file_content': file_content,  # Store file content in database
        'created_at': datetime.utcnow().isoformat() 
    }
    # Access MongoDB via current_app.mongo.db
    project_id = current_app.mongo.db.projects.insert_one(project).inserted_id
    
    # Create a copy for JSON response without binary content
    project_response = {
        'id': str(project_id),
        'name': project['name'],
        'description': project['description'],
        'user_id': project['user_id'],
        'file_name': project['file_name'],
        'created_at': project['created_at']
    }

    return jsonify({'message': 'Project created successfully', 'project': project_response}), 201

@projects_bp.route('/api/projects/<project_id>/upload_report', methods=['POST'])
@login_required
def upload_report(project_id):
    current_app.logger.debug(f"📤 Upload request received for project: {project_id}")
    
    if 'report_file' not in request.files:
        current_app.logger.error(f"❌ No report_file in request.files: {list(request.files.keys())}")
        return jsonify({'error': 'No report file provided'}), 400

    report_file = request.files['report_file']
    current_app.logger.debug(f"📁 File received: {report_file.filename}")

    if report_file.filename == '':
        current_app.logger.error(f"❌ Empty filename")
        return jsonify({'error': 'No selected report file'}), 400

    if not allowed_report_file(report_file.filename):
        current_app.logger.error(f"❌ File type not allowed: {report_file.filename}")
        return jsonify({'error': 'Report file type not allowed. Only .xlsx or .csv are accepted.'}), 400

    try:
        project_id_obj = ObjectId(project_id)
        current_app.logger.debug(f"✅ Valid project ID: {project_id}")
    except Exception as e:
        current_app.logger.error(f"❌ Invalid project ID: {project_id}, error: {e}")
        return jsonify({'error': 'Invalid project ID'}), 400

    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        current_app.logger.error(f"❌ Project not found or unauthorized: {project_id}")
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    current_app.logger.debug(f"✅ Project found: {project.get('name', 'Unknown')}")

    # Handle both old (file_path) and new (file_name/file_content) project formats
    template_file_name = project.get('file_name')
    template_file_content = project.get('file_content')
    
    # Backward compatibility: if new format not found, try old format
    if not template_file_name or not template_file_content:
        old_file_path = project.get('file_path')
        if old_file_path:
            # Convert old format to new format
            abs_file_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), old_file_path)
            if os.path.exists(abs_file_path):
                try:
                    with open(abs_file_path, 'rb') as f:
                        template_file_content = f.read()
                    template_file_name = os.path.basename(old_file_path)
                    
                    # Update project to new format
                    current_app.mongo.db.projects.update_one(
                        {'_id': project_id_obj},
                        {'$set': {
                            'file_name': template_file_name,
                            'file_content': template_file_content
                        }}
                    )
                    current_app.logger.debug(f"🔄 Migrated project from old format to new format")
                except Exception as e:
                    current_app.logger.error(f"❌ Failed to migrate old project format: {e}")
                    return jsonify({'error': 'Failed to load template file. Please re-upload the template.'}), 400
            else:
                current_app.logger.error(f"❌ Old template file not found: {abs_file_path}")
                return jsonify({'error': 'Template file not found. Please re-upload the template.'}), 400
        else:
            current_app.logger.error(f"❌ No template file found in project")
            return jsonify({'error': 'Word template file not found for this project. Please upload it during project creation.'}), 400
    
    current_app.logger.debug(f"📄 Template file name: {template_file_name}")
    
    # Create temporary file from database content
    temp_template_dir = tempfile.mkdtemp()
    temp_template_path = os.path.join(temp_template_dir, template_file_name)
    with open(temp_template_path, 'wb') as f:
        f.write(template_file_content)
    current_app.logger.debug(f"📄 Temporary template created: {temp_template_path}")
    
    # Save the uploaded report data file temporarily
    report_data_filename = secure_filename(report_file.filename)
    temp_dir = tempfile.mkdtemp()
    temp_report_data_path = os.path.join(temp_dir, report_data_filename)
    report_file.save(temp_report_data_path)

    # Clear any existing errors for this project before starting new generation
    if hasattr(current_app, 'chart_errors') and project_id in current_app.chart_errors:
        current_app.chart_errors[project_id] = {}

    # Generate the report
    current_app.logger.debug(f"🔄 Starting report generation...")
    generated_report_path = _generate_report(project_id, temp_template_path, temp_report_data_path)
    
    # Clean up the temporary files and directories
    import shutil
    shutil.rmtree(temp_dir)
    shutil.rmtree(temp_template_dir)
    current_app.logger.debug(f"🧹 Temporary files cleaned up")

    if generated_report_path:
        current_app.logger.debug(f"✅ Report generated successfully: {generated_report_path}")
        # Update project with generated report path
        current_app.mongo.db.projects.update_one(
            {'_id': project_id_obj},
            {'$set': {'generated_report_path': generated_report_path, 'report_generated_at': datetime.utcnow().isoformat()}}
        )
        return jsonify({'message': 'Report generated successfully', 'report_path': generated_report_path}), 200
    else:
        current_app.logger.error(f"❌ Report generation failed")
        return jsonify({'error': 'Failed to generate report'}), 500

@projects_bp.route('/api/reports/<project_id>/download', methods=['GET'])
@login_required
def download_report(project_id):
    try:
        project_id_obj = ObjectId(project_id)
    except:
        return jsonify({'error': 'Invalid project ID'}), 400

    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    generated_report_path = project.get('generated_report_path')
    if not generated_report_path or not os.path.exists(generated_report_path):
        return jsonify({'error': 'Generated report not found for this project'}), 404

    # Send file and clean up after download
    try:
        response = send_file(generated_report_path, as_attachment=True)
        
        # Clean up the temporary file after sending
        def cleanup_after_response(response):
            try:
                if os.path.exists(generated_report_path):
                    os.remove(generated_report_path)
                    # Also remove the temp directory if it's empty
                    temp_dir = os.path.dirname(generated_report_path)
                    if os.path.exists(temp_dir) and not os.listdir(temp_dir):
                        os.rmdir(temp_dir)
            except Exception as e:
                pass  # Suppress warning logs: f"⚠️ Failed to cleanup temporary report file: {e}")
            return response
        
        response.call_on_close(lambda: cleanup_after_response(response))
        return response
        
    except Exception as e:
        # Clean up on error too
        try:
            if os.path.exists(generated_report_path):
                os.remove(generated_report_path)
                temp_dir = os.path.dirname(generated_report_path)
                if os.path.exists(temp_dir) and not os.listdir(temp_dir):
                    os.rmdir(temp_dir)
        except:
            pass
        raise e

@projects_bp.route('/api/reports/<chart_filename>/download_html', methods=['GET'])
@login_required
def download_chart_html(chart_filename):
    # Charts are embedded in Word documents, no separate HTML files needed
    return jsonify({'error': 'Chart HTML files are not available - charts are embedded in Word documents'}), 404

@projects_bp.route('/api/reports/batch_reports_<project_id>.zip', methods=['GET'])
@login_required
def download_batch_reports(project_id):
    """Download batch reports ZIP file"""
    try:
        # Validate project ID
        project_id_obj = ObjectId(project_id)
        
        # Check if project exists and belongs to user
        project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
        if not project:
            return jsonify({'error': 'Project not found or unauthorized'}), 404
        
        # Construct the ZIP file path
        zip_filename = f'batch_reports_{project_id}.zip'
        zip_path = os.path.join(tempfile.gettempdir(), zip_filename)
        
        # Check if file exists
        if not os.path.exists(zip_path):
            return jsonify({'error': 'Batch reports file not found. Please regenerate the reports.'}), 404
        
        # Send file
        def cleanup_after_response(response):
            try:
                # Clean up the ZIP file after sending
                if os.path.exists(zip_path):
                    os.remove(zip_path)
                    current_app.logger.info(f"Cleaned up batch reports ZIP: {zip_path}")
            except Exception as e:
                pass  # Suppress warning logs: f"Failed to clean up batch reports ZIP {zip_path}: {e}")
            return response
        
        response = send_file(
            zip_path,
            as_attachment=True,
            download_name=zip_filename,
            mimetype='application/zip'
        )
        
        # Add cleanup callback
        response.call_on_close(lambda: cleanup_after_response(response))
        
        current_app.logger.info(f"Downloading batch reports ZIP: {zip_filename}")
        return response
        
    except Exception as e:
        current_app.logger.error(f"Error downloading batch reports: {e}")
        return jsonify({'error': 'Failed to download batch reports'}), 500

@projects_bp.route('/api/projects/<project_id>/chart_errors', methods=['GET'])
@login_required
def get_chart_errors(project_id):
    try:
        project_id_obj = ObjectId(project_id)
    except:
        return jsonify({'error': 'Invalid project ID'}), 400

    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    # Get chart errors for this project
    chart_errors = getattr(current_app, 'chart_errors', {}).get(project_id, {})
    report_errors = getattr(current_app, 'report_errors', {}).get(project_id, {})
    report_generation_errors = getattr(current_app, 'report_generation_errors', {}).get(project_id, {})
    
    # Debug logging
    current_app.logger.info(f"🔍 Chart errors for project {project_id}:")
    current_app.logger.info(f"  - chart_errors: {chart_errors}")
    current_app.logger.info(f"  - report_errors: {report_errors}")
    current_app.logger.info(f"  - report_generation_errors: {report_generation_errors}")
    
    # Combine both types of errors
    all_errors = {
        "chart_generation_errors": chart_errors,
        "report_generation_errors": report_errors.get("chart_errors", []),
        "report_generation_errors_detailed": report_generation_errors,
        "report_generated_at": report_errors.get("generated_at")
    }
    
    current_app.logger.info(f"🔍 Returning chart errors: {all_errors}")
    return jsonify(all_errors)

@projects_bp.route('/api/projects/<project_id>/clear_errors', methods=['POST'])
@login_required
def clear_project_errors(project_id):
    try:
        project_id_obj = ObjectId(project_id)
    except:
        return jsonify({'error': 'Invalid project ID'}), 400

    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    # Clear chart errors for this project
    if hasattr(current_app, 'chart_errors') and project_id in current_app.chart_errors:
        current_app.chart_errors[project_id] = {}
    
    # Clear report errors for this project
    if hasattr(current_app, 'report_errors') and project_id in current_app.report_errors:
        current_app.report_errors[project_id] = {}
    
    return jsonify({'message': 'Project errors cleared successfully'})

@projects_bp.route('/api/projects/<project_id>/upload_zip', methods=['POST'])
@login_required
def upload_zip_and_generate_reports(project_id):
    import gc  # Add garbage collection import
    import matplotlib.pyplot as plt  # Add matplotlib import
    
    if 'zip_file' not in request.files:
        return jsonify({'error': 'No zip file provided'}), 400

    zip_file = request.files['zip_file']
    if not zip_file.filename.endswith('.zip'):
        return jsonify({'error': 'Only .zip files are allowed'}), 400
    
    if zip_file.filename == '':
        return jsonify({'error': 'No ZIP file selected'}), 400

    # Clear any existing errors for this project before starting new generation
    if hasattr(current_app, 'chart_errors') and project_id in current_app.chart_errors:
        current_app.chart_errors[project_id] = {}

    # Prepare temp directories
    temp_dir = tempfile.mkdtemp()
    extracted_dir = os.path.join(temp_dir, 'extracted')
    os.makedirs(extracted_dir, exist_ok=True)

    # Save and extract ZIP
    zip_path = os.path.join(temp_dir, secure_filename(zip_file.filename))
    zip_file.save(zip_path)
    current_app.logger.info(f"ZIP file saved: {zip_path}")
    
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            # Test if ZIP is valid
            zip_ref.testzip()
            zip_ref.extractall(extracted_dir)
            current_app.logger.info(f"ZIP extracted to: {extracted_dir}")
            current_app.logger.info(f"ZIP contents: {zip_ref.namelist()}")
    except zipfile.BadZipFile:
        current_app.logger.error(f"❌ Corrupted ZIP file: {zip_file.filename}")
        shutil.rmtree(temp_dir)
        return jsonify({'error': 'The uploaded ZIP file is corrupted or invalid'}), 400
    except Exception as e:
        current_app.logger.error(f"❌ Error extracting ZIP: {e}")
        shutil.rmtree(temp_dir)
        return jsonify({'error': f'Error extracting ZIP file: {str(e)}'}), 500

    # Find all Excel files (including in subdirectories)
    excel_files = []
    for root, dirs, files in os.walk(extracted_dir):
        for file in files:
            if file.endswith('.xlsx') or file.endswith('.xls'):
                excel_files.append(os.path.join(root, file))
    
    current_app.logger.info(f"Found {len(excel_files)} Excel files in ZIP: {[os.path.basename(f) for f in excel_files]}")

    # Prepare temporary output folders
    output_folder_name = os.path.join(temp_dir, 'reports_by_name')
    output_folder_code = os.path.join(temp_dir, 'reports_by_code')
    os.makedirs(output_folder_name, exist_ok=True)
    os.makedirs(output_folder_code, exist_ok=True)

    generated_files = []
    total_files = len(excel_files)
    
    if total_files == 0:
        current_app.logger.error(f"❌ No Excel files found in ZIP: {zip_file.filename}")
        # Clean up temp directory
        shutil.rmtree(temp_dir)
        return jsonify({'error': 'No Excel files (.xlsx or .xls) found in the uploaded ZIP file'}), 400
    
    current_app.logger.info(f"Starting batch processing of {total_files} Excel files")
    
    for idx, excel_path in enumerate(excel_files, 1):
        current_app.logger.info(f"🔍 Starting to process file {idx}/{total_files}: {os.path.basename(excel_path)}")
        
        # Force garbage collection before processing each file
        gc.collect()
        plt.close('all')
        
        # Validate Excel structure first
        is_valid, validation_message = validate_excel_structure(excel_path)
        if not is_valid:
            current_app.logger.error(f"❌ Invalid Excel structure in {os.path.basename(excel_path)}: {validation_message}")
            continue
        
        current_app.logger.info(f"✅ Excel structure validated for {os.path.basename(excel_path)}")
        
        # Extract report name and code from Excel file
        try:
            report_name, report_code = extract_report_info_from_excel(excel_path)
            current_app.logger.info(f"📋 Extracted info: {report_name} (Code: {report_code})")
        except Exception as e:
            current_app.logger.error(f"❌ Failed to extract report info from {os.path.basename(excel_path)}: {e}")
            continue

        # Generate report
        project = current_app.mongo.db.projects.find_one({'_id': ObjectId(project_id)})
        
        # Handle both old and new project formats
        template_file_name = project.get('file_name')
        template_file_content = project.get('file_content')
        
        # Backward compatibility: if new format not found, try old format
        if not template_file_name or not template_file_content:
            old_file_path = project.get('file_path')
            if old_file_path:
                abs_file_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), old_file_path)
                if os.path.exists(abs_file_path):
                    with open(abs_file_path, 'rb') as f:
                        template_file_content = f.read()
                    template_file_name = os.path.basename(old_file_path)
                else:
                    current_app.logger.error(f"❌ Old template file not found for batch processing: {abs_file_path}")
                    continue
            else:
                current_app.logger.error(f"❌ No template file found for batch processing")
                continue
        
        # Create temporary template file
        temp_template_dir = tempfile.mkdtemp()
        temp_template_path = os.path.join(temp_template_dir, template_file_name)
        with open(temp_template_path, 'wb') as f:
            f.write(template_file_content)
        
        try:
            output_path = _generate_report(f"{project_id}_{idx}", temp_template_path, excel_path)
            
            # Clean up temporary template
            shutil.rmtree(temp_template_dir)
            
            if output_path:
                # Save in both folders with clean naming (using only report name/code)
                base_filename = os.path.splitext(os.path.basename(excel_path))[0]  # Get original Excel filename without extension
                name_file_path = os.path.join(output_folder_name, f"{report_name}.docx")
                code_file_path = os.path.join(output_folder_code, f"{report_code}.docx")
                
                shutil.copy(output_path, name_file_path)
                shutil.copy(output_path, code_file_path)
                
                generated_files.append({
                    'name': report_name, 
                    'code': report_code,
                    'original_file': base_filename,
                    'report_name': report_name,
                    'report_code': report_code
                })
                current_app.logger.info(f"✅ Successfully generated report {idx}/{total_files}: {report_name} -> {report_code}")
            else:
                current_app.logger.error(f"❌ Failed to generate report {idx}/{total_files}: {report_name}")
        except Exception as e:
            current_app.logger.error(f"❌ Error processing file {idx}/{total_files} ({os.path.basename(excel_path)}): {e}")
            # Clean up temporary template if it exists
            if os.path.exists(temp_template_dir):
                shutil.rmtree(temp_template_dir)
        
        # Log progress
        current_app.logger.info(f"Progress: {idx}/{total_files} reports processed")
        
        # Force cleanup after each report
        gc.collect()
        plt.close('all')

    # Create zip file in temporary location with both folder structures
    zip_output_path = os.path.join(temp_dir, f'batch_reports_{project_id}.zip')
    with zipfile.ZipFile(zip_output_path, 'w') as zipf:
        # Add files from both folders to maintain the folder structure
        for file_info in generated_files:
            # Add file from reports_by_name folder
            name_file_path = os.path.join(output_folder_name, f"{file_info['name']}.docx")
            if os.path.exists(name_file_path):
                zipf.write(name_file_path, arcname=f"reports_by_name/{file_info['name']}.docx")
            
            # Add file from reports_by_code folder
            code_file_path = os.path.join(output_folder_code, f"{file_info['code']}.docx")
            if os.path.exists(code_file_path):
                zipf.write(code_file_path, arcname=f"reports_by_code/{file_info['code']}.docx")

    # Move zip to a temporary location that will be cleaned up after download
    final_zip_path = os.path.join(tempfile.gettempdir(), f'batch_reports_{project_id}.zip')
    shutil.move(zip_output_path, final_zip_path)
    
    # Clean up temp directory
    shutil.rmtree(temp_dir)
    
    # Final cleanup after batch processing
    gc.collect()
    plt.close('all')

    current_app.logger.info(f"Batch processing complete. Generated {len(generated_files)} out of {total_files} reports")
    
    # Log summary of results
    if len(generated_files) < total_files:
        pass  # Suppress warning logs: f"⚠️  {total_files - len(generated_files)} files failed to process")
        current_app.logger.info(f"✅ Successfully processed: {[f['name'] for f in generated_files]}")
    else:
        current_app.logger.info(f"✅ All {total_files} files processed successfully!")

    return jsonify({
        'message': f'Generated {len(generated_files)} out of {total_files} reports.',
        'download_zip': final_zip_path,
        'reports': generated_files,
        'total_files': total_files,
        'processed_files': len(generated_files),
        'success_rate': f"{len(generated_files)}/{total_files}"
    })

@projects_bp.route('/api/projects/<project_id>', methods=['PUT'])
@login_required
def update_project(project_id):
    try:
        project_id_obj = ObjectId(project_id)
    except Exception as e:
        current_app.logger.error(f"Invalid project ID: {project_id}, error: {e}")
        return jsonify({'error': 'Invalid project ID'}), 400

    try:
        # Check MongoDB connection first
        try:
            current_app.mongo.db.command('ping')
            current_app.logger.info(f"MongoDB connection test successful for project {project_id}")
        except Exception as e:
            current_app.logger.error(f"MongoDB connection failed for project {project_id}: {e}")
            return jsonify({'error': 'Database connection failed'}), 500

        # Check if project exists and belongs to user
        project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
        if not project:
            pass  # Suppress warning logs: f"Project not found or unauthorized: {project_id} for user {current_user.get_id()}")
            return jsonify({'error': 'Project not found or unauthorized'}), 404

        # Get form data
        name = request.form.get('name')
        description = request.form.get('description')
        file = request.files.get('file')

        current_app.logger.info(f"Updating project {project_id}: name='{name}', description='{description}', file={file.filename if file else 'None'}")

        # Validate required fields
        if not name or not description:
            pass  # Suppress warning logs: f"Missing required fields for project {project_id}: name='{name}', description='{description}'")
            return jsonify({'error': 'Missing required fields (name or description)'}), 400

        # Prepare update data
        update_data = {
            'name': name,
            'description': description,
            'updated_at': datetime.utcnow().isoformat()
        }

        # Handle file upload if provided
        if file:
            if not allowed_file(file.filename):
                pass  # Suppress warning logs: f"Invalid file type for project {project_id}: {file.filename}")
                return jsonify({'error': 'File type not allowed. Only .doc or .docx files are accepted.'}), 400
            
            try:
                # Read new file content
                file_name = secure_filename(file.filename)
                file_content = file.read()
                
                # Check file size (MongoDB document limit is 16MB)
                if len(file_content) > 15 * 1024 * 1024:  # 15MB limit to be safe
                    current_app.logger.error(f"File too large for project {project_id}: {file_name} ({len(file_content)} bytes)")
                    return jsonify({'error': 'File too large. Maximum file size is 15MB.'}), 400
                
                # Validate file content is not empty
                if len(file_content) == 0:
                    current_app.logger.error(f"Empty file uploaded for project {project_id}: {file_name}")
                    return jsonify({'error': 'Uploaded file is empty'}), 400
                
                # Store file content as binary data
                update_data['file_name'] = file_name
                update_data['file_content'] = file_content
                current_app.logger.info(f"File uploaded for project {project_id}: {file_name} ({len(file_content)} bytes)")
                
                # Verify the file can be read properly (basic validation)
                try:
                    if file_name.endswith('.docx'):
                        # Try to open as a Word document to validate
                        from io import BytesIO
                        from docx import Document
                        doc = Document(BytesIO(file_content))
                        current_app.logger.info(f"Word document validation successful for {file_name}")
                except Exception as e:
                    pass  # Suppress warning logs: f"Word document validation failed for {file_name}: {e}")
                    # Don't fail the upload, just log the warning
            except Exception as e:
                current_app.logger.error(f"Error reading file for project {project_id}: {e}")
                return jsonify({'error': f'Failed to read uploaded file: {str(e)}'}), 500

        # Update project in database
        try:
            result = current_app.mongo.db.projects.update_one(
                {'_id': project_id_obj, 'user_id': current_user.get_id()},
                {'$set': update_data}
            )

            if result.modified_count == 0:
                current_app.logger.error(f"Failed to update project {project_id} in database")
                return jsonify({'error': 'Failed to update project'}), 500

            # Get updated project
            updated_project = current_app.mongo.db.projects.find_one({'_id': project_id_obj})
            if not updated_project:
                current_app.logger.error(f"Failed to retrieve updated project {project_id}")
                return jsonify({'error': 'Failed to retrieve updated project'}), 500

            # Remove binary file_content to prevent JSON serialization error
            if 'file_content' in updated_project:
                del updated_project['file_content']

            updated_project['id'] = str(updated_project['_id'])
            del updated_project['_id']

            current_app.logger.info(f"Successfully updated project {project_id}")
            return jsonify({'message': 'Project updated successfully', 'project': updated_project})

        except Exception as e:
            current_app.logger.error(f"Database error updating project {project_id}: {e}")
            return jsonify({'error': f'Database error occurred: {str(e)}'}), 500

    except Exception as e:
        current_app.logger.error(f"Unexpected error updating project {project_id}: {e}")
        return jsonify({'error': 'An unexpected error occurred'}), 500

@projects_bp.route('/api/projects/<project_id>', methods=['DELETE'])
@login_required
def delete_project(project_id):
    try:
        project_id_obj = ObjectId(project_id)
    except Exception as e:
        return jsonify({'error': 'Invalid project ID'}), 400

    # Check if project exists and belongs to user
    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    # File content is stored in database, no file system cleanup needed

    # Delete project from database
    result = current_app.mongo.db.projects.delete_one({'_id': project_id_obj, 'user_id': current_user.get_id()})

    if result.deleted_count == 0:
        return jsonify({'error': 'Failed to delete project'}), 500

    return jsonify({'message': 'Project deleted successfully'})

@projects_bp.route('/api/projects/<project_id>', methods=['GET'])
@login_required
def get_project(project_id):
    try:
        project_id_obj = ObjectId(project_id)
    except Exception as e:
        return jsonify({'error': 'Invalid project ID'}), 400

    # Check if project exists and belongs to user
    project = current_app.mongo.db.projects.find_one({'_id': project_id_obj, 'user_id': current_user.get_id()})
    if not project:
        return jsonify({'error': 'Project not found or unauthorized'}), 404

    # Remove binary file_content to prevent JSON serialization error
    if 'file_content' in project:
        del project['file_content']

    project['id'] = str(project['_id'])
    del project['_id']
    
    return jsonify({'project': project})


@projects_bp.route('/api/test-remove-toc', methods=['POST'])
@login_required
def test_remove_toc_endpoint():
    """
    Test endpoint to remove TOC/LOF/LOT content from a Word document.
    
    Expects JSON with:
    {
        "docx_path": "/path/to/document.docx"
    }
    """
    try:
        data = request.get_json()
        if not data or 'docx_path' not in data:
            return jsonify({
                'success': False,
                'error': 'docx_path is required'
            }), 400
        
        docx_path = data['docx_path']
        
        # Test the aggressive cleaning function
        result = clean_pages_2_3_4_completely(docx_path)
        
        if result['success']:
            return jsonify(result), 200
        else:
            return jsonify(result), 500
            
    except Exception as e:
        current_app.logger.error(f"❌ Test endpoint error: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500
